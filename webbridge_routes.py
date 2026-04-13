# webbridge_routes.py — Second-half routes for webbridge.py.
# Contains: auth, user, suggest, job runner, porting, criteria and report endpoints.
# This module is imported at the bottom of webbridge.py after all shared state is defined.
# Circular import is safe because webbridge is already in sys.modules by the time this runs.

import os
import sys
import re
import json
import threading
import time
import uuid
import io
import hashlib
import logging
import heapq
import difflib
import secrets
from csv import DictWriter
from datetime import datetime
from functools import wraps
import requests
from flask import request, send_from_directory, jsonify, abort, Response, stream_with_context
from werkzeug.middleware.dispatcher import DispatcherMiddleware
from werkzeug.security import generate_password_hash, check_password_hash

# ---------------------------------------------------------------------------
# __main__ / module-name fix — same pattern as webbridge_cv.py
if 'webbridge' not in sys.modules:
    _main = sys.modules.get('__main__')
    if _main is not None and os.path.basename(os.path.normpath(getattr(_main, '__file__', ''))) == 'webbridge.py':
        sys.modules['webbridge'] = _main
# ---------------------------------------------------------------------------

from webbridge import (
    app, logger, genai,
    BASE_DIR, OUTPUT_DIR, SEARCH_XLS_DIR, REPORT_TEMPLATES_DIR,
    BUCKET_COMPANIES, BUCKET_JOB_TITLES,
    SECTORS_INDEX,
    CSE_PAGE_SIZE, CSE_PAGE_DELAY,
    GOOGLE_CSE_API_KEY, GOOGLE_CSE_CX, SEARCH_RESULTS_TARGET,
    GEMINI_API_KEY, GEMINI_SUGGEST_MODEL,
    SINGAPORE_CONTEXT, SEARCH_RULES,
    CV_TRANSLATION_MAX_CHARS, LANG_DETECTION_SAMPLE_LENGTH, CV_ANALYSIS_MAX_CHARS,
    MAX_COMMENT_LENGTH, COMMENT_TRUNCATE_LENGTH,
    ASSESSMENT_EXCELLENT_THRESHOLD, ASSESSMENT_GOOD_THRESHOLD, ASSESSMENT_MODERATE_THRESHOLD,
    CITY_TO_COUNTRY_DATA,
    _CV_ANALYZE_SEMAPHORE, _SINGLE_FILE_MAX,
    _rate, _check_user_rate, _csrf_required, _require_admin,
    _is_pdf_bytes,
    _extract_json_object, _extract_confirmed_skills,
    translate_text_pipeline,
    _infer_region_from_country,
    _find_best_sector_match_for_text, _map_keyword_to_sector_label,
    _compute_search_target,
    _should_overwrite_existing, _ensure_rating_metadata_columns, _ensure_search_indexes,
    _persist_jskillset, _fetch_jskillset, _fetch_jskillset_from_process,
    _sync_login_jskillset_to_process, _sync_criteria_jskillset_to_process,
    _increment_cse_query_count, _increment_gemini_query_count, _load_rate_limits,
    _make_flask_limit,
    _pg_connect, _ensure_admin_columns,
    dedupe,
    _normalize_seniority_single, _map_gemini_seniority_to_dropdown,
    _gemini_talent_pool_suggestion,
    _token_set, _build_sectors_token_index, _is_pharma_company, _sectors_allow_pharma,
    _nllb_available, nllb_translate,
    log_identity, log_infrastructure, log_financial, log_security, log_error, log_approval,
    read_all_logs,
    _APP_LOGGER_AVAILABLE,
    _load_search_provider_config,
    _load_llm_provider_config,
)

@app.post("/login")
@_rate(_make_flask_limit("login"))
@_check_user_rate("login")
@_csrf_required
def login_account():
    data = request.get_json(force=True, silent=True) or {}
    username = (data.get("username") or "").strip()
    password = (data.get("password") or "").strip()
    _ip = request.headers.get("X-Forwarded-For", request.remote_addr or "")
    if not (username and password):
        return jsonify({"error":"username and password required"}), 400

    try:
        import common_auth
        hash_password_fn = getattr(common_auth, "hash_password", None)
        verify_password_fn = getattr(common_auth, "verify_password", None)
    except Exception:
        hash_password_fn = None
        verify_password_fn = None

    try:
        import psycopg2
        pg_host=os.getenv("PGHOST","localhost")
        pg_port=int(os.getenv("PGPORT","5432"))
        pg_user=os.getenv("PGUSER","postgres")
        pg_password=os.getenv("PGPASSWORD", "")
        pg_db=os.getenv("PGDATABASE","candidate_db")
        conn=psycopg2.connect(host=pg_host, port=pg_port, user=pg_user, password=pg_password, dbname=pg_db)
        cur=conn.cursor()
        cur.execute("SELECT password, userid, cemail, fullname, role_tag, COALESCE(token,0) FROM login WHERE username=%s", (username,))
        row=cur.fetchone()
        cur.close(); conn.close()
        if not row:
            log_security("login_failed", username=username, ip_address=_ip,
                         detail="User not found", severity="warning")
            return jsonify({"error":"Invalid credentials"}), 401
        stored_pw, userid, cemail, fullname, role_tag, token_val = row
        stored_pw = stored_pw or ""

        if verify_password_fn:
            ok = False
            try:
                ok = bool(verify_password_fn(stored_pw, password))
            except Exception:
                ok = False
            if not ok:
                log_security("login_failed", username=username, ip_address=_ip,
                             detail="Password mismatch", severity="warning")
                return jsonify({"error":"Invalid credentials"}), 401
        else:
            def _local_hash_password(p: str) -> str:
                import hashlib
                salt = os.getenv("PASSWORD_SALT", "")
                return hashlib.sha256((salt + p).encode("utf-8")).hexdigest()
            hashed = hash_password_fn(password) if hash_password_fn else _local_hash_password(password)
            if stored_pw != hashed and stored_pw != password:
                log_security("login_failed", username=username, ip_address=_ip,
                             detail="Password mismatch", severity="warning")
                return jsonify({"error":"Invalid credentials"}), 401

        log_identity(userid=str(userid or ""), username=username,
                     ip_address=_ip, mfa_status="N/A")
        resp = jsonify({"ok": True, "userid": userid or "", "username": username, "cemail": cemail or "", "fullname": fullname or "", "role_tag": role_tag or "", "token": int(token_val or 0)})
        # httponly=False: AutoSourcing.html (and other pages) read the username
        # cookie via document.cookie to identify the logged-in user.  This
        # matches the behaviour of chatbot_api.py which also sets httponly=False.
        _cookie_opts = dict(max_age=2592000, path="/", httponly=False, samesite="lax",
                            secure=os.getenv("FORCE_HTTPS", "0") == "1")
        resp.set_cookie("username", username, **_cookie_opts)
        resp.set_cookie("userid", str(userid or ""), **_cookie_opts)
        return resp, 200
    except Exception as e:
        log_error(source="login", message=str(e), severity="error",
                  username=username, endpoint="/login")
        return jsonify({"error": str(e)}), 500

@app.post("/logout")
@_csrf_required
def logout_account():
    resp = jsonify({"ok": True})
    resp.delete_cookie("username", path="/")
    resp.delete_cookie("userid", path="/")
    return resp

@app.post("/register")
@_rate(_make_flask_limit("register"))
@_check_user_rate("register")
@_csrf_required
def register_account():
    data = request.get_json(force=True, silent=True) or {}

    fullname   = (data.get("fullname") or "").strip()
    corporation = (data.get("corporation") or "").strip()
    cemail     = (data.get("cemail") or "").strip()
    username   = (data.get("username") or "").strip()
    password   = data.get("password") or ""
    userid     = (data.get("userid") or "").strip()
    created_at = (data.get("created_at") or "").strip()

    if not (fullname and cemail and username and password):
        return jsonify({"error": "fullname, cemail, username, password are required"}), 400

    if not userid:
        userid = str(uuid.uuid4().int % 9000000 + 1000000)

    try:
        import common_auth
        hash_password_fn = getattr(common_auth, "hash_password", None)
    except Exception:
        hash_password_fn = None

    if hash_password_fn:
        try:
            hashed_pw = hash_password_fn(password)
        except Exception:
            hashed_pw = None
    else:
        hashed_pw = None

    if not hashed_pw:
        def _local_hash_password(p: str) -> str:
            import hashlib
            salt = os.getenv("PASSWORD_SALT", "")
            return hashlib.sha256((salt + p).encode("utf-8")).hexdigest()
        hashed_pw = _local_hash_password(password)

    try:
        import psycopg2
        from psycopg2 import sql as pgsql
        pg_host=os.getenv("PGHOST","localhost")
        pg_port=int(os.getenv("PGPORT","5432"))
        pg_user=os.getenv("PGUSER","postgres")
        pg_password=os.getenv("PGPASSWORD", "")
        pg_db=os.getenv("PGDATABASE","candidate_db")

        conn=psycopg2.connect(host=pg_host, port=pg_port, user=pg_user, password=pg_password, dbname=pg_db)
        cur=conn.cursor()

        cur.execute("""
            SELECT column_name
            FROM information_schema.columns
            WHERE table_schema='public' AND table_name='login'
        """)
        login_cols = {r[0].lower() for r in cur.fetchall()}

        if "username" in login_cols and "cemail" in login_cols:
            cur.execute("SELECT 1 FROM login WHERE username=%s OR cemail=%s LIMIT 1", (username, cemail))
            if cur.fetchone():
                cur.close(); conn.close()
                return jsonify({"error": "Username or email already registered"}), 409
        elif "username" in login_cols:
            cur.execute("SELECT 1 FROM login WHERE username=%s LIMIT 1", (username,))
            if cur.fetchone():
                cur.close(); conn.close()
                return jsonify({"error": "Username already registered"}), 409
        elif "cemail" in login_cols:
            cur.execute("SELECT 1 FROM login WHERE cemail=%s LIMIT 1", (cemail,))
            if cur.fetchone():
                cur.close(); conn.close()
                return jsonify({"error": "Email already registered"}), 409

        insert_cols = []
        insert_vals = []

        for col, val in [
            ("userid", userid),
            ("username", username),
            ("password", hashed_pw),
            ("fullname", fullname),
            ("cemail", cemail)
        ]:
            if col in login_cols:
                insert_cols.append(col)
                insert_vals.append(val)

        if "corporation" in login_cols and corporation:
            insert_cols.append("corporation"); insert_vals.append(corporation)
        if "created_at" in login_cols and created_at:
            insert_cols.append("created_at"); insert_vals.append(created_at)
        if "role_tag" in login_cols:
            insert_cols.append("role_tag"); insert_vals.append("")
        elif "roletag" in login_cols:
            insert_cols.append("roletag"); insert_vals.append("")
        if "token" in login_cols:
            insert_cols.append("token"); insert_vals.append(0)

        if not insert_cols:
            cur.close(); conn.close()
            return jsonify({"error": "No compatible columns found for registration"}), 500

        col_sql = pgsql.SQL(", ").join(pgsql.Identifier(c) for c in insert_cols)
        placeholders = pgsql.SQL(", ".join(["%s"] * len(insert_cols)))
        cur.execute(pgsql.SQL("INSERT INTO login ({}) VALUES ({})").format(col_sql, placeholders), insert_vals)
        conn.commit()
        cur.close(); conn.close()

        return jsonify({"ok": True, "message": "Registration successful", "username": username, "userid": userid}), 200
    except Exception as e:
        logger.error(f"[Register] {e}")
        return jsonify({"error": str(e)}), 500


# ── Sales-rep self-registration ──────────────────────────────────────────────
# Stores the profile in a dedicated `employee` table (created on first use).

def _safe_cookie_value(s: str) -> str:
    """Strip characters that are illegal in HTTP Set-Cookie values to prevent header injection."""
    import re as _re
    return _re.sub(r'[\x00-\x1f\x7f;,\\ "\'=]', '', str(s or ""))[:256]


def _ensure_employee_table(conn):
    """Create the employee table if it does not already exist."""
    cur = conn.cursor()
    cur.execute("""
        CREATE TABLE IF NOT EXISTS employee (
            id                   SERIAL PRIMARY KEY,
            full_name            TEXT NOT NULL,
            username             TEXT UNIQUE NOT NULL,
            password             TEXT NOT NULL,
            nationality          TEXT,
            location             TEXT,
            skillsets            TEXT,
            industrial_vertical  TEXT,
            language_skills      TEXT,
            travel_availability  TEXT,
            commission           NUMERIC DEFAULT 0,
            ownership            INTEGER DEFAULT 0,
            created_at           TIMESTAMPTZ DEFAULT NOW()
        )
    """)
    # Idempotently add commission/ownership to tables created before this migration.
    for ddl in [
        "ALTER TABLE employee ADD COLUMN IF NOT EXISTS commission NUMERIC DEFAULT 0",
        "ALTER TABLE employee ADD COLUMN IF NOT EXISTS ownership INTEGER DEFAULT 0",
    ]:
        try:
            cur.execute("SAVEPOINT sp_emp_col")
            cur.execute(ddl)
            cur.execute("RELEASE SAVEPOINT sp_emp_col")
        except Exception:
            cur.execute("ROLLBACK TO SAVEPOINT sp_emp_col")
    conn.commit()
    cur.close()


@app.post("/employee/register")
@_csrf_required
def employee_register():
    data = request.get_json(force=True, silent=True) or {}

    full_name            = (data.get("full_name") or "").strip()
    username             = (data.get("username") or "").strip()
    password             = data.get("password") or ""
    nationality          = (data.get("nationality") or "").strip()
    location             = (data.get("location") or "").strip()
    skillsets            = (data.get("skillsets") or "").strip()
    industrial_vertical  = (data.get("industrial_vertical") or "").strip()
    language_skills      = (data.get("language_skills") or "").strip()
    travel_availability  = (data.get("travel_availability") or "").strip()

    if not (full_name and username and password and nationality and location
            and skillsets and industrial_vertical and language_skills and travel_availability):
        return jsonify({"error": "All fields are required."}), 400

    if len(password) < 8 or not (any(c.isalpha() for c in password) and any(c.isdigit() for c in password)):
        return jsonify({"error": "Password must be at least 8 characters and contain both letters and numbers."}), 400

    import re as _re
    if not _re.match(r'^[a-zA-Z0-9_\-\.]+$', username):
        return jsonify({"error": "Username may only contain letters, numbers, underscores, hyphens and dots."}), 400

    # Hash the password using the same mechanism as the main /register route
    try:
        import common_auth
        hash_password_fn = getattr(common_auth, "hash_password", None)
    except Exception:
        hash_password_fn = None

    if hash_password_fn:
        try:
            hashed_pw = hash_password_fn(password)
        except Exception:
            hashed_pw = None
    else:
        hashed_pw = None

    if not hashed_pw:
        hashed_pw = generate_password_hash(password)

    try:
        conn = _pg_connect()
        _ensure_employee_table(conn)
        cur = conn.cursor()
        cur.execute(
            """INSERT INTO employee
                   (full_name, username, password, nationality, location,
                    skillsets, industrial_vertical, language_skills, travel_availability)
               VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s)
               RETURNING username, full_name""",
            (
                full_name[:100], username[:50], hashed_pw,
                nationality[:80], location[:100], skillsets[:1000],
                industrial_vertical[:200], language_skills[:200], travel_availability[:100],
            )
        )
        db_row = cur.fetchone()
        conn.commit()
        cur.close()
        conn.close()
        # Use DB-returned values for cookies so the value originates from the database
        db_username  = db_row[0] if db_row else ""
        db_full_name = db_row[1] if db_row else ""
        _cookie_opts = dict(max_age=86400, path="/", httponly=False, samesite="lax",
                            secure=os.getenv("FORCE_HTTPS", "0") == "1")
        resp = jsonify({"ok": True, "message": "Sales rep registered successfully.", "username": db_username, "full_name": db_full_name})
        resp.set_cookie("emp_username", db_username, **_cookie_opts)
        resp.set_cookie("emp_full_name", db_full_name, **_cookie_opts)
        return resp, 201
    except Exception as e:
        if hasattr(e, 'pgcode') and e.pgcode == '23505':
            return jsonify({"error": "That username is already taken. Please choose another."}), 409
        logger.error(f"[employee_register] {e}")
        return jsonify({"error": "Registration failed due to an internal error. Please try again."}), 500


@app.post("/employee/login")
@_csrf_required
def employee_login():
    """Authenticate a sales rep against the employee table and set a session cookie."""
    data = request.get_json(force=True, silent=True) or {}
    username = (data.get("username") or "").strip()
    password = data.get("password") or ""

    if not (username and password):
        return jsonify({"error": "Username and password are required."}), 400

    try:
        conn = _pg_connect()
        _ensure_employee_table(conn)
        cur = conn.cursor()
        cur.execute(
            "SELECT full_name, password, username FROM employee WHERE username = %s LIMIT 1",
            (username,)
        )
        row = cur.fetchone()
        cur.close()
        conn.close()
    except Exception as e:
        logger.error(f"[employee_login] DB error: {e}")
        return jsonify({"error": "Login failed due to an internal error. Please try again."}), 500

    if not row:
        return jsonify({"error": "Invalid username or password."}), 401

    full_name, stored_hash, db_username = row
    ok = False
    try:
        ok = check_password_hash(stored_hash or "", password)
    except Exception:
        pass
    # Fallback: check common_auth.verify_password if available
    if not ok:
        try:
            import common_auth
            verify_fn = getattr(common_auth, "verify_password", None)
            if verify_fn:
                ok = bool(verify_fn(stored_hash or "", password))
        except Exception:
            pass

    if not ok:
        return jsonify({"error": "Invalid username or password."}), 401

    _cookie_opts = dict(max_age=86400, path="/", httponly=False, samesite="lax",
                        secure=os.getenv("FORCE_HTTPS", "0") == "1")
    # Use DB-sourced username and full_name for cookies (not the raw user-supplied input)
    resp = jsonify({"ok": True, "username": db_username or "", "full_name": full_name or ""})
    resp.set_cookie("emp_username", db_username or "", **_cookie_opts)
    resp.set_cookie("emp_full_name", full_name or "", **_cookie_opts)
    return resp, 200


@app.post("/employee/logout")
@_csrf_required
def employee_logout():
    resp = jsonify({"ok": True})
    resp.delete_cookie("emp_username", path="/")
    resp.delete_cookie("emp_full_name", path="/")
    return resp, 200


@app.get("/employee/check-client")
def employee_check_client():
    """Check whether a corporation name already exists in the login table.
    Returns exact match, fuzzy/partial match, or no match."""
    emp_username = (request.cookies.get("emp_username") or "").strip()
    if not emp_username:
        return jsonify({"error": "Unauthorized"}), 401
    corp = request.args.get("corporation", "").strip()
    if not corp:
        return jsonify({"error": "corporation is required"}), 400
    try:
        conn = _pg_connect()
        with conn.cursor() as cur:
            # 1. Exact match (case-insensitive)
            cur.execute(
                "SELECT corporation FROM login WHERE LOWER(corporation) = LOWER(%s)"
                " AND corporation IS NOT NULL LIMIT 1",
                (corp,))
            row = cur.fetchone()
            if row:
                return jsonify({"exists": True, "match": row[0]})
            # 2. Partial match — query is contained in a corporation name
            cur.execute(
                "SELECT corporation FROM login WHERE LOWER(corporation) LIKE LOWER(%s)"
                " AND corporation IS NOT NULL ORDER BY corporation LIMIT 1",
                (f"%{corp}%",))
            row = cur.fetchone()
            if row:
                return jsonify({"exists": False, "fuzzy": True, "match": row[0]})
            # 3. Corporation name is contained in the query, or word-level match
            words = [w for w in corp.split() if len(w) >= 3]
            if words:
                conditions = " OR ".join(
                    "LOWER(corporation) LIKE LOWER(%s)" for _ in words)
                params = [f"%{w}%" for w in words]
                cur.execute(
                    f"SELECT corporation FROM login WHERE ({conditions})"
                    " AND corporation IS NOT NULL ORDER BY corporation LIMIT 1",
                    params)
                row = cur.fetchone()
                if row:
                    return jsonify({"exists": False, "fuzzy": True, "match": row[0]})
        return jsonify({"exists": False, "fuzzy": False})
    except Exception as e:
        logger.error(f"[employee_check_client] DB error: {e}")
        return jsonify({"error": "Server error"}), 500


@app.get("/employee/dashboard-data")
def employee_dashboard_data():
    """
    Return the sales rep's dashboard payload:
      - employee profile (full_name, username, …)
      - client list: corporations from login table where bd = emp_username
      - transaction logs: financial credit entries for each of those corporations' users
    """
    emp_username = (request.cookies.get("emp_username") or "").strip()
    if not emp_username:
        return jsonify({"error": "Unauthorized"}), 401

    try:
        conn = _pg_connect()
        _ensure_employee_table(conn)
        cur = conn.cursor()

        # 1. Employee profile
        cur.execute(
            """SELECT full_name, username, nationality, location, skillsets,
                      industrial_vertical, language_skills, travel_availability,
                      COALESCE(commission, 0), COALESCE(ownership, 0)
               FROM employee WHERE username = %s LIMIT 1""",
            (emp_username,)
        )
        emp_row = cur.fetchone()
        if not emp_row:
            cur.close(); conn.close()
            return jsonify({"error": "Employee not found."}), 404

        employee = {
            "full_name":           emp_row[0] or "",
            "username":            emp_row[1] or "",
            "nationality":         emp_row[2] or "",
            "location":            emp_row[3] or "",
            "skillsets":           emp_row[4] or "",
            "industrial_vertical": emp_row[5] or "",
            "language_skills":     emp_row[6] or "",
            "travel_availability": emp_row[7] or "",
            "commission":          float(emp_row[8]),
            "ownership":           int(emp_row[9]),
        }

        # 2. Clients: rows in login table where bd = emp_username
        # All usernames assigned to this BD (used for transaction log filtering,
        # regardless of whether they have a corporation set)
        cur.execute(
            """SELECT DISTINCT username FROM login
               WHERE bd = %s AND username IS NOT NULL AND username != ''""",
            (emp_username,)
        )
        all_client_usernames: set = {r[0] for r in cur.fetchall()}

        # Clients with a named corporation (used for the Clients display table)
        cur.execute(
            """SELECT DISTINCT corporation, username
               FROM login
               WHERE bd = %s AND corporation IS NOT NULL AND corporation != ''
               ORDER BY corporation""",
            (emp_username,)
        )
        client_rows = cur.fetchall()
        cur.close()
        conn.close()
    except Exception as e:
        logger.error(f"[employee_dashboard_data] DB error: {e}")
        return jsonify({"error": "Failed to load dashboard data."}), 500

    # Build lookups for the log joining step
    # Map username -> corporation (for labelling log entries)
    username_to_corp: dict = {r[1]: r[0] for r in client_rows if r[1]}
    # Map corporation -> last_credited_date
    corp_credit_map: dict = {r[0]: {"last_credited_date": None} for r in client_rows if r[0]}

    # 3. Financial transaction logs — read the same files as /admin/logs
    transactions: list = []
    try:
        fin_logs = read_all_logs().get("financial", [])

        for entry in fin_logs:
            uname = entry.get("username") or ""
            if uname not in all_client_usernames:
                continue
            corp = username_to_corp.get(uname, "")
            ts = entry.get("timestamp") or ""
            txn_type = (entry.get("transaction_type") or "").lower()

            # Update last_credited_date for the corporation
            if txn_type == "credit" and corp in corp_credit_map:
                existing = corp_credit_map[corp]["last_credited_date"]
                if not existing or ts > existing:
                    corp_credit_map[corp]["last_credited_date"] = ts

            transactions.append({
                "timestamp":          ts,
                "username":           uname,
                "userid":             entry.get("userid") or "",
                "corporation":        corp,
                "transaction_type":   entry.get("transaction_type") or "",
                "transaction_amount": entry.get("transaction_amount"),
                "token_before":       entry.get("token_before"),
                "token_after":        entry.get("token_after"),
                "token_cost_sgd":     entry.get("token_cost_sgd"),
                "revenue_sgd":        entry.get("revenue_sgd"),
                "credits_spent":      entry.get("credits_spent"),
                "token_usage":        entry.get("token_usage"),
                "feature":            entry.get("feature") or "",
            })

        # Sort transactions newest-first
        transactions.sort(key=lambda x: x["timestamp"] or "", reverse=True)
    except Exception as log_err:
        logger.warning(f"[employee_dashboard_data] log read error: {log_err}")

    # 4. Assemble client list with last_credited_date
    clients = [
        {
            "corporation":       r[0],
            "username":          r[1],
            "last_credited_date": corp_credit_map.get(r[0], {}).get("last_credited_date"),
        }
        for r in client_rows if r[0]
    ]

    return jsonify({
        "employee":     employee,
        "clients":      clients,
        "transactions": transactions,
    }), 200


@app.get("/sales_rep_dashboard.html")
def sales_rep_dashboard_html():
    return send_from_directory(BASE_DIR, "sales_rep_dashboard.html")


@app.get("/admin/sales-rep")
@_require_admin
def admin_sales_rep():
    """
    Return one aggregated row per sales rep (bd) with:
      full_name (from employee table), username (bd value),
      total_clients (distinct corporations), tokens_credited (sum of credit txns),
      total_revenue (SGD from spend txns).
    """
    try:
        conn = _pg_connect()
        cur = conn.cursor()
        _ensure_admin_columns(cur)
        conn.commit()
        _ensure_employee_table(conn)

        # All login rows that have a bd value set.
        cur.execute(
            """SELECT username, COALESCE(corporation, '') AS corporation,
                      COALESCE(bd, '') AS bd
               FROM login
               WHERE bd IS NOT NULL AND bd != ''"""
        )
        login_rows = cur.fetchall()

        # Employee full_name, commission, ownership by username.
        cur.execute("SELECT username, full_name, COALESCE(commission,0), COALESCE(ownership,0) FROM employee")
        emp_info = {r[0]: {"full_name": r[1] or r[0], "commission": float(r[2]), "ownership": int(r[3])}
                    for r in cur.fetchall()}

        cur.close()
        conn.close()
    except Exception as e:
        logger.error(f"[admin_sales_rep] DB error: {e}")
        return jsonify({"error": "Failed to query database."}), 500

    # Build per-bd structures.
    bd_corp_sets: dict = {}       # bd -> set of corporation names
    username_to_info: dict = {}   # login username -> { bd, corporation }
    for username, corporation, bd in login_rows:
        if not bd:
            continue
        bd_corp_sets.setdefault(bd, set())
        if corporation:
            bd_corp_sets[bd].add(corporation)
        username_to_info[username] = {"bd": bd, "corporation": corporation}

    bd_acc: dict = {bd: {"tokens_credited": 0, "total_revenue": 0.0, "total_tokens_consumed": 0}
                    for bd in bd_corp_sets}

    # Aggregate from financial logs.
    try:
        fin_logs = read_all_logs().get("financial", [])
        for entry in fin_logs:
            uname = entry.get("username") or ""
            info = username_to_info.get(uname)
            if not info:
                continue
            bd = info["bd"]
            if bd not in bd_acc:
                continue
            txn_type = (entry.get("transaction_type") or "").lower()
            amt = float(entry.get("transaction_amount") or 0)
            if txn_type == "credit":
                bd_acc[bd]["tokens_credited"] += amt
            elif txn_type == "spend":
                bd_acc[bd]["total_tokens_consumed"] += abs(amt)
                rev = float(entry.get("revenue_sgd") or 0)
                if rev > 0:
                    bd_acc[bd]["total_revenue"] += rev
                else:
                    cost = float(entry.get("token_cost_sgd") or 0.10)
                    bd_acc[bd]["total_revenue"] += abs(amt) * cost
    except Exception as log_err:
        logger.warning(f"[admin_sales_rep] log read error: {log_err}")

    result = [
        {
            "full_name":              emp_info.get(bd, {}).get("full_name", bd),
            "username":               bd,
            "total_clients":          len(bd_corp_sets[bd]),
            "tokens_credited":        round(bd_acc[bd]["tokens_credited"]),
            "total_tokens_consumed":  round(bd_acc[bd]["total_tokens_consumed"]),
            "total_revenue":          round(bd_acc[bd]["total_revenue"], 2),
            "commission":             emp_info.get(bd, {}).get("commission", 0),
            "ownership":              emp_info.get(bd, {}).get("ownership", 0),
        }
        for bd in bd_corp_sets
    ]
    result.sort(key=lambda x: (x["full_name"] or "").lower())

    return jsonify({"sales_rep": result}), 200


@app.patch("/admin/sales-rep/<username>")
@_require_admin
def admin_sales_rep_update(username):
    """Update commission rate and ownership period for a sales rep in the employee table."""
    body = request.get_json(silent=True) or {}
    commission = body.get("commission")
    ownership  = body.get("ownership")
    if commission is None and ownership is None:
        return jsonify({"error": "No fields to update."}), 400
    try:
        commission = float(commission) if commission is not None else None
        ownership  = int(ownership)   if ownership  is not None else None
    except (TypeError, ValueError):
        return jsonify({"error": "Invalid commission or ownership value."}), 400
    try:
        conn = _pg_connect()
        cur  = conn.cursor()
        _ensure_employee_table(conn)
        if commission is not None and ownership is not None:
            cur.execute(
                "UPDATE employee SET commission=%s, ownership=%s WHERE username=%s",
                (commission, ownership, username)
            )
        elif commission is not None:
            cur.execute(
                "UPDATE employee SET commission=%s WHERE username=%s",
                (commission, username)
            )
        else:
            cur.execute(
                "UPDATE employee SET ownership=%s WHERE username=%s",
                (ownership, username)
            )
        conn.commit()
        cur.close()
        conn.close()
    except Exception as e:
        logger.error(f"[admin_sales_rep_update] DB error: {e}")
        return jsonify({"error": "Failed to update."}), 500
    return jsonify({"ok": True}), 200


@app.get("/admin/sales-rep/<username>/transactions")
@_require_admin
def admin_sales_rep_transactions(username):
    """Return all financial transaction log entries for a given BD (sales rep) username.
    Accepts optional `from` and `to` query-string params (YYYY-MM-DD) for date filtering.
    """
    date_from = (request.args.get("from") or "").strip()
    date_to   = (request.args.get("to")   or "").strip()

    try:
        conn = _pg_connect()
        cur  = conn.cursor()
        # All login-table usernames assigned to this BD
        cur.execute(
            """SELECT DISTINCT username FROM login
               WHERE bd = %s AND username IS NOT NULL AND username != ''""",
            (username,)
        )
        all_client_usernames: set = {r[0] for r in cur.fetchall()}
        # Map username -> corporation for log-entry labelling
        cur.execute(
            """SELECT DISTINCT username, COALESCE(corporation, '') FROM login
               WHERE bd = %s AND username IS NOT NULL AND username != ''""",
            (username,)
        )
        username_to_corp: dict = {r[0]: r[1] for r in cur.fetchall()}
        cur.close()
        conn.close()
    except Exception as e:
        logger.error(f"[admin_sales_rep_transactions] DB error: {e}")
        return jsonify({"error": "Failed to query database."}), 500

    transactions: list = []
    try:
        fin_logs = read_all_logs().get("financial", [])
        for entry in fin_logs:
            uname = entry.get("username") or ""
            if uname not in all_client_usernames:
                continue
            ts = entry.get("timestamp") or ""
            # Date filter
            if date_from and ts[:10] < date_from:
                continue
            if date_to and ts[:10] > date_to:
                continue
            transactions.append({
                "timestamp":          ts,
                "username":           uname,
                "userid":             entry.get("userid") or "",
                "corporation":        username_to_corp.get(uname, "") or entry.get("corporation") or "",
                "transaction_type":   entry.get("transaction_type") or "",
                "transaction_amount": entry.get("transaction_amount"),
                "token_before":       entry.get("token_before"),
                "token_after":        entry.get("token_after"),
                "token_cost_sgd":     entry.get("token_cost_sgd"),
                "revenue_sgd":        entry.get("revenue_sgd"),
                "credits_spent":      entry.get("credits_spent"),
                "token_usage":        entry.get("token_usage"),
                "feature":            entry.get("feature") or "",
            })
        transactions.sort(key=lambda x: x["timestamp"] or "", reverse=True)
    except Exception as log_err:
        logger.warning(f"[admin_sales_rep_transactions] log read error: {log_err}")

    return jsonify({"transactions": transactions}), 200


@app.get("/token-config")
def token_config():
    """Return the token credit/deduction configuration from rate_limits.json.
    Used by AutoSourcing.html and SourcingVerify.html (Flask-served pages) to
    read dynamic token rates without reaching across to the Node.js server.
    Requires a valid username cookie so the endpoint is not publicly enumerable.
    """
    username = (request.cookies.get("username") or "").strip()
    if not username:
        return jsonify({"error": "Unauthorized"}), 401
    try:
        cfg = _load_rate_limits()
        t = cfg.get("tokens", {})
        s = cfg.get("system", {})
        return jsonify({
            "appeal_approve_credit":     t.get("appeal_approve_credit",     1),
            "verified_selection_deduct": t.get("verified_selection_deduct", 2),
            "rebate_credit_per_profile": t.get("rebate_credit_per_profile", 1),
            "analytic_token_cost":       t.get("analytic_token_cost",       1),
            "initial_token_display":     t.get("initial_token_display",     5000),
            "sourcing_rate_base":        t.get("sourcing_rate_base",        1),
            "sourcing_rate_best_mode":   t.get("sourcing_rate_best_mode",   1.5),
            "sourcing_rate_over50":      t.get("sourcing_rate_over50",      2),
            "sourcing_rate_best_over50": t.get("sourcing_rate_best_over50", 2.5),
            "jd_upload_max_count":       s.get("jd_upload_max_count",       5),
            "jd_upload_max_bytes":       s.get("jd_upload_max_bytes",       6291456),
            "jd_analysis_token_cost":    t.get("jd_analysis_token_cost",    1),
            "token_cost_sgd":            t.get("token_cost_sgd",            0.10),
        }), 200
    except Exception as e:
        logger.error(f"[token-config] {e}")
        return jsonify({"error": "Failed to load token config"}), 500


@app.get("/user/resolve")
def user_resolve():
    username = (request.args.get("username") or "").strip()
    if not username:
        return jsonify({"error": "username required"}), 400
    try:
        import psycopg2
        pg_host=os.getenv("PGHOST","localhost")
        pg_port=int(os.getenv("PGPORT","5432"))
        pg_user=os.getenv("PGUSER","postgres")
        pg_password=os.getenv("PGPASSWORD", "")
        pg_db=os.getenv("PGDATABASE","candidate_db")
        conn=psycopg2.connect(host=pg_host, port=pg_port, user=pg_user, password=pg_password, dbname=pg_db)
        cur=conn.cursor()
        cur.execute("SELECT userid, fullname, role_tag, COALESCE(token,0), COALESCE(target_limit,10), COALESCE(useraccess,'') FROM login WHERE username=%s", (username,))
        row = cur.fetchone()
        if not row:
            cur.close(); conn.close()
            return jsonify({"error":"not found"}), 404
        userid, fullname, login_role_tag, token_val, target_limit_val, useraccess_val = row
        # Use login.role_tag as the authoritative current session role for the recruiter.
        # sourcing.role_tag is per-candidate (for matching) and must not override the recruiter's
        # current active role — old sourcing records from previous searches would cause the session
        # badge to show a stale role even after the recruiter has started a new search.
        resolved_role_tag = login_role_tag or ""
        cur.close(); conn.close()
        return jsonify({"userid": userid or "", "fullname": fullname or "", "role_tag": resolved_role_tag, "token": int(token_val or 0), "target_limit": int(target_limit_val or 10), "useraccess": (useraccess_val or "").strip()}), 200
    except Exception as e:
        return jsonify({"error": str(e)}), 500


# Module-level flag: ALTER TABLE to add last_result_count runs at most once per
# server process so the idempotency guard column is created without per-request DDL.
_token_guard_column_ensured = False

# Module-level flag: ALTER TABLE to add role_tag_session column runs at most once
# per server process for both login and sourcing tables.
_role_tag_session_column_ensured = False

@app.post("/user/token_update")
@_csrf_required
def user_token_update():
    """
    POST /user/token_update
    Sets the token column in the login table to the supplied value.
    Used to persist the current "tokens left" figure after each
    token-consuming operation so the login table always reflects the
    most up-to-date balance.

    Body JSON: { "userid": "<id>", "token": <number>, "result_count": <int|optional> }

    When result_count is supplied the endpoint acts as an idempotent guard:
    the update only fires if result_count differs from the stored
    last_result_count, preventing the feedback loop where the same search
    result count is deducted repeatedly on every page refresh or new-tab load.
    """
    global _token_guard_column_ensured, _role_tag_session_column_ensured
    data = request.get_json(force=True, silent=True) or {}
    userid = (data.get("userid") or "").strip()
    token_val = data.get("token")
    delta_val = data.get("delta")
    # Optional caller-supplied feature tag for granular transaction logging.
    # Falls back to "token_update" when not provided for backward compatibility.
    caller_feature = (data.get("feature") or "").strip() or "token_update"
    # Snapshot the current per-token SGD cost for audit logging.
    _rl_cfg = _load_rate_limits()
    _token_cost_sgd = float((_rl_cfg.get("tokens") or {}).get("token_cost_sgd", 0.10))

    # Delta mode: increment/decrement by a relative amount (used by rebate flow to restore +1 token).
    if delta_val is not None and token_val is None:
        try:
            delta_int = int(delta_val)
        except (TypeError, ValueError):
            return jsonify({"error": "delta must be a number"}), 400
        if not userid:
            return jsonify({"error": "userid is required"}), 400
        try:
            import psycopg2
            pg_host = os.getenv("PGHOST", "localhost")
            pg_port = int(os.getenv("PGPORT", "5432"))
            pg_user = os.getenv("PGUSER", "postgres")
            pg_password = os.getenv("PGPASSWORD", "")
            pg_db = os.getenv("PGDATABASE", "candidate_db")
            conn = psycopg2.connect(host=pg_host, port=pg_port, user=pg_user,
                                    password=pg_password, dbname=pg_db)
            cur = conn.cursor()
            # Skip token deduction (negative delta) for BYOK users
            if delta_int < 0:
                cur.execute(
                    "SELECT COALESCE(token, 0) AS t, useraccess FROM login WHERE userid = %s",
                    (userid,)
                )
                _byok_row = cur.fetchone()
                if _byok_row and (_byok_row[1] or "").strip().lower() == "byok":
                    cur.close()
                    conn.close()
                    return jsonify({"ok": True, "token": int(_byok_row[0])}), 200
            cur.execute(
                "UPDATE login SET token = COALESCE(token, 0) + %s WHERE userid = %s RETURNING token, username",
                (delta_int, userid)
            )
            row = cur.fetchone()
            conn.commit()
            cur.close(); conn.close()
            if not row:
                return jsonify({"error": "user not found"}), 404
            new_token, _uname = int(row[0]), (row[1] or "")
            if _APP_LOGGER_AVAILABLE:
                txn_type_d = "credit" if delta_int > 0 else "spend"
                log_financial(
                    username=_uname, userid=userid, feature=caller_feature,
                    transaction_type=txn_type_d,
                    token_before=new_token - delta_int, token_after=new_token,
                    transaction_amount=abs(delta_int),
                    token_usage=abs(delta_int) if txn_type_d == "spend" else 0,
                    token_cost_sgd=_token_cost_sgd,
                )
            return jsonify({"ok": True, "token": new_token}), 200
        except Exception as e:
            logger.error(f"[TokenUpdate/delta] {e}")
            return jsonify({"error": str(e)}), 500

    if not userid or token_val is None:
        return jsonify({"error": "userid and token are required"}), 400
    try:
        token_int = int(token_val)
    except (TypeError, ValueError):
        return jsonify({"error": "token must be a number"}), 400
    result_count_int = None
    rc = data.get("result_count")
    if rc is not None:
        try:
            result_count_int = int(rc)
        except (TypeError, ValueError):
            pass
    role_tag = (data.get("role_tag") or "").strip()
    _token_before_tx: int | None = None  # captured before the UPDATE for financial logging
    _username_tx: str = ""               # captured for financial logging
    try:
        import psycopg2
        pg_host = os.getenv("PGHOST", "localhost")
        pg_port = int(os.getenv("PGPORT", "5432"))
        pg_user = os.getenv("PGUSER", "postgres")
        pg_password = os.getenv("PGPASSWORD", "")
        pg_db = os.getenv("PGDATABASE", "candidate_db")
        conn = psycopg2.connect(
            host=pg_host, port=pg_port, user=pg_user,
            password=pg_password, dbname=pg_db
        )
        cur = conn.cursor()
        # Skip token deduction for BYOK users (absolute set path)
        cur.execute("SELECT COALESCE(token, 0) AS t, useraccess FROM login WHERE userid = %s", (userid,))
        _byok_check = cur.fetchone()
        if _byok_check and (_byok_check[1] or "").strip().lower() == "byok":
            cur.close()
            conn.close()
            return jsonify({"ok": True, "token": int(_byok_check[0]), "skipped": True}), 200
        try:
            if result_count_int is not None:
                # Ensure idempotency columns exist — run at most once per process
                if not _token_guard_column_ensured:
                    cur.execute(
                        "ALTER TABLE login ADD COLUMN IF NOT EXISTS last_result_count INTEGER"
                    )
                    cur.execute(
                        "ALTER TABLE login ADD COLUMN IF NOT EXISTS last_deducted_role_tag TEXT"
                    )
                    _token_guard_column_ensured = True
                # Ensure session tracking columns exist — run at most once per process
                if not _role_tag_session_column_ensured:
                    cur.execute("ALTER TABLE login ADD COLUMN IF NOT EXISTS session TIMESTAMPTZ")
                    cur.execute("ALTER TABLE sourcing ADD COLUMN IF NOT EXISTS session TIMESTAMPTZ")
                    cur.execute("ALTER TABLE sourcing ALTER COLUMN session DROP DEFAULT")
                    _role_tag_session_column_ensured = True
                # Read current token, stored result count, stored role_tag, login role_tag, session, and username.
                # role_tag and role_tag_session are read so that we can auto-generate the session
                # timestamp for rows where role_tag is already set but role_tag_session is NULL
                # (e.g. rows that pre-existed before the role_tag_session column was added).
                cur.execute(
                    "SELECT token, last_result_count, last_deducted_role_tag,"
                    " role_tag, session, username FROM login WHERE userid = %s",
                    (userid,)
                )
                existing = cur.fetchone()
                if not existing:
                    conn.commit()
                    return jsonify({"error": "user not found"}), 404
                current_token, stored_count, _stored_role_tag_raw, login_role_tag, login_session_ts, login_username = existing
                stored_role_tag = (_stored_role_tag_raw or "").strip()
                _token_before_tx = int(current_token) if current_token is not None else None
                _username_tx = login_username or ""
                # Auto-backfill: if role_tag is already set in login but role_tag_session is NULL,
                # generate a session timestamp now and transfer it to sourcing where role_tag matches.
                # This ensures every role_tag entry is tied to a valid session reference even for
                # rows that existed before the role_tag_session column was introduced.
                if (login_role_tag or "").strip() and login_session_ts is None:
                    cur.execute(
                        "UPDATE login SET session = NOW() WHERE userid = %s RETURNING session",
                        (userid,)
                    )
                    ts_row = cur.fetchone()
                    login_session_ts = ts_row[0] if ts_row else None
                    if login_session_ts is not None and login_username:
                        cur.execute(
                            "UPDATE sourcing SET session = %s WHERE username = %s AND role_tag = %s",
                            (login_session_ts, login_username, login_role_tag)
                        )
                        logger.info(
                            f"[TokenUpdate] Auto-backfilled role_tag_session='{login_session_ts}' "
                            f"for user='{login_username}' (role_tag='{login_role_tag}')"
                        )
                # Backend idempotency guard: skip if same result_count was already persisted.
                # When role_tag is also provided, require that the stored role_tag also matches;
                # a NULL/empty stored role_tag with a provided role_tag is treated as a new session.
                if stored_count is not None and stored_count == result_count_int:
                    if (not role_tag) or (stored_role_tag and stored_role_tag == role_tag):
                        conn.commit()
                        return jsonify({"ok": True, "token": int(current_token) if current_token is not None else 0, "skipped": True}), 200
                # New deduction — persist updated balance, result count, and role_tag
                if role_tag:
                    cur.execute(
                        "UPDATE login SET token = %s, last_result_count = %s, last_deducted_role_tag = %s WHERE userid = %s RETURNING token",
                        (token_int, result_count_int, role_tag, userid)
                    )
                else:
                    cur.execute(
                        "UPDATE login SET token = %s, last_result_count = %s WHERE userid = %s RETURNING token",
                        (token_int, result_count_int, userid)
                    )
            else:
                # Legacy path: no result_count supplied.
                # Uses a session+role_tag guard: if login.session == sourcing.session
                # AND role_tags match, the deduction for this session was already
                # processed — skip to prevent repeated deductions on page refresh.
                if not _role_tag_session_column_ensured:
                    cur.execute("ALTER TABLE login ADD COLUMN IF NOT EXISTS session TIMESTAMPTZ")
                    cur.execute("ALTER TABLE sourcing ADD COLUMN IF NOT EXISTS session TIMESTAMPTZ")
                    cur.execute("ALTER TABLE sourcing ALTER COLUMN session DROP DEFAULT")
                    _role_tag_session_column_ensured = True
                cur.execute(
                    "SELECT role_tag, session, username, token FROM login WHERE userid = %s",
                    (userid,)
                )
                _legacy_row = cur.fetchone()
                if _legacy_row:
                    _legacy_role_tag, _legacy_session_ts, _legacy_username, _legacy_token = _legacy_row
                    _token_before_tx = int(_legacy_token) if _legacy_token is not None else None
                    _username_tx = _legacy_username or ""
                    # Auto-backfill: if role_tag is set but session is NULL, generate now
                    if (_legacy_role_tag or "").strip() and _legacy_session_ts is None:
                        cur.execute(
                            "UPDATE login SET session = NOW() WHERE userid = %s RETURNING session",
                            (userid,)
                        )
                        _legacy_ts_row = cur.fetchone()
                        _legacy_new_ts = _legacy_ts_row[0] if _legacy_ts_row else None
                        if _legacy_new_ts is not None:
                            _legacy_session_ts = _legacy_new_ts
                            if _legacy_username:
                                cur.execute(
                                    "UPDATE sourcing SET session = %s WHERE username = %s AND role_tag = %s",
                                    (_legacy_new_ts, _legacy_username, _legacy_role_tag)
                                )
                                logger.info(
                                    f"[TokenUpdate] Auto-backfilled role_tag_session='{_legacy_new_ts}' "
                                    f"for user='{_legacy_username}' (role_tag='{_legacy_role_tag}') via legacy path"
                                )
                    # Session+role_tag guard: skip deduction when both tables have
                    # the same session timestamp and role_tag (already processed).
                    if (_legacy_session_ts is not None and (_legacy_role_tag or "").strip()
                            and _legacy_username):
                        cur.execute(
                            "SELECT session FROM sourcing"
                            " WHERE username = %s AND role_tag = %s LIMIT 1",
                            (_legacy_username, _legacy_role_tag)
                        )
                        _src_row = cur.fetchone()
                        _src_session = _src_row[0] if _src_row else None
                        if _src_session is not None and _src_session == _legacy_session_ts:
                            conn.commit()
                            return jsonify({"ok": True,
                                            "token": int(_legacy_token) if _legacy_token is not None else 0,
                                            "skipped": True}), 200
                cur.execute(
                    "UPDATE login SET token = %s WHERE userid = %s RETURNING token",
                    (token_int, userid)
                )
            row = cur.fetchone()
            conn.commit()
        finally:
            cur.close()
            conn.close()
        if not row:
            return jsonify({"error": "user not found"}), 404
        new_token = int(row[0])
        # Log token spend/credit transaction
        if _APP_LOGGER_AVAILABLE:
            delta = (new_token - _token_before_tx) if _token_before_tx is not None else None
            if delta is None:
                txn_type = "adjustment"
            elif delta < 0:
                txn_type = "spend"
            elif delta > 0:
                txn_type = "credit"
            else:
                txn_type = "adjustment"
            log_financial(
                username=_username_tx,
                userid=userid,
                feature=caller_feature,
                transaction_type=txn_type,
                token_before=_token_before_tx,
                token_after=new_token,
                transaction_amount=abs(delta) if delta is not None else None,
                token_usage=abs(delta) if (delta is not None and txn_type == "spend") else 0,
                token_cost_sgd=_token_cost_sgd,
            )
        return jsonify({"ok": True, "token": new_token}), 200
    except Exception as e:
        logger.error(f"[TokenUpdate] {e}")
        return jsonify({"error": str(e)}), 500


# ==================== Fetch Skills Endpoint ====================

@app.route("/user/fetch_skills", methods=["GET"])
def user_fetch_skills():
    """
    GET /user/fetch_skills?username=<username>
    Returns the user's skill list from the login table (jskillset or skills column).
    Response: { "skills": ["Python", "C++", ...] }
    """
    username = (request.args.get("username") or "").strip()
    if not username:
        return jsonify({"error": "username required"}), 400
    try:
        skills = _fetch_jskillset(username)
        return jsonify({"skills": skills}), 200
    except Exception as e:
        logger.error(f"[fetch_skills] Error for user='{username}': {e}")
        return jsonify({"error": str(e)}), 500

# ==================== Role Tag Update Endpoint ====================

@app.route("/user/update_role_tag", methods=["POST", "GET"])
def user_update_role_tag():
    """
    POST/GET /user/update_role_tag
    Updates role_tag in both login and sourcing tables for the given username.
    The sourcing table is the authoritative source for role-based job title assessment.

    Session tracking:
    - A timestamp (role_tag_session) is generated and stored in login when role_tag is set.
    - The same timestamp is transferred to sourcing only after validating that the
      role_tag value matches in both tables, ensuring cross-table traceability.
    """
    global _role_tag_session_column_ensured
    if request.method == "POST":
        data = request.get_json(force=True, silent=True) or {}
        username = (data.get("username") or "").strip()
        role_tag = (data.get("role_tag") or "").strip()
    else:
        username = (request.args.get("username") or "").strip()
        role_tag = (request.args.get("role_tag") or "").strip()
    if not username:
        return jsonify({"error": "username required"}), 400
    conn = None
    cur = None
    try:
        import psycopg2
        pg_host = os.getenv("PGHOST", "localhost")
        pg_port = int(os.getenv("PGPORT", "5432"))
        pg_user = os.getenv("PGUSER", "postgres")
        pg_password = os.getenv("PGPASSWORD", "")
        pg_db = os.getenv("PGDATABASE", "candidate_db")
        conn = psycopg2.connect(
            host=pg_host, port=pg_port, user=pg_user,
            password=pg_password, dbname=pg_db
        )
        cur = conn.cursor()
        # Ensure role_tag_session column exists in login and sourcing (once per process).
        # NOTE: This flag mirrors the _token_guard_column_ensured pattern; it is intentionally
        # not protected by a lock for the same reason — IF NOT EXISTS makes the DDL idempotent,
        # so concurrent first-time executions are safe.
        if not _role_tag_session_column_ensured:
            cur.execute("ALTER TABLE login ADD COLUMN IF NOT EXISTS session TIMESTAMPTZ")
            cur.execute("ALTER TABLE sourcing ADD COLUMN IF NOT EXISTS session TIMESTAMPTZ")
            cur.execute("ALTER TABLE sourcing ALTER COLUMN session DROP DEFAULT")
            _role_tag_session_column_ensured = True
        # Step 1: Update login — set role_tag and generate session timestamp atomically
        cur.execute(
            "UPDATE login SET role_tag=%s, session=NOW() WHERE username=%s",
            (role_tag, username)
        )
        if cur.rowcount == 0:
            conn.rollback()
            return jsonify({"error": "User not found"}), 404
        # Step 2: Read back the persisted role_tag and session timestamp from login
        cur.execute(
            "SELECT role_tag, session FROM login WHERE username=%s",
            (username,)
        )
        login_row = cur.fetchone()
        login_role_tag = login_row[0] if login_row else None
        login_session_ts = login_row[1] if login_row else None
        # Step 3: Update sourcing role_tag for all records of this user
        cur.execute("ALTER TABLE sourcing ADD COLUMN IF NOT EXISTS role_tag TEXT DEFAULT ''")
        cur.execute("UPDATE sourcing SET role_tag=%s WHERE username=%s AND (role_tag IS NULL OR role_tag='')", (role_tag, username))
        # Step 4: Validate that role_tag matches in both login and sourcing, then transfer
        # the session timestamp from login to sourcing for consistency and traceability.
        if login_role_tag == role_tag and login_session_ts is not None:
            cur.execute(
                "UPDATE sourcing SET session=%s WHERE username=%s AND role_tag=%s",
                (login_session_ts, username, role_tag)
            )
        conn.commit()
        logger.info(
            f"[UpdateRoleTag] Set role_tag='{role_tag}' session_ts='{login_session_ts}' "
            f"for user='{username}' in login and sourcing tables"
        )
        # login_session_ts may be a datetime object or a plain string depending on
        # the column type and psycopg2 type-casting; handle both safely.
        if login_session_ts is not None:
            session_val = (login_session_ts.isoformat()
                           if hasattr(login_session_ts, 'isoformat')
                           else str(login_session_ts))
        else:
            session_val = None
        return jsonify({"ok": True, "username": username, "role_tag": role_tag,
                        "session": session_val}), 200
    except Exception as e:
        logger.exception(f"[UpdateRoleTag] Failed for user='{username}': {e}")
        if conn:
            try:
                conn.rollback()
            except Exception:
                pass
        return jsonify({"error": str(e)}), 500
    finally:
        if cur:
            try:
                cur.close()
            except Exception:
                pass
        if conn:
            try:
                conn.close()
            except Exception:
                pass

# ==================== VSkillset Integration Endpoints ====================

@app.get("/user/jskillset")
def get_user_jskillset():
    """
    GET /user/jskillset?username=<username>
    Returns the user's jskillset from the login table.
    Response: { "jskillset": ["Python", "C++", ...] }
    """
    username = (request.args.get("username") or "").strip()
    if not username:
        return jsonify({"error": "username required"}), 400
    
    try:
        skills = _fetch_jskillset(username)
        return jsonify({"jskillset": skills}), 200
    except Exception as e:
        logger.error(f"[get_user_jskillset] Error: {e}")
        return jsonify({"error": str(e)}), 500

@app.post("/vskillset/infer")
@_rate(_make_flask_limit("vskillset_infer"))
@_check_user_rate("vskillset_infer")
def vskillset_infer():
    """
    POST /vskillset/infer
    Body: { 
        linkedinurl: "<url>", 
        skills: ["Python", "C++", ...], 
        assessment_level: "L1"|"L2", 
        username: "<optional>" 
    }
    
    Uses Gemini to evaluate each skill based on experience/cv.
    Returns: { 
        results: [ 
            { skill: "Python", probability: 85, category: "High", reason: "..." },
            ...
        ], 
        persisted: true 
    }
    """
    data = request.get_json(force=True, silent=True) or {}
    linkedinurl = (data.get("linkedinurl") or "").strip()
    skills = data.get("skills", [])
    assessment_level = (data.get("assessment_level") or "L2").upper()
    username = (data.get("username") or "").strip()
    force_regen = bool(data.get("force", False))
    
    if not linkedinurl or not skills:
        return jsonify({"error": "linkedinurl and skills required"}), 400
    
    if not isinstance(skills, list) or len(skills) == 0:
        return jsonify({"error": "skills must be a non-empty array"}), 400
    
    try:
        import psycopg2
        from psycopg2 import sql
        pg_host = os.getenv("PGHOST", "localhost")
        pg_port = int(os.getenv("PGPORT", "5432"))
        pg_user = os.getenv("PGUSER", "postgres")
        pg_password = os.getenv("PGPASSWORD", "")
        pg_db = os.getenv("PGDATABASE", "candidate_db")
        
        conn = psycopg2.connect(host=pg_host, port=pg_port, user=pg_user, password=pg_password, dbname=pg_db)
        cur = conn.cursor()
        
        # Normalize linkedin URL
        normalized = linkedinurl.lower().strip().rstrip('/')
        if not normalized.startswith('http'):
            normalized = 'https://' + normalized
        
        # Idempotency guard: if vskillset already exists in DB, return it without
        # re-running Gemini. Pass force=true in the request body to override this.
        if not force_regen:
            try:
                cur.execute("""
                    SELECT vskillset FROM process
                    WHERE LOWER(TRIM(TRAILING '/' FROM linkedinurl)) = %s
                       OR normalized_linkedin = %s
                    LIMIT 1
                """, (normalized, normalized))
                vs_row = cur.fetchone()
                if vs_row and vs_row[0]:
                    existing_vs = vs_row[0]
                    if isinstance(existing_vs, str):
                        existing_vs = json.loads(existing_vs)
                    if isinstance(existing_vs, list) and len(existing_vs) > 0:
                        high_skills = [i["skill"] for i in existing_vs if isinstance(i, dict) and i.get("category") == "High"]
                        cur.close()
                        conn.close()
                        logger.info(f"[vskillset_infer] Returning existing vskillset ({len(existing_vs)} items) for {linkedinurl[:50]} — use force=true to regenerate")
                        return jsonify({
                            "results": existing_vs,
                            "persisted": True,
                            "skipped": True,
                            "high_skills": high_skills,
                            "confirmed_skills": [i["skill"] for i in existing_vs if isinstance(i, dict) and i.get("source") == "confirmed"],
                            "inferred_skills":  [i["skill"] for i in existing_vs if isinstance(i, dict) and i.get("source") == "inferred"],
                        }), 200
            except Exception as _e:
                logger.warning(f"[vskillset_infer] Idempotency check failed ({_e}); proceeding with generation")
        
        # Fetch experience and cv from process table
        experience_text = ""
        cv_text = ""
        
        # Try by normalized_linkedin first, then linkedinurl
        cur.execute("""
            SELECT experience, cv 
            FROM process 
            WHERE LOWER(TRIM(TRAILING '/' FROM linkedinurl)) = %s 
               OR normalized_linkedin = %s
            LIMIT 1
        """, (normalized, normalized))
        row = cur.fetchone()
        
        if row:
            experience_text = (row[0] or "").strip()
            cv_text = (row[1] or "").strip()
        
        # Use experience as primary, cv as fallback
        profile_context = experience_text if experience_text else cv_text
        
        if not profile_context:
            cur.close()
            conn.close()
            return jsonify({
                "error": "No experience or CV data found for this profile",
                "results": [],
                "persisted": False
            }), 404
        
        # STEP 1: Extractive pass - mark skills explicitly mentioned in experience text as confirmed/High
        explicitly_confirmed = _extract_confirmed_skills(profile_context, skills)
        confirmed_set = set(s.lower() for s in explicitly_confirmed)
        confirmed_results = [
            {
                "skill": skill,
                "probability": 100,
                "category": "High",
                "reason": "Explicitly mentioned in experience text",
                "source": "confirmed"
            }
            for skill in explicitly_confirmed
        ]
        logger.info(f"[vskillset_infer] Extractive pass: {len(confirmed_results)}/{len(skills)} skills confirmed from text")

        # STEP 2: Only send unconfirmed skills to Gemini for inference
        unconfirmed_skills = [s for s in skills if s.lower() not in confirmed_set]
        inferred_results = []

        if unconfirmed_skills:
            prompt = f"""SYSTEM:
You are an expert technical recruiter evaluating candidate skillsets based on their work experience.

TASK:
For each skill in the list below, evaluate the candidate's likely proficiency based on their experience.
These skills were NOT found explicitly in the experience text, so use contextual inference from
job titles, companies, products, sector, and experience patterns.
Assign a probability score (0-100) and categorize as Low (<40), Medium (40-74), or High (75-100).

CANDIDATE PROFILE:
{profile_context[:3000]}

SKILLS TO INFER (not found explicitly in experience text):
{json.dumps(unconfirmed_skills, ensure_ascii=False)}

OUTPUT FORMAT (JSON):
{{
  "evaluations": [
    {{
      "skill": "skill_name",
      "probability": 0-100,
      "category": "Low|Medium|High",
      "reason": "Brief explanation based on companies and roles"
    }}
  ]
}}

Return ONLY the JSON object, no other text."""

            raw_text = (unified_llm_call_text(prompt) or "").strip()
            _increment_gemini_query_count(username)

            parsed = _extract_json_object(raw_text)

            if not parsed or "evaluations" not in parsed:
                logger.warning(f"[vskillset_infer] Gemini returned invalid JSON: {raw_text[:200]}")
                # Fallback: create basic inferred results for unconfirmed skills
                for skill in unconfirmed_skills:
                    inferred_results.append({
                        "skill": skill,
                        "probability": 50,
                        "category": "Medium",
                        "reason": "Unable to parse Gemini response",
                        "source": "inferred"
                    })
            else:
                inferred_results = parsed["evaluations"]

            # Ensure all required fields are present and annotate source
            for item in inferred_results:
                if "probability" not in item:
                    item["probability"] = 50
                if "category" not in item:
                    prob = item.get("probability", 50)
                    if prob >= 75:
                        item["category"] = "High"
                    elif prob >= 40:
                        item["category"] = "Medium"
                    else:
                        item["category"] = "Low"
                if "reason" not in item:
                    item["reason"] = "No reasoning provided"
                item["source"] = "inferred"

        # STEP 3: Merge confirmed + inferred results
        results = confirmed_results + inferred_results
        
        # Persist to database
        # 1. Store full annotated results in vskillset column (JSON)
        # 2. Store only High skills in skillset column as comma-separated string
        
        vskillset_json = json.dumps(results, ensure_ascii=False)
        high_skills = [item["skill"] for item in results if item["category"] == "High"]
        # Ensure all skills are strings before joining
        skillset_str = ", ".join([str(s) for s in high_skills if s])
        
        # Check if vskillset column exists
        cur.execute("""
            SELECT column_name 
            FROM information_schema.columns
            WHERE table_schema='public' AND table_name='process' 
              AND column_name IN ('vskillset', 'skillset')
        """)
        available_cols = {r[0] for r in cur.fetchall()}
        
        # Update process table
        updates = []
        if 'vskillset' in available_cols:
            updates.append("vskillset = %s")
        
        if updates:
            update_sql = sql.SQL("UPDATE process SET {} WHERE LOWER(TRIM(TRAILING '/' FROM linkedinurl)) = %s").format(sql.SQL(", ".join(updates)))
            update_values = []
            if 'vskillset' in available_cols:
                update_values.append(vskillset_json)
            update_values.append(normalized)
            cur.execute(update_sql, tuple(update_values))
        
        # Skillset: merge new High skills into existing value (add only; never remove or replace)
        if 'skillset' in available_cols and high_skills:
            cur.execute(
                "SELECT skillset FROM process WHERE LOWER(TRIM(TRAILING '/' FROM linkedinurl)) = %s",
                (normalized,)
            )
            _sk_row = cur.fetchone()
            _existing_sk = (_sk_row[0] or "") if _sk_row else ""
            _existing_parts = [s.strip() for s in _existing_sk.split(",") if s.strip()]
            _existing_set = {s.lower() for s in _existing_parts}
            _new_high = [s for s in high_skills if s.lower() not in _existing_set]
            if _new_high:
                _merged_sk = ", ".join(_existing_parts + _new_high)
                cur.execute(
                    "UPDATE process SET skillset = %s"
                    " WHERE LOWER(TRIM(TRAILING '/' FROM linkedinurl)) = %s",
                    (_merged_sk, normalized)
                )
                logger.info(f"[vskillset_infer] Merged {len(_new_high)} new High skills into skillset for {linkedinurl[:50]}")
            else:
                logger.info(f"[vskillset_infer] No new High skills for {linkedinurl[:50]} — skillset unchanged")
        
        conn.commit()
        
        cur.close()
        conn.close()
        
        return jsonify({
            "results": results,
            "persisted": True,
            "confirmed_skills": [item["skill"] for item in results if item.get("source") == "confirmed"],
            "inferred_skills": [item["skill"] for item in results if item.get("source") == "inferred"],
            "high_skills": high_skills
        }), 200
        
    except Exception as e:
        logger.error(f"[vskillset_infer] Error: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({"error": str(e), "persisted": False}), 500

@app.get("/process/skillsets")
def get_process_skillsets():
    """
    GET /process/skillsets?linkedin=<linkedinurl>
    Returns the persisted skillset and vskillset for a candidate.
    Response: { 
        "skillset": ["Python", "C++", ...], 
        "vskillset": [ 
            { "skill": "Python", "probability": 85, "category": "High", "reason": "..." },
            ...
        ] 
    }
    """
    linkedinurl = (request.args.get("linkedin") or "").strip()
    if not linkedinurl:
        return jsonify({"error": "linkedin parameter required"}), 400
    
    try:
        import psycopg2
        from psycopg2 import sql as pgsql
        pg_host = os.getenv("PGHOST", "localhost")
        pg_port = int(os.getenv("PGPORT", "5432"))
        pg_user = os.getenv("PGUSER", "postgres")
        pg_password = os.getenv("PGPASSWORD", "")
        pg_db = os.getenv("PGDATABASE", "candidate_db")
        
        conn = psycopg2.connect(host=pg_host, port=pg_port, user=pg_user, password=pg_password, dbname=pg_db)
        cur = conn.cursor()
        
        # Normalize linkedin URL
        normalized = linkedinurl.lower().strip().rstrip('/')
        if not normalized.startswith('http'):
            normalized = 'https://' + normalized
        
        # Check which columns exist
        cur.execute("""
            SELECT column_name 
            FROM information_schema.columns
            WHERE table_schema='public' AND table_name='process' 
              AND column_name IN ('vskillset', 'skillset')
        """)
        available_cols = {r[0] for r in cur.fetchall()}
        
        # Build SELECT query based on available columns
        select_cols = []
        if 'vskillset' in available_cols:
            select_cols.append('vskillset')
        if 'skillset' in available_cols:
            select_cols.append('skillset')
        
        if not select_cols:
            cur.close()
            conn.close()
            return jsonify({"skillset": [], "vskillset": []}), 200
        
        query = pgsql.SQL("SELECT {} FROM process WHERE LOWER(TRIM(TRAILING '/' FROM linkedinurl)) = %s LIMIT 1").format(pgsql.SQL(", ").join(pgsql.Identifier(c) for c in select_cols))
        cur.execute(query, (normalized,))
        row = cur.fetchone()
        
        cur.close()
        conn.close()
        
        if not row:
            return jsonify({"skillset": [], "vskillset": []}), 200
        
        result = {}
        col_idx = 0
        
        if 'vskillset' in available_cols:
            vskillset_raw = row[col_idx]
            col_idx += 1
            if vskillset_raw:
                if isinstance(vskillset_raw, str):
                    try:
                        result["vskillset"] = json.loads(vskillset_raw)
                    except (json.JSONDecodeError, ValueError):
                        result["vskillset"] = []
                elif isinstance(vskillset_raw, list):
                    result["vskillset"] = vskillset_raw
                else:
                    result["vskillset"] = []
            else:
                result["vskillset"] = []
        else:
            result["vskillset"] = []
        
        if 'skillset' in available_cols:
            skillset_raw = row[col_idx]
            if skillset_raw:
                if isinstance(skillset_raw, str):
                    try:
                        parsed = json.loads(skillset_raw)
                        if isinstance(parsed, list):
                            result["skillset"] = parsed
                        else:
                            result["skillset"] = [s.strip() for s in skillset_raw.split(',') if s.strip()]
                    except (json.JSONDecodeError, ValueError):
                        result["skillset"] = [s.strip() for s in skillset_raw.split(',') if s.strip()]
                elif isinstance(skillset_raw, list):
                    result["skillset"] = skillset_raw
                else:
                    result["skillset"] = []
            else:
                result["skillset"] = []
        else:
            result["skillset"] = []
        
        return jsonify(result), 200
        
    except Exception as e:
        logger.error(f"[get_process_skillsets] Error: {e}")
        return jsonify({"error": str(e)}), 500

# ==================== End VSkillset Integration ====================

# ... rest of file unchanged beyond this point ...

# Suggestion code: caching, supplemental lists, enforcement with sector-aware filtering
SUGGEST_CACHE = {}
SUGGEST_CACHE_LOCK = threading.Lock()
MAX_SUGGESTIONS_PER_TAG = int(os.getenv("MAX_SUGGESTIONS_PER_TAG", "6"))
COMPANY_SUGGESTIONS_LIMIT = int(os.getenv("COMPANY_SUGGESTIONS_LIMIT", "30"))

def _clean_list(items, limit=20):
    out=[]; seen=set()
    for x in items or []:
        if not isinstance(x,str): continue
        t=re.sub(r'\s+',' ',x).strip()
        if not t: continue
        k=t.lower()
        if k in seen: continue
        t=re.sub(r'[;,/]+$','',t)
        seen.add(k); out.append(t)
        if len(out)>=limit: break
    return out

# Regex to match trailing corporate entity type suffixes (legal entity words, not brand words).
_CORP_SUFFIX_RE = re.compile(
    r'(\s*,?\s*'
    r'(?:Co\.\s*,?\s*Ltd\.?|Co\.?,?|Ltd\.?|Inc\.?|K\.K\.?|Corp\.?|Corporation|GmbH'
    r'|S\.A\.?|N\.V\.?|B\.V\.?|A\.G\.?|PLC|Plc|L\.L\.C\.?|LLC|Company,?)'
    r'\s*[,.]?\s*)$',
    re.IGNORECASE
)

def _strip_corp_suffix(name: str) -> str:
    """Strip trailing corporate entity suffixes (Co., Ltd., Inc., K.K., etc.) from a company name."""
    if not name:
        return name
    result = name.strip().rstrip(',').rstrip('.').strip()
    # Run up to 3 passes to handle compound suffixes such as "Co., Ltd." — each pass strips
    # one suffix token (e.g. pass 1 removes "Ltd.", pass 2 removes the trailing "Co.").
    for _ in range(3):
        new = _CORP_SUFFIX_RE.sub('', result).strip().rstrip(',').rstrip('.').strip()
        if new == result:
            break
        result = new
    # If stripping removed the entire name (edge case), fall back to the original.
    return result if result else name.strip()

def _country_to_region(country: str):
    c=(country or "").strip().lower()
    if not c: return None
    apac={"singapore","japan","taiwan","hong kong","china","south korea","korea","vietnam","thailand","malaysia","indonesia","philippines","australia","new zealand","india"}
    emea={"united kingdom","uk","england","ireland","germany","france","spain","italy","portugal","belgium","netherlands","switzerland","austria","poland","czech republic","czechia","sweden","norway","denmark","finland"}
    amer={"united states","usa","us","canada","mexico","brazil","argentina","chile","colombia"}
    if c in apac: return "apac"
    if c in emea: return "emea"
    if c in amer: return "americas"
    return None

COMPANY_REGION_PRESENCE = {
    "iqvia": {"apac","emea","americas"},
    "labcorp drug development": {"apac","emea","americas"},
    "labcorp": {"apac","emea","americas"},
    "ppd": {"apac","emea","americas"},
    "parexel": {"apac","emea","americas"},
    "icon": {"apac","emea","americas"},
    "syneos health": {"apac","emea","americas"},
    "novotech": {"apac"},
    "tigermed": {"apac"},
    "pfizer": {"apac","emea","americas"},
    "roche": {"apac","emea","americas"},
    "novartis": {"apac","emea","americas"},
    "johnson & johnson": {"apac","emea","americas"},
    "merck": {"apac","emea","americas"},
    "gsk": {"apac","emea","americas"},
    "sanofi": {"apac","emea","americas"},
    "astrazeneca": {"apac","emea","americas"},
    "bayer": {"apac","emea","americas"},
}

def _has_local_presence(company: str, country: str) -> bool:
    if not country: return True
    region=_country_to_region(country)
    k=(company or "").strip().lower()
    pres=COMPANY_REGION_PRESENCE.get(k)
    if pres:
        if region and region in pres: return True
        if country.strip().lower() in pres: return True
        return False
    return True

CRO_COMPETITORS = ["IQVIA","Labcorp Drug Development","Labcorp","ICON","Parexel","PPD","Syneos Health","Novotech","Tigermed"]
CRA_ADJACENT_ROLES = ["Clinical Trial Associate","Site Manager","Clinical Research Coordinator","Clinical Operations Lead","Study Start-Up Specialist","Clinical Project Manager"]

_BANNED_GENERIC_COMPANY_PHRASES = {
    "gaming studio","game studio","tech company","technology company","software company","pharma company","pharmaceutical company",
    "biotech company","marketing agency","consulting firm","it services provider","design agency","media company","advertising agency",
    "creative studio","blockchain company","web3 company","healthcare company","medical company","diagnostics company","clinical research company",
    "research organization","manufacturing company","energy company","data company"
}

def _is_real_company(name: str) -> bool:
    if not name or not isinstance(name, str):
        return False
    n = name.strip()
    if len(n) < 3:
        return False
    lower = n.lower()
    if lower in _BANNED_GENERIC_COMPANY_PHRASES:
        return False
    if re.search(r'[A-Z]', n):
        return True
    if '&' in n:
        return True
    return False

# Country/region words that Gemini appends to company names (e.g. "Electronic Arts China")
# Strip these trailing tokens so search results are based on the clean brand name.
_COMPANY_COUNTRY_SUFFIX_RE = re.compile(
    r'\s+(?:china|india|japan|korea|taiwan|singapore|malaysia|indonesia|thailand|vietnam|'
    r'philippines|australia|germany|france|uk|us|usa|emea|apac|latam|anz|mea|'
    r'asia(?:\s+pacific)?|pacific|americas|europe|international|global|limited|ltd\.?|'
    r'pte\.?\s*ltd\.?|inc\.?|corp(?:oration)?\.?|llc\.?|co\.?\s*ltd\.?|holdings?)$',
    re.IGNORECASE
)

# Parenthetical regional/status suffixes e.g. "(Japan)", "(Asia Pacific)", "(Merged)"
_COMPANY_PAREN_SUFFIX_RE = re.compile(r'\s*\([^)]+\)\s*$')

# Trailing "&" company-form patterns e.g. "& Co., Inc.", "& Co.", "& Sons"
_COMPANY_AMPERSAND_SUFFIX_RE = re.compile(
    r'\s*&\s*(?:co\.?\s*(?:,\s*)?(?:inc\.?|ltd\.?|llc\.?|plc\.?)?|sons?|partners?|associates?)\s*$',
    re.IGNORECASE
)

# Industry/entity-type descriptor words that Gemini appends to brand names
# e.g. "Takeda Pharmaceutical Company" → "Takeda", "Roche Diagnostics" → "Roche"
# Deliberately narrow: only strip words that are clearly generic descriptors, not brand differentiators.
_COMPANY_INDUSTRY_SUFFIX_RE = re.compile(
    r'\s+(?:pharmaceutical(?:s|(?:\s+company)?)?|diagnostics|biotech(?:nology)?|'
    r'life\s+sciences?|healthcare|health\s*care)$',
    re.IGNORECASE
)

def _strip_company_country_suffix(name: str) -> str:
    """
    Remove trailing suffixes from a company name to return the core brand name.
    Steps applied in order:
      1. Parenthetical regional/status suffixes: "(Japan)", "(Asia Pacific)"
      2. Ampersand company-form patterns: "& Co., Inc.", "& Sons"
      3. Industry/entity-type descriptors: "Pharmaceutical Company", "Diagnostics", "HealthCare"
         (applied up to 2 passes to handle chains like "Pharmaceutical Company")
      4. Country/region/legal-entity suffixes: "China", "Japan", "Ltd", "Inc", "Holdings"
         (applied up to 2 passes)
    Falls back to the original name if stripping reduces it to < 3 chars.
    """
    if not name:
        return name
    original = name.strip()
    s = original

    # Step 1: parenthetical suffixes
    s = _COMPANY_PAREN_SUFFIX_RE.sub('', s).strip()

    # Step 2: ampersand company-form patterns
    s = _COMPANY_AMPERSAND_SUFFIX_RE.sub('', s).strip()

    # Steps 3+4: interleave industry-type and country/legal suffix stripping.
    # Up to _MAX_SUFFIX_STRIP_PASSES passes so chained descriptors like
    # "Pharmaceuticals Corporation" are fully unwound in a single call.
    _MAX_PASSES = 3
    for _ in range(_MAX_PASSES):
        prev = s
        s2 = _COMPANY_INDUSTRY_SUFFIX_RE.sub('', s).strip()
        if s2 != s:
            s = s2
        s2 = _COMPANY_COUNTRY_SUFFIX_RE.sub('', s).strip()
        if s2 != s:
            s = s2
        if s == prev:
            break

    # Keep original if stripping makes it too short (< 3 chars)
    return s if len(s) >= 3 else original

def _supplement_companies(existing, country: str, limit: int, sectors=None):
    """
    Add companies from BUCKET_COMPANIES until we reach the desired limit,
    but do NOT include pharma companies unless sectors explicitly allow pharma.
    """
    pool=[]
    seen=set(x.lower() for x in existing)
    allow_pharma = _sectors_allow_pharma(sectors)
    for bucket, data in BUCKET_COMPANIES.items():
        for group in ("global","apac"):
            for c in data.get(group,[]) or []:
                cl=c.strip()
                if not cl: continue
                if cl.lower() in seen: continue
                # Skip pharma unless allowed by sectors
                if not allow_pharma and _is_pharma_company(cl):
                    continue
                if not _has_local_presence(cl, country):
                    continue
                pool.append(cl)
                seen.add(cl.lower())
                if len(existing)+len(pool) >= limit:
                    break
            if len(existing)+len(pool) >= limit:
                break
        if len(existing)+len(pool) >= limit:
            break
    return existing + pool[:max(0, limit-len(existing))]

def _enforce_company_limit(raw_list, country: str, limit: int, sectors=None):
    """
    Clean raw_list of strings into a limited list of companies.
    If result is shorter than limit, supplement from bucket list but avoid pharma unless sectors allow it.
    """
    cleaned=[]
    seen=set()
    allow_pharma = _sectors_allow_pharma(sectors)
    for c in raw_list or []:
        if not isinstance(c,str): continue
        t=_strip_company_country_suffix(c.strip())
        if not t: continue
        k=t.lower()
        if k in seen: continue
        if not _is_real_company(t):
            continue
        # Skip pharma unless allowed
        if not allow_pharma and _is_pharma_company(t):
            continue
        if not _has_local_presence(t, country):
            continue
        seen.add(k); cleaned.append(t)
        if len(cleaned) >= limit:
            break
    if len(cleaned) < limit:
        cleaned = _supplement_companies(cleaned, country, limit, sectors)
    return cleaned[:limit]

def _gemini_suggestions(job_titles, companies, industry, languages=None, sectors=None, country: str = None, products: list = None):
    languages = languages or []
    sectors = sectors or []
    products = products or []
    locality_hint = "Prioritize Singapore/APAC relevance where naturally applicable." if SINGAPORE_CONTEXT else ""
    
    # Add country-specific filtering instruction
    country_filter_hint = ""
    if country:
        country_filter_hint = f"\n- When suggesting companies, ONLY recommend companies with a legal entity or registered presence in {country}.\n- Exclude companies that do not operate in {country}.\n"

    # Add strict sector rule when sectors are provided to prevent cross-sector leakage
    sector_strict_hint = ""
    if sectors:
        sector_strict_hint = (
            "\n- STRICT SECTOR RULE for company.related: ONLY include companies whose PRIMARY BUSINESS and CORE"
            " OPERATIONS are direct competitors in the specified sector(s). EXCLUDE any company from a different"
            " industry that merely uses or purchases products/services in those sectors. Examples of what to exclude:\n"
            "  * For Gaming / Technology sectors: do NOT include pharma, healthcare, finance, insurance, or"
            " manufacturing companies, even if they use software or hire engineers internally.\n"
            "  * For Healthcare / Clinical Research sectors: do NOT include gaming, tech, or retail companies.\n"
            "  * For Industrial & Manufacturing sectors: do NOT include pure software, gaming, or financial services companies.\n"
            "  Competitors must share the same product/service focus as the job context.\n"
        )

    # Add product-based competitor hint when products are present.
    # Only applied when no companies are provided: when companies exist, they already
    # give Gemini strong competitor context, and adding products could create conflicting signals.
    product_hint = ""
    if products and not companies:
        product_hint = (
            f"\n- PRODUCT CONTEXT: The JD references these products/technologies: {', '.join(products[:10])}.\n"
            "  When no companies are explicitly mentioned, prioritize direct competitors that manufacture or sell these SAME products.\n"
            "  For example: if the JD mentions 'Aircon' or 'HVAC', suggest companies like Daikin, Carrier, Trane, Mitsubishi Electric, LG Electronics, etc.\n"
            "  Do NOT suggest companies from unrelated industries just to fill the list.\n"
        )

    input_obj = {
        "sectors": sectors,
        "jobTitles": job_titles,
        "companies": companies,
        "languages": languages,
        "location": (country or "").strip()
    }
    company_limit = COMPANY_SUGGESTIONS_LIMIT
    job_limit = MAX_SUGGESTIONS_PER_TAG
    prompt = (
        "SYSTEM:\nYou are a sourcing assistant. Produce concise, boolean-friendly suggestions.\n"
        "Return STRICT JSON ONLY in the form:\n"
        "{\"job\":{\"related\":[...]}, \"company\":{\"related\":[]}}\n"
        f"Hard requirements:\n"
        f"- Provide EXACTLY {job_limit} distinct, real, professional job title variants in job.related (if context allows; otherwise fill remaining with closest relevant titles).\n"
        f"- Provide EXACTLY {company_limit} distinct, real, company or organization names in company.related.\n"
        "- Company names MUST be real, brand-level entities (e.g., 'Ubisoft', 'Electronic Arts', 'Epic Games').\n"
        "- DO NOT output generic placeholders (e.g., 'Gaming Studio', 'Tech Company', 'Pharma Company', 'Consulting Firm', 'Marketing Agency').\n"
        + country_filter_hint
        + sector_strict_hint
        + product_hint
        + "- No duplicates, no commentary, no extra keys.\n"
        "- If insufficient context, fill remaining slots with well-known global or APAC companies relevant to the sectors/location.\n"
        "- Maintain JSON key order as shown.\n"
        f"{locality_hint}\n\nINPUT(JSON): {json.dumps(input_obj, ensure_ascii=False)}\n\nJSON:"
    )
    try:
        text = (unified_llm_call_text(prompt) or "").strip()
        start=text.find('{'); end=text.rfind('}')
        if start!=-1 and end!=-1 and end>start:
            parsed=json.loads(text[start:end+1])
            out={"job":{"related":[]}, "company":{"related":[]}}
            if isinstance(parsed,dict):
                jr=parsed.get("job",{}).get("related",[])
                cr=parsed.get("company",{}).get("related",[])
                jr_clean=_clean_list([s for s in jr if isinstance(s,str)], job_limit)
                if len(jr_clean) < job_limit:
                    heuristic_extra=_heuristic_job_suggestions(job_titles or jr_clean, industry, languages, sectors) or []
                    for h in heuristic_extra:
                        if h not in jr_clean and len(jr_clean) < job_limit:
                            jr_clean.append(h)
                # Pass sectors to enforce function so it can avoid adding pharma unless allowed
                cr_enforced=_enforce_company_limit(cr, country, company_limit, sectors)
                out["job"]["related"]=jr_clean[:job_limit]
                out["company"]["related"]=cr_enforced[:company_limit]
            return out
    except Exception as e:
        logger.warning(f"[Gemini Suggest] Failure: {e}")
    return None

def _heuristic_job_suggestions(job_titles, companies, industry, languages=None, sectors=None):
    out=set()
    languages = languages or []
    sectors = sectors or []
    for jt in job_titles:
        base=jt.strip()
        if not base: continue
        if "Senior" not in base and len(out)<MAX_SUGGESTIONS_PER_TAG: out.add(f"Senior {base}")
        if "Lead" not in base and len(out)<MAX_SUGGESTIONS_PER_TAG: out.add(f"Lead {base}")
        if industry=="Gaming" and "Game" not in base and len(out)<MAX_SUGGESTIONS_PER_TAG: out.add(f"Game {base}")
        if "Manager" not in base and not base.endswith("Manager") and len(out)<MAX_SUGGESTIONS_PER_TAG: out.add(f"{base} Manager")
        if len(out)>=MAX_SUGGESTIONS_PER_TAG: break
    if languages and len(out)<MAX_SUGGESTIONS_PER_TAG:
        for lang in languages:
            for role in [f"{lang} Translator", f"{lang} Interpreter", f"{lang} Localization", f"{lang} Linguist"]:
                if len(out)>=MAX_SUGGESTIONS_PER_TAG: break
                out.add(role)
            if len(out)>=MAX_SUGGESTIONS_PER_TAG: break
    if len(out)<MAX_SUGGESTIONS_PER_TAG:
        sect_join=" ".join(sectors).lower()
        jt_join=" ".join(job_titles).lower()
        if ("clinical research" in sect_join) or ("cra" in jt_join) or ("clinical research associate" in jt_join):
            for jt in CRA_ADJACENT_ROLES:
                if len(out)>=COMPANY_SUGGESTIONS_LIMIT: break
                out.add(jt)
    return dedupe(list(out))[:MAX_SUGGESTIONS_PER_TAG]

def _heuristic_company_suggestions(companies, languages=None, sectors=None, country: str = None):
    out=set()
    sectors = sectors or []
    for c in companies:
        base=c.strip()
        if not base: continue
        if base.endswith("Inc") or base.endswith("Inc."):
            cleaned=base.replace("Inc.","").replace("Inc","").strip()
            if cleaned: out.add(cleaned)
        if "Labs" not in base and len(out)<COMPANY_SUGGESTIONS_LIMIT: out.add(f"{base} Labs")
        if "Studio" not in base and len(out)<COMPANY_SUGGESTIONS_LIMIT: out.add(f"{base} Studio")
        if len(out)>=COMPANY_SUGGESTIONS_LIMIT: break
    if len(out)<COMPANY_SUGGESTIONS_LIMIT:
        sect_join=" ".join(sectors).lower()
        comp_join=" ".join(companies).lower()
        cro_context=("clinical research" in sect_join) or any(k in comp_join for k in ["iqvia","ppd","labcorp","parexel","icon","syneos","novotech","tigermed"])
        if cro_context:
            for cro in CRO_COMPETITORS:
                if len(out)>=COMPANY_SUGGESTIONS_LIMIT: break
                if _has_local_presence(cro, country):
                    out.add(cro)
    filtered=[c for c in out if _has_local_presence(c, country)]
    final=_enforce_company_limit(filtered, country, COMPANY_SUGGESTIONS_LIMIT)
    return final[:COMPANY_SUGGESTIONS_LIMIT]

def _prioritize_cross_sector(sets):
    freq={}
    for s in sets:
        for c in s: freq[c]=freq.get(c,0)+1
    cross=[c for c,f in freq.items() if f>1]; single=[c for c,f in freq.items() if f==1]
    ordered=[]; seen=set()
    for s in sets:
        for c in s:
            if c in cross and c not in seen: ordered.append(c); seen.add(c)
    for s in sets:
        for c in s:
            if c in single and c not in seen: ordered.append(c); seen.add(c)
    return ordered

def _heuristic_multi_sector(selected, user_job_title, user_company, languages=None):
    languages = languages or []
    # Use canonical bucket mapping to map selected sector labels to BUCKET_COMPANIES keys
    buckets=[_canon_sector_bucket(x) for x in selected] or ["other"]
    per_sets=[]
    for b in buckets:
        entries=BUCKET_COMPANIES.get(b, {})
        vals=entries.get("global", [])
        if SINGAPORE_CONTEXT:
            vals=list(dict.fromkeys(entries.get("apac", []) + vals))
        per_sets.append(set(vals))
    companies=_prioritize_cross_sector(per_sets)
    jobs=[]; seen=set()
    for b in buckets:
        for t in BUCKET_JOB_TITLES.get(b, []):
            k=t.lower()
            if k not in seen:
                seen.add(k); jobs.append(t)
    if not jobs:
        jobs=BUCKET_JOB_TITLES["other"][:]
    if languages:
        for lang in languages:
            for role in [f"{lang} Translator", f"{lang} Interpreter", f"{lang} Localization", f"{lang} Linguist"]:
                if role.lower() not in seen:
                    jobs.insert(0, role); seen.add(role.lower())
    companies=_enforce_company_limit(companies, None, 20)
    return {"job":{"related":jobs[:15]}, "company":{"related":companies[:20]}}

# Ensure canon mapping includes financial keywords
def _normalize_sector_name(s: str):
    s=(s or "").strip().lower()
    rep={"pharmaceutical":"pharmaceuticals","pharma":"pharmaceuticals","biotech":"biotechnology","med device":"medical devices",
         "medical device":"medical devices","devices":"medical devices","medtech":"medical devices","diagnostic":"diagnostics",
         "health tech":"healthtech","health tech.":"healthtech","healthcare tech":"healthtech","web3":"web3 & blockchain",
         "blockchain":"web3 & blockchain","ai":"ai & data","data":"ai & data","cyber security":"cybersecurity"}
    return rep.get(s, s).replace("&amp;","&").strip()

def _canon_sector_bucket(name: str):
    s=_normalize_sector_name(name)
    if not s:
        return "other"
    # Financial mappings
    if any(k in s for k in ["financial", "finance", "bank", "banking", "insurance", "investment", "asset", "asset management", "asset-management", "fintech", "wealth"]):
        return "financial_services"
    if any(k in s for k in ["pharmaceutical","pharmaceuticals","biotech","biotechnology"]): return "pharma_biotech"
    if "medical device" in s or "medtech" in s or "devices" in s: return "medical_devices"
    if "diagnostic" in s: return "diagnostics"
    if "healthtech" in s or "health tech" in s: return "healthtech"
    if "clinical_research" in s or "clinical research" in s: return "clinical_research"
    if "software" in s or "saas" in s or "technology" in s or "ai & data" in s or "ai" in s: return "technology"
    if "cybersecurity" in s: return "cybersecurity"
    if "automotive" in s or "manufactur" in s or "industrial" in s: return "manufacturing"
    if "energy" in s or "renewable" in s: return "energy"
    if "gaming" in s: return "gaming"
    if "web3" in s or "blockchain" in s: return "web3"
    return "other"

def _bucket_to_sector_label(bucket_name: str):
    """
    Map bucket names (from BUCKET_COMPANIES) to sectors.json labels.
    Returns a sector label that should exist in SECTORS_INDEX, or None.
    """
    bucket_to_label = {
        "pharma_biotech": "Healthcare > Pharmaceuticals",
        "medical_devices": "Healthcare > Medical Devices",
        "diagnostics": "Healthcare > Diagnostics",
        "clinical_research": "Healthcare > Clinical Research",
        "healthtech": "Healthcare > HealthTech",
        "technology": "Technology",
        "manufacturing": "Industrial & Manufacturing",
        "energy": "Energy & Environment",
        "gaming": "Media, Gaming & Entertainment > Gaming",
        "web3": "Emerging & Cross-Sector > Web3 & Blockchain",
        "financial_services": "Financial Services",
        "cybersecurity": "Technology > Cybersecurity",
        "other": None
    }
    
    label = bucket_to_label.get(bucket_name)
    # Verify the label exists in SECTORS_INDEX before returning
    if label and label in SECTORS_INDEX:
        return label
    
    # If exact match not found, try to find a partial match in SECTORS_INDEX
    if label:
        label_lower = label.lower()
        for idx_label in SECTORS_INDEX:
            idx_label_lower = idx_label.lower()
            if label_lower in idx_label_lower or idx_label_lower in label_lower:
                return idx_label
    
    return None

@app.post("/suggest")
def suggest():
    data = request.get_json(force=True, silent=True) or {}
    job_titles = data.get("jobTitles") or []
    companies = data.get("companies") or []
    industry = data.get("industry") or "Non-Gaming"
    languages = data.get("languages") or []
    sectors = data.get("sectors") or data.get("selectedSectors") or []
    country = (data.get("country") or "").strip()
    products = data.get("products") or []  # Product references extracted from JD
    key = (tuple(sorted([jt.strip().lower() for jt in job_titles])),
           tuple(sorted([c.strip().lower() for c in companies])),
           industry.lower(),
           tuple(sorted([str(x).lower() for x in languages])),
           tuple(sorted([str(x).lower() for x in sectors])),
           country.lower())
    with SUGGEST_CACHE_LOCK:
        cached=SUGGEST_CACHE.get(key)
    if cached:
        return jsonify(cached)

    user_jobs_clean = [jt.strip() for jt in job_titles if isinstance(jt, str) and jt.strip()]

    gem=_gemini_suggestions(job_titles, companies, industry, languages, sectors, country, products)
    if gem:
        existing_companies = {c.lower() for c in companies if isinstance(c, str) and c.strip()}
        gem_job_raw = [s for s in gem.get("job", {}).get("related", []) if isinstance(s, str) and s.strip()]
        gem_comp_raw = [s for s in gem.get("company", {}).get("related", []) if isinstance(s, str) and s.strip()]
        gem_job_filtered = [s for s in gem_job_raw if not any(s.strip().lower() == uj.lower() for uj in user_jobs_clean)]
        gem_comp_filtered = [s for s in gem_comp_raw if s.strip().lower() not in existing_companies]
        combined_jobs = list(gem_job_filtered)
        for uj in reversed(user_jobs_clean):
            if not any(uj.lower() == existing.lower() for existing in combined_jobs):
                combined_jobs.insert(0, uj)
        final_job_list = _clean_list(combined_jobs, MAX_SUGGESTIONS_PER_TAG)[:MAX_SUGGESTIONS_PER_TAG]
        final_company_list = gem_comp_filtered[:COMPANY_SUGGESTIONS_LIMIT]
        payload = {
            "job": {"related": final_job_list},
            "company": {"related": final_company_list},
            "engine": "gemini"
        }
    else:
        heuristic_jobs = _heuristic_job_suggestions(job_titles, industry, languages, sectors) or []
        heuristic_companies = _heuristic_company_suggestions(companies, languages, sectors, country) or []
        combined_jobs = list(heuristic_jobs)
        for uj in reversed(user_jobs_clean):
            if not any(uj.lower() == existing.lower() for existing in combined_jobs):
                combined_jobs.insert(0, uj)
        final_job_list = _clean_list(combined_jobs, MAX_SUGGESTIONS_PER_TAG)[:MAX_SUGGESTIONS_PER_TAG]
        final_company_list = heuristic_companies[:COMPANY_SUGGESTIONS_LIMIT]
        payload = {
            "job": {"related": final_job_list},
            "company": {"related": final_company_list},
            "engine": "heuristic"
        }

    with SUGGEST_CACHE_LOCK:
        SUGGEST_CACHE[key]=payload
    return jsonify(payload)

@app.post("/sector_suggest")
def sector_suggest():
    data = request.get_json(force=True, silent=True) or {}
    sectors_list = data.get("selectedSectors") or ([data.get("selectedSector")] if data.get("selectedSector") else [])
    sectors_list=[s for s in sectors_list if isinstance(s,str) and s.strip()]
    user_company=(data.get("userCompany") or "").strip()
    user_job_title=(data.get("userJobTitle") or "").strip()
    languages = data.get("languages") or []
    if not sectors_list and not user_company and not user_job_title and not languages:
        return jsonify({"job":{"related":[]}, "company":{"related":[]}}), 200
    normalized=[]
    for s in sectors_list:
        parts=[p.strip() for p in re.split(r'>', s) if p.strip()]
        normalized.append(parts[-1] if parts else s)
    normalized=[n for n in normalized if n]
    gem=_gemini_multi_sector(normalized, user_job_title, user_company, languages)
    if gem and (gem.get("job",{}).get("related") or gem.get("company",{}).get("related")):
        comp_rel = gem.get("company",{}).get("related") or []
        gem["company"]["related"] = [_strip_corp_suffix(c) for c in comp_rel if c]
        return jsonify(gem), 200
    result = _heuristic_multi_sector(normalized, user_job_title, user_company, languages)
    comp_rel = result.get("company",{}).get("related") or []
    result["company"]["related"] = [_strip_corp_suffix(c) for c in comp_rel if c]
    return jsonify(result), 200

JOBS = {}
JOBS_LOCK = threading.Lock()
PERSIST_JOBS_TO_FILES = os.getenv("PERSIST_JOBS_TO_FILES", "1") == "1"
JOB_FILE_PREFIX="job_"; JOB_FILE_SUFFIX=".json"
_USERNAME_SAFE_RE = re.compile(r'[^A-Za-z0-9_-]')
def _job_file(job_id: str, username: str = "") -> str:
    safe_username = _USERNAME_SAFE_RE.sub('', username or "")
    suffix = f"_{safe_username}" if safe_username else ""
    return os.path.join(OUTPUT_DIR, f"{JOB_FILE_PREFIX}{job_id}{suffix}{JOB_FILE_SUFFIX}")
def persist_job(job_id: str):
    if not PERSIST_JOBS_TO_FILES: return
    try:
        with JOBS_LOCK:
            job=JOBS.get(job_id)
            if not job: return
            username = job.get("username") or ""
            tmp=_job_file(job_id, username)+".tmp"
            with open(tmp,"w",encoding="utf-8") as f: json.dump(job,f,ensure_ascii=False,indent=2)
            os.replace(tmp,_job_file(job_id, username))
    except Exception as e:
        logger.warning(f"[Persist] {e}")
def add_message(job_id: str, text: str):
    with JOBS_LOCK:
        job=JOBS.get(job_id)
        if not job: return
        job['messages'].append(text)
        job['status_html']="<br>".join(job['messages'][-12:])
    persist_job(job_id)

# ... [Job helper functions] ...
LINKEDIN_PROFILE_RE = re.compile(r'(?:^|\.)linkedin\.com/(?:in|pub)/', re.I)
CLEAN_LINKEDIN_SUFFIX_RE = re.compile(r'\s*\|\s*LinkedIn.*$', re.I)
MULTI_SPACE_RE = re.compile(r'\s+')

def is_linkedin_profile(url: str) -> bool:
    return bool(url and LINKEDIN_PROFILE_RE.search(url))

def parse_linkedin_title(title: str):
    if not title: return None, None, None
    cleaned=CLEAN_LINKEDIN_SUFFIX_RE.sub('', title).strip()
    cleaned=cleaned.replace('–','-').replace('—','-')
    if '-' not in cleaned: return None, None, None
    name_part, rest = cleaned.split('-', 1)
    name=name_part.strip()
    if len(name.split())>9 or len(name)<2: return None, None, None
    if not re.search(r'[A-Za-z]', name): return None, None, None
    rest=rest.strip()
    company=""; jobtitle=rest
    at_idx=rest.lower().find(" at ")
    if at_idx!=-1:
        jobtitle=rest[:at_idx].strip()
        company=rest[at_idx+4:].strip()
    name=MULTI_SPACE_RE.sub(' ',name)
    jobtitle=MULTI_SPACE_RE.sub(' ',jobtitle)
    company=MULTI_SPACE_RE.sub(' ',company)
    return name or None, jobtitle or None, company or None

# ── LLM provider adapters ─────────────────────────────────────────────────────

def openai_call_text(prompt: str, api_key: str, model: str = "gpt-4o-mini",
                     system_prompt: str = None,
                     temperature: float = None,
                     max_output_tokens: int = None) -> str | None:
    """Call OpenAI Chat Completions API; returns text content or None on failure."""
    try:
        import openai as _openai  # type: ignore
    except ImportError:
        logger.warning("[OpenAI] openai package not installed; pip install openai")
        return None
    try:
        client = _openai.OpenAI(api_key=api_key.strip())
        messages = []
        if system_prompt:
            messages.append({"role": "system", "content": system_prompt})
        messages.append({"role": "user", "content": prompt})
        kwargs: dict = {"model": model, "messages": messages}
        if temperature is not None:
            kwargs["temperature"] = temperature
        if max_output_tokens is not None:
            kwargs["max_tokens"] = max_output_tokens
        resp = client.chat.completions.create(**kwargs)
        return resp.choices[0].message.content
    except Exception as exc:
        logger.warning(f"[OpenAI] call failed: {exc}")
        return None


def anthropic_call_text(prompt: str, api_key: str, model: str = "claude-3-5-haiku-20241022",
                        system_prompt: str = None,
                        temperature: float = None,
                        max_output_tokens: int = None) -> str | None:
    """Call Anthropic Messages API; returns text content or None on failure."""
    try:
        import anthropic as _anthropic  # type: ignore
    except ImportError:
        logger.warning("[Anthropic] anthropic package not installed; pip install anthropic")
        return None
    try:
        client = _anthropic.Anthropic(api_key=api_key.strip())
        kwargs: dict = {"model": model,
                        "max_tokens": max_output_tokens if max_output_tokens is not None else 4096,
                        "messages": [{"role": "user", "content": prompt}]}
        if system_prompt:
            kwargs["system"] = system_prompt
        if temperature is not None:
            kwargs["temperature"] = temperature
        resp = client.messages.create(**kwargs)
        return resp.content[0].text
    except Exception as exc:
        logger.warning(f"[Anthropic] call failed: {exc}")
        return None


def gemini_call_text(prompt: str, api_key: str, model: str = "gemini-2.5-flash-lite",
                     temperature: float = None,
                     max_output_tokens: int = None) -> str | None:
    """Call Gemini GenerativeModel; returns text or None on failure."""
    try:
        import google.generativeai as _genai  # type: ignore
        _genai.configure(api_key=api_key.strip())
        m = _genai.GenerativeModel(model)
        gen_cfg = {}
        if temperature is not None:
            gen_cfg["temperature"] = temperature
        if max_output_tokens is not None:
            gen_cfg["max_output_tokens"] = max_output_tokens
        resp = m.generate_content(prompt, generation_config=gen_cfg if gen_cfg else None)
        return resp.text
    except Exception as exc:
        logger.warning(f"[Gemini] call failed: {exc}")
        return None


def unified_llm_call_text(prompt: str, system_prompt: str = None,
                           temperature: float = None,
                           max_output_tokens: int = None) -> str | None:
    """Route an LLM text call through the active provider from llm_provider_config.json.
    Priority: active_provider field → Gemini fallback.
    Returns the text response or None if no provider is configured / all fail."""
    cfg = _load_llm_provider_config()
    active = cfg.get("active_provider", "gemini")

    if active == "openai":
        oai = cfg.get("openai", {})
        key = (oai.get("api_key") or "").strip()
        model = oai.get("model", "gpt-4o-mini")
        if key and oai.get("enabled") == "enabled":
            result = openai_call_text(prompt, key, model, system_prompt,
                                      temperature=temperature,
                                      max_output_tokens=max_output_tokens)
            if result is not None:
                return result

    if active == "anthropic":
        ant = cfg.get("anthropic", {})
        key = (ant.get("api_key") or "").strip()
        model = ant.get("model", "claude-3-5-haiku-20241022")
        if key and ant.get("enabled") == "enabled":
            result = anthropic_call_text(prompt, key, model, system_prompt,
                                         temperature=temperature,
                                         max_output_tokens=max_output_tokens)
            if result is not None:
                return result

    # Gemini path (default / fallback)
    gem = cfg.get("gemini", {})
    gem_key = (gem.get("api_key") or "").strip() or (GEMINI_API_KEY or "").strip()
    gem_model = gem.get("model", GEMINI_SUGGEST_MODEL)
    if gem_key:
        return gemini_call_text(prompt, gem_key, gem_model,
                                temperature=temperature,
                                max_output_tokens=max_output_tokens)

    return None

def google_cse_search_page(query: str, api_key: str, cx: str, num: int, start_index: int, gl_hint: str = None):
    if not api_key or not cx: return [], 0
    endpoint="https://www.googleapis.com/customsearch/v1"
    params={"key":api_key,"cx":cx,"q":query,"num":min(num,10),"start":start_index}
    if gl_hint: params["gl"]=gl_hint
    try:
        r=requests.get(endpoint, params=params, timeout=30)
        r.raise_for_status()
        data=r.json()
        items=data.get("items",[]) or []
        total_str=(data.get("searchInformation") or {}).get("totalResults","0") or "0"
        try:
            estimated_total=int(str(total_str).replace(",",""))
        except (ValueError, TypeError):
            estimated_total=0
        out=[]
        for it in items:
            out.append({"link":it.get("link") or "","title":it.get("title") or "","snippet":it.get("snippet") or "","displayLink":it.get("displayLink") or ""})
        return out, estimated_total
    except Exception as e:
        logger.warning(f"[CSE] page fetch failed: {e}")
        return [], 0

def serper_search_page(query: str, api_key: str, num: int, gl_hint: str = None, page: int = 1):
    """Fetch one page of results from the Serper.dev Google Search API.

    Returns the same ``(results, estimated_total)`` tuple as
    ``google_cse_search_page`` so callers are interchangeable.
    Serper does not support cursor-based pagination via a start index; it
    uses a ``page`` parameter instead.  The caller (``unified_search_page``)
    manages page increments.
    """
    if not api_key:
        return [], 0
    endpoint = "https://google.serper.dev/search"
    payload = {"q": query, "num": min(num, 10), "page": page}
    if gl_hint:
        payload["gl"] = gl_hint
    try:
        r = requests.post(
            endpoint,
            headers={"X-API-KEY": api_key, "Content-Type": "application/json"},
            json=payload,
            timeout=30,
        )
        r.raise_for_status()
        data = r.json()
        organic = data.get("organic") or []
        total_str = str((data.get("searchParameters") or {}).get("totalResults", "0") or "0")
        try:
            estimated_total = int(total_str.replace(",", ""))
        except (ValueError, TypeError):
            estimated_total = 0
        out = []
        for it in organic:
            out.append({
                "link": it.get("link") or "",
                "title": it.get("title") or "",
                "snippet": it.get("snippet") or "",
                "displayLink": it.get("displayLink") or (it.get("link") or ""),
            })
        return out, estimated_total
    except Exception as e:
        logger.warning(f"[Serper] page fetch failed: {e}")
        return [], 0

def dataforseo_search_page(query: str, login: str, password: str, num: int = 10, gl_hint: str = None, page: int = 1):
    """Fetch one page of results from the DataforSEO Google Organic Live API.

    Uses HTTP Basic Auth (RFC 7617, UTF-8 encoded) with an explicit Authorization
    header to avoid any character-encoding issues with intermediary libraries.
    Returns the same ``(results, estimated_total)`` tuple as the other adapters
    so callers are interchangeable.  DataforSEO supports an ``offset`` parameter
    for pagination (0-based), which ``unified_search_page`` maps from ``start_index``.
    """
    import base64
    # Strip any accidental whitespace that could corrupt the auth header
    login = (login or "").strip()
    password = (password or "").strip()
    if not login or not password:
        return [], 0
    endpoint = "https://api.dataforseo.com/v3/serp/google/organic/live/regular"
    offset = max(0, (page - 1) * num)
    task = {
        "keyword": query,
        "language_code": "en",
        "location_code": 2840,
        "device": "desktop",
        "depth": min(num, 100),
        "offset": offset,
    }
    # Build Basic Auth header manually using UTF-8 (RFC 7617) to avoid
    # requests' internal latin-1 encoding which can corrupt non-ASCII chars
    credentials = base64.b64encode(f"{login}:{password}".encode("utf-8")).decode("ascii")
    try:
        r = requests.post(
            endpoint,
            headers={
                "Authorization": f"Basic {credentials}",
                "Content-Type": "application/json",
            },
            json=[task],
            timeout=30,
        )
        if not r.ok:
            # Log the response body so we can see DataforSEO's actual error message
            try:
                err_body = r.json()
            except Exception:
                err_body = r.text[:300]
            logger.warning(f"[DataforSEO] page fetch failed: {r.status_code} — {err_body}")
            return [], 0
        data = r.json()
        tasks = data.get("tasks") or []
        if not tasks:
            return [], 0
        task_result = (tasks[0].get("result") or [{}])[0]
        items = task_result.get("items") or []
        estimated_total = task_result.get("items_count") or 0
        out = []
        for it in items:
            if it.get("type") != "organic":
                continue
            out.append({
                "link": it.get("url") or "",
                "title": it.get("title") or "",
                "snippet": it.get("description") or "",
                "displayLink": it.get("domain") or (it.get("url") or ""),
            })
        return out, estimated_total
    except Exception as e:
        logger.warning(f"[DataforSEO] page fetch failed: {e}")
        return [], 0


def unified_search_page(query: str, num: int, start_index: int, gl_hint: str = None,
                        user_provider: str = None, user_serper_key: str = None,
                        user_dfs_login: str = None, user_dfs_password: str = None):
    """Search wrapper that routes to the configured active provider.

    Per-user provider (from Option A service config) takes priority over the global
    admin config.  Priority: per-user Serper → per-user DataforSEO → admin Serper →
    admin DataforSEO → Google CSE (fallback).

    Returns ``(results, estimated_total)`` identical to the individual adapters.
    """
    page = max(1, ((start_index - 1) // max(num, 1)) + 1)

    # Per-user search provider takes priority over the global admin config
    if user_provider == 'serper' and user_serper_key:
        return serper_search_page(query, user_serper_key, num, gl_hint=gl_hint, page=page)
    if user_provider == 'dataforseo' and user_dfs_login and user_dfs_password:
        return dataforseo_search_page(query, user_dfs_login, user_dfs_password, num, gl_hint=gl_hint, page=page)

    cfg = _load_search_provider_config()

    serper_cfg = cfg.get("serper", {})
    serper_key = serper_cfg.get("api_key", "")
    if serper_cfg.get("enabled", "disabled") == "enabled" and serper_key:
        return serper_search_page(query, serper_key, num, gl_hint=gl_hint, page=page)

    dfs_cfg = cfg.get("dataforseo", {})
    dfs_login = (dfs_cfg.get("login") or "").strip()
    dfs_password = (dfs_cfg.get("password") or "").strip()
    if dfs_cfg.get("enabled", "disabled") == "enabled" and dfs_login and dfs_password:
        return dataforseo_search_page(query, dfs_login, dfs_password, num, gl_hint=gl_hint, page=page)

    # Fall back to Google CSE
    return google_cse_search_page(query, GOOGLE_CSE_API_KEY, GOOGLE_CSE_CX, num, start_index, gl_hint=gl_hint)

def _is_private_host(url: str) -> bool:
    """Return True if the URL resolves to a private/loopback/reserved IP — used to block SSRF."""
    import socket
    import ipaddress
    from urllib.parse import urlparse
    try:
        host = urlparse(url).hostname
        if not host:
            return True
        for addrinfo in socket.getaddrinfo(host, None):
            ip = ipaddress.ip_address(addrinfo[4][0])
            if (ip.is_private or ip.is_loopback or ip.is_reserved
                    or ip.is_link_local or ip.is_multicast):
                return True
        return False
    except Exception:
        return True


def get_linkedin_profile_picture(linkedin_url: str, display_name: str = None):
    """
    Retrieve LinkedIn profile picture URL using scraping and Google Custom Search.
    Returns profile picture URL or None if not found.

    Priority:
    1. Try to fetch og:image meta tag directly from LinkedIn profile
    2. Google CSE text search for the LinkedIn profile page — extract
       pagemap.cse_thumbnail or pagemap.metatags[og:image] (most reliable,
       Google caches the metadata even for authenticated pages)
    3. Google CSE image search as a last-resort fallback
    4. Return None if no valid image found

    Security Note: Validates LinkedIn URLs to prevent SSRF attacks.
    """
    if not linkedin_url:
        return None

    # SECURITY: Validate LinkedIn URL to prevent SSRF
    # Must be a valid LinkedIn profile URL
    if not re.match(r'^https?://([a-z]+\.)?linkedin\.com/in/[a-zA-Z0-9\-._~%]+/?$', linkedin_url, re.IGNORECASE):
        logger.warning(f"[Profile Pic] Invalid LinkedIn URL format: {linkedin_url}")
        return None

    profile_pic_url = None

    # Method 1: Try to fetch og:image meta tag directly from LinkedIn profile
    try:
        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
        }
        response = requests.get(linkedin_url, headers=headers, timeout=10)
        
        # Handle rate limiting and forbidden responses
        if response.status_code == 429:
            logger.warning(f"[Profile Pic] Rate limited by LinkedIn: {linkedin_url}")
            # Continue to fallback method
        elif response.status_code == 403:
            logger.warning(f"[Profile Pic] Forbidden by LinkedIn (may require auth): {linkedin_url}")
            # Continue to fallback method
        elif response.status_code == 200:
            # Parse HTML to find og:image meta tag
            # Note: LinkedIn may actively block scraping - this is best-effort
            og_image_match = re.search(r'<meta[^>]*property=["\']og:image["\'][^>]*content=["\']([^"\']+)["\']', response.text, re.IGNORECASE)
            if not og_image_match:
                # Try reverse order (content before property)
                og_image_match = re.search(r'<meta[^>]*content=["\']([^"\']+)["\'][^>]*property=["\']og:image["\']', response.text, re.IGNORECASE)
            
            if og_image_match:
                profile_pic_url = og_image_match.group(1)
                logger.info(f"[Profile Pic] Found og:image from LinkedIn profile: {profile_pic_url}")
                
                # Validate that it's not a placeholder or default image
                if profile_pic_url and not any(placeholder in profile_pic_url.lower() for placeholder in ['default', 'placeholder', 'ghost']):
                    return profile_pic_url
                else:
                    logger.info(f"[Profile Pic] og:image appears to be placeholder, trying fallback")
    except Exception as e:
        logger.warning(f"[Profile Pic] Failed to fetch og:image from LinkedIn (may be blocked): {e}")

    # Method 2 & 3: search fallback (Serper.dev, DataforSEO, or Google CSE)
    _search_cfg = _load_search_provider_config()
    _serper_cfg = _search_cfg.get("serper", {})
    _serper_active = _serper_cfg.get("enabled", "disabled") == "enabled" and bool(_serper_cfg.get("api_key"))
    _dfs_cfg = _search_cfg.get("dataforseo", {})
    _dfs_active = (_dfs_cfg.get("enabled", "disabled") == "enabled"
                   and bool(_dfs_cfg.get("login")) and bool(_dfs_cfg.get("password")))
    _cse_available = bool(GOOGLE_CSE_API_KEY and GOOGLE_CSE_CX)
    if not profile_pic_url and (_serper_active or _dfs_active or _cse_available):
        try:
            from urllib.parse import urlparse

            # Extract URL slug (e.g. john-doe-12345)
            match = re.search(r'linkedin\.com/in/([^/?#]+)', linkedin_url)
            if not match:
                logger.warning(f"[Profile Pic] Could not extract profile slug from URL: {linkedin_url}")
                return None

            profile_slug = match.group(1).rstrip('/')

            # ── Method 2: Text search — Google caches pagemap metadata including og:image ──
            # This works even when LinkedIn requires login to view the profile page.
            # Build query: prefer display name (more specific), fall back to slug.
            def _run_text_search(query_str: str) -> str | None:
                """Run a unified text search and extract the best profile picture URL."""
                try:
                    items, _ = unified_search_page(query_str, 5, 1)
                    for item in items:
                        pagemap = item.get("pagemap", {})
                        # Priority: cse_thumbnail (Google's cached thumbnail — no auth needed)
                        thumbnails = pagemap.get("cse_thumbnail") or []
                        if thumbnails and thumbnails[0].get("src"):
                            src = thumbnails[0]["src"]
                            logger.info(f"[Profile Pic] cse_thumbnail found: {src}")
                            return src
                        # Fallback: og:image from metatags
                        for mt in pagemap.get("metatags", []):
                            og = mt.get("og:image") or mt.get("twitter:image")
                            if og:
                                logger.info(f"[Profile Pic] og:image found via metatags: {og}")
                                return og
                except Exception as exc:
                    logger.warning(f"[Profile Pic] text search failed ({query_str!r}): {exc}")
                return None

            # Try display name first if provided
            if display_name and display_name.strip():
                clean_name = display_name.strip()
                profile_pic_url = _run_text_search(f'"{clean_name}" site:linkedin.com/in')
            # Fall back to URL slug
            if not profile_pic_url:
                profile_pic_url = _run_text_search(f'site:linkedin.com/in "{profile_slug}"')

            # ── Method 3: Image search — last resort (Google CSE only) ──
            if not profile_pic_url and _cse_available and not _serper_active:
                try:
                    endpoint = "https://www.googleapis.com/customsearch/v1"
                    # Build image query: display name is more useful than URL slug here
                    img_query = (
                        f'"{display_name.strip()}" site:linkedin.com/in'
                        if display_name and display_name.strip()
                        else f'site:linkedin.com/in "{profile_slug}"'
                    )
                    params = {
                        "key": GOOGLE_CSE_API_KEY,
                        "cx": GOOGLE_CSE_CX,
                        "q": img_query,
                        "searchType": "image",
                        "num": 5,
                    }
                    r = requests.get(endpoint, params=params, timeout=15)
                    r.raise_for_status()
                    items = r.json().get("items", [])
                    for item in items:
                        image_url = item.get("link", "")
                        context_link = item.get("image", {}).get("contextLink", "")
                        if not image_url:
                            continue
                        # SECURITY: context link must be from linkedin.com
                        try:
                            parsed = urlparse(context_link)
                            if not (parsed.netloc == "linkedin.com" or parsed.netloc.endswith(".linkedin.com")):
                                continue
                        except Exception:
                            continue
                        # Prefer square-ish, reasonably sized images
                        width = item.get("image", {}).get("width", 0)
                        height = item.get("image", {}).get("height", 0)
                        if width and height:
                            aspect = width / height if height else 0
                            if 0.7 <= aspect <= 1.4 and width < 1200:
                                profile_pic_url = image_url
                                logger.info(f"[Profile Pic] Image search hit: {image_url}")
                                break
                    # Last resort: take first image result regardless of dimensions
                    if not profile_pic_url and items:
                        profile_pic_url = items[0].get("link")
                        logger.info(f"[Profile Pic] Using first image result: {profile_pic_url}")
                except Exception as exc:
                    logger.warning(f"[Profile Pic] CSE image search failed: {exc}")

        except Exception as e:
            logger.warning(f"[Profile Pic] Search fallback failed: {e}")

    # Final validation: ensure URL is not empty or broken
    if profile_pic_url:
        try:
            # SECURITY: reject URLs that resolve to private/loopback addresses (SSRF)
            if _is_private_host(profile_pic_url):
                logger.warning(f"[Profile Pic] SSRF: blocked private-host URL: {profile_pic_url}")
                return None
            # Quick HEAD request to verify image exists; allow redirects and treat 2xx/3xx as valid
            try:
                head_response = requests.head(profile_pic_url, timeout=8, allow_redirects=True)
                status = head_response.status_code
                if 200 <= status < 400:
                    return profile_pic_url
                # Some CDNs reject HEAD but serve GET — fall through to return URL if content-type ok
                if status == 405:
                    # Method Not Allowed: server does not support HEAD, trust CSE result
                    return profile_pic_url
                logger.warning(f"[Profile Pic] Image URL returned status {status}: {profile_pic_url}")
                return None
            except requests.exceptions.Timeout:
                # If HEAD times out the URL is likely unreachable; return None
                logger.warning(f"[Profile Pic] HEAD request timed out for: {profile_pic_url}")
                return None
        except Exception as e:
            logger.warning(f"[Profile Pic] Failed to validate image URL: {e}")
            return None
    
    return None

def fetch_image_bytes_from_url(image_url: str, max_size_mb=5):
    """
    Fetch image bytes from a URL and return as bytes suitable for bytea storage.
    Returns bytes or None if fetch failed or image too large.
    
    Args:
        image_url: The URL of the image to fetch
        max_size_mb: Maximum allowed image size in MB (default: 5MB)
    """
    if not image_url:
        return None
    
    try:
        # SECURITY: block SSRF — reject URLs that resolve to private/loopback addresses
        if _is_private_host(image_url):
            logger.warning(f"[Fetch Image Bytes] SSRF: blocked private-host URL: {image_url}")
            return None
        response = requests.get(image_url, timeout=15, stream=True)
        response.raise_for_status()
        
        # Check content type
        content_type = response.headers.get('Content-Type', '')
        if not content_type.startswith('image/'):
            logger.warning(f"[Fetch Image Bytes] Invalid content type: {content_type}")
            return None
        
        # Check content length
        content_length = response.headers.get('Content-Length')
        if content_length and int(content_length) > max_size_mb * 1024 * 1024:
            logger.warning(f"[Fetch Image Bytes] Image too large: {content_length} bytes")
            return None
        
        # Read the image data
        image_bytes = response.content
        
        # Verify size after download
        if len(image_bytes) > max_size_mb * 1024 * 1024:
            logger.warning(f"[Fetch Image Bytes] Downloaded image too large: {len(image_bytes)} bytes")
            return None
        
        logger.info(f"[Fetch Image Bytes] Successfully fetched {len(image_bytes)} bytes from {image_url}")
        return image_bytes
        
    except Exception as e:
        logger.warning(f"[Fetch Image Bytes] Failed to fetch image from {image_url}: {e}")
        return None

def _dedupe_links(records):
    seen=set(); out=[]
    for r in records:
        link=r.get("link")
        if not link or link in seen: continue
        seen.add(link); out.append(r)
    return out

def _infer_primary_job_title(job_titles):
    if job_titles and isinstance(job_titles,list) and job_titles:
        return job_titles[0]
    return ""

def _perform_cse_queries(job_id, queries, target_limit, country,
                         user_provider=None, user_serper_key=None,
                         user_dfs_login=None, user_dfs_password=None):
    results=[]
    m_cc=re.search(r'site:([a-z]{2})\.linkedin\.com/in', " ".join(queries), re.I)
    country_code_hint = m_cc.group(1).lower() if m_cc else None

    # Determine active search provider label for job status messages
    # Per-user provider takes priority over admin config for label determination
    if user_provider == 'serper' and user_serper_key:
        _provider_label = "Serper (user)"
    elif user_provider == 'dataforseo' and user_dfs_login and user_dfs_password:
        _provider_label = "DataforSEO (user)"
    else:
        _sp = _load_search_provider_config()
        _serper_on = (
            _sp.get("serper", {}).get("enabled", "disabled") == "enabled"
            and bool(_sp.get("serper", {}).get("api_key"))
        )
        _dfs_on = (
            _sp.get("dataforseo", {}).get("enabled", "disabled") == "enabled"
            and bool(_sp.get("dataforseo", {}).get("login"))
            and bool(_sp.get("dataforseo", {}).get("password"))
        )
        _provider_label = "Serper" if _serper_on else ("DataforSEO" if _dfs_on else "CSE")

    global_collected = 0

    for q in queries:
        # Global stop-loss: target already reached, no need to fire more queries.
        still_needed = target_limit - global_collected
        if still_needed <= 0:
            add_message(job_id, f"Target reached: {global_collected}/{target_limit} — skipping remaining queries")
            break

        # Each query tries to collect however many are still needed to reach the
        # overall target, so shortfalls from earlier queries are automatically filled.
        gathered=0; start_index=1; pages_fetched=0
        effective_target = still_needed
        add_message(job_id, f"Running {_provider_label}: {q} target={effective_target} (need {still_needed} more to reach {target_limit})")

        while gathered < effective_target:
            remaining = effective_target - gathered
            page_size = min(CSE_PAGE_SIZE, remaining)

            page, estimated_total = unified_search_page(
                q, page_size, start_index, gl_hint=country_code_hint,
                user_provider=user_provider, user_serper_key=user_serper_key,
                user_dfs_login=user_dfs_login, user_dfs_password=user_dfs_password,
            )
            pages_fetched+=1

            # Per-query stop-loss: if Google reports fewer total results than we
            # are requesting from this query, cap the query target to what Google
            # says is actually available.  This prevents wasting API quota on
            # pages that will always return empty, but does NOT prevent subsequent
            # queries from running to make up the shortfall.
            if pages_fetched == 1 and estimated_total > 0 and estimated_total < effective_target:
                effective_target = estimated_total
                add_message(job_id, f"  Stop-loss: {_provider_label} reports ~{estimated_total} results for this query — capping to {effective_target}")

            if not page:
                add_message(job_id, f"  No results page start={start_index}")
                break

            results.extend(page); gathered+=len(page); global_collected+=len(page)
            if len(page) < page_size: break
            start_index += len(page)

            # Safety break — prevents runaway pagination on unexpectedly large indices
            if pages_fetched >= 20: break
            time.sleep(CSE_PAGE_DELAY)

        add_message(job_id, f"{_provider_label} done (collected {gathered}). pages={pages_fetched}")
    return _dedupe_links(results)

def _infer_seniority_from_titles(job_titles):
    if not job_titles: return None
    joined=" ".join([t or "" for t in job_titles])
    # Coordinator always maps to Associate (Junior) — checked first to prevent misclassification
    if re.search(r"\bCoordinator\b", joined, flags=re.I): return "Associate"
    if re.search(r"\bAssociate\b", joined, flags=re.I): return "Associate"
    if re.search(r"\bManager\b", joined, flags=re.I): return "Manager"
    if re.search(r"\bDirector\b", joined, flags=re.I): return "Director"
    return None

_SPECIALS = "<>àÀáÁâÂãÃäÄåÅæÆçÇèÈéÉêÊëËìÌíÍîÎïÏðÐñÑòÒóÓôÔõÖøØùÙúÚûÛüÜýÝÿŸšŠžŽłŁßþÞœŒ~"
_SPECIALS_RE = re.compile("[" + re.escape(_SPECIALS) + "]")

def _sanitize_for_excel(val: str) -> str:
    if not isinstance(val, str): return val or ""
    try:
        import unicodedata
        s=unicodedata.normalize("NFKC", val)
    except Exception:
        s=val
    s=(s.replace("–","-").replace("—","-").replace("’","'").replace("‘","'").replace("“",'"').replace("”",'"'))
    s=_SPECIALS_RE.sub("", s)
    try:
        import unicodedata
        s="".join(ch for ch in s if unicodedata.category(ch)[0]!="C" and ch not in {"\u200b","\u200c","\u200d","\ufeff"})
    except Exception:
        s=s.replace("\u200b","").replace("\u200c","").replace("\u200d","").replace("\ufeff","")
    s=re.sub(r"\s+"," ",s).strip()
    if len(s)>512: s=s[:512]
    return s

def _aggregate_company_dropdown(meta):
    if not isinstance(meta, dict):
        return []
    user = meta.get('user_companies') or []
    auto = meta.get('auto_suggest_companies') or []
    sectors = meta.get('selected_sectors') or []
    languages = meta.get('languages') or []
    sector_companies=[]
    try:
        if sectors:
            norm=[]
            for s in sectors:
                if not isinstance(s,str): continue
                parts=[p.strip() for p in re.split(r'>', s) if p.strip()]
                norm.append(parts[-1] if parts else s)
            sector_payload=_heuristic_multi_sector(norm,"","",languages)
            sector_companies = sector_payload.get('company',{}).get('related',[]) if sector_payload else []
    except Exception as e:
        logger.warning(f"[Dropdown] Sector heuristic failed: {e}")
    merged=[]; seen=set()
    for source in (user, auto, sector_companies):
        for c in source:
            if not isinstance(c,str): continue
            t=c.strip()
            if not t: continue
            k=t.lower()
            if k in seen: continue
            seen.add(k); merged.append(t)
            if len(merged)>=200: break
        if len(merged)>=200: break
    return merged

def _extract_company_from_jobtitle(job_title_raw: str, existing_company: str, company_list):
    if not job_title_raw or existing_company:
        return existing_company, job_title_raw
    seps=r"[\s\-\|,/@]"
    candidates=sorted([c for c in (company_list or []) if isinstance(c,str) and c.strip()], key=lambda x: len(x), reverse=True)
    for comp in candidates:
        pat=re.compile(rf"(^|{seps}+)" rf"({re.escape(comp)})" rf"({seps}+|$)", re.IGNORECASE)
        m=pat.search(job_title_raw)
        if not m: continue
        start_company,end_company=m.span(2)
        cleaned=job_title_raw[:start_company]+job_title_raw[end_company:]
        cleaned=re.sub(rf"({seps}+)", " ", cleaned).strip(" -|,/@").strip()
        return comp, cleaned
    return existing_company, job_title_raw

def _gemini_extract_company_from_jobtitle(job_title_raw: str, candidates=None):
    if not job_title_raw: return None, job_title_raw
    try:
        context={"jobTitle": job_title_raw.strip(), "knownCandidates": (candidates or [])[:30]}
        prompt=("Extract inline employer/company from jobTitle strictly if present. "
                "Return JSON {\"company\":\"\",\"jobTitleWithoutCompany\":\"\"}. "
                f"INPUT:\n{json.dumps(context,ensure_ascii=False)}\nOUTPUT:")
        text = (unified_llm_call_text(prompt) or "").strip()
        start=text.find('{'); end=text.rfind('}')
        if start==-1 or end==-1 or end<=start: return None, job_title_raw
        obj=json.loads(text[start:end+1])
        company=(obj.get("company") or "").strip()
        jt_wo=(obj.get("jobTitleWithoutCompany") or "").strip()
        if not company: return None, job_title_raw
        if not jt_wo: jt_wo=job_title_raw
        return company, jt_wo
    except Exception as e:
        logger.warning(f"[Gemini Title->Company] {e}")
        return None, job_title_raw


# ---------------------------------------------------------------------------
# Criteria-file helpers — must be defined BEFORE `import webbridge_cv` because
# webbridge_cv.py imports these names from this module at its top level.
CRITERIA_OUTPUT_DIR = os.getenv(
    "CRITERIA_OUTPUT_DIR",
    r"F:\Recruiting Tools\Autosourcing\output\Criteras"
)


def _get_criteria_filepath(username, role_tag):
    """Return the full path for the criteria JSON file for the given user/role_tag.
    Returns None if either argument is empty.
    """
    username = (username or "").strip()
    role_tag = (role_tag or "").strip()
    if not username or not role_tag:
        return None
    safe_role = re.sub(r'[<>:"/\\|?*\.]', '_', role_tag).strip('_')
    safe_user = re.sub(r'[<>:"/\\|?*\.]', '_', username).strip('_')
    if not safe_role or not safe_user:
        return None
    return os.path.join(CRITERIA_OUTPUT_DIR, f"{safe_role} {safe_user}.json")


def _read_search_criteria(username, role_tag):
    """Load and return the criteria dict from the saved JSON file, or None if not found."""
    filepath = _get_criteria_filepath(username, role_tag)
    if not filepath:
        return None
    try:
        with open(filepath, "r", encoding="utf-8") as fh:
            data = json.load(fh)
        return data.get("criteria") or None
    except FileNotFoundError:
        return None
    except Exception:
        logger.warning(f"[load_search_criteria] Failed to read {filepath}", exc_info=True)
        return None


# ---------------------------------------------------------------------------
# Import second-half routes (job runner, sourcing, CV processing, bulk assess).
# webbridge_cv.py is a sibling module that imports shared state from this file;
# the circular import is safe because all names below are defined before this
# import statement is reached.

# ---------------------------------------------------------------------------
# Import CV-processing routes.
# webbridge_cv.py imports shared state from webbridge (and webbridge_routes);
# the circular import is safe because all names are defined before this line.
import webbridge_cv  # registers routes with `app`
from webbridge_cv import _gemini_multi_sector, _core_assess_profile  # backward refs

@app.get("/")
def index():
    html_file=os.path.join(BASE_DIR, "AutoSourcing.html")
    if os.path.isfile(html_file): return send_from_directory(BASE_DIR, "AutoSourcing.html")
    return "AutoSourcing WebBridge is running! (AutoSourcing.html not found)", 200

@app.get("/AutoSourcing.html")
def autosourcing_explicit(): return send_from_directory(BASE_DIR, "AutoSourcing.html")

@app.get("/sales_rep_register.html")
def sales_rep_register_html():
    return send_from_directory(BASE_DIR, "sales_rep_register.html")

@app.get('/favicon.ico')
def favicon():
    path=os.path.join(BASE_DIR, 'favicon.ico')
    if not os.path.isfile(path): abort(404)
    return send_from_directory(BASE_DIR, 'favicon.ico', mimetype='image/vnd.microsoft.icon')

# --- START: New Endpoint to serve data_sorter.json ---
@app.get("/data_sorter.json")
def get_data_sorter_json():
    """
    Serve data_sorter.json if present in static folder.
    This allows frontend or other services to access reference lists (JobFamilyRoles, GeoCountries)
    even when data_sorter.py is not active or directly reachable.
    """
    try:
        # Check standard static location relative to BASE_DIR
        static_folder = os.path.join(BASE_DIR, "static")
        filename = "data_sorter.json"
        file_path = os.path.join(static_folder, filename)
        
        if os.path.isfile(file_path):
            return send_from_directory(static_folder, filename, mimetype='application/json')
        else:
            # Fallback check in base dir just in case
            if os.path.isfile(os.path.join(BASE_DIR, filename)):
                return send_from_directory(BASE_DIR, filename, mimetype='application/json')
            
            return jsonify({"error": "data_sorter.json not found"}), 404
    except Exception as e:
        logger.warning(f"Failed to serve data_sorter.json: {e}")
        return jsonify({"error": str(e)}), 500
# --- END: New Endpoint ---

# --- START: Integration of data_sorter.py ---
try:
    import data_sorter
    if hasattr(data_sorter, 'app'):
        # Ensure session compatibility if needed by sharing secret key
        try:
            data_sorter.app.secret_key = app.secret_key
        except Exception:
            pass

        # Mount data_sorter app at /data_sorter prefix
        app.wsgi_app = DispatcherMiddleware(app.wsgi_app, {
            '/data_sorter': data_sorter.app.wsgi_app
        })
        logger.info("Integrated data_sorter app mounted at /data_sorter")
    else:
        logger.warning("data_sorter module found but has no 'app' attribute.")
except ImportError:
    logger.warning("data_sorter.py not found. Skipping integration.")
except Exception as e:
    logger.warning(f"Failed to integrate data_sorter: {e}")
# --- END: Integration ---

def _startup_backfill_role_tag_session():
    """
    One-time startup backfill: for every login row where role_tag is set but
    role_tag_session is NULL, generate a timestamp (NOW()) and transfer it to
    all matching sourcing rows (WHERE username matches AND role_tag matches).

    This handles rows that existed before the role_tag_session column was
    introduced via ALTER TABLE … ADD COLUMN IF NOT EXISTS (which sets NULL for
    pre-existing rows).  Called once when the server process starts.
    """
    try:
        import psycopg2
        pg_host = os.getenv("PGHOST", "localhost")
        pg_port = int(os.getenv("PGPORT", "5432"))
        pg_user = os.getenv("PGUSER", "postgres")
        pg_password = os.getenv("PGPASSWORD", "")
        pg_db = os.getenv("PGDATABASE", "candidate_db")
        conn = psycopg2.connect(
            host=pg_host, port=pg_port, user=pg_user,
            password=pg_password, dbname=pg_db
        )
        cur = conn.cursor()
        try:
            # Ensure columns exist before touching them
            cur.execute("ALTER TABLE login ADD COLUMN IF NOT EXISTS session TIMESTAMPTZ")
            cur.execute("ALTER TABLE sourcing ADD COLUMN IF NOT EXISTS session TIMESTAMPTZ")
            cur.execute("ALTER TABLE sourcing ALTER COLUMN session DROP DEFAULT")
            # Find all login rows with role_tag set but role_tag_session NULL
            cur.execute(
                "SELECT username, role_tag FROM login"
                " WHERE role_tag IS NOT NULL AND role_tag <> '' AND session IS NULL"
            )
            rows = cur.fetchall()
            count = 0
            for username, role_tag in rows:
                if not username:
                    continue
                # Generate a timestamp for this row and write it to login
                cur.execute(
                    "UPDATE login SET session = NOW()"
                    " WHERE username = %s AND role_tag = %s AND session IS NULL"
                    " RETURNING session",
                    (username, role_tag)
                )
                ts_row = cur.fetchone()
                if ts_row and ts_row[0] is not None:
                    # Transfer the same timestamp to sourcing for matching rows
                    cur.execute(
                        "UPDATE sourcing SET session = %s"
                        " WHERE username = %s AND role_tag = %s",
                        (ts_row[0], username, role_tag)
                    )
                    count += 1
            conn.commit()
            if count:
                logger.info(f"[Startup] Backfilled role_tag_session for {count} user(s) missing a session timestamp.")
            else:
                logger.info("[Startup] role_tag_session backfill: no rows needed backfilling (all sessions already set or no role_tag entries found).")
            global _role_tag_session_column_ensured
            _role_tag_session_column_ensured = True
        finally:
            cur.close()
            conn.close()
    except Exception as e:
        logger.error(f"[Startup] role_tag_session backfill failed: {e}")


_startup_backfill_role_tag_session()

# ── API Porting routes ─────────────────────────────────────────────────────────
import re as _re

_PORTING_INPUT_DIR = os.path.normpath(
    os.getenv("PORTING_INPUT_DIR", os.path.join(os.path.dirname(os.path.abspath(__file__)), "porting_input"))
)
_PORTING_MAPPINGS_DIR = os.path.normpath(
    os.getenv("PORTING_MAPPINGS_DIR", os.path.join(os.path.dirname(os.path.abspath(__file__)), "porting_mappings"))
)
_PROCESS_TABLE_FIELDS = [
    'id','name','company','jobtitle','country','linkedinurl','username','userid',
    'product','sector','jobfamily','geographic','seniority','skillset',
    'sourcingstatus','email','mobile','office','role_tag','experience','cv',
    'education','exp','rating','pic','tenure','comment','vskillset',
    'compensation','lskillset','jskillset',
]

def _porting_safe_name(s):
    return _re.sub(r'[^a-zA-Z0-9_\-]', '_', str(s))

def _porting_get_key() -> bytes:
    """Return a stable 32-byte encryption key.

    Priority:
    1. PORTING_SECRET env var (set by the operator for production use).
    2. Persisted key file  <porting_input>/porting.key  (auto-created on first run).
    """
    secret = os.getenv("PORTING_SECRET", "").strip()
    if secret:
        return (secret + "!" * 32)[:32].encode()[:32]
    # Auto-generate / reuse a persistent random key so restarts stay compatible.
    key_path = os.path.join(_PORTING_INPUT_DIR, "porting.key")
    os.makedirs(_PORTING_INPUT_DIR, exist_ok=True)
    if os.path.exists(key_path):
        with open(key_path, "rb") as fh:
            raw = fh.read()
        if len(raw) >= 32:
            return raw[:32]
        logger.warning("[porting] porting.key is shorter than 32 bytes — regenerating.")
    raw = os.urandom(32)
    import tempfile
    fd, tmp = tempfile.mkstemp(dir=_PORTING_INPUT_DIR)
    try:
        os.write(fd, raw)
        os.close(fd)
        try:
            os.chmod(tmp, 0o600)
        except OSError:
            pass
        os.replace(tmp, key_path)
    except Exception:
        try:
            os.close(fd)
        except OSError:
            pass
        try:
            os.unlink(tmp)
        except OSError:
            pass
        raise
    return raw


def _porting_encrypt(data: bytes) -> bytes:
    """AES-256-GCM encrypt.  Returns nonce(12) + ciphertext + tag(16).
    Auto-installs the 'cryptography' package if it is not already present."""
    try:
        from cryptography.hazmat.primitives.ciphers.aead import AESGCM
    except ImportError:
        import subprocess
        import sys as _sys
        logger.info("[porting] 'cryptography' not found — installing…")
        result = subprocess.run(
            [_sys.executable, "-m", "pip", "install", "cryptography"],
            capture_output=True, text=True,
        )
        if result.returncode != 0:
            logger.error("[porting] pip install cryptography failed: %s", result.stderr)
            raise RuntimeError(
                "The 'cryptography' package is required for encryption. "
                "Install it with: pip install cryptography"
            ) from None
        from cryptography.hazmat.primitives.ciphers.aead import AESGCM
    key = _porting_get_key()
    nonce = os.urandom(12)
    ct = AESGCM(key).encrypt(nonce, data, None)
    return nonce + ct

def _porting_login_required():
    """Return (username, None) or (None, error_response)."""
    username = (request.cookies.get("username") or "").strip()
    if not username:
        return None, (jsonify({"error": "Authentication required"}), 401)
    return username, None

@app.post("/api/porting/upload")
def porting_upload():
    username, err = _porting_login_required()
    if err:
        return err
    try:
        body = request.get_json(silent=True) or {}
        upload_type = body.get("type", "")
        content = body.get("content", "")
        filename = body.get("filename", "")
        if not upload_type or not content:
            return jsonify({"error": "Missing type or content"}), 400
        if upload_type not in ("file", "text"):
            return jsonify({"error": 'type must be "file" or "text"'}), 400
        import base64
        if upload_type == "file":
            raw = base64.b64decode(content)
        else:
            raw = content.encode("utf-8")
        if len(raw) > 1024 * 1024:
            return jsonify({"error": "Content too large (max 1 MB)"}), 413
        safe_fname = os.path.basename(str(filename)).replace(" ", "_") if filename else (
            "upload.env" if upload_type == "file" else "api_keys.txt"
        )
        safe_fname = _re.sub(r'[^a-zA-Z0-9_\-\.]', '_', safe_fname)
        safe_fname = f"{_porting_safe_name(username)}_{int(__import__('time').time()*1000)}_{safe_fname}"
        os.makedirs(_PORTING_INPUT_DIR, exist_ok=True)
        encrypted = _porting_encrypt(raw)
        dest = os.path.join(_PORTING_INPUT_DIR, safe_fname + ".enc")
        with open(dest, "wb") as fh:
            fh.write(encrypted)
        return jsonify({"ok": True, "stored": safe_fname + ".enc"})
    except Exception as exc:
        logger.exception("[porting/upload]")
        return jsonify({"error": "Upload failed", "detail": str(exc)}), 500

@app.post("/api/porting/map")
def porting_map():
    username, err = _porting_login_required()
    if err:
        return err
    try:
        body = request.get_json(silent=True) or {}
        names = body.get("names", [])
        if not isinstance(names, list) or not names:
            return jsonify({"error": "names must be a non-empty array"}), 400
        fields_str = ", ".join(_PROCESS_TABLE_FIELDS)
        names_str = ", ".join(f'"{str(n)}"' for n in names)
        prompt = (
            f'You are a database field mapping assistant.\n'
            f'Available target fields (PostgreSQL "process" table): {fields_str}\n\n'
            f'Map each of the following external API field names to the SINGLE best-matching target field.\n'
            f'If there is no reasonable match, use null.\n'
            f'Return ONLY a JSON object (no markdown, no explanation) where each key is the input name and '
            f'each value is the matching target field name or null.\n\n'
            f'Input names: {names_str}'
        )
        raw = (unified_llm_call_text(prompt) or "").strip()
        if not raw:
            return jsonify({"error": "No LLM provider configured."}), 500
        _increment_gemini_query_count(username)
        raw = _re.sub(r'^```(?:json)?\s*', '', raw, flags=_re.IGNORECASE)
        raw = _re.sub(r'\s*```$', '', raw).strip()
        try:
            mapping = json.loads(raw)
        except Exception:
            return jsonify({"error": "LLM returned invalid JSON", "raw": raw}), 500
        cleaned = {k: (v if v and v in _PROCESS_TABLE_FIELDS else None) for k, v in mapping.items()}
        return jsonify({"ok": True, "mapping": cleaned})
    except Exception as exc:
        logger.exception("[porting/map]")
        return jsonify({"error": "Mapping failed", "detail": str(exc)}), 500

@app.post("/api/porting/confirm")
def porting_confirm():
    username, err = _porting_login_required()
    if err:
        return err
    try:
        body = request.get_json(silent=True) or {}
        mapping = body.get("mapping")
        if not mapping or not isinstance(mapping, dict):
            return jsonify({"error": "mapping is required"}), 400
        for k, v in mapping.items():
            if v is not None and v not in _PROCESS_TABLE_FIELDS:
                return jsonify({"error": f"Invalid target field: {v}"}), 400
        os.makedirs(_PORTING_MAPPINGS_DIR, exist_ok=True)
        path_out = os.path.join(_PORTING_MAPPINGS_DIR, _porting_safe_name(username) + ".json")
        with open(path_out, "w", encoding="utf-8") as fh:
            json.dump({"username": username, "mapping": mapping}, fh, indent=2)
        return jsonify({"ok": True})
    except Exception as exc:
        logger.exception("[porting/confirm]")
        return jsonify({"error": "Confirm failed", "detail": str(exc)}), 500

@app.get("/api/porting/mapping")
def porting_get_mapping():
    username, err = _porting_login_required()
    if err:
        return err
    try:
        path_in = os.path.join(_PORTING_MAPPINGS_DIR, _porting_safe_name(username) + ".json")
        if not os.path.isfile(path_in):
            return jsonify({"mapping": None})
        with open(path_in, encoding="utf-8") as fh:
            data = json.load(fh)
        return jsonify({"mapping": data.get("mapping")})
    except Exception as exc:
        logger.exception("[porting/mapping]")
        return jsonify({"error": "Could not load mapping", "detail": str(exc)}), 500

@app.post("/api/porting/export")
def porting_export():
    username, err = _porting_login_required()
    if err:
        return err
    try:
        path_map = os.path.join(_PORTING_MAPPINGS_DIR, _porting_safe_name(username) + ".json")
        if not os.path.isfile(path_map):
            return jsonify({"error": "No confirmed mapping found. Please complete the mapping step first."}), 400
        with open(path_map, encoding="utf-8") as fh:
            mapping = json.load(fh).get("mapping", {})
        cols = [c for c in _PROCESS_TABLE_FIELDS if c not in ("cv", "pic")]
        conn = _pg_connect()
        try:
            cur = conn.cursor()
            col_sql = ", ".join(f'"{c}"' for c in cols)
            cur.execute(f'SELECT {col_sql} FROM "process" WHERE username = %s', (username,))
            rows = cur.fetchall()
            cur.close()
        finally:
            conn.close()
        if not rows:
            return jsonify({"error": "No data found for this user in the process table."}), 404
        reverse_map = {proc: ext for ext, proc in mapping.items() if proc}
        exported = [
            {reverse_map.get(col, col): (row[i] if row[i] is not None else None) for i, col in enumerate(cols)}
            for row in rows
        ]
        json_str = json.dumps(exported, indent=2, default=str)
        body_req = request.get_json(silent=True) or {}
        target_url = body_req.get("targetUrl", "")
        if target_url:
            try:
                import urllib.parse as _up
                import urllib.request as _ur
                import ipaddress as _ipaddr
                import socket as _sock
                parsed = _up.urlparse(target_url)
                if parsed.scheme not in ("http", "https"):
                    raise ValueError("targetUrl must use http or https scheme")
                _host = parsed.hostname or ""
                if not _host:
                    raise ValueError("targetUrl must include a hostname")
                # Block requests to loopback / link-local / private ranges (SSRF guard)
                try:
                    _resolved = _sock.getaddrinfo(_host, None, proto=_sock.IPPROTO_TCP)
                    for _af, _st, _pr, _cn, _sa in _resolved:
                        _ip = _ipaddr.ip_address(_sa[0])
                        if _ip.is_loopback or _ip.is_private or _ip.is_link_local or _ip.is_reserved:
                            raise ValueError(f"targetUrl resolves to a disallowed address: {_sa[0]}")
                except _sock.gaierror:
                    raise ValueError("targetUrl hostname could not be resolved")
                # Reconstruct URL from parsed (and validated) components so the
                # request is made to a known-safe value, not the raw user string.
                _safe_url = _up.urlunparse((
                    parsed.scheme, parsed.netloc,
                    parsed.path, parsed.params, parsed.query, ""
                ))
                req_obj = _ur.Request(
                    _safe_url,
                    data=json_str.encode("utf-8"),
                    headers={"Content-Type": "application/json"},
                    method="POST",
                )
                with _ur.urlopen(req_obj, timeout=15):
                    pass
            except Exception as push_err:
                logger.warning(f"[porting/export] push to {target_url} failed: {push_err}")
        from flask import make_response
        resp = make_response(json_str)
        resp.headers["Content-Type"] = "application/json"
        resp.headers["Content-Disposition"] = f'attachment; filename="porting_export_{_porting_safe_name(username)}.json"'
        log_approval(action="export_pdf_triggered", username=username,
                     detail=f"Data export triggered; {len(exported)} row(s)")
        return resp
    except Exception as exc:
        logger.exception("[porting/export]")
        log_error(source="porting_export", message=str(exc), severity="error",
                  username=username, endpoint="/api/porting/export")
        return jsonify({"error": "Export failed", "detail": str(exc)}), 500


# ── BYOK (Bring Your Own Keys) routes ─────────────────────────────────────────
_BYOK_REQUIRED_KEYS = [
    'GEMINI_API_KEY', 'GOOGLE_CSE_API_KEY', 'GOOGLE_API_KEY',
    'GOOGLE_CSE_CX', 'GOOGLE_CLIENT_ID', 'GOOGLE_CLIENT_SECRET',
]

def _byok_path(username: str) -> str:
    byok_dir = os.path.join(_PORTING_INPUT_DIR, 'byok')
    os.makedirs(byok_dir, exist_ok=True)
    return os.path.join(byok_dir, _porting_safe_name(username) + '.enc')


@app.post("/api/porting/byok/activate")
def byok_activate():
    username, err = _porting_login_required()
    if err:
        return err
    try:
        body = request.get_json(silent=True) or {}
        keys = {}
        missing = []
        for k in _BYOK_REQUIRED_KEYS:
            val = str(body.get(k, '')).strip()
            if not val:
                missing.append(k)
            else:
                keys[k] = val
        if missing:
            return jsonify({"error": f"Missing required keys: {', '.join(missing)}"}), 400
        raw = json.dumps({'username': username, 'keys': keys}).encode('utf-8')
        encrypted = _porting_encrypt(raw)
        dest = _byok_path(username)
        with open(dest, 'wb') as fh:
            fh.write(encrypted)
        log_infrastructure("byok_activated", username=username,
                           detail="BYOK keys activated", status="success")
        return jsonify({"ok": True, "byok_active": True})
    except Exception as exc:
        logger.exception("[porting/byok/activate]")
        log_error(source="byok_activate", message=str(exc), severity="error",
                  username=username, endpoint="/api/porting/byok/activate")
        return jsonify({"error": "BYOK activation failed", "detail": str(exc)}), 500


@app.get("/api/porting/byok/status")
def byok_status():
    username, err = _porting_login_required()
    if err:
        return err
    try:
        active = os.path.isfile(_byok_path(username))
        return jsonify({"byok_active": active})
    except Exception as exc:
        logger.exception("[porting/byok/status]")
        return jsonify({"error": "Could not check BYOK status", "detail": str(exc)}), 500


@app.get("/api/porting/credentials/status")
def porting_credentials_status():
    """Return whether the user has any uploaded credential files on file."""
    username, err = _porting_login_required()
    if err:
        return err
    try:
        safe_prefix = _porting_safe_name(username) + "_"
        has_creds = any(
            f.startswith(safe_prefix) and f.endswith(".enc")
            for f in os.listdir(_PORTING_INPUT_DIR)
        ) if os.path.isdir(_PORTING_INPUT_DIR) else False
        return jsonify({"credentials_on_file": has_creds})
    except Exception as exc:
        logger.exception("[porting/credentials/status]")
        return jsonify({"error": "Could not check credential status", "detail": str(exc)}), 500


@app.delete("/api/porting/byok/deactivate")
def byok_deactivate():
    username, err = _porting_login_required()
    if err:
        return err
    try:
        dest = _byok_path(username)
        if os.path.isfile(dest):
            os.remove(dest)
        log_infrastructure(
            "byok_deactivated",
            username=username,
            detail="BYOK keys file removed",
            status="success",
            key_type="ALL",
            deactivation_reason="manual",
        )
        return jsonify({"ok": True, "byok_active": False})
    except Exception as exc:
        logger.exception("[porting/byok/deactivate]")
        return jsonify({"error": "Could not deactivate BYOK", "detail": str(exc)}), 500


@app.post("/api/porting/byok/validate")
def byok_validate():
    """Validate BYOK keys by probing live Google Cloud APIs + checking credential formats.
    Steps:
      1. Gemini API  — list models (validates GEMINI_API_KEY + billing)
      2. Custom Search API — single query (validates GOOGLE_CSE_API_KEY + GOOGLE_CSE_CX)
      3. GOOGLE_API_KEY format check
      4. OAuth client credential format check
    Returns a structured results array without storing anything."""
    username, err = _porting_login_required()
    if err:
        return err
    try:
        import re
        import urllib.request as _ureq
        import urllib.parse as _uparse

        body = request.get_json(silent=True) or {}
        keys = {}
        missing = []
        for k in _BYOK_REQUIRED_KEYS:
            raw = body.get(k)
            if not isinstance(raw, (str, int, float)):
                missing.append(k); continue
            val = str(raw).strip()
            if not val or len(val) > 512:
                missing.append(k)
            else:
                keys[k] = val
        if missing:
            return jsonify({"error": f"Missing required keys: {', '.join(missing)}"}), 400

        def _probe(url, timeout=8):
            """GET url; returns (http_status_or_None, body_text)."""
            try:
                with _ureq.urlopen(url, timeout=timeout) as resp:
                    return resp.status, resp.read().decode('utf-8', errors='replace')
            except Exception as exc:
                if hasattr(exc, 'code'):
                    try:
                        return exc.code, exc.read().decode('utf-8', errors='replace')
                    except Exception:
                        return exc.code, ''
                return None, str(exc)

        def _err_msg(body_text, fallback):
            try:
                return json.loads(body_text).get('error', {}).get('message', fallback)
            except Exception:
                return fallback

        results = []

        # ── Step 1: Gemini API (GEMINI_API_KEY + billing) ────────────────────────
        gemini_url = (
            "https://generativelanguage.googleapis.com/v1beta/models?key="
            + _uparse.quote(keys['GEMINI_API_KEY'], safe='')
        )
        status, body_text = _probe(gemini_url)
        if status == 200:
            results.append({'step': 'gemini', 'label': 'Gemini API', 'status': 'ok',
                            'detail': 'API key is valid and billing is active.'})
        elif status == 403:
            results.append({'step': 'gemini', 'label': 'Gemini API', 'status': 'error',
                            'detail': _err_msg(body_text, 'Gemini API is not enabled or billing is inactive on this project.'),
                            'consoleUrl': 'https://console.cloud.google.com/apis/library/generativelanguage.googleapis.com'})
        elif status == 400:
            results.append({'step': 'gemini', 'label': 'Gemini API', 'status': 'error',
                            'detail': _err_msg(body_text, 'Invalid GEMINI_API_KEY.'),
                            'consoleUrl': 'https://console.cloud.google.com/apis/credentials'})
        else:
            detail = f'Unexpected HTTP {status}' if status else f'Could not reach Google APIs: {body_text}'
            results.append({'step': 'gemini', 'label': 'Gemini API', 'status': 'warn', 'detail': detail})

        # ── Step 2: Custom Search API (GOOGLE_CSE_API_KEY + GOOGLE_CSE_CX) ───────
        cse_url = (
            "https://customsearch.googleapis.com/customsearch/v1?key="
            + _uparse.quote(keys['GOOGLE_CSE_API_KEY'], safe='')
            + "&cx=" + _uparse.quote(keys['GOOGLE_CSE_CX'], safe='')
            + "&q=test&num=1"
        )
        status, body_text = _probe(cse_url)
        if status == 200:
            results.append({'step': 'cse', 'label': 'Custom Search API', 'status': 'ok',
                            'detail': 'CSE API key and Search Engine ID are valid.'})
        elif status == 403:
            results.append({'step': 'cse', 'label': 'Custom Search API', 'status': 'error',
                            'detail': _err_msg(body_text, 'Custom Search API is not enabled or billing is required.'),
                            'consoleUrl': 'https://console.cloud.google.com/apis/library/customsearch.googleapis.com'})
        elif status == 400:
            results.append({'step': 'cse', 'label': 'Custom Search API', 'status': 'error',
                            'detail': _err_msg(body_text, 'Invalid GOOGLE_CSE_API_KEY or GOOGLE_CSE_CX Search Engine ID.'),
                            'consoleUrl': 'https://console.cloud.google.com/apis/credentials'})
        else:
            detail = f'Unexpected HTTP {status}' if status else f'Could not reach Custom Search API: {body_text}'
            results.append({'step': 'cse', 'label': 'Custom Search API', 'status': 'warn', 'detail': detail})

        # ── Step 3: GOOGLE_API_KEY format ─────────────────────────────────────────
        google_api_key_ok = bool(re.fullmatch(r'AIza[0-9A-Za-z\-_]{35}', keys['GOOGLE_API_KEY']))
        results.append({
            'step': 'google_api_key', 'label': 'GOOGLE_API_KEY Format',
            'status': 'ok' if google_api_key_ok else 'warn',
            'detail': ('Key format is valid (AIza… 39-character format).' if google_api_key_ok
                       else 'Key format looks unusual — expected a 39-character key starting with "AIza".'),
            **({'consoleUrl': 'https://console.cloud.google.com/apis/credentials'} if not google_api_key_ok else {}),
        })

        # ── Step 4: OAuth client credentials ──────────────────────────────────────
        client_id_ok = bool(re.fullmatch(r'\d+-[a-zA-Z0-9]+\.apps\.googleusercontent\.com', keys['GOOGLE_CLIENT_ID']))
        client_secret_ok = bool(re.match(r'^(GOCSPX-[A-Za-z0-9_\-]{28,}|[A-Za-z0-9_\-]{24,})$', keys['GOOGLE_CLIENT_SECRET']))
        if not client_id_ok:
            results.append({'step': 'oauth', 'label': 'OAuth Client Credentials', 'status': 'error',
                            'detail': 'GOOGLE_CLIENT_ID must have the format <numbers>-<id>.apps.googleusercontent.com',
                            'consoleUrl': 'https://console.cloud.google.com/apis/credentials'})
        elif not client_secret_ok:
            results.append({'step': 'oauth', 'label': 'OAuth Client Credentials', 'status': 'warn',
                            'detail': 'GOOGLE_CLIENT_SECRET format looks unusual (expected "GOCSPX-…"). Verify it was copied from Google Cloud Console → Credentials → OAuth 2.0 Client.',
                            'consoleUrl': 'https://console.cloud.google.com/apis/credentials'})
        else:
            results.append({'step': 'oauth', 'label': 'OAuth Client Credentials', 'status': 'ok',
                            'detail': 'Client ID and Client Secret formats are valid.'})

        all_ok = all(r['status'] in ('ok', 'warn') for r in results)
        overall_status = "success" if all_ok else "fail"
        failed_steps = [r['label'] for r in results if r['status'] == 'error']
        log_infrastructure("byok_validation", username=username,
                           detail="; ".join(failed_steps) if failed_steps else "All checks passed",
                           status=overall_status)
        return jsonify({'ok': all_ok, 'results': results})
    except Exception as exc:
        logger.exception("[porting/byok/validate]")
        log_error(source="byok_validate", message=str(exc), severity="error",
                  username=username, endpoint="/api/porting/byok/validate")
        return jsonify({"error": "Validation failed", "detail": str(exc)}), 500


# ── Per-User Service Config (Option A) ───────────────────────────────────────
# Encrypted per-user storage for Search Engine / LLM / Email Verification keys.
# Keys are AES-256-GCM encrypted using the same _porting_get_key() helper.
# Each user's config is stored at <PORTING_INPUT_DIR>/user-services/<username>.enc
# Format matches server.js: IV(16) + tag(16) + ciphertext so server.js can decrypt it.

def _svc_config_path(username: str) -> str:
    svc_dir = os.path.join(_PORTING_INPUT_DIR, 'user-services')
    os.makedirs(svc_dir, exist_ok=True)
    return os.path.join(svc_dir, _porting_safe_name(username) + '.enc')


def _svc_config_encrypt(data: bytes) -> bytes:
    """AES-256-GCM encrypt in Node.js-compatible format: IV(16) + tag(16) + ciphertext."""
    from cryptography.hazmat.primitives.ciphers import Cipher, algorithms, modes
    key = _porting_get_key()
    iv = os.urandom(16)  # 16-byte IV to match server.js
    cipher = Cipher(algorithms.AES(key), modes.GCM(iv))
    encryptor = cipher.encryptor()
    ct = encryptor.update(data) + encryptor.finalize()
    tag = encryptor.tag
    return iv + tag + ct  # IV(16) + tag(16) + ciphertext


def _svc_config_decrypt(data: bytes) -> bytes:
    """AES-256-GCM decrypt in Node.js-compatible format: IV(16) + tag(16) + ciphertext."""
    if len(data) < 33:  # 16-byte IV + 16-byte tag + at least 1 byte
        raise ValueError("Encrypted data too short")
    from cryptography.hazmat.primitives.ciphers import Cipher, algorithms, modes
    key = _porting_get_key()
    iv = data[:16]
    tag = data[16:32]
    ct = data[32:]
    cipher = Cipher(algorithms.AES(key), modes.GCM(iv, tag))
    decryptor = cipher.decryptor()
    return decryptor.update(ct) + decryptor.finalize()


def _porting_decrypt(data: bytes) -> bytes:
    """AES-256-GCM decrypt.  Expects nonce(12) + ciphertext + tag(16)."""
    if len(data) < 28:  # 12-byte nonce + 16-byte tag minimum
        raise ValueError("Encrypted data too short")
    from cryptography.hazmat.primitives.ciphers.aead import AESGCM
    key = _porting_get_key()
    nonce = data[:12]
    ct = data[12:]
    return AESGCM(key).decrypt(nonce, ct, None)


@app.get("/api/user-service-config/status")
def user_svc_config_status():
    username, err = _porting_login_required()
    if err:
        return err
    try:
        path = _svc_config_path(username)
        if not os.path.isfile(path):
            return jsonify({"active": False})
        with open(path, 'rb') as fh:
            raw = fh.read()
        try:
            decrypted = _svc_config_decrypt(raw)
            stored = json.loads(decrypted.decode('utf-8'))
            providers = {
                'search': stored.get('search', {}).get('provider', 'google_cse'),
                'llm': stored.get('llm', {}).get('provider', 'gemini'),
                'email_verif': stored.get('email_verif', {}).get('provider', 'default'),
            }
            return jsonify({"active": True, "providers": providers})
        except Exception:
            return jsonify({"active": False})
    except Exception as exc:
        logger.exception("[user-service-config/status]")
        return jsonify({"error": "Could not retrieve service config status"}), 500


@app.get("/api/user-service-config/search-keys")
def user_svc_config_search_keys():
    """Return decrypted search credentials for the authenticated user.
    Used by AutoSourcing.html to inject per-user search keys into the /start_job payload."""
    username, err = _porting_login_required()
    if err:
        return err
    try:
        path = _svc_config_path(username)
        if not os.path.isfile(path):
            return jsonify({"provider": "google_cse"})
        with open(path, 'rb') as fh:
            raw = fh.read()
        try:
            decrypted = _svc_config_decrypt(raw)
            cfg = json.loads(decrypted.decode('utf-8'))
        except Exception:
            logger.warning("[user-service-config/search-keys] decrypt/parse failed for %s", username)
            return jsonify({"provider": "google_cse"})
        search = cfg.get('search', {})
        result = {"provider": search.get('provider', 'google_cse')}
        if search.get('provider') == 'serper' and search.get('SERPER_API_KEY'):
            result['SERPER_API_KEY'] = search['SERPER_API_KEY']
        if search.get('provider') == 'dataforseo' and search.get('DATAFORSEO_LOGIN'):
            result['DATAFORSEO_LOGIN'] = search['DATAFORSEO_LOGIN']
        if search.get('provider') == 'dataforseo' and search.get('DATAFORSEO_PASSWORD'):
            result['DATAFORSEO_PASSWORD'] = search['DATAFORSEO_PASSWORD']
        return jsonify(result)
    except Exception as exc:
        logger.exception("[user-service-config/search-keys]")
        return jsonify({"provider": "google_cse"})


@app.post("/api/user-service-config/validate")
def user_svc_config_validate():
    """Validate provided keys by calling each service's API. Does NOT store anything."""
    username, err = _porting_login_required()
    if err:
        return err
    try:
        import urllib.request as _ureq2
        import urllib.parse as _uparse2
        import base64

        body = request.get_json(silent=True) or {}
        search = body.get('search') or {}
        llm = body.get('llm') or {}
        email_verif = body.get('email_verif') or {}

        def _probe_get(url, headers=None, timeout=8):
            req = _ureq2.Request(url, headers=headers or {})
            try:
                with _ureq2.urlopen(req, timeout=timeout) as resp:
                    return resp.status, resp.read().decode('utf-8', errors='replace')
            except Exception as exc2:
                if hasattr(exc2, 'code'):
                    try:
                        return exc2.code, exc2.read().decode('utf-8', errors='replace')
                    except Exception:
                        return exc2.code, ''
                return None, str(exc2)

        def _probe_post(url, data, headers=None, timeout=8):
            req = _ureq2.Request(url, data=data, headers=headers or {}, method='POST')
            try:
                with _ureq2.urlopen(req, timeout=timeout) as resp:
                    return resp.status, resp.read().decode('utf-8', errors='replace')
            except Exception as exc2:
                if hasattr(exc2, 'code'):
                    try:
                        return exc2.code, exc2.read().decode('utf-8', errors='replace')
                    except Exception:
                        return exc2.code, ''
                return None, str(exc2)

        results = []

        # ── Search Engine ──────────────────────────────────────────────────────
        sp = (search.get('provider') or '').strip()
        if sp == 'google_cse' or not sp:
            results.append({'label': 'Search Engine', 'status': 'ok',
                            'detail': 'Using platform Google CSE — no custom key required.'})
        elif sp == 'serper':
            key = (search.get('SERPER_API_KEY') or '').strip()
            if not key:
                results.append({'label': 'Serper.dev', 'status': 'error',
                                'detail': 'SERPER_API_KEY is required.'})
            else:
                payload = json.dumps({'q': 'test', 'num': 1}).encode('utf-8')
                status, _ = _probe_post('https://google.serper.dev/search', payload,
                                        headers={'X-API-KEY': key, 'Content-Type': 'application/json'})
                if status == 200:
                    results.append({'label': 'Serper.dev', 'status': 'ok', 'detail': 'API key is valid.'})
                elif status in (401, 403):
                    results.append({'label': 'Serper.dev', 'status': 'error',
                                    'detail': f'Authentication failed (HTTP {status}). Check your SERPER_API_KEY.'})
                else:
                    results.append({'label': 'Serper.dev', 'status': 'warn',
                                    'detail': f'Unexpected HTTP {status} — key may be valid but quota or plan issue possible.'
                                    if status else 'Could not reach Serper API.'})
        elif sp == 'dataforseo':
            login = (search.get('DATAFORSEO_LOGIN') or '').strip()
            pwd   = (search.get('DATAFORSEO_PASSWORD') or '').strip()
            if not login or not pwd:
                results.append({'label': 'DataforSEO', 'status': 'error',
                                'detail': 'DATAFORSEO_LOGIN and DATAFORSEO_PASSWORD are both required.'})
            else:
                auth = base64.b64encode(f'{login}:{pwd}'.encode('utf-8')).decode('ascii')
                status, _ = _probe_get(
                    'https://api.dataforseo.com/v3/appendix/user_data',
                    headers={'Authorization': f'Basic {auth}'}
                )
                if status == 200:
                    results.append({'label': 'DataforSEO', 'status': 'ok', 'detail': 'Credentials are valid.'})
                elif status in (401, 403):
                    results.append({'label': 'DataforSEO', 'status': 'error',
                                    'detail': f'Authentication failed (HTTP {status}). Check login/password.'})
                else:
                    results.append({'label': 'DataforSEO', 'status': 'warn',
                                    'detail': f'Unexpected HTTP {status}.' if status else 'Could not reach DataforSEO API.'})

        # ── LLM ───────────────────────────────────────────────────────────────
        lp = (llm.get('provider') or '').strip()
        if lp == 'gemini' or not lp:
            results.append({'label': 'LLM', 'status': 'ok',
                            'detail': 'Using platform Gemini — no custom key required.'})
        elif lp == 'openai':
            key = (llm.get('OPENAI_API_KEY') or '').strip()
            if not key:
                results.append({'label': 'OpenAI', 'status': 'error', 'detail': 'OPENAI_API_KEY is required.'})
            else:
                status, _ = _probe_get('https://api.openai.com/v1/models',
                                       headers={'Authorization': f'Bearer {key}'})
                if status == 200:
                    results.append({'label': 'OpenAI', 'status': 'ok', 'detail': 'API key is valid.'})
                elif status == 401:
                    results.append({'label': 'OpenAI', 'status': 'error',
                                    'detail': 'Authentication failed. Check your OPENAI_API_KEY.'})
                else:
                    results.append({'label': 'OpenAI', 'status': 'warn',
                                    'detail': f'Unexpected HTTP {status}.' if status else 'Could not reach OpenAI API.'})
        elif lp == 'anthropic':
            key = (llm.get('ANTHROPIC_API_KEY') or '').strip()
            if not key:
                results.append({'label': 'Anthropic', 'status': 'error', 'detail': 'ANTHROPIC_API_KEY is required.'})
            else:
                payload = json.dumps({
                    'model': 'claude-3-haiku-20240307', 'max_tokens': 1,
                    'messages': [{'role': 'user', 'content': 'hi'}]
                }).encode('utf-8')
                status, _ = _probe_post('https://api.anthropic.com/v1/messages', payload, headers={
                    'x-api-key': key, 'anthropic-version': '2023-06-01',
                    'Content-Type': 'application/json',
                })
                if status == 401:
                    results.append({'label': 'Anthropic', 'status': 'error',
                                    'detail': 'Authentication failed. Check your ANTHROPIC_API_KEY.'})
                elif status:
                    results.append({'label': 'Anthropic', 'status': 'ok',
                                    'detail': f'API key accepted (HTTP {status}).'})
                else:
                    results.append({'label': 'Anthropic', 'status': 'warn',
                                    'detail': 'Could not reach Anthropic API.'})

        # ── Email Verification ────────────────────────────────────────────────
        ep = (email_verif.get('provider') or '').strip()
        if ep == 'default' or not ep:
            results.append({'label': 'Email Verification', 'status': 'ok',
                            'detail': 'Using platform default verification — no custom key required.'})
        elif ep == 'neverbounce':
            key = (email_verif.get('NEVERBOUNCE_API_KEY') or '').strip()
            if not key:
                results.append({'label': 'NeverBounce', 'status': 'error',
                                'detail': 'NEVERBOUNCE_API_KEY is required.'})
            else:
                status, _ = _probe_get(
                    f'https://api.neverbounce.com/v4/account/info?key={_uparse2.quote(key, safe="")}')
                if status == 200:
                    results.append({'label': 'NeverBounce', 'status': 'ok', 'detail': 'API key is valid.'})
                elif status in (401, 403):
                    results.append({'label': 'NeverBounce', 'status': 'error',
                                    'detail': f'Authentication failed (HTTP {status}).'})
                else:
                    results.append({'label': 'NeverBounce', 'status': 'warn',
                                    'detail': f'Unexpected HTTP {status}.' if status else 'Could not reach NeverBounce API.'})
        elif ep == 'zerobounce':
            key = (email_verif.get('ZEROBOUNCE_API_KEY') or '').strip()
            if not key:
                results.append({'label': 'ZeroBounce', 'status': 'error',
                                'detail': 'ZEROBOUNCE_API_KEY is required.'})
            else:
                status, body_text = _probe_get(
                    f'https://api.zerobounce.net/v2/getcredits?api_key={_uparse2.quote(key, safe="")}')
                if status == 200:
                    try:
                        credits = json.loads(body_text).get('Credits')
                    except Exception:
                        credits = None
                    try:
                        credits_num = int(credits) if credits is not None else None
                    except (TypeError, ValueError):
                        credits_num = None
                    if credits_num is not None and credits_num > 0:
                        results.append({'label': 'ZeroBounce', 'status': 'ok',
                                        'detail': f'API key valid. Credits remaining: {credits_num}.'})
                    elif credits_num == 0:
                        results.append({'label': 'ZeroBounce', 'status': 'warn',
                                        'detail': 'API key valid but account has 0 credits.'})
                    else:
                        results.append({'label': 'ZeroBounce', 'status': 'ok', 'detail': 'API key accepted.'})
                elif status in (400, 401):
                    results.append({'label': 'ZeroBounce', 'status': 'error',
                                    'detail': f'Authentication failed (HTTP {status}).'})
                else:
                    results.append({'label': 'ZeroBounce', 'status': 'warn',
                                    'detail': f'Unexpected HTTP {status}.' if status else 'Could not reach ZeroBounce API.'})
        elif ep == 'bouncer':
            key = (email_verif.get('BOUNCER_API_KEY') or '').strip()
            if not key:
                results.append({'label': 'Bouncer', 'status': 'error',
                                'detail': 'BOUNCER_API_KEY is required.'})
            else:
                # Use the verify endpoint for validation — the /account endpoint returns
                # 403 on some Bouncer plans even with a valid key.  A test verification
                # call returns 401 only when the key itself is invalid; any other code
                # (200 = success, 402 = no credits, 429 = rate-limited) means the key works.
                # timeout=5 in the query string is the Bouncer per-request MX lookup limit.
                status, _ = _probe_get(
                    'https://api.usebouncer.com/v1.1/email/verify?email=test%40usebouncer.com&timeout=5',
                    headers={'x-api-key': key})
                if status == 401:
                    results.append({'label': 'Bouncer', 'status': 'error',
                                    'detail': 'Authentication failed — invalid API key (HTTP 401).'})
                elif status == 402:
                    results.append({'label': 'Bouncer', 'status': 'warn',
                                    'detail': 'API key is valid but the account has no credits (HTTP 402).'})
                elif status == 429:
                    results.append({'label': 'Bouncer', 'status': 'warn',
                                    'detail': 'API key is valid but the request was rate-limited (HTTP 429).'})
                elif status == 200:
                    results.append({'label': 'Bouncer', 'status': 'ok', 'detail': 'API key is valid.'})
                elif status is not None:
                    # Any non-401 HTTP response means the key was accepted by Bouncer.
                    results.append({'label': 'Bouncer', 'status': 'ok',
                                    'detail': f'API key accepted (HTTP {status}).'})
                else:
                    results.append({'label': 'Bouncer', 'status': 'warn',
                                    'detail': 'Could not reach Bouncer API — please try again.'})

        has_error = any(r['status'] == 'error' for r in results)
        return jsonify({'ok': not has_error, 'results': results})
    except Exception as exc:
        logger.exception("[user-service-config/validate]")
        return jsonify({"error": "Validation failed"}), 500


@app.post("/api/user-service-config/activate")
def user_svc_config_activate():
    """Encrypt and store per-user service config."""
    username, err = _porting_login_required()
    if err:
        return err
    try:
        body = request.get_json(silent=True) or {}
        # Store the whole payload (providers + keys) encrypted
        raw = json.dumps({
            'username': username,
            'search': body.get('search') or {},
            'llm': body.get('llm') or {},
            'email_verif': body.get('email_verif') or {},
        }).encode('utf-8')
        encrypted = _svc_config_encrypt(raw)
        dest = _svc_config_path(username)
        with open(dest, 'wb') as fh:
            fh.write(encrypted)
        log_infrastructure("user_svc_config_activated", username=username,
                           detail="Per-user service config activated", status="success")
        return jsonify({"ok": True, "active": True})
    except Exception as exc:
        logger.exception("[user-service-config/activate]")
        return jsonify({"error": "Activation failed"}), 500


@app.delete("/api/user-service-config/deactivate")
def user_svc_config_deactivate():
    """Delete the per-user service config (keys are wiped)."""
    username, err = _porting_login_required()
    if err:
        return err
    try:
        path = _svc_config_path(username)
        if os.path.isfile(path):
            os.remove(path)
        log_infrastructure("user_svc_config_deactivated", username=username,
                           detail="Per-user service config removed", status="success")
        return jsonify({"ok": True, "active": False})
    except Exception as exc:
        logger.exception("[user-service-config/deactivate]")
        return jsonify({"error": "Deactivation failed"}), 500


@app.get("/load_search_criteria")
def load_search_criteria():
    """Return the saved search criteria and profile list for the given username and role_tag.

    Query params:
        username  – recruiter username
        role_tag  – recruiter's active role tag
    Returns the criteria JSON object (including profiles list), or 404 if no file exists.
    """
    username = (request.args.get("username") or "").strip()
    role_tag = (request.args.get("role_tag") or "").strip()
    if not username or not role_tag:
        return jsonify({"error": "username and role_tag are required"}), 400
    filepath = _get_criteria_filepath(username, role_tag)
    if not filepath:
        return jsonify({"error": "No criteria file found"}), 404
    try:
        with open(filepath, "r", encoding="utf-8") as fh:
            record = json.load(fh)
    except FileNotFoundError:
        return jsonify({"error": "No criteria file found"}), 404
    except Exception as exc:
        logger.warning(f"[load_search_criteria] Failed to read {filepath}: {exc}")
        return jsonify({"error": "Failed to read criteria file"}), 500
    return jsonify({
        "ok": True,
        "criteria": record.get("criteria") or {},
        "name": record.get("name") or record.get("profiles") or [],
    }), 200


@app.post("/save_search_criteria")
def save_search_criteria():
    """Save the search category breakdown criteria to a JSON file on the server.

    Expected payload:
        username   – recruiter username (e.g. "orlha")
        role_tag   – current role tag    (e.g. "Site Activation Manager")
        criteria   – object with keys: Job Title, Seniority, Sector, Country,
                     Company, Skillset, Tenure
    File is written to  <CRITERIA_OUTPUT_DIR>/<role_tag> <username>.json
    """
    try:
        body = request.get_json(force=True, silent=True) or {}
        username = (body.get("username") or "").strip()
        role_tag = (body.get("role_tag") or "").strip()
        criteria = body.get("criteria") or {}

        if not username or not role_tag:
            return jsonify({"error": "username and role_tag are required"}), 400

        filepath = _get_criteria_filepath(username, role_tag)
        if not filepath:
            return jsonify({"error": "Invalid role_tag or username after sanitization"}), 400
        filename = os.path.basename(filepath)

        os.makedirs(CRITERIA_OUTPUT_DIR, exist_ok=True)

        # Fetch profile names from the sourcing table for this user/role search
        profile_names = []
        try:
            _pconn = _pg_connect()
            try:
                _pcur = _pconn.cursor()
                _pcur.execute(
                    "SELECT DISTINCT name FROM sourcing "
                    "WHERE username=%s AND role_tag=%s AND name IS NOT NULL AND name != ''",
                    (username, role_tag)
                )
                # Strip the "님" honorific suffix and surrounding whitespace
                profile_names = [
                    row[0].replace("님", "").strip()
                    for row in _pcur.fetchall()
                    if row[0]
                ]
                _pcur.close()
            finally:
                _pconn.close()
        except Exception as _pe:
            logger.warning(f"[save_search_criteria] Could not fetch profile names: {_pe}")

        record = {
            "role_tag": role_tag,
            "username": username,
            "saved_at": datetime.utcnow().isoformat() + "Z",
            "name": profile_names,
            "criteria": {
                "Job Title":  criteria.get("Job Title") or [],
                "Seniority":  criteria.get("Seniority") or "",
                "Sector":     criteria.get("Sector") or [],
                "Country":    criteria.get("Country") or "",
                "Company":    criteria.get("Company") or [],
                "Skillset":   criteria.get("Skillset") or [],
                "Tenure":     criteria.get("Tenure"),
            }
        }

        with open(filepath, "w", encoding="utf-8") as fh:
            json.dump(record, fh, ensure_ascii=False, indent=2)

        logger.info(f"[save_search_criteria] Written to {filepath} with {len(profile_names)} profile(s)")
        return jsonify({"ok": True, "file": filename, "name": len(profile_names)}), 200

    except Exception as exc:
        logger.exception("[save_search_criteria]")
        return jsonify({"error": str(exc)}), 500


def _find_criteria_file_for_candidate(candidate_name: str):
    """Scan CRITERIA_OUTPUT_DIR for a JSON file whose 'name' list contains candidate_name.
    Returns the file path and parsed record, or (None, None) if not found.
    """
    if not candidate_name or not os.path.isdir(CRITERIA_OUTPUT_DIR):
        return None, None
    norm = candidate_name.replace("님", "").strip().lower()
    try:
        for fname in os.listdir(CRITERIA_OUTPUT_DIR):
            if not fname.endswith(".json"):
                continue
            fpath = os.path.join(CRITERIA_OUTPUT_DIR, fname)
            try:
                with open(fpath, "r", encoding="utf-8") as fh:
                    rec = json.load(fh)
                names = rec.get("name") or rec.get("profiles") or []
                for n in names:
                    if n.replace("님", "").strip().lower() == norm:
                        return fpath, rec
            except Exception:
                continue
    except Exception:
        pass
    return None, None


def _criteria_record_to_pdf_bytes(record: dict) -> bytes:
    """Convert a criteria JSON record to a minimal PDF.
    Falls back to reportlab if available, otherwise writes a raw minimal PDF.
    """
    # -- Flatten content to lines -----------------------------------------------
    lines = []
    lines.append(("title", f"Search Criteria Report"))
    lines.append(("gap", ""))
    lines.append(("key", f"Role: {record.get('role_tag', '')}"))
    lines.append(("key", f"Generated: {record.get('saved_at', '')}"))
    lines.append(("gap", ""))
    lines.append(("section", "CRITERIA"))
    criteria = record.get("criteria") or {}
    for k, v in criteria.items():
        if isinstance(v, list):
            v_str = ", ".join(str(x) for x in v) if v else "—"
        else:
            v_str = str(v) if v is not None else "—"
        lines.append(("item", f"{k}: {v_str}"))
    lines.append(("gap", ""))
    lines.append(("section", "SOURCED PROFILES"))
    profiles = record.get("name") or record.get("profiles") or []
    for p in profiles:
        lines.append(("item", f"  • {p}"))

    # -- Try reportlab first -----------------------------------------------------
    try:
        from reportlab.pdfgen import canvas as rl_canvas
        from reportlab.lib.pagesizes import A4
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
        import io as _io
        buf = _io.BytesIO()
        c = rl_canvas.Canvas(buf, pagesize=A4)
        w, h = A4
        y = h - 50
        for kind, text in lines:
            if kind == "gap":
                y -= 10; continue
            if kind == "title":
                c.setFont("Helvetica-Bold", 16)
            elif kind == "section":
                c.setFont("Helvetica-Bold", 12)
                y -= 4
            else:
                c.setFont("Helvetica", 11)
            # Encode to latin-1 to avoid non-Latin characters crashing Helvetica
            safe = text.encode("latin-1", errors="replace").decode("latin-1")
            c.drawString(40, y, safe)
            y -= 16
            if y < 60:
                c.showPage()
                y = h - 50
        c.save()
        return buf.getvalue()
    except ImportError:
        pass

    # -- Raw minimal PDF fallback (Latin-1 only) ---------------------------------
    import struct as _struct
    text_parts = []
    for kind, text in lines:
        safe = text.encode("latin-1", errors="replace").decode("latin-1")
        size = 16 if kind == "title" else (12 if kind == "section" else 11)
        bold = "-Bold" if kind in ("title", "section") else ""
        text_parts.append((safe, size, bold))

    def _pdf_str(s):
        return s.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")

    stream_lines = ["BT", "/F1 16 Tf", "40 792 Td"]
    y = 742
    for safe, size, bold in text_parts:
        if not safe.strip():
            y -= 10
            stream_lines.append(f"0 -{10} Td")
            continue
        fname = f"/F{'B' if bold else '1'}"
        stream_lines.append(f"{fname} {size} Tf")
        stream_lines.append(f"({_pdf_str(safe)}) Tj")
        y -= 16
        stream_lines.append(f"0 -{16} Td")
    stream_lines.append("ET")
    stream = "\n".join(stream_lines).encode("latin-1")

    def _obj(n, d, s=None):
        out = f"{n} 0 obj\n{d}\n"
        if s is not None:
            out += f"stream\n"
            out = out.encode("latin-1") + s + b"\nendstream\n"
            return out + b"endobj\n"
        return (out + "endobj\n").encode("latin-1")

    o1 = _obj(1, "<< /Type /Catalog /Pages 2 0 R >>")
    o2 = _obj(2, "<< /Type /Pages /Kids [3 0 R] /Count 1 >>")
    o3 = _obj(3, f"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 595 842] /Contents 4 0 R /Resources << /Font << /F1 5 0 R /FB 6 0 R >> >> >>")
    slen = len(stream)
    o4 = _obj(4, f"<< /Length {slen} >>", stream)
    o5 = _obj(5, "<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>")
    o6 = _obj(6, "<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica-Bold >>")

    header = b"%PDF-1.4\n"
    body = o1 + o2 + o3 + o4 + o5 + o6
    offsets = []
    pos = len(header)
    for chunk in (o1, o2, o3, o4, o5, o6):
        offsets.append(pos)
        pos += len(chunk)

    xref_offset = len(header) + len(body)
    xref = f"xref\n0 7\n0000000000 65535 f \n"
    for off in offsets:
        xref += f"{off:010d} 00000 n \n"
    trailer = f"trailer\n<< /Size 7 /Root 1 0 R >>\nstartxref\n{xref_offset}\n%%EOF\n"
    return header + body + (xref + trailer).encode("latin-1")


@app.get("/sourcing/has_criteria_json")
def has_criteria_json():
    """Check whether a criteria JSON file exists for the given candidate.

    Query params:
        linkedin  – candidate LinkedIn URL
        name      – candidate name (used for lookup)
    Returns { "exists": true/false }
    """
    name = (request.args.get("name") or "").strip()
    linkedin = (request.args.get("linkedin") or "").strip()
    if not name and not linkedin:
        return jsonify({"exists": False}), 200
    # If no name provided, try to look it up from sourcing table
    if not name and linkedin:
        try:
            conn = _pg_connect()
            try:
                cur = conn.cursor()
                cur.execute(
                    "SELECT name FROM sourcing WHERE linkedinurl=%s AND name IS NOT NULL LIMIT 1",
                    (linkedin,)
                )
                row = cur.fetchone()
                cur.close()
                if row and row[0]:
                    name = row[0].replace("님", "").strip()
            finally:
                conn.close()
        except Exception:
            pass
    if not name:
        return jsonify({"exists": False}), 200
    fpath, _ = _find_criteria_file_for_candidate(name)
    return jsonify({"exists": fpath is not None}), 200


@app.get("/sourcing/download_criteria_pdf")
def download_criteria_pdf():
    """Download the criteria JSON file for the given candidate as a PDF.

    Query params:
        linkedin  – candidate LinkedIn URL
        name      – candidate name (used for file lookup)
    Returns the PDF file as an attachment, or 404 if no criteria file found.
    """
    name = (request.args.get("name") or "").strip()
    linkedin = (request.args.get("linkedin") or "").strip()
    if not name and not linkedin:
        return "name or linkedin required", 400
    # Look up name from sourcing table if not supplied
    if not name and linkedin:
        try:
            conn = _pg_connect()
            try:
                cur = conn.cursor()
                cur.execute(
                    "SELECT name FROM sourcing WHERE linkedinurl=%s AND name IS NOT NULL LIMIT 1",
                    (linkedin,)
                )
                row = cur.fetchone()
                cur.close()
                if row and row[0]:
                    name = row[0].replace("님", "").strip()
            finally:
                conn.close()
        except Exception as exc:
            logger.warning(f"[download_criteria_pdf] DB lookup failed: {exc}")
    if not name:
        return "No candidate name found", 404
    fpath, record = _find_criteria_file_for_candidate(name)
    if not record:
        return "No criteria file found for this candidate", 404
    try:
        pdf_bytes = _criteria_record_to_pdf_bytes(record)
    except Exception as exc:
        logger.exception("[download_criteria_pdf] PDF generation failed")
        return f"PDF generation failed: {exc}", 500
    safe_role = re.sub(r'[<>:"/\\|?*]', '_', record.get("role_tag", "criteria"))
    safe_name = re.sub(r'[<>:"/\\|?*]', '_', name)
    fname = f"{safe_role} {safe_name}.pdf"
    from flask import Response as _Response
    return _Response(
        pdf_bytes,
        mimetype="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="{fname}"'},
    )


# ---------------------------------------------------------------------------
# Assessment Report Generation (criteria JSON + bulk assessment results → PDF)
# ---------------------------------------------------------------------------

def _enrich_assessment_with_db_vskillset(result: dict, linkedin_url: str = "", process_id=None) -> dict:
    """Attach vskillset and/or category_appraisals from the DB when they are missing from result.

    Two enrichment passes:
    1. vskillset — read from the DB `vskillset` column and attach when absent from result.
       This fixes LinkedIn / SourcingVerify reports where the individual assessment file was
       written before vskillset was added by the caller.
    2. category_appraisals / scoring fields — merged from the DB `rating` column when the
       result's category_appraisals is empty or absent (happens when bulk assessment ran
       before CV data was written to DB, causing _core_assess_profile to exit early with no
       active criteria and no weight breakdown).
    """
    if not isinstance(result, dict):
        return result

    needs_vskillset = not result.get("vskillset")
    needs_appraisals = not result.get("category_appraisals")

    if not needs_vskillset and not needs_appraisals:
        return result
    if not linkedin_url and not process_id:
        return result

    try:
        conn = _pg_connect()
        try:
            cur = conn.cursor()
            if process_id:
                cur.execute(
                    "SELECT vskillset, rating FROM process WHERE id = %s LIMIT 1",
                    (process_id,)
                )
            else:
                cur.execute(
                    "SELECT vskillset, rating FROM process WHERE linkedinurl = %s"
                    " ORDER BY rating_updated_at DESC NULLS LAST LIMIT 1",
                    (linkedin_url,)
                )
            row = cur.fetchone()
            cur.close()
            if row:
                result = dict(result)
                # Pass 1: attach vskillset from DB vskillset column
                if needs_vskillset and row[0]:
                    db_vs = row[0]
                    if isinstance(db_vs, str):
                        try:
                            db_vs = json.loads(db_vs)
                        except Exception:
                            db_vs = None
                    if db_vs:
                        result["vskillset"] = db_vs
                # Pass 2: merge scoring fields from DB rating when appraisals are missing.
                # Fields merged: the full scoring breakdown (category_appraisals, criteria,
                # total_score, stars) and display fields (assessment_level, comments,
                # overall_comment). Fields intentionally excluded: vskillset (handled above),
                # is_level2 (re-derived), file (disk path, not relevant to report).
                if needs_appraisals and row[1]:
                    db_rating = row[1]
                    if isinstance(db_rating, str):
                        try:
                            db_rating = json.loads(db_rating)
                        except Exception:
                            db_rating = None
                    if isinstance(db_rating, dict) and db_rating.get("category_appraisals"):
                        for _merge_key in ("category_appraisals", "criteria", "total_score",
                                           "stars", "assessment_level", "comments",
                                           "overall_comment"):
                            if db_rating.get(_merge_key) and not result.get(_merge_key):
                                result[_merge_key] = db_rating[_merge_key]
        finally:
            conn.close()
    except Exception as _enrich_exc:
        logger.debug(f"[enrich_assessment] DB enrichment failed (non-fatal): {_enrich_exc}")
    return result


def _find_assessment_for_candidate(linkedin_url: str):
    """Find the latest assessment result for a candidate.

    Checks (in order):
      1. OUTPUT_DIR/assessments/assessment_{sha256[:16]}*.json  (individual assessment — may include username suffix)
      2. OUTPUT_DIR/bulk_*_results*.json                        (most-recent bulk run)

    Returns the assessment result dict, or None if not found.
    """
    if not linkedin_url:
        return None
    # 1. Individual assessment file (written by gemini_assess_profile / _assess_and_persist)
    # Match both legacy `assessment_{hash}.json` and new `assessment_{hash}_{username}.json`.
    _hash_prefix = "assessment_" + hashlib.sha256(linkedin_url.encode("utf-8")).hexdigest()[:16]
    assess_dir = os.path.join(OUTPUT_DIR, "assessments")
    _assess_path_legacy = os.path.join(assess_dir, _hash_prefix + ".json")
    try:
        # Try exact legacy name first for speed
        if os.path.exists(_assess_path_legacy):
            with open(_assess_path_legacy, "r", encoding="utf-8") as fh:
                result = json.load(fh)
            return _enrich_assessment_with_db_vskillset(result, linkedin_url=linkedin_url)
        # Scan for files with username suffix (new naming)
        if os.path.isdir(assess_dir):
            candidates = [f for f in os.listdir(assess_dir) if f.startswith(_hash_prefix) and f.endswith(".json")]
            if candidates:
                # Pick most recently modified
                candidates.sort(key=lambda f: os.path.getmtime(os.path.join(assess_dir, f)), reverse=True)
                with open(os.path.join(assess_dir, candidates[0]), "r", encoding="utf-8") as fh:
                    result = json.load(fh)
                return _enrich_assessment_with_db_vskillset(result, linkedin_url=linkedin_url)
    except Exception:
        pass
    # 2. Bulk results files — scan newest first
    if not os.path.isdir(OUTPUT_DIR):
        return None
    try:
        bulk_files = sorted(
            [f for f in os.listdir(OUTPUT_DIR) if f.endswith("_results.json")],
            key=lambda f: os.path.getmtime(os.path.join(OUTPUT_DIR, f)),
            reverse=True,
        )
        norm_url = linkedin_url.strip().lower()
        for fname in bulk_files:
            fpath = os.path.join(OUTPUT_DIR, fname)
            try:
                with open(fpath, "r", encoding="utf-8") as fh:
                    records = json.load(fh)
                if not isinstance(records, list):
                    continue
                for rec in records:
                    if isinstance(rec, dict) and rec.get("linkedinurl", "").strip().lower() == norm_url:
                        result = rec.get("result")
                        if result and not result.get("_skipped"):
                            # Attach vskillset from the record sibling if not already present
                            if "vskillset" not in result and rec.get("vskillset"):
                                result = dict(result)
                                result["vskillset"] = rec.get("vskillset")
                            return _enrich_assessment_with_db_vskillset(result, linkedin_url=linkedin_url)
            except Exception:
                continue
    except Exception:
        pass
    # 3. DB fallback — read the `rating` column from the process table.
    #    This covers candidates assessed via the individual path before file-writing
    #    was added, and acts as a safety net when the on-disk file is missing.
    try:
        conn = _pg_connect()
        try:
            cur = conn.cursor()
            cur.execute(
                "SELECT rating FROM process WHERE linkedinurl = %s AND rating IS NOT NULL"
                " ORDER BY rating_updated_at DESC NULLS LAST LIMIT 1",
                (linkedin_url,)
            )
            row = cur.fetchone()
            cur.close()
            if row and row[0]:
                rating = row[0]
                if isinstance(rating, str):
                    try:
                        rating = json.loads(rating)
                    except Exception:
                        rating = None
                if isinstance(rating, dict) and not rating.get("_skipped"):
                    return _enrich_assessment_with_db_vskillset(rating, linkedin_url=linkedin_url)
        finally:
            conn.close()
    except Exception:
        pass
    return None


def _lines_to_pdf_bytes(lines: list) -> bytes:
    """Convert a list of (kind, text) tuples to PDF bytes.

    Kinds: 'title', 'section', 'key', 'item', 'gap', 'table'
    For 'table' kind, text is a list of rows (list of strings). The first row
    is rendered as a bold header.  Cell text is automatically word-wrapped so
    content never overflows the page boundary.

    Tries reportlab first; falls back to a raw minimal PDF.
    """
    try:
        from reportlab.pdfgen import canvas as rl_canvas
        from reportlab.lib.pagesizes import A4
        import io as _io

        buf = _io.BytesIO()
        c = rl_canvas.Canvas(buf, pagesize=A4)
        page_w, page_h = A4
        MARGIN_L = 40
        MARGIN_R = 30
        CONTENT_W = page_w - MARGIN_L - MARGIN_R  # usable width
        y = page_h - 50

        def _safe(s):
            return str(s).encode("latin-1", errors="replace").decode("latin-1")

        def _wrap_text(text, font, size, avail_w):
            """Word-wrap text into lines that fit within avail_w pixels."""
            safe = _safe(str(text))
            if not safe.strip():
                return [""]
            words = safe.split()
            result, cur = [], ""
            for word in words:
                test = (cur + " " + word).strip() if cur else word
                if c.stringWidth(test, font, size) <= avail_w:
                    cur = test
                else:
                    if cur:
                        result.append(cur)
                    # Single word too long: truncate to fit
                    while word and c.stringWidth(word, font, size) > avail_w and len(word) > 1:
                        word = word[:-1]
                    cur = word
            if cur:
                result.append(cur)
            return result or [""]

        def _new_page():
            nonlocal y
            c.showPage()
            y = page_h - 50

        def _ensure_space(needed):
            nonlocal y
            if y - needed < 60:
                _new_page()

        def _draw_table(rows):
            """Draw a table with word-wrapped cells and dynamic row heights."""
            nonlocal y
            if not rows:
                return
            col_count = max(len(r) for r in rows)
            if col_count == 0:
                return

            # Column widths by column count
            if col_count == 4:
                # CATEGORY | WEIGHT | RATING/STATUS | COMMENT
                col_w = [105, 45, 82, CONTENT_W - 105 - 45 - 82]
            elif col_count == 3:
                col_w = [100, 75, CONTENT_W - 100 - 75]
            elif col_count == 2:
                col_w = [130, CONTENT_W - 130]
            else:
                col_w = [CONTENT_W / col_count] * col_count

            x_pos = [MARGIN_L]
            for cw in col_w[:-1]:
                x_pos.append(x_pos[-1] + cw)

            FONT_SZ = 9
            LINE_H = FONT_SZ + 3
            PAD_H = 3   # horizontal padding per side
            PAD_V = 4   # vertical padding per side

            # Draw outer border around the entire table
            c.setStrokeColorRGB(0.55, 0.55, 0.55)
            c.setLineWidth(0.7)

            for r_idx, row in enumerate(rows):
                is_hdr = r_idx == 0
                font_n = "Helvetica-Bold" if is_hdr else "Helvetica"

                # Pad row to col_count
                padded = list(row) + [""] * (col_count - len(row))

                # Word-wrap each cell
                cell_wrapped = []
                for ci in range(col_count):
                    avail = col_w[ci] - 2 * PAD_H
                    cell_wrapped.append(_wrap_text(str(padded[ci]), font_n, FONT_SZ, avail))

                n_lines = max(len(cl) for cl in cell_wrapped)
                row_h = n_lines * LINE_H + 2 * PAD_V
                row_h = max(row_h, LINE_H + 2 * PAD_V)

                _ensure_space(row_h + 4)

                # Background
                if is_hdr:
                    c.setFillColorRGB(0.18, 0.36, 0.56)  # professional dark blue
                    c.rect(MARGIN_L, y - row_h, CONTENT_W, row_h, fill=1, stroke=0)
                    c.setFillColorRGB(1, 1, 1)
                elif r_idx % 2 == 1:
                    c.setFillColorRGB(0.95, 0.97, 1.0)
                    c.rect(MARGIN_L, y - row_h, CONTENT_W, row_h, fill=1, stroke=0)
                    c.setFillColorRGB(0.1, 0.1, 0.1)
                else:
                    c.setFillColorRGB(1, 1, 1)
                    c.rect(MARGIN_L, y - row_h, CONTENT_W, row_h, fill=1, stroke=0)
                    c.setFillColorRGB(0.1, 0.1, 0.1)

                # Cell text
                c.setFont(font_n, FONT_SZ)
                for ci, (cl, xp, cw) in enumerate(zip(cell_wrapped, x_pos, col_w)):
                    text_y = y - PAD_V - LINE_H + 2
                    for ln in cl:
                        c.drawString(xp + PAD_H, text_y, ln)
                        text_y -= LINE_H

                # Row bottom border
                c.setStrokeColorRGB(0.70, 0.70, 0.70)
                c.setLineWidth(0.4)
                c.line(MARGIN_L, y - row_h, MARGIN_L + CONTENT_W, y - row_h)

                c.setFillColorRGB(0, 0, 0)
                y -= row_h

            # Outer border (left + right verticals)
            c.setStrokeColorRGB(0.55, 0.55, 0.55)
            c.setLineWidth(0.6)
            y -= 4

        for kind, text in lines:
            if kind == "gap":
                y -= 10
                continue

            if kind == "table":
                _draw_table(text)
                continue

            # Regular text kinds
            if kind == "title":
                font_n, font_s = "Helvetica-Bold", 16
                indent = MARGIN_L
                _ensure_space(font_s + 14)
                c.setFont(font_n, font_s)
                c.drawString(indent, y, _safe(str(text)))
                y -= font_s + 4
                # Decorative underline
                c.setStrokeColorRGB(0.18, 0.36, 0.56)
                c.setLineWidth(1.5)
                c.line(MARGIN_L, y, MARGIN_L + CONTENT_W, y)
                y -= 6
                continue
            elif kind == "section":
                font_n, font_s = "Helvetica-Bold", 11
                indent = MARGIN_L
                y -= 4
            elif kind == "key":
                font_n, font_s = "Helvetica-Bold", 10
                indent = MARGIN_L
            else:  # "item"
                font_n, font_s = "Helvetica", 9
                indent = MARGIN_L + 8

            avail_w = CONTENT_W - (indent - MARGIN_L)
            for ln in _wrap_text(str(text), font_n, font_s, avail_w):
                _ensure_space(font_s + 6)
                c.setFont(font_n, font_s)
                c.drawString(indent, y, ln)
                y -= font_s + 3

        c.save()
        return buf.getvalue()
    except ImportError:
        pass

    # Raw minimal PDF fallback
    def _pdf_str(s):
        return s.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")

    stream_lines = ["BT", "/F1 16 Tf", "40 792 Td"]
    for kind, text in lines:
        if kind == "table":
            rows = text
            for r_idx, row in enumerate(rows):
                row_text = " | ".join(str(c) for c in row)
                safe = row_text.encode("latin-1", errors="replace").decode("latin-1")
                fname = "/FB" if r_idx == 0 else "/F1"
                stream_lines.append(f"{fname} 9 Tf")
                stream_lines.append(f"({_pdf_str(safe)}) Tj")
                stream_lines.append("0 -14 Td")
            continue
        safe = str(text).encode("latin-1", errors="replace").decode("latin-1")
        if not safe.strip() or kind == "gap":
            stream_lines.append("0 -10 Td")
            continue
        size = 16 if kind == "title" else (12 if kind == "section" else 11)
        bold = kind in ("title", "section", "key")
        fname = "/FB" if bold else "/F1"
        stream_lines.append(f"{fname} {size} Tf")
        stream_lines.append(f"({_pdf_str(safe)}) Tj")
        stream_lines.append("0 -16 Td")
    stream_lines.append("ET")
    stream = "\n".join(stream_lines).encode("latin-1")

    def _obj(n, d, s=None):
        out = f"{n} 0 obj\n{d}\n"
        if s is not None:
            out = out.encode("latin-1") + b"stream\n" + s + b"\nendstream\nendobj\n"
            return out
        return (out + "endobj\n").encode("latin-1")

    o1 = _obj(1, "<< /Type /Catalog /Pages 2 0 R >>")
    o2 = _obj(2, "<< /Type /Pages /Kids [3 0 R] /Count 1 >>")
    o3 = _obj(3, "<< /Type /Page /Parent 2 0 R /MediaBox [0 0 595 842] /Contents 4 0 R /Resources << /Font << /F1 5 0 R /FB 6 0 R >> >> >>")
    slen = len(stream)
    o4 = _obj(4, f"<< /Length {slen} >>", stream)
    o5 = _obj(5, "<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>")
    o6 = _obj(6, "<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica-Bold >>")

    header = b"%PDF-1.4\n"
    body = o1 + o2 + o3 + o4 + o5 + o6
    offsets = []
    pos = len(header)
    for chunk in (o1, o2, o3, o4, o5, o6):
        offsets.append(pos)
        pos += len(chunk)

    xref_offset = len(header) + len(body)
    xref = "xref\n0 7\n0000000000 65535 f \n"
    for off in offsets:
        xref += f"{off:010d} 00000 n \n"
    trailer = f"trailer\n<< /Size 7 /Root 1 0 R >>\nstartxref\n{xref_offset}\n%%EOF\n"
    return header + body + (xref + trailer).encode("latin-1")


def _build_report_lines(candidate_name: str, criteria_record: dict, assessment_result: dict) -> list:
    """Build the (kind, text) lines for a full assessment report PDF.

    All data sections are rendered as structured tables so that no text
    overflows the page boundary.  Table cells word-wrap automatically.
    """
    role_tag = criteria_record.get("role_tag", "")
    saved_at = criteria_record.get("saved_at", "")
    criteria = criteria_record.get("criteria") or {}

    lines = []

    # ── Title ─────────────────────────────────────────────────────────────────
    lines.append(("title", "Assessment Report"))
    lines.append(("gap", ""))

    # ── Candidate information table ───────────────────────────────────────────
    lines.append(("table", [
        ["FIELD", "VALUE"],
        ["Candidate", candidate_name],
        ["Role", role_tag],
        ["Date", saved_at],
    ]))
    lines.append(("gap", ""))

    if assessment_result:
        # ── Assessment Summary ─────────────────────────────────────────────────
        stars = max(0, min(int(assessment_result.get("stars", 0) or 0), 5))
        star_str = ("*" * stars) + ("." * (5 - stars)) + f" ({stars}/5)"
        overall = assessment_result.get("overall_comment", "") or ""

        lines.append(("section", "ASSESSMENT SUMMARY"))
        summary_rows = [
            ["FIELD", "VALUE"],
            ["Overall Score", str(assessment_result.get("total_score", "-"))],
            ["Stars", star_str],
            ["Assessment Level", str(assessment_result.get("assessment_level", "-"))],
        ]
        if overall:
            summary_rows.append(["Overall Comment", overall])
        lines.append(("table", summary_rows))
        lines.append(("gap", ""))

        # ── Score Breakdown ────────────────────────────────────────────────────
        breakdown = assessment_result.get("criteria") or {}
        if breakdown:
            lines.append(("section", "SCORE BREAKDOWN"))
            bd_rows = [["CATEGORY", "SCORE"]]
            for cat, score in breakdown.items():
                bd_rows.append([str(cat), str(score)])
            lines.append(("table", bd_rows))
            lines.append(("gap", ""))

        # ── Category Appraisals (4-column with wrapped COMMENT) ────────────────
        appraisals = assessment_result.get("category_appraisals") or {}
        if appraisals:
            lines.append(("section", "CATEGORY APPRAISALS"))
            ap_rows = [["CATEGORY", "WEIGHT", "RATING / STATUS", "COMMENT"]]
            for cat, appraisal in appraisals.items():
                if isinstance(appraisal, dict):
                    weight = appraisal.get("weight_percent", "")
                    rating = appraisal.get("rating", "")
                    status = appraisal.get("status", "")
                    comment = appraisal.get("comment", "")
                    rating_status = f"{str(rating)} / {str(status)}" if status else str(rating)
                    ap_rows.append([
                        str(cat),
                        f"{weight}%" if weight not in (None, "") else "-",
                        rating_status,
                        str(comment),
                    ])
                else:
                    ap_rows.append([str(cat), "-", "-", str(appraisal)])
            lines.append(("table", ap_rows))
            lines.append(("gap", ""))

        # ── Skill Comments ─────────────────────────────────────────────────────
        comments_raw = assessment_result.get("comments")
        if comments_raw:
            lines.append(("section", "SKILL COMMENTS"))
            if isinstance(comments_raw, str):
                # Narrative string — render each paragraph as a row in a
                # single-column table so text stays within page boundaries.
                paras = [p.strip() for p in comments_raw.split("\n") if p.strip()]
                if paras:
                    sc_rows = [["COMMENTS"]]
                    for para in paras:
                        sc_rows.append([para])
                    lines.append(("table", sc_rows))
            elif isinstance(comments_raw, (list, tuple)):
                sc_rows = [["SKILL", "STATUS", "COMMENT"]]
                for entry in comments_raw:
                    if isinstance(entry, dict):
                        skill = str(entry.get("skill") or entry.get("category", ""))
                        match = str(entry.get("match") or entry.get("status", ""))
                        note = str(entry.get("comment") or entry.get("note", ""))
                        sc_rows.append([skill, match, note])
                    else:
                        sc_rows.append([str(entry), "", ""])
                lines.append(("table", sc_rows))
            lines.append(("gap", ""))

    # ── Search Criteria (2-column table) ──────────────────────────────────────
    if criteria:
        lines.append(("section", "SEARCH CRITERIA"))
        crit_rows = [["CRITERIA", "VALUE"]]
        for k, v in criteria.items():
            if isinstance(v, list):
                v_str = ", ".join(str(x) for x in v) if v else "-"
            else:
                v_str = str(v) if v is not None else "-"
            crit_rows.append([str(k), v_str])
        lines.append(("table", crit_rows))

    return lines


@app.get("/sourcing/has_report")
def has_report():
    """Check whether a full assessment report can be generated for the given candidate.

    Requires both a criteria JSON file AND a completed assessment result.
    Query params:
        linkedin  – candidate LinkedIn URL
        name      – candidate name
    Returns { "exists": true/false }
    """
    name = (request.args.get("name") or "").strip()
    linkedin = (request.args.get("linkedin") or "").strip()
    if not name and not linkedin:
        return jsonify({"exists": False}), 200
    # Look up name from DB if missing
    if not name and linkedin:
        try:
            conn = _pg_connect()
            try:
                cur = conn.cursor()
                cur.execute(
                    "SELECT name FROM sourcing WHERE linkedinurl=%s AND name IS NOT NULL LIMIT 1",
                    (linkedin,)
                )
                row = cur.fetchone()
                cur.close()
                if row and row[0]:
                    name = row[0].replace("님", "").strip()
            finally:
                conn.close()
        except Exception:
            pass
    if not name:
        return jsonify({"exists": False}), 200
    # criteria JSON must exist
    fpath, _ = _find_criteria_file_for_candidate(name)
    if not fpath:
        return jsonify({"exists": False}), 200
    # assessment result must also exist
    assessment = _find_assessment_for_candidate(linkedin) if linkedin else None
    return jsonify({"exists": assessment is not None}), 200


@app.get("/sourcing/check_reassess")
def check_reassess():
    """Check whether the Reassess File button should be visible for a candidate.

    Logic: query the `process` table for the given LinkedIn URL and inspect the
    `rating` JSON column's `category_appraisals` for the 7 Category Breakdown
    categories (Company, Country, Job Title, Sector, Seniority, Skillset, Tenure).

    Returns JSON:
        {
          "show_reassess": true,   # true = show the button (at least one category rating missing)
          "has_rating":    false,  # true = a completed assessment exists in process.rating
          "missing_fields": ["Seniority", "Tenure"]  # list of categories missing a rating
        }

    The button should be shown when show_reassess is true (i.e. at least one
    Category Breakdown rating is NULL / missing in category_appraisals).
    """
    linkedin = (request.args.get("linkedin") or "").strip().rstrip("/")
    if not linkedin:
        return jsonify({"show_reassess": False, "has_rating": False, "missing_fields": []}), 200

    # Normalise the URL the same way the DB does
    normalized = linkedin.lower().rstrip("/")

    # The 7 required Category Breakdown categories
    _REQUIRED_CATEGORIES = ["Company", "Country", "Job Title", "Sector", "Seniority", "Skillset", "Tenure"]

    try:
        conn = _pg_connect()
        try:
            cur = conn.cursor()
            cur.execute(
                """
                SELECT rating
                FROM process
                WHERE LOWER(TRIM(TRAILING '/' FROM linkedinurl)) = %s
                   OR normalized_linkedin = %s
                ORDER BY id DESC
                LIMIT 1
                """,
                (normalized, normalized),
            )
            row = cur.fetchone()
            cur.close()
        finally:
            conn.close()

        if not row:
            # No process record at all — show the button so the user can upload a CV
            return jsonify({"show_reassess": True, "has_rating": False, "missing_fields": _REQUIRED_CATEGORIES}), 200

        rating_raw = row[0]

        # Parse the rating JSON column
        has_rating = False
        category_appraisals = {}
        if rating_raw is not None:
            r_str = str(rating_raw).strip() if not isinstance(rating_raw, str) else rating_raw.strip()
            if r_str and r_str not in ("{}", "null", "[]", ""):
                has_rating = True
                try:
                    rating_data = json.loads(r_str) if isinstance(r_str, str) else rating_raw
                    if isinstance(rating_data, dict):
                        category_appraisals = rating_data.get("category_appraisals") or {}
                except Exception:
                    pass

        if not has_rating:
            # No assessment yet — show the button for all categories
            return jsonify({"show_reassess": True, "has_rating": False, "missing_fields": _REQUIRED_CATEGORIES}), 200

        # Check each required category for a non-null rating value in category_appraisals
        missing = []
        for cat_name in _REQUIRED_CATEGORIES:
            appraisal = category_appraisals.get(cat_name)
            if not isinstance(appraisal, dict):
                missing.append(cat_name)
                continue
            rating_val = appraisal.get("rating")
            if rating_val is None or rating_val == "":
                missing.append(cat_name)

        show_reassess = len(missing) > 0

        return jsonify({
            "show_reassess": show_reassess,
            "has_rating": has_rating,
            "missing_fields": missing,
        }), 200

    except Exception as e:
        # On any DB error, default to showing the button so the user is never blocked
        return jsonify({"show_reassess": True, "has_rating": False, "missing_fields": [], "error": str(e)}), 200


def _build_report_docx(candidate_name: str, criteria_record: dict, assessment_result: dict,
                       candidate_jobtitle: str = "") -> bytes:
    """Generate a well-structured Word document (.docx) assessment report.

    Uses python-docx tables for every data section so text is always contained
    within column boundaries — no overflow or alignment issues.

    Returns the raw .docx bytes.
    """
    import io as _io
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Cm
        from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
    except ImportError:
        raise RuntimeError("python-docx is not installed; cannot generate DOCX report")

    role_tag = criteria_record.get("role_tag", "")
    saved_at = criteria_record.get("saved_at", "")
    criteria = criteria_record.get("criteria") or {}

    # Format date as YYYY-MM-DD only (strip time component)
    date_str = str(saved_at)
    if "T" in date_str:
        date_str = date_str.split("T")[0]
    elif len(date_str) > 10:
        date_str = date_str[:10]

    # Use candidate's current job title when available, else fall back to role_tag
    display_role = candidate_jobtitle or role_tag

    doc = Document()

    # ── Criteria key → human-readable label (also used as appraisal category lookup) ──
    _CRIT_LABEL = {
        "job_titles": "Job Title",
        "job_title": "Job Title",
        "jobtitle_role_tag": "Job Title",
        "jobtitle": "Job Title",
        "country": "Country",
        "countries": "Country",
        "company": "Company",
        "companies": "Company",
        "sector": "Sector",
        "sectors": "Sector",
        "tenure": "Tenure",
        "min_tenure": "Tenure",
        "skills": "Skillset",
        "skillset": "Skillset",
        "seniority": "Seniority",
    }

    # ── Page margins ──────────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # ── Styles helpers ────────────────────────────────────────────────────────
    DARK_BLUE = RGBColor(0x12, 0x36, 0x5E)   # header bg approximated as font colour
    HDR_BG = "123660"                         # dark navy hex for table header shading
    ALT_BG = "EBF0F8"                         # light blue for alternating rows
    WHITE_BG = "FFFFFF"

    def _shade_cell(cell, hex_color):
        """Apply solid background shading to a table cell."""
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), hex_color)
        tcPr.append(shd)

    def _set_col_widths(table, widths_cm):
        """Set absolute column widths."""
        for row in table.rows:
            for idx, cell in enumerate(row.cells):
                if idx < len(widths_cm):
                    cell.width = Cm(widths_cm[idx])

    def _add_table(headers, rows_data, col_widths_cm=None):
        """Add a formatted table with a dark header row and alternating rows."""
        all_rows = [headers] + rows_data
        tbl = doc.add_table(rows=len(all_rows), cols=len(headers))
        tbl.style = "Table Grid"
        tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

        for r_idx, row_vals in enumerate(all_rows):
            is_hdr = r_idx == 0
            row = tbl.rows[r_idx]
            for c_idx, val in enumerate(row_vals):
                cell = row.cells[c_idx]
                cell.text = str(val) if val is not None else ""
                para = cell.paragraphs[0]
                run = para.runs[0] if para.runs else para.add_run(cell.text)
                run.text = str(val) if val is not None else ""
                run.font.size = Pt(9)
                run.font.bold = is_hdr
                if is_hdr:
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                    _shade_cell(cell, HDR_BG)
                elif r_idx % 2 == 0:
                    _shade_cell(cell, ALT_BG)
                else:
                    _shade_cell(cell, WHITE_BG)
                cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP

        if col_widths_cm:
            _set_col_widths(tbl, col_widths_cm)
        return tbl

    def _add_section_heading(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(text.upper())
        run.font.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = DARK_BLUE
        # Underline the section heading
        p2 = doc.add_paragraph()
        p2.paragraph_format.space_before = Pt(0)
        p2.paragraph_format.space_after = Pt(4)
        run2 = p2.add_run("─" * 60)
        run2.font.size = Pt(7)
        run2.font.color.rgb = DARK_BLUE

    # ── Title ─────────────────────────────────────────────────────────────────
    title_p = doc.add_paragraph()
    title_run = title_p.add_run("Assessment Report")
    title_run.font.bold = True
    title_run.font.size = Pt(18)
    title_run.font.color.rgb = DARK_BLUE
    title_p.paragraph_format.space_after = Pt(6)

    # ── Candidate info table ──────────────────────────────────────────────────
    _add_section_heading("Candidate Information")
    _add_table(
        ["Field", "Value"],
        [
            ["Candidate", candidate_name],
            ["Role", display_role],
            ["Date", date_str],
        ],
        col_widths_cm=[4.5, 12.5],
    )
    doc.add_paragraph()

    if assessment_result:
        # ── Assessment Summary ─────────────────────────────────────────────────
        stars = max(0, min(int(assessment_result.get("stars", 0) or 0), 5))
        star_str = ("★" * stars) + ("☆" * (5 - stars)) + f"  ({stars}/5)"
        overall = str(assessment_result.get("overall_comment", "") or "")
        # Ensure overall score shows as e.g. "96%" without duplicating the % symbol
        raw_score = str(assessment_result.get("total_score", "-")).rstrip("%")
        score_display = f"{raw_score}%" if raw_score != "-" else "-"

        _add_section_heading("Assessment Summary")
        summary_data = [
            ["Overall Score", score_display],
            ["Stars", star_str],
            ["Level", str(assessment_result.get("assessment_level", "-"))],
        ]
        if overall:
            summary_data.append(["Overall Comment", overall])
        _add_table(["Field", "Value"], summary_data, col_widths_cm=[4.5, 12.5])
        doc.add_paragraph()

        # ── Category Appraisals (build early so weights are available for criteria table) ──
        appraisals = assessment_result.get("category_appraisals") or {}

        def _get_weight_for_criteria(crit_key):
            """Return weight_percent string from category_appraisals for a criteria key."""
            ap_label = _CRIT_LABEL.get(crit_key.lower(), "")
            for ap_cat, ap_val in appraisals.items():
                if (ap_label and ap_cat.lower() == ap_label.lower()) or ap_cat.lower() == crit_key.lower():
                    if isinstance(ap_val, dict):
                        w = ap_val.get("weight_percent", "")
                        if w not in (None, ""):
                            return f"{w}%"
            return ""

        if criteria:
            _add_section_heading("Search Criteria")
            crit_data = []
            for k, v in criteria.items():
                if isinstance(v, list):
                    v_str = ", ".join(str(x) for x in v) if v else "-"
                else:
                    v_str = str(v) if v is not None else "-"
                display_name = _CRIT_LABEL.get(k.lower(), str(k))
                weight_val = _get_weight_for_criteria(k)
                crit_data.append([display_name, v_str, weight_val])
            _add_table(
                ["Criteria", "Value", "Weight"],
                crit_data,
                col_widths_cm=[4.0, 10.5, 2.5],
            )
            doc.add_paragraph()

        # ── Category Appraisals ────────────────────────────────────────────────
        if appraisals:
            _add_section_heading("Category Appraisals")
            # Build reverse mapping: display_name → criteria breakdown score
            criteria_breakdown = assessment_result.get("criteria") or {}
            _label_to_keys = {}
            for _k, _v in _CRIT_LABEL.items():
                _label_to_keys.setdefault(_v, []).append(_k)

            def _get_score_for_category(display_name):
                """Return actual computed score for a category from the criteria breakdown."""
                dn_lower = display_name.lower()
                try:
                    # Try direct lowercase match
                    if dn_lower in criteria_breakdown:
                        return str(round(float(criteria_breakdown[dn_lower]), 1))
                    # Try mapped internal keys
                    for _key in _label_to_keys.get(display_name, []):
                        if _key in criteria_breakdown:
                            return str(round(float(criteria_breakdown[_key]), 1))
                except (ValueError, TypeError):
                    pass
                return "-"

            ap_data = []
            for cat, appraisal in appraisals.items():
                if isinstance(appraisal, dict):
                    score_val = _get_score_for_category(str(cat))
                    rating = str(appraisal.get("rating", "") or "")
                    status = str(appraisal.get("status", "") or "")
                    comment = str(appraisal.get("comment", "") or "")
                    rating_status = f"{rating} / {status}" if status else rating
                    ap_data.append([str(cat), score_val, rating_status, comment])
                else:
                    ap_data.append([str(cat), "-", "-", str(appraisal)])
            _add_table(
                ["Category", "Score", "Rating / Status", "Comment"],
                ap_data,
                col_widths_cm=[3.5, 1.8, 3.2, 8.5],
            )
            doc.add_paragraph()

        # ── Verified Skillset (below Category Appraisals) ─────────────────────
        vskillset = assessment_result.get("vskillset")
        if vskillset:
            _add_section_heading("Verified Skillset")
            # vskillset may be a list of dicts or a single dict
            if isinstance(vskillset, dict):
                vskillset = [vskillset]
            if isinstance(vskillset, list) and vskillset:
                vs_data = []
                for item in vskillset:
                    if isinstance(item, dict):
                        skill = str(item.get("skill", ""))
                        prob = str(item.get("probability", ""))
                        if prob:
                            prob = prob.rstrip("%") + "%"
                        cat = str(item.get("category", ""))
                        reason = str(item.get("reason", ""))
                        vs_data.append([skill, prob, cat, reason])
                    else:
                        vs_data.append([str(item), "", "", ""])
                _add_table(
                    ["Skill", "Probability", "Category", "Reason"],
                    vs_data,
                    col_widths_cm=[3.0, 2.0, 2.0, 10.0],
                )
                doc.add_paragraph()

        # ── Conclusion (formerly Skill Comments) ───────────────────────────────
        comments_raw = assessment_result.get("comments")
        if comments_raw:
            _add_section_heading("Conclusion")
            if isinstance(comments_raw, str):
                paras = [p.strip() for p in comments_raw.split("\n") if p.strip()]
                if paras:
                    _add_table(["Comments"], [[p] for p in paras], col_widths_cm=[17.0])
            elif isinstance(comments_raw, (list, tuple)):
                sc_data = []
                for entry in comments_raw:
                    if isinstance(entry, dict):
                        skill = str(entry.get("skill") or entry.get("category", ""))
                        match = str(entry.get("match") or entry.get("status", ""))
                        note = str(entry.get("comment") or entry.get("note", ""))
                        sc_data.append([skill, match, note])
                    else:
                        sc_data.append([str(entry), "", ""])
                _add_table(["Skill", "Status", "Comment"], sc_data, col_widths_cm=[4.0, 3.0, 10.0])
            doc.add_paragraph()

    elif criteria:
        # No assessment yet — still show Search Criteria
        _add_section_heading("Search Criteria")
        crit_data = []
        for k, v in criteria.items():
            if isinstance(v, list):
                v_str = ", ".join(str(x) for x in v) if v else "-"
            else:
                v_str = str(v) if v is not None else "-"
            display_name = _CRIT_LABEL.get(k.lower(), str(k))
            crit_data.append([display_name, v_str])
        _add_table(["Criteria", "Value"], crit_data, col_widths_cm=[5.5, 11.5])

    buf = _io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _find_assessment_for_candidate_by_id(process_id: int):
    """Find the latest assessment for a no-LinkedIn candidate by process table id.

    Scans bulk result files for a record whose 'id' or 'process_id' matches.
    Falls back to the DB `rating` column keyed by process primary key.
    """
    if not process_id:
        return None
    # Scan bulk result files
    if os.path.isdir(OUTPUT_DIR):
        try:
            bulk_files = sorted(
                [f for f in os.listdir(OUTPUT_DIR) if f.endswith("_results.json")],
                key=lambda f: os.path.getmtime(os.path.join(OUTPUT_DIR, f)),
                reverse=True,
            )
            for fname in bulk_files:
                fpath = os.path.join(OUTPUT_DIR, fname)
                try:
                    with open(fpath, "r", encoding="utf-8") as fh:
                        records = json.load(fh)
                    if not isinstance(records, list):
                        continue
                    for rec in records:
                        if not isinstance(rec, dict):
                            continue
                        rec_id = rec.get("id") or rec.get("process_id")
                        if rec_id and int(rec_id) == process_id:
                            result = rec.get("result")
                            if result and not result.get("_skipped"):
                                if "vskillset" not in result and rec.get("vskillset"):
                                    result = dict(result)
                                    result["vskillset"] = rec.get("vskillset")
                                return _enrich_assessment_with_db_vskillset(result, process_id=process_id)
                except Exception:
                    continue
        except Exception:
            pass
    # DB fallback — read rating column by process primary key
    try:
        conn = _pg_connect()
        try:
            cur = conn.cursor()
            cur.execute(
                "SELECT rating FROM process WHERE id = %s AND rating IS NOT NULL"
                " ORDER BY rating_updated_at DESC NULLS LAST LIMIT 1",
                (process_id,)
            )
            row = cur.fetchone()
            cur.close()
            if row and row[0]:
                rating = row[0]
                if isinstance(rating, str):
                    try:
                        rating = json.loads(rating)
                    except Exception:
                        rating = None
                if isinstance(rating, dict) and not rating.get("_skipped"):
                    return _enrich_assessment_with_db_vskillset(rating, process_id=process_id)
        finally:
            conn.close()
    except Exception:
        pass
    return None


def _find_assessment_for_candidate_by_name(candidate_name: str):
    """Find the latest assessment for a candidate by name (last-resort fallback).

    Only used when neither LinkedIn URL nor process_id is available.
    """
    if not candidate_name:
        return None
    norm_name = candidate_name.strip().lower()
    if os.path.isdir(OUTPUT_DIR):
        try:
            bulk_files = sorted(
                [f for f in os.listdir(OUTPUT_DIR) if f.endswith("_results.json")],
                key=lambda f: os.path.getmtime(os.path.join(OUTPUT_DIR, f)),
                reverse=True,
            )
            for fname in bulk_files:
                fpath = os.path.join(OUTPUT_DIR, fname)
                try:
                    with open(fpath, "r", encoding="utf-8") as fh:
                        records = json.load(fh)
                    if not isinstance(records, list):
                        continue
                    for rec in records:
                        if not isinstance(rec, dict):
                            continue
                        rec_name = (rec.get("name") or "").strip().lower()
                        if rec_name and rec_name == norm_name:
                            result = rec.get("result")
                            if result and not result.get("_skipped"):
                                if "vskillset" not in result and rec.get("vskillset"):
                                    result = dict(result)
                                    result["vskillset"] = rec.get("vskillset")
                                # Enrich using the process_id embedded in the record (if any)
                                _rec_pid = rec.get("process_id") or rec.get("id")
                                return _enrich_assessment_with_db_vskillset(
                                    result, process_id=_rec_pid if _rec_pid else None
                                )
                except Exception:
                    continue
        except Exception:
            pass
    return None


@app.get("/sourcing/download_report")
def download_report():
    """Generate and download a formal assessment report as a Word document (.docx).

    Combines the criteria JSON and bulk/individual assessment results into one document.
    The generated file is also saved to REPORT_TEMPLATES_DIR for record-keeping.

    Query params:
        linkedin   – candidate LinkedIn URL (optional for no-LinkedIn records)
        name       – candidate name
        process_id – process table primary key (used when linkedin is empty)
    """
    name = (request.args.get("name") or "").strip()
    linkedin = (request.args.get("linkedin") or "").strip()
    process_id_str = (request.args.get("process_id") or "").strip()
    process_id = int(process_id_str) if process_id_str.isdigit() else None
    if not name and not linkedin and not process_id:
        return "name or linkedin required", 400
    # Look up name and linkedin from DB by process_id when LinkedIn URL is absent
    if process_id and not linkedin:
        try:
            conn = _pg_connect()
            try:
                cur = conn.cursor()
                cur.execute(
                    "SELECT name, linkedinurl FROM process WHERE id = %s LIMIT 1",
                    (process_id,)
                )
                row = cur.fetchone()
                cur.close()
                if row:
                    if not name and row[0]:
                        name = row[0].replace("님", "").strip()
                    if not linkedin and row[1]:
                        linkedin = row[1].strip()
            finally:
                conn.close()
        except Exception as exc:
            logger.warning(f"[download_report] process_id lookup failed: {exc}")
    # Look up name from DB if missing (LinkedIn known)
    if not name and linkedin:
        try:
            conn = _pg_connect()
            try:
                cur = conn.cursor()
                cur.execute(
                    "SELECT name FROM sourcing WHERE linkedinurl=%s AND name IS NOT NULL LIMIT 1",
                    (linkedin,)
                )
                row = cur.fetchone()
                cur.close()
                if row and row[0]:
                    name = row[0].replace("님", "").strip()
            finally:
                conn.close()
        except Exception as exc:
            logger.warning(f"[download_report] DB lookup failed: {exc}")
    if not name:
        return "No candidate name found", 404
    # Fetch candidate's current job title from DB for the report
    candidate_jobtitle = ""
    if linkedin:
        try:
            conn = _pg_connect()
            try:
                cur = conn.cursor()
                cur.execute(
                    "SELECT jobtitle FROM sourcing WHERE linkedinurl=%s AND jobtitle IS NOT NULL"
                    " ORDER BY id DESC LIMIT 1",
                    (linkedin,)
                )
                row = cur.fetchone()
                cur.close()
                if row and row[0]:
                    candidate_jobtitle = row[0].strip()
            finally:
                conn.close()
        except Exception as exc:
            logger.warning(f"[download_report] jobtitle lookup failed: {exc}")
    elif process_id:
        try:
            conn = _pg_connect()
            try:
                cur = conn.cursor()
                cur.execute(
                    "SELECT jobtitle FROM process WHERE id = %s AND jobtitle IS NOT NULL LIMIT 1",
                    (process_id,)
                )
                row = cur.fetchone()
                cur.close()
                if row and row[0]:
                    candidate_jobtitle = row[0].strip()
            finally:
                conn.close()
        except Exception as exc:
            logger.warning(f"[download_report] process_id jobtitle lookup failed: {exc}")
    # When process_id is known, look up the candidate's current role_tag+username from
    # the process table and use that to find the correct criteria file — this prevents
    # the name-based scan from returning the wrong criteria file (e.g. cloud) when the
    # candidate's name also appears in a file for a different role.
    fpath = None
    criteria_record = None
    _rt_username = None
    _rt_role_tag = None
    if process_id or linkedin:
        try:
            _rt_conn = _pg_connect()
            try:
                _rt_cur = _rt_conn.cursor()
                if process_id:
                    _rt_cur.execute(
                        "SELECT role_tag, username FROM process WHERE id = %s LIMIT 1",
                        (process_id,)
                    )
                else:
                    _rt_cur.execute(
                        "SELECT role_tag, username FROM process WHERE linkedinurl = %s LIMIT 1",
                        (linkedin,)
                    )
                _rt_row = _rt_cur.fetchone()
                _rt_cur.close()
                if _rt_row:
                    _rt_role_tag = (_rt_row[0] or "").strip()
                    _rt_username = (_rt_row[1] or "").strip()
            finally:
                _rt_conn.close()
        except Exception as _rt_exc:
            logger.warning(f"[download_report] role_tag lookup failed: {_rt_exc}")
    if _rt_role_tag and _rt_username and os.path.isdir(CRITERIA_OUTPUT_DIR):
        # Scan criteria files: prefer exact role_tag match, then substring/prefix match
        try:
            for _cfname in os.listdir(CRITERIA_OUTPUT_DIR):
                if not _cfname.endswith(".json"):
                    continue
                _cfpath = os.path.join(CRITERIA_OUTPUT_DIR, _cfname)
                try:
                    with open(_cfpath, "r", encoding="utf-8") as _cfh:
                        _cfrec = json.load(_cfh)
                    _cf_rt = (_cfrec.get("role_tag") or "").strip()
                    _cf_un = (_cfrec.get("username") or "").strip()
                    if _cf_un == _rt_username and (
                        _cf_rt == _rt_role_tag
                        or any(
                            r.strip().lower() == _cf_rt.lower()
                            for r in _rt_role_tag.split(",")
                        )
                        or any(
                            r.strip().lower() == _rt_role_tag.lower()
                            for r in _cf_rt.split(",")
                        )
                    ):
                        criteria_record = _cfrec
                        fpath = _cfpath
                        break
                except Exception:
                    continue
        except Exception as _scan_exc:
            logger.warning(f"[download_report] criteria scan failed: {_scan_exc}")
    # Fallback: scan criteria files by candidate name (last resort)
    if not criteria_record:
        fpath, criteria_record = _find_criteria_file_for_candidate(name)
    if not criteria_record:
        return "No criteria file found for this candidate", 404
    assessment_result = None
    if linkedin:
        assessment_result = _find_assessment_for_candidate(linkedin)
    if assessment_result is None and process_id:
        assessment_result = _find_assessment_for_candidate_by_id(process_id)
    if assessment_result is None and name:
        assessment_result = _find_assessment_for_candidate_by_name(name)
    try:
        docx_bytes = _build_report_docx(name, criteria_record, assessment_result or {}, candidate_jobtitle=candidate_jobtitle)
    except Exception as exc:
        logger.exception("[download_report] DOCX generation failed")
        return f"Report generation failed: {exc}", 500
    safe_name = re.sub(r'[<>:"/\\|?*]', '_', name)
    safe_role = re.sub(r'[<>:"/\\|?*]', '_', criteria_record.get("role_tag", "report"))
    fname = f"{safe_name} {safe_role}.docx"
    # Persist to REPORT_TEMPLATES_DIR (already created at startup)
    try:
        out_path = os.path.join(REPORT_TEMPLATES_DIR, fname)
        with open(out_path, "wb") as fh:
            fh.write(docx_bytes)
        logger.info(f"[download_report] Saved report to {out_path}")
    except Exception as exc:
        logger.warning(f"[download_report] Could not save to templates dir: {exc}")
    from flask import Response as _Response
    return _Response(
        docx_bytes,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{fname}"'},
    )