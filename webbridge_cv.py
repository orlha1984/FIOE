# webbridge_cv.py — Second-half routes for webbridge.py.
# Contains: job runner, sourcing, process management, CV processing, and bulk assessment.
# This module is imported at the bottom of webbridge.py after all shared state is defined.
# Circular import is safe because webbridge is already in sys.modules by the time this runs.

import logging
import os
import sys
import re
import json
import threading
import time
import uuid
import io
import hashlib
from csv import DictWriter
from datetime import datetime
from flask import request, send_from_directory, jsonify, abort, Response, stream_with_context

# Sanitise username for use in filenames (allow only alphanumeric, _ and -)
_CV_USERNAME_SAFE_RE = re.compile(r'[^A-Za-z0-9_-]')

# Structured activity logger
try:
    from app_logger import log_agentic, log_approval, log_error as _log_error_cv
    _CV_LOGGER_AVAILABLE = True
except ImportError:
    _CV_LOGGER_AVAILABLE = False
    def log_agentic(**_kw): pass
    def log_approval(**_kw): pass
    def _log_error_cv(**_kw): pass

# ---------------------------------------------------------------------------
# __main__ / module-name fix
# When webbridge.py is invoked directly (`python webbridge.py`) Python registers
# it as '__main__', NOT as 'webbridge'.  Without this alias, the import below
# would trigger a second, independent load of webbridge.py, running all
# initialisation twice and eventually hitting a circular-import error when the
# second load again tries to `import webbridge_cv`.
# By registering the already-running __main__ under the 'webbridge' key before
# the import, Python reuses the partially-loaded module instead of re-loading.
if 'webbridge' not in sys.modules:
    _main = sys.modules.get('__main__')
    if _main is not None and os.path.basename(os.path.normpath(getattr(_main, '__file__', ''))) == 'webbridge.py':
        sys.modules['webbridge'] = _main
# ---------------------------------------------------------------------------

# Import shared state and utilities from webbridge (master module).
# All names below are defined in webbridge.py before webbridge_routes is imported,
# so they are fully initialized by the time this module runs.
from webbridge import (
    app, logger, genai,
    BASE_DIR, OUTPUT_DIR, SEARCH_XLS_DIR,
    GOOGLE_CSE_API_KEY, GOOGLE_CSE_CX, SEARCH_RESULTS_TARGET,
    GEMINI_API_KEY, GEMINI_SUGGEST_MODEL,
    SINGAPORE_CONTEXT, SEARCH_RULES,
    CV_TRANSLATION_MAX_CHARS, LANG_DETECTION_SAMPLE_LENGTH, CV_ANALYSIS_MAX_CHARS,
    MAX_COMMENT_LENGTH, COMMENT_TRUNCATE_LENGTH,
    ASSESSMENT_EXCELLENT_THRESHOLD, ASSESSMENT_GOOD_THRESHOLD, ASSESSMENT_MODERATE_THRESHOLD,
    CITY_TO_COUNTRY_DATA,
    _CV_ANALYZE_SEMAPHORE, _SINGLE_FILE_MAX,
    _rate, _check_user_rate, _csrf_required,
    _is_pdf_bytes,
    _extract_json_object, _extract_confirmed_skills,
    translate_text_pipeline,
    _infer_region_from_country,
    _find_best_sector_match_for_text, _map_keyword_to_sector_label,
    _resolve_sector_from_ml_master,
    _resolve_jobfamily_seniority_from_ml_master,
    _compute_search_target,
    _should_overwrite_existing, _ensure_rating_metadata_columns, _ensure_search_indexes,
    _persist_jskillset, _fetch_jskillset, _fetch_jskillset_from_process,
    _sync_login_jskillset_to_process,
    _sync_criteria_jskillset_to_process,
    _increment_cse_query_count,
    _increment_gemini_query_count,
    _load_rate_limits,
    _make_flask_limit,
)

# Items that moved to webbridge_routes.py (originally lines 5152-9431 of webbridge.py).
# webbridge_routes is imported at the bottom of webbridge.py, so it is always
# in sys.modules and fully initialised up to this point by the time this line runs.
from webbridge_routes import (
    JOBS, JOBS_LOCK,
    COMPANY_SUGGESTIONS_LIMIT,
    _clean_list, _enforce_company_limit, _heuristic_job_suggestions,
    _infer_primary_job_title, _infer_seniority_from_titles, _perform_cse_queries,
    is_linkedin_profile, parse_linkedin_title,
    get_linkedin_profile_picture, fetch_image_bytes_from_url,
    add_message, persist_job,
    _sanitize_for_excel, _aggregate_company_dropdown,
    _extract_company_from_jobtitle, _gemini_extract_company_from_jobtitle,
    _role_tag_session_column_ensured,
    CRITERIA_OUTPUT_DIR, _get_criteria_filepath, _read_search_criteria,
    unified_llm_call_text,
)


def _bulk_assess_flask_limit():
    """Return the Flask-Limiter rate string for bulk_assess, read from admin config.

    Uses the dedicated 'bulk_assess' key so the bulk assessment endpoint can be
    granted a higher limit than 'start_job' without affecting sourcing job limits.
    This prevents rate limiting from causing inconsistent vskillset coverage when
    many records need to be assessed in a single DB Dockin – DB Analytic session.
    """
    try:
        cfg  = _load_rate_limits()
        feat = cfg.get("defaults", {}).get("bulk_assess", {})
        req  = int(feat.get("requests", 30))
        win  = int(feat.get("window_seconds", 60))
        return f"{req} per {win} second"
    except Exception:
        return "30 per minute"


def _geography_flask_limit():
    """Return the Flask-Limiter rate string for geography lookup, read from rate_limits.json.
    Keeps the Flask-Limiter hard cap in sync with the admin-configured geography budget
    so the server-side limit matches whatever is set in the admin panel."""
    try:
        cfg  = _load_rate_limits()
        feat = cfg.get("defaults", {}).get("geography", {})
        req  = int(feat.get("requests", 200))
        win  = int(feat.get("window_seconds", 60))
        return f"{req} per {win} second"
    except Exception:
        return "200 per minute"


def _job_runner(job_id, queries, fallback_queries, auto_expand, manual_urls, search_results_only, country, dynamic_target, job_titles,
                user_search_provider=None, user_serper_key=None, user_dfs_login=None, user_dfs_password=None):
    global SEARCH_RESULTS_TARGET
    add_message(job_id, "Starting search pipeline...")
    rows=[]; urls=[]
    cse_queries_fired = 0
    if isinstance(dynamic_target, int) and dynamic_target > 0:
        SEARCH_RESULTS_TARGET = dynamic_target
    target_limit = SEARCH_RESULTS_TARGET
    primary_job_title=_infer_primary_job_title(job_titles)
    try:
        executed_primary=False
        if (search_results_only or auto_expand):
            # Allow search to proceed when a per-user search provider is configured
            # even if the global CSE keys are not set.
            _has_user_search = (
                (user_search_provider == 'serper' and user_serper_key) or
                (user_search_provider == 'dataforseo' and user_dfs_login and user_dfs_password)
            )
            if not _has_user_search and (not GOOGLE_CSE_API_KEY or not GOOGLE_CSE_CX):
                add_message(job_id, "ERROR: GOOGLE_CSE_API_KEY/CX not set. Cannot run search.")
            else:
                executed_primary=True
                primary_q = queries or ["site:linkedin.com/in"]
                cse_results=_perform_cse_queries(
                    job_id, primary_q, target_limit, country,
                    user_provider=user_search_provider, user_serper_key=user_serper_key,
                    user_dfs_login=user_dfs_login, user_dfs_password=user_dfs_password,
                )
                cse_queries_fired += len(primary_q)
                urls=[r["link"] for r in cse_results]
                with JOBS_LOCK:
                    JOBS[job_id]['urls']=urls
                    JOBS[job_id]['progress']['total']=len(urls)
                persist_job(job_id)
                processed=0
                for item in cse_results:
                    link=item.get("link","") or ""
                    title=item.get("title") or ""
                    domain_part=re.sub(r'^https?://','',link+'/').split('/')[0].lower()
                    if is_linkedin_profile(link):
                        name, jobtitle, company = parse_linkedin_title(title)
                        if name or jobtitle or company:
                            rows.append({"Name":name or "", "Company":company or "", "JobTitle":jobtitle or "", "Country":country or "", "LinkedInURL":link})
                    processed+=1
                    with JOBS_LOCK:
                        JOBS[job_id]['progress']['processed']=processed
                    if processed % 15 == 0: persist_job(job_id)
        # Run fallback (job-title-only) queries whenever primary has not yet reached the
        # target limit -- not just when it returned zero rows.  This ensures remaining
        # slots are always filled with the best available candidates, regardless of how
        # many the company-filtered queries found.
        if executed_primary and fallback_queries and len(rows) < target_limit:
            # still_needed_fb is always positive here because the condition above guards it
            still_needed_fb = target_limit - len(rows)
            add_message(job_id, f"Primary collected {len(rows)}/{target_limit}. Running fallback (job-title only) queries for remaining {still_needed_fb}...")
            cse_results=_perform_cse_queries(
                job_id, fallback_queries, still_needed_fb, country,
                user_provider=user_search_provider, user_serper_key=user_serper_key,
                user_dfs_login=user_dfs_login, user_dfs_password=user_dfs_password,
            )
            cse_queries_fired += len(fallback_queries)
            # Collect already-seen URLs so duplicates from fallback are discarded
            existing_urls = {r.get("LinkedInURL","") for r in rows}
            new_urls=[r["link"] for r in cse_results if r.get("link") not in existing_urls]
            with JOBS_LOCK:
                JOBS[job_id]['urls']=list(dict.fromkeys(JOBS[job_id]['urls']+new_urls))
                JOBS[job_id]['progress']['total']=len(JOBS[job_id]['urls'])
            persist_job(job_id)
            processed=0
            for item in cse_results:
                link=item.get("link"); title=item.get("title")
                if not link or link in existing_urls:
                    continue
                if is_linkedin_profile(link):
                    name, jobtitle, company = parse_linkedin_title(title)
                    if name or jobtitle or company:
                        rows.append({"Name":name or "", "Company":company or "", "JobTitle":jobtitle or "", "Country":country or "", "LinkedInURL":link})
                        existing_urls.add(link)
                processed+=1
                with JOBS_LOCK:
                    JOBS[job_id]['progress']['processed']=processed
                if processed % 15 == 0: persist_job(job_id)
        if manual_urls and not (search_results_only or auto_expand):
            add_message(job_id, f"Processing manual URLs: {len(manual_urls)}")
            for u in manual_urls:
                domain_part=re.sub(r'^https?://','',u).split('/')[0].lower()
                if is_linkedin_profile(u):
                    rows.append({"Name":"","Company":"","JobTitle":primary_job_title,"Country":country or "","LinkedInURL":u})
            with JOBS_LOCK:
                JOBS[job_id]['urls']=list(dict.fromkeys(JOBS[job_id]['urls']+manual_urls))
                JOBS[job_id]['progress']['total']=len(JOBS[job_id]['urls'])
                JOBS[job_id]['progress']['processed']=len(JOBS[job_id]['urls'])
            persist_job(job_id)
        dedup=[]; seen=set()
        for r in rows:
            key=(r.get("LinkedInURL",""), r.get("Name","").lower(), r.get("JobTitle","").lower())
            if key in seen: continue
            seen.add(key); dedup.append(r)
        rows=dedup
        with JOBS_LOCK:
            meta=(JOBS.get(job_id) or {}).get('meta',{})
        seniority_effective=meta.get('seniority') or _infer_seniority_from_titles(job_titles)
        if seniority_effective:
            before=len(rows)
            srules=(SEARCH_RULES or {}).get("seniority_rules") or {}
            excl_rule=None
            for k in srules.keys():
                if str(k).strip().lower()==seniority_effective.lower():
                    excl_rule=srules.get(k); break
            if excl_rule and isinstance(excl_rule,dict):
                tokens_raw=excl_rule.get("xrayExclusion","")
                tokens=[]
                m=re.search(r"\((.*?)\)", tokens_raw)
                if m:
                    tokens=[t.strip().strip('"\'') for t in m.group(1).split("OR")]
                elif tokens_raw:
                    m2=re.search(r"-\s*([A-Za-z ]+)", tokens_raw)
                    if m2: tokens=[m2.group(1).strip()]
                filtered=[]
                lowers=[t.lower() for t in tokens if t]
                if lowers:
                    for row in rows:
                        jt=(row.get("JobTitle") or "").lower()
                        if not any(x in jt for x in lowers):
                            filtered.append(row)
                    rows=filtered
            after=len(rows)
            add_message(job_id, f"Seniority filter '{seniority_effective}' applied: kept {after}/{before} rows.")
        # Persist CSE query count for the user who triggered this job
        if cse_queries_fired > 0:
            with JOBS_LOCK:
                job_username = (JOBS.get(job_id) or {}).get('username', '')
            if job_username:
                _increment_cse_query_count(job_username, cse_queries_fired)
        csv_name,xlsx_name=_write_outputs(job_id, rows)
        # CSV/XLSX both saved to SEARCH_XLS_DIR now
        csv_ok=bool(csv_name) and os.path.exists(os.path.join(SEARCH_XLS_DIR, csv_name))
        xlsx_ok=bool(xlsx_name) and os.path.exists(os.path.join(SEARCH_XLS_DIR, xlsx_name))
        with JOBS_LOCK:
            job=JOBS.get(job_id)
            if job:
                job['output_csv']=csv_name
                job['output_xlsx']=xlsx_name
                job['done']=bool(csv_ok or xlsx_ok)
                job['messages'].append(f"Job complete. Wrote {len(rows)} rows.")
                job['status_html']="<br>".join(job['messages'][-12:])
        persist_job(job_id)
    except Exception as e:
        add_message(job_id, f"Pipeline error: {e}")
        with JOBS_LOCK:
            job=JOBS.get(job_id)
            if job:
                job['done']=False
                job['status_html']="<br>".join(job['messages'][-12:])
        persist_job(job_id)

def _write_outputs(job_id, rows):
    with JOBS_LOCK:
        job_meta=(JOBS.get(job_id) or {}).get('meta',{})
        job_top=(JOBS.get(job_id) or {})
    dropdown_companies=_aggregate_company_dropdown(job_meta)
    processed=[]
    for r in rows:
        link=r.get("LinkedInURL","")
        country_val=r.get("Country","")
        if not is_linkedin_profile(link): country_val=""
        raw_name=r.get("Name",""); raw_company=r.get("Company",""); raw_job=r.get("JobTitle","")
        moved_company, adjusted_job=_extract_company_from_jobtitle(raw_job, raw_company, dropdown_companies)
        if not moved_company:
            g_company,g_job=_gemini_extract_company_from_jobtitle(raw_job, dropdown_companies)
            if g_company:
                moved_company=g_company.strip()
                adjusted_job=(g_job or raw_job).strip()
        name_val=_sanitize_for_excel(raw_name)
        job_val=_sanitize_for_excel(adjusted_job)
        company_val=(moved_company or "").strip()
        processed.append({"Name":name_val,"Company":company_val,"JobTitle":job_val,"Country":country_val,"LinkedInURL":link})
    csv_name=f"{job_id}_results.csv"
    csv_path=os.path.join(SEARCH_XLS_DIR, csv_name)
    # Ensure target dir exists
    os.makedirs(SEARCH_XLS_DIR, exist_ok=True)
    with open(csv_path,"w",encoding="utf-8",newline="") as f:
        w=DictWriter(f, fieldnames=["Name","Company","JobTitle","Country","LinkedInURL"])
        w.writeheader()
        for pr in processed: w.writerow(pr)
        try:
            f.flush(); os.fsync(f.fileno())
        except Exception:
            pass
    xlsx_name=None
    try:
        import openpyxl
        from openpyxl import Workbook
        from openpyxl.worksheet.datavalidation import DataValidation
        wb=Workbook(); ws=wb.active
        ws.title="Results"
        ws.append(["Name","Company","JobTitle","Country","LinkedInURL"])
        for pr in processed: ws.append([pr["Name"],pr["Company"],pr["JobTitle"],pr["Country"],pr["LinkedInURL"]])
        company_dropdown=dropdown_companies
        if company_dropdown:
            cs=wb.create_sheet(title="Companies")
            cs.append(["Companies"])
            for c in company_dropdown: cs.append([c])
            last_row_comp=len(company_dropdown)+1
            dv=DataValidation(type="list", formula1=f"=Companies!$A$2:$A${last_row_comp}", allow_blank=True,
                              showErrorMessage=True, error="Select a company from the dropdown list.", errorTitle="Invalid Company")
            ws.add_data_validation(dv)
            last_results_row=ws.max_row
            dv.add(f"B2:B{last_results_row}")
        xlsx_name=f"{job_id}_results.xlsx"
        # Save to SEARCH_XLS_DIR
        xlsx_full=os.path.join(SEARCH_XLS_DIR, xlsx_name)
        wb.save(xlsx_full)
        for _ in range(20):
            try:
                if os.path.exists(xlsx_full) and os.path.getsize(xlsx_full)>0:
                    break
            except Exception:
                pass
            time.sleep(0.05)
        try:
            import psycopg2
            from psycopg2 import sql
            excel_path=xlsx_full
            wb_ing=openpyxl.load_workbook(excel_path, read_only=True, data_only=True)
            sheet=wb_ing["Results"]
            headers=[cell.value for cell in next(sheet.iter_rows(min_row=1, max_row=1))]
            expected=["Name","Company","JobTitle","Country","LinkedInURL"]
            if headers!=expected:
                logger.warning(f"[Ingest] Header mismatch. Expected {expected}, got {headers}. Skipping ingestion.")
            else:
                data_rows=[]
                for row in sheet.iter_rows(min_row=2, values_only=True):
                    if all((val is None or str(val).strip()=="" ) for val in row): continue
                    data_rows.append(row)
                if data_rows:
                    pg_host=os.getenv("PGHOST","localhost")
                    pg_port=int(os.getenv("PGPORT","5432"))
                    pg_user=os.getenv("PGUSER","postgres")
                    pg_password=os.getenv("PGPASSWORD", "")
                    pg_db=os.getenv("PGDATABASE","candidate_db")
                    logger.info(f"[Ingest] Connecting to Postgres host={pg_host} port={pg_port} db={pg_db} user={pg_user}")
                    conn=psycopg2.connect(host=pg_host, port=pg_port, user=pg_user, password=pg_password, dbname=pg_db)
                    conn.autocommit=False
                    cur=conn.cursor()
                    
                    # Check if pic column exists in sourcing table
                    cur.execute("""
                        SELECT column_name 
                        FROM information_schema.columns 
                        WHERE table_schema='public' AND table_name='sourcing' AND column_name='pic'
                    """)
                    has_pic_column = bool(cur.fetchone())
                    
                    active_userid=(job_meta.get('userid') or job_top.get('userid') or '').strip()
                    active_username=(job_meta.get('username') or job_top.get('username') or '').strip()
                    
                    if has_pic_column:
                        # Include pic column in insert
                        insert_stmt=sql.SQL("INSERT INTO sourcing (userid, username, name, company, jobtitle, country, linkedinurl, pic) VALUES (%s, %s, %s, %s, %s, %s, %s, %s) ON CONFLICT DO NOTHING")
                        batch_rows = []
                        for r in data_rows:
                            # r is a tuple: (name, company, jobtitle, country, linkedinurl)
                            # LinkedInURL is at index 4 (0-based indexing, 5th column)
                            linkedin_url = r[4]
                            # Retrieve profile picture and convert to bytea
                            pic_bytes = None
                            try:
                                pic_url = get_linkedin_profile_picture(linkedin_url, display_name=r[0]) if linkedin_url else None
                                if pic_url:
                                    pic_bytes = fetch_image_bytes_from_url(pic_url)
                                    if pic_bytes:
                                        import psycopg2
                                        pic_bytes = psycopg2.Binary(pic_bytes)
                                    else:
                                        # Byte fetch failed — store URL as bytes so client can try direct load
                                        import psycopg2
                                        pic_bytes = psycopg2.Binary(pic_url.encode('utf-8'))
                            except Exception as pic_err:
                                logger.warning(f"[Ingest] Failed to get profile pic for {linkedin_url}: {pic_err}")
                            batch_rows.append((active_userid, active_username, r[0], r[1], r[2], r[3], linkedin_url, pic_bytes))
                    else:
                        # No pic column, use original insert
                        insert_stmt=sql.SQL("INSERT INTO sourcing (userid, username, name, company, jobtitle, country, linkedinurl) VALUES (%s, %s, %s, %s, %s, %s, %s) ON CONFLICT DO NOTHING")
                        batch_rows=[(active_userid, active_username, r[0], r[1], r[2], r[3], r[4]) for r in data_rows]
                    
                    batch_size=500
                    total_inserted=0
                    for i in range(0,len(batch_rows),batch_size):
                        batch=batch_rows[i:i+batch_size]
                        cur.executemany(insert_stmt, batch)
                        total_inserted+=len(batch)
                    conn.commit()
                    logger.info(f"[Ingest] Inserted {total_inserted} rows into sourcing (userid='{active_userid}' username='{active_username}').")
                    # Transfer role_tag from login table into sourcing table for this user
                    if active_username:
                        try:
                            cur.execute("SELECT role_tag, session FROM login WHERE username=%s LIMIT 1", (active_username,))
                            rt_row = cur.fetchone()
                            login_role_tag = rt_row[0] if rt_row and rt_row[0] else ""
                            login_session_ts = rt_row[1] if rt_row else None
                            if login_role_tag:
                                cur.execute("ALTER TABLE sourcing ADD COLUMN IF NOT EXISTS role_tag TEXT DEFAULT ''")
                                cur.execute("UPDATE sourcing SET role_tag=%s WHERE username=%s AND (role_tag IS NULL OR role_tag='')", (login_role_tag, active_username))
                                # Transfer session timestamp from login to sourcing after validating role_tag matches.
                                if login_session_ts is not None:
                                    cur.execute("ALTER TABLE sourcing ADD COLUMN IF NOT EXISTS session TIMESTAMPTZ")
                                    cur.execute("ALTER TABLE sourcing ALTER COLUMN session DROP DEFAULT")
                                    cur.execute(
                                        "UPDATE sourcing SET session=%s WHERE username=%s AND role_tag=%s",
                                        (login_session_ts, active_username, login_role_tag)
                                    )
                                conn.commit()
                                logger.info(f"[Ingest] Transferred role_tag='{login_role_tag}' session_ts='{login_session_ts}' from login to sourcing for user='{active_username}'.")
                        except Exception as e_rt:
                            logger.warning(f"[Ingest] Failed to transfer role_tag to sourcing: {e_rt}")
                    cur.close(); conn.close()
                else:
                    logger.info(f"[Ingest] No data rows to insert from {xlsx_name}.")
        except Exception as e:
            logger.warning(f"[Ingest] PostgreSQL ingestion failed: {e}")
    except Exception as e:
        logger.warning(f"[Excel Dropdown] XLSX generation failed: {e}")
    return csv_name, xlsx_name

def _gemini_multi_sector(selected, user_job_title, user_company, languages=None):
    languages = languages or []
    region_hint="SG/APAC" if SINGAPORE_CONTEXT else None
    input_obj={"selectedSectors": selected, "userJobTitle": user_job_title or None, "userCompany": user_company or None, "languages": languages, "regionHint": region_hint}
    prompt=("SYSTEM:\nYou are a sourcing assistant integrated into an application. Generate concise suggestions.\n"
            "Return ONLY JSON: {\"job\":{\"related\":[...]},\"company\":{\"related\":[...]}}\n"
            f"- Provide EXACTLY 15 real job titles (or fill with closest relevant if fewer) in job.related.\n"
            f"- Provide EXACTLY {COMPANY_SUGGESTIONS_LIMIT} real company names (brand-level) in company.related.\n"
            "- STRICT SECTOR RULE for company.related: ONLY include companies whose PRIMARY BUSINESS and CORE\n"
            "  OPERATIONS fall within the selected sector(s). EXCLUDE any company from a different industry\n"
            "  that merely uses or purchases services in those sectors. Examples of what to exclude:\n"
            "  * For Gaming / Technology sectors: do NOT include pharma, healthcare, finance, insurance, or\n"
            "    manufacturing companies, even if they use software or hire engineers internally.\n"
            "  * For Healthcare / Clinical Research sectors: do NOT include gaming, tech, or retail companies.\n"
            "  Competitors must share the same product/service focus as companies already in the sector.\n"
            "- NO generic placeholders (e.g., 'Tech Company', 'Gaming Studio').\n"
            "- NO commentary or extra keys.\n"
            f"INPUT:\n{json.dumps(input_obj,ensure_ascii=False)}\nJSON:")
    try:
        resp_text = unified_llm_call_text(prompt)
        parsed=_extract_json_object(resp_text or "")
        if not isinstance(parsed, dict): return None
        job=parsed.get("job",{}) if isinstance(parsed.get("job"),dict) else {}
        comp=parsed.get("company",{}) if isinstance(parsed.get("company"),dict) else {}
        jr=job.get("related") if isinstance(job.get("related"),list) else []
        cr=comp.get("related") if isinstance(comp.get("related"),list) else []
        jr_clean=_clean_list([s for s in jr if isinstance(s,str)], 15)
        if len(jr_clean) < 15:
            extra=_heuristic_job_suggestions(jr_clean, "Non-Gaming", languages, selected) or []
            for e in extra:
                if e not in jr_clean and len(jr_clean) < 15:
                    jr_clean.append(e)
        cr_enforced=_enforce_company_limit(cr, None, COMPANY_SUGGESTIONS_LIMIT)
        return {"job":{"related":jr_clean[:15]}, "company":{"related":cr_enforced[:COMPANY_SUGGESTIONS_LIMIT]}}
    except Exception as e:
        logger.warning(f"[Gemini Multi-Sector] {e}")
    return None

@app.post("/start_job")
@_rate(_make_flask_limit("start_job"))
@_check_user_rate("start_job")
def start_job():
    global _role_tag_session_column_ensured
    data=request.get_json(force=True, silent=True) or {}
    queries=data.get('queries') or []
    fallback_queries=data.get('fallbackQueries') or []
    auto_expand=bool(data.get('autoExpand'))
    manual_urls=data.get('manualUrls') or []
    search_results_only=bool(data.get('searchResultsOnly'))
    country=(data.get("country") or "").strip()
    languages=data.get("languages") or []
    language_query=(data.get("languageQuery") or "").strip()
    auto_suggest_companies=data.get("autoSuggestedCompanyNames") or []
    user_companies=data.get("companyNames") or []
    job_titles=data.get("jobTitles") or []
    current_role=bool(data.get("currentRole"))
    selected_sectors=data.get("selectedSectors") or data.get("sectors") or []
    seniority=(data.get("seniority") or "").strip()
    deep_mode=bool(data.get("deepMode"))
    xray_platform_queries=data.get("xrayPlatformQueries") or []
    channel_gaming=bool(data.get("channelGaming"))
    channel_media=bool(data.get("channelMedia"))
    channel_technology=bool(data.get("channelTechnology"))
    if deep_mode and xray_platform_queries:
        for q in xray_platform_queries:
            if q not in queries:
                queries.append(q)
    channel_count=int(channel_gaming)+int(channel_media)+int(channel_technology)
    platform_count=len(xray_platform_queries)
    
    # Check if user provided an explicit target limit
    user_target_raw = data.get("userTarget")
    dynamic_target = 0
    if user_target_raw is not None:
        try:
            dynamic_target = int(user_target_raw)
        except Exception:
            dynamic_target = 0

    # Dynamically adjust the global SEARCH_RESULTS_TARGET to reflect the user's Target Limit
    if dynamic_target > 0:
        _wb_mod = sys.modules.get('webbridge') or sys.modules.get('__main__')
        if _wb_mod is not None:
            _wb_mod.SEARCH_RESULTS_TARGET = dynamic_target

    if dynamic_target <= 0:
        dynamic_target=_compute_search_target(job_titles, country, user_companies, auto_suggest_companies,
                                              selected_sectors, languages, current_role, seniority or None,
                                              channel_count, platform_count)
                                          
    job_id=uuid.uuid4().hex[:10]
    userid=(data.get('userid') or '').strip()
    username=(data.get('username') or '').strip()

    # Per-user search provider credentials injected by AutoSourcing.html from Option A config
    _user_search_provider = (data.get('_userSearchProvider') or '').strip().lower() or None
    _user_serper_key      = (data.get('_serperApiKey')       or '').strip() or None
    _user_dfs_login       = (data.get('_dfsLogin')           or '').strip() or None
    _user_dfs_password    = (data.get('_dfsPassword')        or '').strip() or None

    # --- PATCH START: Automatically update role_tag in login and sourcing tables based on job_titles ---
    # The requirement is that autosourcing.html search title must pass automatically to login.role_tag
    # and also be transferred to sourcing.role_tag for all records of this user.
    try:
        if username and job_titles:
            # Construct the tag string, same logic as frontend: joined by commas
            # job_titles is a list of strings
            role_tag_val = ", ".join([str(t).strip() for t in job_titles if t]).strip()
            if role_tag_val:
                import psycopg2
                pg_host_l = os.getenv("PGHOST", "localhost")
                pg_port_l = int(os.getenv("PGPORT", "5432"))
                pg_user_l = os.getenv("PGUSER", "postgres")
                pg_password_l = os.getenv("PGPASSWORD", "")
                pg_db_l = os.getenv("PGDATABASE", "candidate_db")
                conn_l = psycopg2.connect(host=pg_host_l, port=pg_port_l, user=pg_user_l, password=pg_password_l, dbname=pg_db_l)
                cur_l = conn_l.cursor()
                # Ensure role_tag_session column exists (reuse global flag; ADD COLUMN IF NOT EXISTS is idempotent)
                if not _role_tag_session_column_ensured:
                    cur_l.execute("ALTER TABLE login ADD COLUMN IF NOT EXISTS session TIMESTAMPTZ")
                    cur_l.execute("ALTER TABLE sourcing ADD COLUMN IF NOT EXISTS session TIMESTAMPTZ")
                    cur_l.execute("ALTER TABLE sourcing ALTER COLUMN session DROP DEFAULT")
                    _role_tag_session_column_ensured = True
                # Update login table — set role_tag and generate session timestamp
                cur_l.execute("UPDATE login SET role_tag=%s, session=NOW() WHERE username=%s", (role_tag_val, username))
                # Read back the session timestamp
                cur_l.execute("SELECT role_tag, session FROM login WHERE username=%s", (username,))
                _login_row = cur_l.fetchone()
                _login_role_tag = _login_row[0] if _login_row else None
                _login_session_ts = _login_row[1] if _login_row else None
                # Transfer role_tag to sourcing table (authoritative source for assessments)
                cur_l.execute("ALTER TABLE sourcing ADD COLUMN IF NOT EXISTS role_tag TEXT DEFAULT ''")
                cur_l.execute("UPDATE sourcing SET role_tag=%s WHERE username=%s AND (role_tag IS NULL OR role_tag='')", (role_tag_val, username))
                # Transfer session timestamp to sourcing after validating role_tag matches
                if _login_role_tag == role_tag_val and _login_session_ts is not None:
                    cur_l.execute(
                        "UPDATE sourcing SET session=%s WHERE username=%s AND role_tag=%s",
                        (_login_session_ts, username, role_tag_val)
                    )
                conn_l.commit()
                cur_l.close()
                conn_l.close()
                logger.info(f"[StartJob Auto-Update] Set role_tag='{role_tag_val}' session_ts='{_login_session_ts}' for user='{username}' in login and sourcing tables")
    except Exception as e_rt:
        logger.warning(f"[StartJob Auto-Update role_tag] Failed: {e_rt}")
    # --- PATCH END ---

    with JOBS_LOCK:
        JOBS[job_id]={
            'status_html':'Job created. Initializing...',
            'done':False,
            'output_csv':None,
            'output_xlsx':None,
            'progress':{'processed':0,'total':0},
            'urls':[],
            'messages':[],
            'started':time.time(),
            'userid':userid,
            'username':username,
            'meta':{
                'languages':languages,
                'language_query':language_query,
                'auto_suggest_companies':auto_suggest_companies,
                'user_companies':user_companies,
                'fallback_queries':fallback_queries,
                'selected_sectors':selected_sectors,
                'dynamic_target':dynamic_target,
                'seniority':seniority or None,
                'deepMode':deep_mode,
                'platform_queries':xray_platform_queries,
                'included_platforms':{
                    'gaming':channel_gaming,
                    'media':channel_media,
                    'technology':channel_technology
                },
                'channel_count':channel_count,
                'platform_count':platform_count,
                'userid':userid,
                'username':username
            }
        }
    persist_job(job_id)
    threading.Thread(target=_job_runner,
                     args=(job_id, queries, fallback_queries, auto_expand, manual_urls,
                           search_results_only, country, dynamic_target, job_titles,
                           _user_search_provider, _user_serper_key, _user_dfs_login, _user_dfs_password),
                     daemon=True).start()
    # Log agentic intent event
    _agentic_filters = []
    if country: _agentic_filters.append(f"country:{country}")
    if seniority: _agentic_filters.append(f"seniority:{seniority}")
    if selected_sectors: _agentic_filters.extend([f"sector:{s}" for s in selected_sectors[:3]])
    log_agentic(username=username, userid=userid,
                query="; ".join(queries[:3]) if queries else "",
                filters=_agentic_filters,
                result_count=dynamic_target)
    return jsonify({'job_id': job_id}), 200

@app.get('/job_status/<job_id>')
def job_status(job_id):
    with JOBS_LOCK:
        job=JOBS.get(job_id)
    if not job:
        return jsonify({"error":"Unknown job id"}), 404
    return jsonify(job)

@app.get('/download/<filename>')
def download_file(filename):
    # Look in BASE_DIR first (legacy), then SEARCH_XLS_DIR
    file_path_base = os.path.join(BASE_DIR, filename)
    if os.path.exists(file_path_base):
        return send_from_directory(BASE_DIR, filename, as_attachment=True)
    file_path_search = os.path.join(SEARCH_XLS_DIR, filename)
    if os.path.exists(file_path_search):
        return send_from_directory(SEARCH_XLS_DIR, filename, as_attachment=True)
    return {'error':'File not found'}, 404

@app.get("/SourcingVerify.html")
def sourcing_verify_page():
    return send_from_directory(BASE_DIR, "SourcingVerify.html")

# Module-level flag: _ensure_search_indexes runs at most once per process start.
_search_indexes_ensured = False
_search_indexes_lock = threading.Lock()


def _get_pg_conn():
    """Open and return a new psycopg2 connection using env vars."""
    import psycopg2
    return psycopg2.connect(
        host=os.getenv("PGHOST", "localhost"),
        port=int(os.getenv("PGPORT", "5432")),
        user=os.getenv("PGUSER", "postgres"),
        password=os.getenv("PGPASSWORD", ""),
        dbname=os.getenv("PGDATABASE", "candidate_db"),
    )


def _lazy_ensure_search_indexes():
    """Call _ensure_search_indexes exactly once per process (thread-safe)."""
    global _search_indexes_ensured
    if _search_indexes_ensured:
        return
    with _search_indexes_lock:
        if _search_indexes_ensured:
            return
        try:
            conn = _get_pg_conn()
            cur = conn.cursor()
            _ensure_search_indexes(cur, conn)
            cur.close()
            conn.close()
        except Exception as e:
            logger.warning(f"[SearchIdx] lazy ensure failed (non-fatal): {e}")
        finally:
            _search_indexes_ensured = True


def _parse_rating_row(row_dict, rating_raw):
    """Parse a rating JSON blob and attach convenience fields to row_dict in-place."""
    if not rating_raw:
        row_dict["rating"] = ""
        return
    row_dict["rating"] = rating_raw
    try:
        rating_obj = None
        if isinstance(rating_raw, str):
            rating_obj = json.loads(rating_raw)
        elif isinstance(rating_raw, dict):
            rating_obj = rating_raw
        if rating_obj:
            total_score_str = rating_obj.get("total_score", "")
            if total_score_str and "%" in total_score_str:
                row_dict["rating_score"] = total_score_str.replace("%", "").strip()
            stars = rating_obj.get("stars")
            if stars is not None:
                row_dict["rating_stars"] = str(stars)
            assessment_level = rating_obj.get("assessment_level", "")
            if "Level 1" in assessment_level or "L1" in assessment_level:
                row_dict["rating_level"] = "L1"
            elif "Level 2" in assessment_level or "L2" in assessment_level:
                row_dict["rating_level"] = "L2"
    except Exception as e:
        logger.warning(f"[Sourcing List] Failed to parse rating JSON for {row_dict.get('linkedinurl','')}: {e}")


@app.get("/sourcing/list")
def sourcing_list():
    """List sourcing candidates with optional full-text search and server-side sorting.

    Query parameters:
      userid     – required; filters by owner
      q          – optional; full-text + trigram search query
      sort_by    – optional; one of: name, company, jobtitle, rating_score, relevance (default)
      sort_dir   – optional; asc | desc (default desc when sort_by=relevance/rating_score, else asc)
      page       – page number (1-based)
      page_size  – rows per page (max 1000)
      all        – "1" / "true" / "yes" to return all rows without paging
    """
    try:
        _lazy_ensure_search_indexes()

        userid = (request.args.get("userid") or "").strip()
        if not userid:
            return jsonify({"rows": []})

        q = (request.args.get("q") or "").strip()
        sort_by = (request.args.get("sort_by") or "").strip().lower()
        sort_dir_raw = (request.args.get("sort_dir") or "").strip().lower()

        # Validate sort_by against an explicit allowlist to prevent SQL injection
        _ALLOWED_SORT = {"name", "company", "jobtitle", "rating_score", "relevance", ""}
        if sort_by not in _ALLOWED_SORT:
            sort_by = ""
        if sort_dir_raw not in {"asc", "desc", ""}:
            sort_dir_raw = ""

        page = request.args.get("page", type=int)
        page_size = request.args.get("page_size", type=int) or request.args.get("pagesize", type=int)
        all_flag = (request.args.get("all") or "").strip().lower() in {"1", "true", "yes"}
        use_paging = (bool(page and page_size) and not all_flag)
        if use_paging:
            page = max(1, int(page))
            page_size = max(1, min(int(page_size), 1000))
            offset = (page - 1) * page_size

        conn = _get_pg_conn()
        cur = conn.cursor()

        # Detect optional pic column once
        cur.execute("""
            SELECT 1 FROM information_schema.columns
            WHERE table_schema='public' AND table_name='sourcing' AND column_name='pic'
        """)
        has_pic = cur.fetchone() is not None

        # Detect whether process has a skillset column (for search)
        cur.execute("""
            SELECT 1 FROM information_schema.columns
            WHERE table_schema='public' AND table_name='process' AND column_name='skillset'
        """)
        has_skillset = cur.fetchone() is not None

        # Detect whether search_vector columns exist (added by _ensure_search_indexes)
        cur.execute("""
            SELECT table_name FROM information_schema.columns
            WHERE table_schema='public' AND column_name='search_vector'
              AND table_name IN ('sourcing','process')
        """)
        _sv_tables = {r[0] for r in cur.fetchall()}
        has_sourcing_sv = 'sourcing' in _sv_tables
        has_process_sv  = 'process'  in _sv_tables

        # Detect whether pg_trgm extension is available (needed for similarity())
        cur.execute("""
            SELECT 1 FROM pg_extension WHERE extname = 'pg_trgm'
        """)
        has_trgm = cur.fetchone() is not None

        pic_col = "s.pic," if has_pic else ""
        pic_idx = 6 if has_pic else None  # 0-based index of pic in result tuple
        rating_idx = 7 if has_pic else 6

        # ------------------------------------------------------------------
        # Build ORDER BY clause
        # ------------------------------------------------------------------
        # When a search query is active, default sort is by relevance DESC.
        # When sort_by=rating_score, sort using a numeric cast of the JSON field.
        if sort_by == "rating_score":
            # Extract the numeric total_score from the rating JSONB blob.
            # regexp_match (PG 10+) returns the captured group as an array
            # element; no backslash escaping in the replacement string is
            # needed, which avoids the psycopg2 '%' format-specifier issue and
            # the PostgreSQL standard_conforming_strings '\1' literal-string
            # problem that caused "invalid input syntax for type integer: '\1'".
            _rating_score_expr = (
                "(regexp_match(COALESCE(p.rating::text, ''), "
                "'\"total_score\":[[:space:]]*\"?([0-9]+)'))[1]::int"
            )
            _dir = "DESC" if sort_dir_raw != "asc" else "ASC"
            order_sql = f"{_rating_score_expr} {_dir} NULLS LAST"
        elif sort_by == "company":
            order_sql = f"COALESCE(p.company, s.company) {'DESC' if sort_dir_raw == 'desc' else 'ASC'} NULLS LAST"
        elif sort_by == "jobtitle":
            order_sql = f"COALESCE(p.jobtitle, s.jobtitle) {'DESC' if sort_dir_raw == 'desc' else 'ASC'} NULLS LAST"
        elif sort_by == "name":
            order_sql = f"s.name {'DESC' if sort_dir_raw == 'desc' else 'ASC'} NULLS LAST"
        elif q:
            # Default when search active: relevance descending
            order_sql = "relevance_score DESC NULLS LAST"
        else:
            # No search, no explicit sort: default alphabetical by name
            order_sql = "s.name ASC NULLS LAST"

        # ------------------------------------------------------------------
        # Build WHERE clause + params
        # ------------------------------------------------------------------
        where_clauses = ["s.userid = %s"]
        params_where = [userid]

        if q:
            # Combine tsvector full-text search (when columns available) with
            # ILIKE fallback so queries work even before search_vector is populated.
            like_q = f"%{q}%"
            fts_parts = []
            extra_params: list = []
            if has_sourcing_sv:
                fts_parts.append("s.search_vector @@ websearch_to_tsquery('english', %s)")
                extra_params.append(q)
            if has_process_sv:
                fts_parts.append("p.search_vector @@ websearch_to_tsquery('english', %s)")
                extra_params.append(q)
            fts_clause = f"({' OR '.join(fts_parts)}) OR " if fts_parts else ""
            where_clauses.append(f"""(
                {fts_clause}
                s.name      ILIKE %s
                OR s.jobtitle  ILIKE %s
                OR s.company   ILIKE %s
            )""")
            params_where += extra_params + [like_q, like_q, like_q]

        where_sql = " AND ".join(where_clauses)

        # ------------------------------------------------------------------
        # Relevance score expression (used in SELECT and ORDER BY when q set)
        # ------------------------------------------------------------------
        if q:
            skillset_tsvec = f"to_tsvector('english', coalesce(p.skillset, ''))" if has_skillset else "to_tsvector('')"
            if has_trgm:
                relevance_expr = f"""(
                    ts_rank(
                        setweight(to_tsvector('english', coalesce(s.jobtitle, '')), 'A') ||
                        setweight(to_tsvector('english', coalesce(s.company,  '')), 'B') ||
                        setweight(to_tsvector('english', coalesce(s.name,     '')), 'C') ||
                        setweight(to_tsvector('english', coalesce(s.experience,'')), 'D') ||
                        setweight({skillset_tsvec}, 'B'),
                        websearch_to_tsquery('english', %s)
                    ) +
                    0.4 * greatest(
                        similarity(coalesce(s.jobtitle, ''), %s),
                        similarity(coalesce(s.company,  ''), %s),
                        similarity(coalesce(s.name,     ''), %s)
                    )
                ) AS relevance_score"""
                params_score = [q, q, q, q]
            else:
                relevance_expr = f"""
                    ts_rank(
                        setweight(to_tsvector('english', coalesce(s.jobtitle, '')), 'A') ||
                        setweight(to_tsvector('english', coalesce(s.company,  '')), 'B') ||
                        setweight(to_tsvector('english', coalesce(s.name,     '')), 'C') ||
                        setweight(to_tsvector('english', coalesce(s.experience,'')), 'D') ||
                        setweight({skillset_tsvec}, 'B'),
                        websearch_to_tsquery('english', %s)
                    ) AS relevance_score"""
                params_score = [q]
        else:
            relevance_expr = "0.0 AS relevance_score"
            params_score = []

        # ------------------------------------------------------------------
        # COUNT for pagination
        # ------------------------------------------------------------------
        total = None
        if use_paging:
            count_params = list(params_where)
            cur.execute(f"""
                SELECT COUNT(*)
                FROM sourcing s
                LEFT JOIN process p ON s.linkedinurl = p.linkedinurl
                WHERE {where_sql}
            """, count_params)
            total = int(cur.fetchone()[0])

        # ------------------------------------------------------------------
        # Main SELECT
        # ------------------------------------------------------------------
        select_sql = f"""
            SELECT s.name,
                   COALESCE(p.company,  s.company)  AS company,
                   COALESCE(p.jobtitle, s.jobtitle) AS jobtitle,
                   COALESCE(p.country,  s.country)  AS country,
                   s.experience,
                   s.linkedinurl,
                   {pic_col}
                   p.rating,
                   {relevance_expr},
                   COALESCE(s.role_tag, '') AS role_tag
            FROM sourcing s
            LEFT JOIN process p ON s.linkedinurl = p.linkedinurl
            WHERE {where_sql}
            ORDER BY {order_sql}
        """
        final_params = params_score + params_where
        if use_paging:
            select_sql += " LIMIT %s OFFSET %s"
            final_params += [page_size, offset]

        cur.execute(select_sql, final_params)

        # ------------------------------------------------------------------
        # Build result rows
        # ------------------------------------------------------------------
        import base64
        rows = []
        # Column index for relevance_score is after rating; role_tag follows it
        relevance_col_idx = (rating_idx + 1)
        role_tag_col_idx = relevance_col_idx + 1

        def _strip_nim(name: str) -> str:
            """Strip the Korean honorific suffix '님' (U+B2D8) and surrounding whitespace from a name."""
            return (name or "").strip().rstrip('\uB2D8').strip()

        for r in cur.fetchall():
            row_dict = {
                "name": _strip_nim(r[0]),
                "company": r[1] or "",
                "jobtitle": r[2] or "",
                "country": r[3] or "",
                "experience": r[4] or "",
                "linkedinurl": r[5] or "",
                "rating": "",
                "rating_score": "",
                "rating_stars": "",
                "rating_level": "",
                "relevance_score": float(r[relevance_col_idx]) if r[relevance_col_idx] is not None else 0.0,
                "role_tag": r[role_tag_col_idx] or "",
            }
            if has_pic and pic_idx is not None:
                pic_data = r[pic_idx]
                if pic_data:
                    if isinstance(pic_data, (bytes, memoryview)):
                        raw = bytes(pic_data)
                        # Detect if the stored value is actually a URL (fallback path)
                        try:
                            decoded = raw.decode('utf-8', errors='strict')
                            if decoded.startswith(('http://', 'https://', 'data:')):
                                row_dict["pic"] = decoded
                            else:
                                row_dict["pic"] = base64.b64encode(raw).decode("utf-8")
                        except (UnicodeDecodeError, Exception):
                            row_dict["pic"] = base64.b64encode(raw).decode("utf-8")
                    else:
                        row_dict["pic"] = str(pic_data)
                else:
                    row_dict["pic"] = ""
            else:
                row_dict["pic"] = ""

            _parse_rating_row(row_dict, r[rating_idx])
            rows.append(row_dict)

        cur.close()
        conn.close()

        if rows:
            s = rows[0]
            logger.info(
                f"[Sourcing List] {len(rows)} rows | q={q!r} | sort={sort_by}/{sort_dir_raw} | "
                f"rating_score={s.get('rating_score')}, stars={s.get('rating_stars')}, level={s.get('rating_level')}"
            )

        resp = {"rows": rows}
        if use_paging:
            resp.update({"page": page, "page_size": page_size, "total": total})

        response = jsonify(resp)
        response.headers["Cache-Control"] = "no-cache, no-store, must-revalidate"
        response.headers["Pragma"] = "no-cache"
        response.headers["Expires"] = "0"
        return response
    except Exception as e:
        logger.warning(f"[Sourcing List] {e}")
        return jsonify({"error": str(e)}), 500



@app.get("/sourcing/autocomplete")
def sourcing_autocomplete():
    """Return top job-title suggestions matching a prefix/query for autocomplete UIs.

    Query parameters:
      q      – required; search prefix or partial term (min 2 chars)
      userid – optional; restrict suggestions to this user's sourcing pool
      limit  – max suggestions to return (default 10, max 30)

    Uses trigram similarity for fuzzy prefix matching so it tolerates typos.
    Results are ranked by similarity score descending, then alphabetically.
    """
    try:
        q = (request.args.get("q") or "").strip()
        if len(q) < 2:
            return jsonify({"suggestions": []})

        userid = (request.args.get("userid") or "").strip()
        limit = min(int(request.args.get("limit") or 10), 30)

        conn = _get_pg_conn()
        cur = conn.cursor()

        params = []
        user_filter = ""
        if userid:
            user_filter = "WHERE s.userid = %s"
            params.append(userid)

        # Query distinct job titles from both sourcing and process tables,
        # ranked by trigram similarity to the query string.
        cur.execute(f"""
            WITH candidates AS (
                SELECT COALESCE(p.jobtitle, s.jobtitle) AS jobtitle
                FROM sourcing s
                LEFT JOIN process p ON s.linkedinurl = p.linkedinurl
                {user_filter}
                UNION
                SELECT jobtitle FROM process WHERE jobtitle IS NOT NULL AND jobtitle <> ''
            )
            SELECT DISTINCT jobtitle,
                   similarity(jobtitle, %s) AS sim
            FROM candidates
            WHERE jobtitle IS NOT NULL
              AND jobtitle <> ''
              AND (
                  jobtitle ILIKE %s
                  OR similarity(jobtitle, %s) > 0.15
              )
            ORDER BY sim DESC, jobtitle ASC
            LIMIT %s
        """, params + [q, f"%{q}%", q, limit])

        suggestions = [row[0] for row in cur.fetchall()]
        cur.close()
        conn.close()
        return jsonify({"suggestions": suggestions})
    except Exception as e:
        logger.warning(f"[Sourcing Autocomplete] {e}")
        return jsonify({"suggestions": []})

@app.post("/sourcing/update")
@_csrf_required
def sourcing_update():
    data=request.get_json(force=True, silent=True) or {}
    linkedinurl=(data.get("linkedinurl") or "").strip()
    field=(data.get("field") or "").strip().lower()
    value=(data.get("value") or "").strip()
    allowed_fields = {
        "name": "name",
        "company": "company",
        "jobtitle": "jobtitle",
        "country": "country",
        "appeal": "appeal",
        "experience": "experience"
    }
    if not linkedinurl or field not in allowed_fields:
        return jsonify({"error":"Invalid parameters"}), 400
    if field == "appeal" and len(value) > 500:
        value = value[:500]
    if field == "experience" and len(value) > 5000:
        value = value[:5000]
    try:
        import psycopg2
        from psycopg2 import sql
        pg_host=os.getenv("PGHOST","localhost")
        pg_port=int(os.getenv("PGPORT","5432"))
        pg_user=os.getenv("PGUSER","postgres")
        pg_password=os.getenv("PGPASSWORD", "")
        pg_db=os.getenv("PGDATABASE","candidate_db")
        conn=psycopg2.connect(host=pg_host, port=pg_port, user=pg_user, password=pg_password, dbname=pg_db)
        cur=conn.cursor()
        col_identifier = sql.Identifier(allowed_fields[field])
        query = sql.SQL("UPDATE sourcing SET {col}=%s WHERE linkedinurl=%s").format(col=col_identifier)
        cur.execute(query, (value, linkedinurl))
        affected=cur.rowcount
        conn.commit()
        cur.close(); conn.close()
        if affected==0:
            return jsonify({"error":"Row not found"}), 404
        return jsonify({"updated":affected,"field":field,"value":value})
    except Exception as e:
        logger.warning(f"[Sourcing Update] {e}")
        return jsonify({"error":str(e)}), 500

@app.post("/sourcing/delete")
@_csrf_required
def sourcing_delete():
    data=request.get_json(force=True, silent=True) or {}
    arr=data.get("linkedinurls")
    userid=(data.get("userid") or "").strip()
    # include_process=True  → also delete matching rows from the process table (used by rebate)
    include_process = bool(data.get("include_process", False))
    # Appeal safeguard: records with a non-empty appeal value are NEVER deleted regardless of caller.
    APPEAL_GUARD = "(appeal IS NULL OR appeal = '')"

    # Require at least one of: linkedinurls list or userid
    if not userid and (not isinstance(arr, list) or not arr):
        return jsonify({"error":"linkedinurls list or userid required"}), 400

    cleaned = []
    if isinstance(arr, list):
        cleaned = [(x or "").strip() for x in arr if (x or "").strip()]

    try:
        import psycopg2
        from psycopg2 import sql as pgsql
        pg_host=os.getenv("PGHOST","localhost")
        pg_port=int(os.getenv("PGPORT","5432"))
        pg_user=os.getenv("PGUSER","postgres")
        pg_password=os.getenv("PGPASSWORD", "")
        pg_db=os.getenv("PGDATABASE","candidate_db")
        conn=psycopg2.connect(host=pg_host, port=pg_port, user=pg_user, password=pg_password, dbname=pg_db)
        cur=conn.cursor()

        if userid and not cleaned:
            # User-scoped purge: delete all sourcing rows for this user (appeal rows protected)
            cur.execute(
                pgsql.SQL("DELETE FROM sourcing WHERE userid = %s AND {}").format(pgsql.SQL(APPEAL_GUARD)),
                (userid,)
            )
        else:
            # URL-list delete: appeal rows are always protected
            cur.execute(
                pgsql.SQL("DELETE FROM sourcing WHERE linkedinurl = ANY(%s) AND {}").format(pgsql.SQL(APPEAL_GUARD)),
                (cleaned,)
            )
        deleted=cur.rowcount

        if include_process:
            # For rebate: also remove from process table (no appeal guard needed there)
            if userid and not cleaned:
                cur.execute("DELETE FROM process WHERE userid = %s", (userid,))
            else:
                cur.execute("DELETE FROM process WHERE linkedinurl = ANY(%s)", (cleaned,))

        conn.commit()
        cur.close(); conn.close()
        return jsonify({"deleted":deleted})
    except Exception as e:
        logger.warning(f"[Sourcing Delete] {e}")
        return jsonify({"error":str(e)}), 500

@app.post("/sourcing/mark_ineligible")
@_csrf_required
def sourcing_mark_ineligible():
    """Mark unassessed sourcing candidates as ineligible for assessment.

    Called when the user proceeds to AutoSourcing.html with unassessed
    candidates still in Talent Evaluation.  Any subsequent call to
    /process/bulk_assess for one of these URLs will be skipped server-side.

    The backend validates which candidates are truly unassessed by comparing
    names between the sourcing table and the process table: if a candidate's
    name exists in sourcing but NOT in process, they have no CV and no
    assessment.  The explicitly supplied linkedinurls list is also marked
    ineligible so both client-side and server-side lists are unified.

    Request JSON:
        {"userid": "...", "linkedinurls": ["https://..."]}   # urls optional
    """
    data = request.get_json(force=True, silent=True) or {}
    userid = (data.get("userid") or "").strip()
    linkedinurls = data.get("linkedinurls") or []
    if isinstance(linkedinurls, str):
        linkedinurls = [linkedinurls]
    client_urls = [(u or "").strip() for u in linkedinurls if (u or "").strip()]

    try:
        import psycopg2
        from psycopg2 import sql as pgsql
        pg_host = os.getenv("PGHOST", "localhost")
        pg_port = int(os.getenv("PGPORT", "5432"))
        pg_user = os.getenv("PGUSER", "postgres")
        pg_password = os.getenv("PGPASSWORD", "")
        pg_db = os.getenv("PGDATABASE", "candidate_db")
        conn = psycopg2.connect(host=pg_host, port=pg_port, user=pg_user,
                                password=pg_password, dbname=pg_db)
        cur = conn.cursor()

        # Ensure the ineligible column exists (ALTER TABLE is idempotent here)
        cur.execute("""
            DO $$
            BEGIN
                IF NOT EXISTS (
                    SELECT 1 FROM information_schema.columns
                    WHERE table_schema='public' AND table_name='sourcing'
                      AND column_name='assessment_ineligible'
                ) THEN
                    ALTER TABLE sourcing ADD COLUMN assessment_ineligible BOOLEAN DEFAULT FALSE;
                END IF;
            END$$;
        """)
        conn.commit()

        db_validated_count = 0

        # Backend validation: compare sourcing.name vs process.name (case-insensitive).
        # A candidate whose name exists in the sourcing table but NOT in the process
        # table has never had a CV uploaded and is therefore unassessed.
        if userid:
            try:
                cur.execute("""
                    UPDATE sourcing s
                    SET    assessment_ineligible = TRUE
                    WHERE  s.userid = %s
                      AND  COALESCE(s.assessment_ineligible, FALSE) = FALSE
                      AND  NOT EXISTS (
                               SELECT 1 FROM process p
                               WHERE  LOWER(TRIM(p.name)) = LOWER(TRIM(s.name))
                           )
                """, (userid,))
                db_validated_count = cur.rowcount
                conn.commit()
                logger.info(
                    f"[mark_ineligible] DB name-comparison marked {db_validated_count} rows "
                    f"ineligible for userid={userid!r}"
                )
            except Exception as _db_err:
                logger.warning(
                    f"[mark_ineligible] DB name-comparison failed (non-fatal; "
                    f"explicit URL list will still be marked): {_db_err}"
                )
                conn.rollback()

        # Also mark the explicitly supplied URLs from the client
        url_marked = 0
        if client_urls:
            if userid:
                cur.execute(
                    "UPDATE sourcing SET assessment_ineligible = TRUE "
                    "WHERE linkedinurl = ANY(%s) AND userid = %s",
                    (client_urls, userid)
                )
            else:
                cur.execute(
                    "UPDATE sourcing SET assessment_ineligible = TRUE WHERE linkedinurl = ANY(%s)",
                    (client_urls,)
                )
            url_marked = cur.rowcount
            conn.commit()

        total_marked = db_validated_count + url_marked
        cur.close()
        conn.close()
        logger.info(f"[mark_ineligible] Total {total_marked} rows ineligible for userid={userid!r}")
        return jsonify({"marked": total_marked, "db_validated": db_validated_count, "url_marked": url_marked})
    except Exception as e:
        logger.warning(f"[mark_ineligible] {e}")
        return jsonify({"error": str(e)}), 500

@app.post("/process/delete")
@_csrf_required
def process_delete_entry():
    data = request.get_json(force=True, silent=True) or {}
    linkedinurls = data.get("linkedinurls")
    # Support single or list
    if not linkedinurls:
        single = data.get("linkedinurl")
        if single:
            linkedinurls = [single]
    
    username = (data.get("username") or "").strip()
    userid = (data.get("userid") or "").strip()

    if not linkedinurls or not isinstance(linkedinurls, list):
        return jsonify({"error": "linkedinurl or linkedinurls list required"}), 400

    cleaned = [str(x).strip() for x in linkedinurls if str(x).strip()]
    if not cleaned:
         return jsonify({"error": "No valid URLs"}), 400

    try:
        import psycopg2
        from psycopg2 import sql
        pg_host=os.getenv("PGHOST","localhost")
        pg_port=int(os.getenv("PGPORT","5432"))
        pg_user=os.getenv("PGUSER","postgres")
        pg_password=os.getenv("PGPASSWORD", "")
        pg_db=os.getenv("PGDATABASE","candidate_db")
        conn=psycopg2.connect(host=pg_host, port=pg_port, user=pg_user, password=pg_password, dbname=pg_db)
        cur=conn.cursor()

        deleted_total = 0

        # Check for normalized column
        cur.execute("SELECT column_name FROM information_schema.columns WHERE table_schema='public' AND table_name='process' AND column_name='normalized_linkedin'")
        has_normalized = bool(cur.fetchone())

        for url in cleaned:
             # Try exact delete
             clause = "linkedinurl = %s"
             args = [url]
             if userid:
                 clause += " AND userid = %s"
                 args.append(userid)
             
             cur.execute(sql.SQL("DELETE FROM process WHERE {}").format(sql.SQL(clause)), tuple(args))
             cnt = cur.rowcount
             
             if cnt == 0 and has_normalized:
                 norm = _normalize_linkedin_to_path(url)
                 if norm:
                     clause_n = "normalized_linkedin = %s"
                     args_n = [norm]
                     if userid:
                         clause_n += " AND userid = %s"
                         args_n.append(userid)
                     cur.execute(sql.SQL("DELETE FROM process WHERE {}").format(sql.SQL(clause_n)), tuple(args_n))
                     cnt = cur.rowcount
             
             if cnt > 0:
                 deleted_total += 1

        new_token = 0
        if deleted_total > 0:
             # Update token
             cur.execute("SELECT column_name FROM information_schema.columns WHERE table_schema='public' AND table_name='login' AND column_name='userid'")
             login_has_userid = bool(cur.fetchone())
             
             if login_has_userid and userid:
                 cur.execute("UPDATE login SET token = COALESCE(token,0) + %s WHERE userid = %s RETURNING COALESCE(token,0)", (deleted_total, userid))
                 row = cur.fetchone()
                 if row: new_token = row[0]
             elif username:
                 cur.execute("UPDATE login SET token = COALESCE(token,0) + %s WHERE username = %s RETURNING COALESCE(token,0)", (deleted_total, username))
                 row = cur.fetchone()
                 if row: new_token = row[0]
        else:
             # Just fetch current token
             if username:
                 cur.execute("SELECT COALESCE(token,0) FROM login WHERE username=%s", (username,))
                 r = cur.fetchone()
                 if r: new_token = r[0]

        conn.commit()
        cur.close(); conn.close()

        return jsonify({"deleted": deleted, "token_delta": deleted, "new_token": int(new_token)}), 200

    except Exception as e:
        logger.error(f"[Process Delete] {e}")
        return jsonify({"error": str(e)}), 500

@app.post("/process/update")
@_csrf_required
def process_update():
    """
    Update process table fields. Accepts linkedinurl and any allowed field to update.
    Used for updating tenure and profile picture (pic column).
    """
    data = request.get_json(force=True, silent=True) or {}
    linkedinurl = (data.get("linkedinurl") or "").strip()
    
    if not linkedinurl:
        return jsonify({"error": "linkedinurl required"}), 400
    
    # Define allowed fields for update
    allowed_fields = {
        "tenure": "tenure",
        "pic": "pic",
        "name": "name",
        "company": "company",
        "jobtitle": "jobtitle",
        "country": "country"
    }
    
    # Collect fields to update
    updates = {}
    for key, col in allowed_fields.items():
        if key in data:
            updates[col] = data[key]
    
    if not updates:
        return jsonify({"error": "No valid fields to update"}), 400
    
    try:
        import psycopg2
        from psycopg2 import sql
        pg_host = os.getenv("PGHOST", "localhost")
        pg_port = int(os.getenv("PGPORT", "5432"))
        pg_user = os.getenv("PGUSER", "postgres")
        pg_password = os.getenv("PGPASSWORD", "")
        pg_db = os.getenv("PGDATABASE", "candidate_db")
        conn = psycopg2.connect(host=pg_host, port=pg_port, user=pg_user, password=pg_password, dbname=pg_db)
        cur = conn.cursor()
        
        # Check which columns exist in the process table
        cur.execute("""
            SELECT column_name 
            FROM information_schema.columns 
            WHERE table_schema='public' AND table_name='process'
        """)
        existing_cols = {r[0].lower() for r in cur.fetchall()}
        
        # Filter updates to only existing columns
        valid_updates = {k: v for k, v in updates.items() if k.lower() in existing_cols}
        
        if not valid_updates:
            cur.close(); conn.close()
            return jsonify({"error": "No valid columns to update in process table"}), 400
        
        # Build UPDATE query
        set_parts = []
        params = []
        for col, value in valid_updates.items():
            set_parts.append(sql.SQL("{} = %s").format(sql.Identifier(col)))
            params.append(value)
        
        params.append(linkedinurl)
        
        query = sql.SQL("UPDATE process SET {} WHERE linkedinurl = %s").format(
            sql.SQL(", ").join(set_parts)
        )
        
        cur.execute(query, params)
        affected = cur.rowcount
        
        # If no rows affected, try with normalized_linkedin
        if affected == 0 and 'normalized_linkedin' in existing_cols:
            normalized = _normalize_linkedin_to_path(linkedinurl)
            if normalized:
                params[-1] = normalized
                query_norm = sql.SQL("UPDATE process SET {} WHERE normalized_linkedin = %s").format(
                    sql.SQL(", ").join(set_parts)
                )
                cur.execute(query_norm, params)
                affected = cur.rowcount
        
        conn.commit()
        cur.close(); conn.close()
        
        if affected == 0:
            return jsonify({"error": "No matching record found"}), 404
        
        return jsonify({
            "updated": affected,
            "fields": list(valid_updates.keys())
        }), 200
        
    except Exception as e:
        logger.error(f"[Process Update] {e}")
        return jsonify({"error": str(e)}), 500
        
@app.post("/sourcing/save_profile_json")
def sourcing_save_profile_json():
    data = request.get_json(force=True, silent=True) or {}
    linkedinurl = (data.get("linkedinurl") or "").strip()
    userid = (data.get("userid") or "").strip()
    
    if not linkedinurl:
        return jsonify({"error": "linkedinurl required"}), 400

    try:
        import psycopg2
        pg_host=os.getenv("PGHOST","localhost")
        pg_port=int(os.getenv("PGPORT","5432"))
        pg_user=os.getenv("PGUSER","postgres")
        pg_password=os.getenv("PGPASSWORD", "")
        pg_db=os.getenv("PGDATABASE","candidate_db")
        conn=psycopg2.connect(host=pg_host, port=pg_port, user=pg_user, password=pg_password, dbname=pg_db)
        cur=conn.cursor()
        
        # Determine query based on provided URL format
        # Try exact match first, then normalized if needed
        cur.execute("SELECT name FROM sourcing WHERE linkedinurl=%s", (linkedinurl,))
        row = cur.fetchone()
        
        if not row:
             # Try normalized lookup if strict match fails
             from urllib.parse import urlparse
             path = urlparse(linkedinurl).path.lower().rstrip('/')
             if path:
                 cur.execute("SELECT name FROM sourcing WHERE LOWER(linkedinurl) LIKE %s LIMIT 1", (f"%{path}%",))
                 row = cur.fetchone()

        candidate_name = None
        if row and row[0]:
            candidate_name = row[0].strip()

        # Fetch username from login table if userid is provided
        username_str = "unknown"
        if userid:
            cur.execute("SELECT username FROM login WHERE userid=%s", (userid,))
            u_row = cur.fetchone()
            if u_row and u_row[0]:
                username_str = u_row[0].strip()

        cur.close(); conn.close()

        if not candidate_name:
            return jsonify({"error": "Profile not found or name empty"}), 404
             
        # Sanitize filename components to "pname {username}.json"
        safe_username = re.sub(r'[\\/*?:"<>|]', "", username_str)
        filename = f"pname {safe_username}.json"
        out_path = os.path.join(OUTPUT_DIR, filename)
        
        with open(out_path, "w", encoding="utf-8") as f:
            json.dump({"name": candidate_name, "linkedinurl": linkedinurl}, f, ensure_ascii=False, indent=2)
            
        return jsonify({"status": "ok", "file": filename}), 200
        
    except Exception as e:
        logger.error(f"[Save Profile JSON] {e}")
        return jsonify({"error": str(e)}), 500

def _normalize_linkedin_to_path(linkedin_url: str) -> str:
    if not linkedin_url:
        return ""
    s = linkedin_url.split('?', 1)[0].strip()
    path = re.sub(r'^https?://[^/]+', '', s, flags=re.I)
    path = path.lower().rstrip('/')
    return path

@app.post("/sourcing/market_analysis")
def sourcing_market_analysis():
    payload = request.get_json(force=True, silent=True) or {}
    records = payload.get("records")
    if not isinstance(records, list) or not records:
        return jsonify({"error": "records list required"}), 400

    normalized_records = []
    for r in records:
        if not isinstance(r, dict):
            continue
        name = (r.get("name") or "").strip()
        company = (r.get("organisation") or r.get("company") or "").strip()
        jobtitle_val = (r.get("jobtitle") or r.get("role") or "").strip()
        country = (r.get("country") or "").strip()
        linkedinurl = (r.get("snapshot_at") or r.get("linkedinurl") or "").strip()
        username = (r.get("username") or "").strip()
        userid = (r.get("userid") or "").strip()
        role_tag_val = (r.get("role_tag") or r.get("roleTag") or "").strip()
        experience_val = (r.get("experience") or "").strip()
        rating_val = (r.get("rating") or "").strip()
        
        normalized_linkedin = _normalize_linkedin_to_path(linkedinurl)
        normalized = {
            "name": name,
            "company": company,
            "jobtitle": jobtitle_val,
            "country": country,
            "linkedinurl": linkedinurl,
            "normalized_linkedin": normalized_linkedin,
            "username": username,
            "userid": userid,
            "role_tag": role_tag_val,
            "experience": experience_val,
            "rating": rating_val
        }
        normalized_records.append(normalized)

    valid_records = []
    for nr in normalized_records:
        if nr["name"] and nr["company"] and nr["jobtitle"] and nr["country"] and nr["linkedinurl"]:
            valid_records.append(nr)

    if not valid_records:
        return jsonify({"error": "No valid rows to insert into process table after normalization"}), 400

    inserted_process = 0
    try:
        import psycopg2
        from psycopg2 import sql
        pg_host=os.getenv("PGHOST","localhost")
        pg_port=int(os.getenv("PGPORT","5432"))
        pg_user=os.getenv("PGUSER","postgres")
        pg_password=os.getenv("PGPASSWORD", "")
        pg_db=os.getenv("PGDATABASE","candidate_db")
        conn=psycopg2.connect(host=pg_host, port=pg_port, user=pg_user, password=pg_password, dbname=pg_db)
        cur=conn.cursor()

        # Discover available columns in process table
        cur.execute("""
            SELECT column_name
            FROM information_schema.columns
            WHERE table_schema = 'public' AND table_name = 'process'
        """)
        cols = {row[0].lower() for row in cur.fetchall()}

        preferred_title_col = None
        if 'jobtitle' in cols:
            preferred_title_col = 'jobtitle'
        elif 'role' in cols:
            preferred_title_col = 'role'
        else:
            cur.close()
            conn.close()
            msg = "Process table does not have 'jobtitle' or 'role' column."
            logger.error(msg)
            return jsonify({"error": msg}), 500

        role_col = None
        if 'role_tag' in cols:
            role_col = 'role_tag'
        elif 'roletag' in cols:
            role_col = 'roletag'

        experience_col = None
        if 'experience' in cols:
            experience_col = 'experience'

        rating_col = None
        if 'rating' in cols:
            rating_col = 'rating'

        normalized_col = 'normalized_linkedin' if 'normalized_linkedin' in cols else None

        # NEW: If process table exposes an 'id' column, we'll attempt to copy sourcing.id into process.id
        process_has_id = 'id' in cols

        # PATCH: Check for geographic column
        geo_col = 'geographic' if 'geographic' in cols else None
        
        # PATCH: Check for pic column
        pic_col = 'pic' if 'pic' in cols else None

        # Helper to lookup role_tag from sourcing (primary) or process (fallback) if missing
        def _get_role_tag_for_user(c, uname, uid):
            if not uname and not uid: return None
            try:
                # Try sourcing table first (authoritative source for assessments)
                if uname:
                    c.execute("SELECT role_tag FROM sourcing WHERE username=%s AND role_tag IS NOT NULL AND role_tag != '' LIMIT 1", (uname,))
                else:
                    c.execute("SELECT role_tag FROM sourcing WHERE userid=%s AND role_tag IS NOT NULL AND role_tag != '' LIMIT 1", (uid,))
                r = c.fetchone()
                if r and r[0]: return r[0]
                # Fallback to process table
                if uname:
                    c.execute("SELECT role_tag FROM process WHERE username=%s AND role_tag IS NOT NULL AND role_tag != '' LIMIT 1", (uname,))
                else:
                    c.execute("SELECT role_tag FROM process WHERE userid=%s AND role_tag IS NOT NULL AND role_tag != '' LIMIT 1", (uid,))
                r = c.fetchone()
                return r[0] if r and r[0] else None
            except Exception as e_rt:
                logger.warning(f"Failed to lookup role_tag from sourcing/process: {e_rt}")
                return None

        # Build list for INITIAL INSERT - strictly core identity fields only.
        # EXCLUDE rating and experience from here to prevent positional mismatches.
        # If process has id, include it so we can set process.id = sourcing.id
        field_list = []
        if process_has_id:
            field_list.append('id')
        field_list.extend(['name', 'company', preferred_title_col, 'country', 'linkedinurl', 'username', 'userid'])
        if role_col:
            field_list.append(role_col)
        if normalized_col:
            field_list.append(normalized_col)
        if geo_col:
            field_list.append(geo_col)
        if pic_col:
            field_list.append(pic_col)

        placeholders = sql.SQL(', ').join([sql.Placeholder() for _ in field_list])
        insert_sql = sql.SQL("INSERT INTO process ({fields}) VALUES ({placeholders})").format(
            fields=sql.SQL(', ').join([sql.Identifier(f) for f in field_list]),
            placeholders=placeholders
        )

        # We'll keep a small cache mapping linkedinurl -> sourcing.id to avoid repeated sourcing lookups
        sourcing_id_cache = {}

        def get_sourcing_id_by_linkedin(link):
            if not link: return None
            if link in sourcing_id_cache:
                return sourcing_id_cache[link]
            try:
                cur2 = conn.cursor()
                cur2.execute("SELECT id FROM sourcing WHERE linkedinurl = %s LIMIT 1", (link,))
                r = cur2.fetchone()
                if not r:
                    # fallback to LIKE normalized path
                    path = _normalize_linkedin_to_path(link)
                    if path:
                        cur2.execute("SELECT id FROM sourcing WHERE LOWER(linkedinurl) LIKE %s LIMIT 1", (f"%{path.lower()}",))
                        r = cur2.fetchone()
                cur2.close()
                sid = r[0] if r else None
                sourcing_id_cache[link] = sid
                return sid
            except Exception:
                sourcing_id_cache[link] = None
                return None

        batch_rows = []
        for nr in valid_records:
            # Check for existing record to avoid duplicate insert if ON CONFLICT DO NOTHING is not effective (no unique constraint)
            exists = False
            l_val = nr.get('linkedinurl')
            norm_val = nr.get('normalized_linkedin')

            # Attempt to find sourcing.id for this linkedin so we can set process.id accordingly
            sourcing_id = get_sourcing_id_by_linkedin(l_val)

            try:
                # First, try to find existing process row using multiple strategies
                if process_has_id and sourcing_id:
                    # Check by id
                    check_query = sql.SQL("SELECT 1 FROM process WHERE id = %s LIMIT 1")
                    cur.execute(check_query, (sourcing_id,))
                    if cur.fetchone():
                        exists = True
                if not exists:
                    # Next try exact linkedinurl match
                    check_query = sql.SQL("SELECT 1 FROM process WHERE linkedinurl = %s LIMIT 1")
                    cur.execute(check_query, (l_val,))
                    if cur.fetchone():
                        exists = True
                if not exists and normalized_col and norm_val:
                    # Fallback check by normalized
                    check_norm = sql.SQL("SELECT 1 FROM process WHERE {} = %s LIMIT 1").format(sql.Identifier(normalized_col))
                    cur.execute(check_norm, (norm_val,))
                    if cur.fetchone():
                        exists = True
            except Exception as e_check:
                logger.warning(f"[Market Analysis Existence Check] Failed for {l_val}: {e_check}")
                # Assume not exists or let insert handle it if DB error allows

            if not exists:
                vals = []
                # Pre-calculate geographic
                geo_val = None
                if geo_col:
                    country_input = nr.get('country', '')
                    if country_input:
                        geo_val = _infer_region_from_country(country_input)

                # Pre-calculate role_tag if missing
                role_val_final = nr.get('role_tag')
                if not role_val_final and role_col:
                    role_val_final = _get_role_tag_for_user(cur, nr.get('username'), nr.get('userid'))
                
                # Pre-calculate profile picture if needed
                pic_val = None
                if pic_col:
                    linkedin_url = nr.get('linkedinurl')
                    if linkedin_url:
                        try:
                            pic_url = get_linkedin_profile_picture(linkedin_url, display_name=nr.get('name'))
                            if pic_url:
                                pic_bytes = fetch_image_bytes_from_url(pic_url)
                                if pic_bytes:
                                    import psycopg2
                                    pic_val = psycopg2.Binary(pic_bytes)
                                else:
                                    # Byte fetch failed — store URL as bytes so client can try direct load
                                    import psycopg2
                                    pic_val = psycopg2.Binary(pic_url.encode('utf-8'))
                        except Exception as pic_err:
                            logger.warning(f"[Market Analysis] Failed to get profile pic for {linkedin_url}: {pic_err}")

                for f in field_list:
                    if f == 'id':
                        # Use sourcing_id if available; else None so DB can assign serial if permitted
                        vals.append(sourcing_id)
                    elif f == 'name':
                        vals.append(nr.get('name'))
                    elif f == 'company':
                        vals.append(nr.get('company'))
                    elif f == preferred_title_col:
                        vals.append(nr.get('jobtitle'))
                    elif f == 'country':
                        vals.append(nr.get('country'))
                    elif f == 'linkedinurl':
                        vals.append(nr.get('linkedinurl'))
                    elif f == 'username':
                        vals.append(nr.get('username'))
                    elif f == 'userid':
                        vals.append(nr.get('userid'))
                    elif role_col and f == role_col:
                        vals.append(role_val_final or None)
                    elif normalized_col and f == normalized_col:
                        vals.append(nr.get('normalized_linkedin') or None)
                    elif geo_col and f == geo_col:
                        vals.append(geo_val)
                    elif pic_col and f == pic_col:
                        vals.append(pic_val)
                    else:
                        vals.append(None)

                try:
                    cur.execute(insert_sql, tuple(vals))
                    inserted_process += cur.rowcount
                except Exception as e_ins:
                    # Insert failed; likely duplicate key or id collision. Rollback to keep session consistent and continue.
                    logger.warning(f"[Market Analysis Insert] Insert failed for {l_val}: {e_ins}")
                    conn.rollback()

        # Explicitly UPDATE core fields for every record.
        # This covers updates for existing records (where conflict occurred).
        update_fields = ['name', 'company', 'country', 'username', 'userid']
        if role_col:
            update_fields.append(role_col)
        update_fields.append(preferred_title_col)
        if geo_col:
            update_fields.append(geo_col)

        # Iterate over records to perform updates
        for nr in valid_records:
            l_val = nr.get('linkedinurl')
            norm_val = nr.get('normalized_linkedin')
            if not l_val: continue

            # Attempt to get sourcing id again (cache will make this cheap)
            sourcing_id = get_sourcing_id_by_linkedin(l_val)

            # Pre-calculate geographic for update
            geo_val = None
            if geo_col:
                country_input = nr.get('country', '')
                if country_input:
                    geo_val = _infer_region_from_country(country_input)

            # Prepare role_tag for update if missing in incoming data
            role_val_final = nr.get('role_tag')
            if not role_val_final and role_col:
                role_val_final = _get_role_tag_for_user(cur, nr.get('username'), nr.get('userid'))

            # Update core fields if record exists
            set_parts = []
            update_values = []

            for f in update_fields:
                val = None
                if f == 'name': val = nr.get('name')
                elif f == 'company': val = nr.get('company')
                elif f == 'country': val = nr.get('country')
                elif f == 'username': val = nr.get('username')
                elif f == 'userid': val = nr.get('userid')
                elif f == preferred_title_col: val = nr.get('jobtitle')
                elif role_col and f == role_col: val = role_val_final
                elif geo_col and f == geo_col: val = geo_val

                if val is not None:
                    set_parts.append(sql.SQL("{} = %s").format(sql.Identifier(f)))
                    update_values.append(val)

            if set_parts:
                try:
                    if process_has_id and sourcing_id:
                        # Prefer update by process.id when we have sourcing.id mapped
                        update_query = sql.SQL("UPDATE process SET {} WHERE id = %s").format(sql.SQL(', ').join(set_parts))
                        cur.execute(update_query, update_values + [sourcing_id])
                        if cur.rowcount == 0:
                            # fallback to linkedinurl/normalized
                            update_query = sql.SQL("UPDATE process SET {} WHERE linkedinurl = %s").format(sql.SQL(', ').join(set_parts))
                            cur.execute(update_query, update_values + [l_val])
                            if cur.rowcount == 0 and normalized_col and norm_val:
                                update_norm_query = sql.SQL("UPDATE process SET {} WHERE {} = %s").format(
                                    sql.SQL(', ').join(set_parts),
                                    sql.Identifier(normalized_col)
                                )
                                cur.execute(update_norm_query, update_values + [norm_val])
                    else:
                        update_query = sql.SQL("UPDATE process SET {} WHERE linkedinurl = %s").format(sql.SQL(', ').join(set_parts))
                        cur.execute(update_query, update_values + [l_val])
                        if cur.rowcount == 0 and normalized_col and norm_val:
                            update_norm_query = sql.SQL("UPDATE process SET {} WHERE {} = %s").format(
                                sql.SQL(', ').join(set_parts),
                                sql.Identifier(normalized_col)
                            )
                            cur.execute(update_norm_query, update_values + [norm_val])
                except Exception as e_upd_core:
                    logger.warning(f"[Market Analysis Core Update] Failed for {l_val}: {e_upd_core}")

            # Explicit update for Rating
            if rating_col:
                r_val = nr.get('rating')
                if r_val is not None:
                    try:
                        if process_has_id and sourcing_id:
                            query = sql.SQL("UPDATE process SET {} = %s WHERE id = %s").format(sql.Identifier(rating_col))
                            cur.execute(query, (r_val, sourcing_id))
                            # Fallbacks
                            if cur.rowcount == 0:
                                query = sql.SQL("UPDATE process SET {} = %s WHERE linkedinurl = %s").format(sql.Identifier(rating_col))
                                cur.execute(query, (r_val, l_val))
                                if cur.rowcount == 0 and normalized_col and norm_val:
                                    query_norm = sql.SQL("UPDATE process SET {} = %s WHERE {} = %s").format(sql.Identifier(rating_col), sql.Identifier(normalized_col))
                                    cur.execute(query_norm, (r_val, norm_val))
                        else:
                            query = sql.SQL("UPDATE process SET {} = %s WHERE linkedinurl = %s").format(sql.Identifier(rating_col))
                            cur.execute(query, (r_val, l_val))
                            if cur.rowcount == 0 and normalized_col and norm_val:
                                query_norm = sql.SQL("UPDATE process SET {} = %s WHERE {} = %s").format(sql.Identifier(rating_col), sql.Identifier(normalized_col))
                                cur.execute(query_norm, (r_val, norm_val))
                    except Exception as e_upd:
                        logger.warning(f"[Market Analysis Update Patch] Failed to update rating for {l_val}: {e_upd}")

            # Explicit update for Experience
            if experience_col:
                e_val = nr.get('experience')
                if e_val is not None:
                     try:
                        if process_has_id and sourcing_id:
                            query = sql.SQL("UPDATE process SET {} = %s WHERE id = %s").format(sql.Identifier(experience_col))
                            cur.execute(query, (e_val, sourcing_id))
                            if cur.rowcount == 0:
                                query = sql.SQL("UPDATE process SET {} = %s WHERE linkedinurl = %s").format(sql.Identifier(experience_col))
                                cur.execute(query, (e_val, l_val))
                                if cur.rowcount == 0 and normalized_col and norm_val:
                                    query_norm = sql.SQL("UPDATE process SET {} = %s WHERE {} = %s").format(sql.Identifier(experience_col), sql.Identifier(normalized_col))
                                    cur.execute(query_norm, (e_val, norm_val))
                        else:
                            query = sql.SQL("UPDATE process SET {} = %s WHERE linkedinurl = %s").format(sql.Identifier(experience_col))
                            cur.execute(query, (e_val, l_val))
                            if cur.rowcount == 0 and normalized_col and norm_val:
                                query_norm = sql.SQL("UPDATE process SET {} = %s WHERE {} = %s").format(sql.Identifier(experience_col), sql.Identifier(normalized_col))
                                cur.execute(query_norm, (e_val, norm_val))
                     except Exception as e_upd:
                        logger.warning(f"[Market Analysis Update Patch] Failed to update experience for {l_val}: {e_upd}")

        # PATCH: Sync sequence if we inserted IDs
        if process_has_id and inserted_process > 0:
             try:
                 cur.execute("SELECT setval(pg_get_serial_sequence('process', 'id'), (SELECT MAX(id) FROM process))")
                 conn.commit()
             except Exception:
                 conn.rollback()

        conn.commit()
        cur.close(); conn.close()

        return jsonify({
            "inserted_process": inserted_process,
            "received_process": len(batch_rows),
            "used_title_column": preferred_title_col,
            "used_role_column": role_col,
            "used_experience_column": experience_col,
            "used_rating_column": rating_col,
            "used_normalized_column": normalized_col,
            "process_id_mapped_from_sourcing": process_has_id
        }), 200
    except Exception as e:
        logger.warning(f"[Market Analysis Insert -> process] {e}")
        return jsonify({"error": str(e)}), 500

@app.get("/process/geography")
@_rate(_geography_flask_limit)
@_check_user_rate("geography")
def process_geography():
    linkedin = (request.args.get("linkedin") or "").strip()
    if not linkedin:
        return jsonify({"error": "linkedin param required"}), 400
    linkedin_norm = linkedin.split('?')[0].rstrip('/')
    linkedin_path = _normalize_linkedin_to_path(linkedin)

    def _standardize_host(url_str: str) -> str:
        s = (url_str or "").strip()
        if not s:
            return s
        return re.sub(r'^https?://[^/]+', 'https://www.linkedin.com', s, flags=re.I)

    linkedin_norm_www = _standardize_host(linkedin_norm)

    try:
        import psycopg2
        from psycopg2 import sql
        pg_host=os.getenv("PGHOST","localhost")
        pg_port=int(os.getenv("PGPORT","5432"))
        pg_user=os.getenv("PGUSER","postgres")
        pg_password=os.getenv("PGPASSWORD", "")
        pg_db=os.getenv("PGDATABASE","candidate_db")
        conn=psycopg2.connect(host=pg_host, port=pg_port, user=pg_user, password=pg_password, dbname=pg_db)
        cur=conn.cursor()

        cur.execute("""
            SELECT column_name
            FROM information_schema.columns
            WHERE table_schema = 'public' AND table_name = 'process'
        """)
        available_cols = {r[0].lower() for r in cur.fetchall()}

        desired = [
            "id","name","company","jobtitle","country","linkedinurl","username","userid",
            "product","sector","jobfamily","geographic","seniority","skillset","sourcingstatus",
            "email","mobile","office","role_tag","experience","cv","exp","education","tenure","vskillset"
        ]
        selected = [c for c in desired if c in available_cols]
        if not selected:
            cur.close(); conn.close()
            return jsonify({"error": "Process table does not contain expected columns."}), 500

        fields_sql = sql.SQL(', ').join([sql.Identifier(c) for c in selected])

        row = None

        if 'normalized_linkedin' in available_cols and linkedin_path:
            q_norm_exact = sql.SQL("SELECT {fields} FROM process WHERE normalized_linkedin = %s LIMIT 1").format(fields=fields_sql)
            cur.execute(q_norm_exact, (linkedin_path,))
            row = cur.fetchone()

        if not row and 'linkedinurl' in available_cols:
            q_legacy_exact = sql.SQL("SELECT {fields} FROM process WHERE linkedinurl = %s LIMIT 1").format(fields=fields_sql)
            cur.execute(q_legacy_exact, (linkedin_norm,))
            row = cur.fetchone()

        if not row and 'linkedinurl' in available_cols and linkedin_norm_www and linkedin_norm_www != linkedin_norm:
            q_legacy_www = sql.SQL("SELECT {fields} FROM process WHERE linkedinurl = %s LIMIT 1").format(fields=fields_sql)
            cur.execute(q_legacy_www, (linkedin_norm_www,))
            row = cur.fetchone()

        if not row and 'snapshot_at' in available_cols:
            q_snap_exact = sql.SQL("SELECT {fields} FROM process WHERE snapshot_at = %s LIMIT 1").format(fields=fields_sql)
            cur.execute(q_snap_exact, (linkedin_norm,))
            row = cur.fetchone()
            if not row and linkedin_norm_www and linkedin_norm_www != linkedin_norm:
                cur.execute(q_snap_exact, (linkedin_norm_www,))
                row = cur.fetchone()

        if not row and linkedin_path and 'linkedinurl' in available_cols:
            suffix = linkedin_path
            q_like_legacy = sql.SQL("SELECT {fields} FROM process WHERE LOWER(linkedinurl) LIKE %s LIMIT 1").format(fields=fields_sql)
            cur.execute(q_like_legacy, (f"%{suffix.lower()}",))
            row = cur.fetchone()

        if not row and linkedin_path and 'snapshot_at' in available_cols:
            q_like_snap = sql.SQL("SELECT {fields} FROM process WHERE LOWER(snapshot_at) LIKE %s LIMIT 1").format(fields=fields_sql)
            cur.execute(q_like_snap, (f"%{linkedin_path.lower()}",))
            row = cur.fetchone()

        if not row and 'normalized_linkedin' in available_cols and linkedin_path:
            q_like_norm = sql.SQL("SELECT {fields} FROM process WHERE normalized_linkedin LIKE %s LIMIT 1").format(fields=fields_sql)
            cur.execute(q_like_norm, (f"%{linkedin_path}",))
            row = cur.fetchone()

        if not row:
            sourcing_fields = ["name","company","jobtitle","country","experience","linkedinurl"]
            q_src = sql.SQL("SELECT {fields} FROM sourcing WHERE LOWER(linkedinurl) LIKE %s LIMIT 1").format(
                fields=sql.SQL(', ').join([sql.Identifier(f) for f in sourcing_fields])
            )
            cur.execute(q_src, (f"%{linkedin_path.lower()}",))
            srow = cur.fetchone()
            if srow:
                result = {}
                for c in selected:
                    if c == "name": result[c] = srow[0] or ""
                    elif c == "company": result[c] = srow[1] or ""
                    elif c == "jobtitle": result[c] = srow[2] or ""
                    elif c == "country": result[c] = srow[3] or ""
                    elif c == "experience": result[c] = srow[4] or ""
                    elif c == "linkedinurl": result[c] = srow[5] or ""
                    else: result[c] = ""
                
                # PATCH: Infer geographic region if missing or matches country (legacy data)
                country_val = result.get("country", "")
                geo_val = result.get("geographic", "")
                if country_val and (not geo_val or geo_val.strip().lower() == country_val.strip().lower()):
                    inferred = _infer_region_from_country(country_val)
                    if inferred:
                        result["geographic"] = inferred

                cur.close(); conn.close()
                # Add no-cache headers to ensure fresh data after assessments
                response = jsonify(result)
                response.headers['Cache-Control'] = 'no-cache, no-store, must-revalidate'
                response.headers['Pragma'] = 'no-cache'
                response.headers['Expires'] = '0'
                return response, 200

        cur.close(); conn.close()
        if not row:
            # Add no-cache headers even for 404 responses
            response = jsonify(None)
            response.headers['Cache-Control'] = 'no-cache, no-store, must-revalidate'
            response.headers['Pragma'] = 'no-cache'
            response.headers['Expires'] = '0'
            return response, 404

        result = {}
        for idx, col in enumerate(selected):
            val = row[idx]
            if col == 'cv':
                 result[col] = bool(val)
            elif col == 'seniority':
                 # Add "-level" suffix for UI display if not already present
                 seniority_val = val if val is not None else ""
                 if seniority_val and not seniority_val.endswith('-level'):
                     result[col] = seniority_val + '-level'
                 else:
                     result[col] = seniority_val
            else:
                 result[col] = val if val is not None else ""

        # PATCH: Infer geographic region if missing or duplicate of country
        country_val = result.get("country", "")
        geo_val = result.get("geographic", "")
        
        # If geographic is missing OR matches country (e.g. "Singapore"=="Singapore"), try to infer "Asia"
        if country_val and (not geo_val or geo_val.strip().lower() == country_val.strip().lower()):
            inferred = _infer_region_from_country(country_val)
            if inferred:
                result["geographic"] = inferred

        # Add no-cache headers to ensure fresh data after assessments
        response = jsonify(result)
        response.headers['Cache-Control'] = 'no-cache, no-store, must-revalidate'
        response.headers['Pragma'] = 'no-cache'
        response.headers['Expires'] = '0'
        return response, 200
    except Exception as e:
        logger.warning(f"[Process Geography] {e}")
        try:
            return jsonify({"error": str(e)}), 500
        except Exception:
            return jsonify({"error": "Internal error"}), 500

@app.post("/process/upload_cv")
@_rate(_make_flask_limit("upload_cv"))
@_check_user_rate("upload_cv")
def process_upload_cv():
    """
    Uploads a PDF CV file to the 'process' table, storing it in the 'cv' bytea column.
    Trigger analysis after upload.
    Also persists candidate name if provided.
    """
    try:
        if 'cv' not in request.files:
            return jsonify({"error": "No file part"}), 400
        file = request.files['cv']
        if file.filename == '':
            return jsonify({"error": "No selected file"}), 400
        
        # AFFECTED: Extract name from request form
        candidate_name = (request.form.get('name') or '').strip()
        
        if file and file.filename.lower().endswith('.pdf'):
            linkedinurl = request.form.get('linkedinurl', '').strip()
            if not linkedinurl:
                 return jsonify({"error": "linkedinurl required"}), 400

            if (request.content_length or 0) > _SINGLE_FILE_MAX:
                return jsonify({"error": "File too large (max 6 MB)"}), 413

            file_bytes = file.read()

            if len(file_bytes) > _SINGLE_FILE_MAX:
                return jsonify({"error": "File too large (max 6 MB)"}), 413

            if not _is_pdf_bytes(file_bytes):
                return jsonify({"error": "Uploaded file is not a valid PDF"}), 400

            # Validate that the candidate name appears in the PDF text
            if candidate_name:
                try:
                    from pypdf import PdfReader as _PdfReader
                    _reader = _PdfReader(io.BytesIO(file_bytes))
                    _pdf_text = " ".join(
                        (p.extract_text() or "") for p in _reader.pages
                    ).lower()
                    _name_lower = candidate_name.lower()
                    # Check full name or all individual name parts appear in PDF text
                    _name_parts = _name_lower.split()
                    _name_found = _name_lower in _pdf_text or (
                        len(_name_parts) >= 2 and all(part in _pdf_text for part in _name_parts)
                    )
                    if not _name_found:
                        return jsonify({
                            "error": "The profile name in the uploaded PDF is invalid. "
                                     "Please ensure the name matches an entry in the search result.",
                            "name_mismatch": True
                        }), 400
                except ImportError:
                    logger.warning("[Upload CV] pypdf not available for name validation; skipping check")
                except Exception as _e:
                    logger.warning(f"[Upload CV] Name validation failed (non-fatal): {_e}")

            import psycopg2
            from psycopg2 import sql
            pg_host=os.getenv("PGHOST","localhost")
            pg_port=int(os.getenv("PGPORT","5432"))
            pg_user=os.getenv("PGUSER","postgres")
            pg_password=os.getenv("PGPASSWORD", "")
            pg_db=os.getenv("PGDATABASE","candidate_db")
            
            conn=psycopg2.connect(host=pg_host, port=pg_port, user=pg_user, password=pg_password, dbname=pg_db)
            cur=conn.cursor()
            
            binary_cv = psycopg2.Binary(file_bytes)
            normalized = _normalize_linkedin_to_path(linkedinurl)
            
            # --- PATCH START: Insert ID from sourcing into process if exists ---
            sourcing_id = None
            try:
                # Discover if we have 'id' column in process first
                cur.execute("SELECT column_name FROM information_schema.columns WHERE table_schema='public' AND table_name='process' AND column_name='id'")
                has_process_id = bool(cur.fetchone())
                
                if has_process_id:
                    # Try to get id from sourcing table
                    cur.execute("SELECT id FROM sourcing WHERE linkedinurl = %s LIMIT 1", (linkedinurl,))
                    sid_row = cur.fetchone()
                    if not sid_row and normalized:
                        cur.execute("SELECT id FROM sourcing WHERE LOWER(linkedinurl) LIKE %s LIMIT 1", (f"%{normalized}%",))
                        sid_row = cur.fetchone()
                    
                    if sid_row:
                        sourcing_id = sid_row[0]
            except Exception as e_id:
                logger.warning(f"[Upload CV] Failed to lookup sourcing ID: {e_id}")
            # --- PATCH END ---

            # Try updating by existing ID first if we found one
            updated = False
            
            # AFFECTED: Prepare update fields including name if present
            update_fields = ["cv = %s"]
            update_values = [binary_cv]
            if candidate_name:
                update_fields.append("name = %s")
                update_values.append(candidate_name)
            
            update_sql_fragment = ", ".join(update_fields)

            if sourcing_id:
                try:
                    cur.execute(sql.SQL("UPDATE process SET {} WHERE id = %s").format(sql.SQL(update_sql_fragment)), tuple(update_values + [sourcing_id]))
                    if cur.rowcount > 0:
                        updated = True
                except Exception:
                    conn.rollback()

            if not updated:
                cur.execute(sql.SQL("UPDATE process SET {} WHERE linkedinurl = %s").format(sql.SQL(update_sql_fragment)), tuple(update_values + [linkedinurl]))
                if cur.rowcount > 0:
                    updated = True
            
            if not updated and normalized:
                try:
                    cur.execute(sql.SQL("UPDATE process SET {} WHERE normalized_linkedin = %s").format(sql.SQL(update_sql_fragment)), tuple(update_values + [normalized]))
                    if cur.rowcount > 0:
                        updated = True
                except Exception:
                    conn.rollback()
            
            if not updated:
                # Insert new record
                try:
                    cols = ["linkedinurl", "cv"]
                    vals = [linkedinurl, binary_cv]
                    placeholders = ["%s", "%s"]
                    
                    # AFFECTED: Include name in insert
                    if candidate_name:
                        cols.append("name")
                        vals.append(candidate_name)
                        placeholders.append("%s")
                    
                    if sourcing_id:
                        cols.append("id")
                        vals.append(sourcing_id)
                        placeholders.append("%s")
                    
                    if normalized:
                        # check if column exists
                        cur.execute("SELECT column_name FROM information_schema.columns WHERE table_schema='public' AND table_name='process' AND column_name='normalized_linkedin'")
                        if cur.fetchone():
                            cols.append("normalized_linkedin")
                            vals.append(normalized)
                            placeholders.append("%s")

                    insert_q = sql.SQL("INSERT INTO process ({}) VALUES ({})").format(
                        sql.SQL(', ').join(map(sql.Identifier, cols)),
                        sql.SQL(', ').join(map(sql.SQL, placeholders))
                    )
                    cur.execute(insert_q, tuple(vals))
                    
                    # PATCH: Sync sequence if we inserted IDs
                    if sourcing_id:
                        try:
                            cur.execute("SELECT setval(pg_get_serial_sequence('process', 'id'), (SELECT MAX(id) FROM process))")
                        except Exception:
                            pass

                except psycopg2.errors.UniqueViolation:
                    conn.rollback()
                    # Fallback update again just in case race condition
                    cur.execute(sql.SQL("UPDATE process SET {} WHERE linkedinurl = %s").format(sql.SQL(update_sql_fragment)), tuple(update_values + [linkedinurl]))
                except Exception as e:
                    conn.rollback()
                    return jsonify({"error": f"Database error on insert: {str(e)}"}), 500

            conn.commit()
            cur.close(); conn.close()
            
            # Fire and forget analysis in background
            threading.Thread(target=analyze_cv_background, args=(linkedinurl, file_bytes)).start()

            return jsonify({"status": "ok"}), 200
        else:
            return jsonify({"error": "Invalid file type, PDF required"}), 400
    except Exception as e:
        logger.error(f"[Upload CV] {e}")
        return jsonify({"error": str(e)}), 500

@app.post("/process/upload_multiple_cvs")
@_rate(_make_flask_limit("upload_multiple_cvs"))
@_check_user_rate("upload_multiple_cvs")
def process_upload_multiple_cvs():
    """
    Accept multiple CV files from a browser upload (FormData 'files').
    For each file:
      - attempt to fuzzy-match filename to a sourcing row (by name -> linkedinurl)
      - if matched, ensure record exists in 'process' table (auto-accept)
      - update process.cv
      - spawn analyze_cv_background(linkedinurl, bytes) for analysis
    Returns: { uploaded_count: int, errors: [ ... ] }
    """
    try:
        if 'files' not in request.files and not request.files:
            files = []
            for k in request.files:
                files.extend(request.files.getlist(k))
        else:
            files = request.files.getlist('files')

        # Optional role_tag override passed by the frontend (Step 3 selection)
        override_role_tag = (request.form.get('role_tag') or '').strip() or None

        if not files:
            return jsonify({"uploaded_count": 0, "errors": ["No files provided"]}), 400

        # Normalize list, filter allowed extensions
        allowed_ext = ('.pdf', '.doc', '.docx')
        to_process = [f for f in files if f and f.filename and f.filename.lower().endswith(allowed_ext)]
        rejected = [f.filename for f in files if f and f.filename and not f.filename.lower().endswith(allowed_ext)]

        import psycopg2
        from psycopg2 import sql
        pg_host=os.getenv("PGHOST","localhost")
        pg_port=int(os.getenv("PGPORT","5432"))
        pg_user=os.getenv("PGUSER","postgres")
        pg_password=os.getenv("PGPASSWORD", "")
        pg_db=os.getenv("PGDATABASE","candidate_db")
        conn=psycopg2.connect(host=pg_host, port=pg_port, user=pg_user, password=pg_password, dbname=pg_db)
        cur=conn.cursor()

        # Fetch context from sourcing
        cur.execute("SELECT id, name, linkedinurl, company, jobtitle, country, username, userid FROM sourcing WHERE name IS NOT NULL AND name != ''")
        candidates = cur.fetchall()  # list of tuples

        def normalize_name(s):
            # Remove common suffixes like " _ LinkedIn" before normalization
            s = (s or '').lower()
            # Remove file extension
            s = re.sub(r'\.(pdf|doc|docx)$', '', s, flags=re.IGNORECASE)
            # Remove " _ linkedin" suffix (common in LinkedIn profile PDFs)
            s = re.sub(r'\s*_\s*linkedin\s*$', '', s)
            # Remove all non-alphanumeric characters
            return re.sub(r'[^a-z0-9]', '', s)
        
        def clean_name_for_display(s):
            """Clean special characters and artifacts from names for display.
            Removes non-printable characters, special Unicode artifacts, and non-Latin characters.
            Preserves Latin letters (including accented characters like José, François) and common name punctuation."""
            if not s:
                return s
            
            # Define Unicode ranges for allowed characters
            ASCII_MAX = 127  # Standard ASCII (0-127)
            LATIN_EXTENDED_MAX = 591  # Covers Latin-1 Supplement + Latin Extended A/B
            
            # Remove non-printable characters and non-Latin Unicode characters
            # Keep only Latin letters (ASCII + Latin-1 Supplement + Latin Extended blocks),
            # spaces, hyphens, periods, apostrophes, commas
            # This filters out Korean (님), special artifacts (δïÿ), etc.
            cleaned = []
            for char in s:
                if not char.isprintable():
                    continue  # Skip non-printable characters
                
                char_code = ord(char)
                # Allow common punctuation (works in all ranges)
                if char in ' -.\',':
                    cleaned.append(char)
                # Allow ASCII letters (A-Z, a-z)
                elif char_code <= ASCII_MAX and char.isalpha():
                    cleaned.append(char)
                # Allow Latin-1 Supplement and Latin Extended letters (e.g., À, É, ñ)
                elif ASCII_MAX < char_code <= LATIN_EXTENDED_MAX and char.isalpha():
                    cleaned.append(char)
                # Reject everything else (Korean, Chinese, Arabic, special symbols, etc.)
            
            result = ''.join(cleaned)
            # Normalize multiple spaces to single space
            result = re.sub(r'\s+', ' ', result)
            return result.strip()

        candidate_map = {}
        # Map: normalized_name -> list of records
        # record: {id, name, linkedinurl, company, jobtitle, country, username, userid}
        for row in candidates:
            sid, cname, clink, comp, job, ctry, uname, uid = row
            norm = normalize_name(cname)
            if len(norm) < 3: continue
            # Clean name for display to remove special characters
            clean_cname = clean_name_for_display(cname)
            entry = {
                "id": sid, "name": clean_cname, "linkedinurl": clink,
                "company": comp, "jobtitle": job, "country": ctry,
                "username": uname, "userid": uid
            }
            candidate_map.setdefault(norm, []).append(entry)

        # Supplement the name map with records from the process table so that
        # newly docked candidates (inserted by /candidates/bulk but not yet in
        # sourcing) can still be matched by filename.  Sourcing records take
        # priority — we only add a process entry when the normalised name is
        # not already covered by sourcing.  Only include rows without a CV
        # (cv IS NULL) since records that already have a CV don't need matching.
        try:
            cur.execute("SELECT id, name, linkedinurl, company, jobtitle, country, username, userid FROM process WHERE name IS NOT NULL AND name != '' AND cv IS NULL")
            process_rows = cur.fetchall()
            for prow in process_rows:
                pid2, pname, plink, pcomp, pjob, pctry, puname, puid = prow
                pnorm = normalize_name(pname)
                if len(pnorm) < 3 or pnorm in candidate_map:
                    continue  # sourcing entry already covers this name
                clean_pname = clean_name_for_display(pname)
                pentry = {
                    "id": pid2, "name": clean_pname, "linkedinurl": plink,
                    "company": pcomp, "jobtitle": pjob, "country": pctry,
                    "username": puname, "userid": puid
                }
                candidate_map.setdefault(pnorm, []).append(pentry)
        except Exception as _proc_map_err:
            logger.warning(f"[Upload Multiple CVs] Could not supplement name map from process table: {_proc_map_err}")

        uploaded_count = 0
        errors = []
        uploaded_profiles = []  # Track successfully uploaded profiles
        did_insert_explicit_id = False

        # Check process columns
        cur.execute("SELECT column_name FROM information_schema.columns WHERE table_schema='public' AND table_name='process'")
        process_cols = {r[0].lower() for r in cur.fetchall()}
        
        has_process_id = 'id' in process_cols
        has_role_tag = 'role_tag' in process_cols
        
        user_role_tags = {}
        
        for f in to_process:
            fname = f.filename or "unnamed"
            fname_norm = normalize_name(fname)
            matched_entry = None

            # Substring match strategy
            possible = []
            for norm_name, entries in candidate_map.items():
                if norm_name in fname_norm:
                    for e in entries:
                        possible.append((len(norm_name), e))
            if possible:
                possible.sort(key=lambda x: x[0], reverse=True)
                matched_entry = possible[0][1]

            try:
                if matched_entry:
                    file_bytes = f.read()
                    fname_lower = (f.filename or "").lower()
                    if fname_lower.endswith('.pdf') and not _is_pdf_bytes(file_bytes):
                        errors.append(f"{f.filename}: not a valid PDF (magic bytes mismatch)")
                        continue
                    binary_cv = psycopg2.Binary(file_bytes)
                    m_link = matched_entry['linkedinurl']
                    sourcing_id = matched_entry['id']
                    
                    # Logic: Try to find existing process row to update (by ID or URL)
                    pid = None
                    
                    # 1. Try match by ID (preferred if safe)
                    if has_process_id and sourcing_id:
                        cur.execute("SELECT id FROM process WHERE id=%s", (sourcing_id,))
                        r_id = cur.fetchone()
                        if r_id: pid = r_id[0]
                        
                    # 2. Try match by LinkedIn URL if no ID match (case-insensitive, trailing-slash tolerant)
                    if not pid:
                        norm_m_link = m_link.rstrip('/').lower() if m_link else ''
                        if norm_m_link:
                            cur.execute("SELECT id FROM process WHERE LOWER(RTRIM(linkedinurl, '/')) = LOWER(%s)", (norm_m_link,))
                            r_link = cur.fetchone()
                            if r_link: pid = r_link[0]

                    if pid:
                        # Update existing record - also update name to clean any special characters
                        cleaned_name = matched_entry['name']  # Already cleaned from clean_name_for_display
                        cur.execute("UPDATE process SET cv=%s, name=%s WHERE id=%s", (binary_cv, cleaned_name, pid))
                        conn.commit()
                        uploaded_count += 1
                        if m_link:
                            uploaded_profiles.append(m_link)  # Include updated records so verifyBulkUploadedProfiles waits for analyze_cv_background
                        threading.Thread(target=analyze_cv_background, args=(m_link, file_bytes), kwargs={'process_id': pid, 'override_role_tag': override_role_tag}).start()
                    else:
                        # Insert new record into process
                        r_tag = ""
                        u_name = matched_entry['username']
                        if u_name:
                            if u_name not in user_role_tags:
                                # Try sourcing table first (authoritative), fallback to process table
                                cur.execute("SELECT role_tag FROM sourcing WHERE username=%s AND role_tag IS NOT NULL AND role_tag != '' LIMIT 1", (u_name,))
                                rt = cur.fetchone()
                                if not rt or not rt[0]:
                                    cur.execute("SELECT role_tag FROM process WHERE username=%s AND role_tag IS NOT NULL AND role_tag != '' LIMIT 1", (u_name,))
                                    rt = cur.fetchone()
                                user_role_tags[u_name] = rt[0] if rt else ""
                            r_tag = user_role_tags[u_name]

                        # Infer Geographic from sourcing country
                        src_country = matched_entry['country'] or ""
                        geo_val = _infer_region_from_country(src_country)

                        # Base columns
                        ins_cols = ["linkedinurl", "name", "company", "jobtitle", "country", "username", "userid", "cv"]
                        ins_vals = [m_link, matched_entry['name'], matched_entry['company'], matched_entry['jobtitle'], src_country, matched_entry['username'], matched_entry['userid'], binary_cv]
                        
                        # Add Geographic if column exists
                        if 'geographic' in process_cols and geo_val:
                            ins_cols.append("geographic")
                            ins_vals.append(geo_val)

                        if has_role_tag:
                            ins_cols.append("role_tag")
                            ins_vals.append(r_tag)
                        
                        # Attempt Insert with Explicit ID
                        inserted = False
                        new_pid = None
                        if has_process_id and sourcing_id:
                            try:
                                cols_id = ins_cols + ["id"]
                                vals_id = ins_vals + [sourcing_id]
                                placeholders = ["%s"] * len(vals_id)
                                
                                q = sql.SQL("INSERT INTO process ({}) VALUES ({})").format(
                                    sql.SQL(", ").join(map(sql.Identifier, cols_id)),
                                    sql.SQL(", ").join(map(sql.SQL, placeholders))
                                )
                                cur.execute(q, vals_id)
                                conn.commit()
                                inserted = True
                                new_pid = sourcing_id
                                did_insert_explicit_id = True
                            except psycopg2.errors.UniqueViolation:
                                conn.rollback()
                                # Fallback to no-ID insert below
                            except Exception as e:
                                conn.rollback()
                                logger.warning(f"[Bulk Upload] ID Insert failed for {m_link}: {e}")
                        
                        # Fallback: Insert without ID (auto-serial)
                        if not inserted:
                            try:
                                placeholders = ["%s"] * len(ins_vals)
                                q = sql.SQL("INSERT INTO process ({}) VALUES ({}) RETURNING id").format(
                                    sql.SQL(", ").join(map(sql.Identifier, ins_cols)),
                                    sql.SQL(", ").join(map(sql.SQL, placeholders))
                                )
                                cur.execute(q, ins_vals)
                                ret = cur.fetchone()
                                if ret: new_pid = ret[0]
                                conn.commit()
                                inserted = True
                            except psycopg2.errors.UniqueViolation:
                                conn.rollback()
                                # Race condition: Row created in meantime? Try Update
                                try:
                                    cur.execute("UPDATE process SET cv=%s WHERE linkedinurl=%s", (binary_cv, m_link))
                                    conn.commit()
                                    inserted = True
                                except Exception:
                                    conn.rollback()

                        if inserted:
                            uploaded_count += 1
                            uploaded_profiles.append(m_link)  # Track uploaded profile
                            threading.Thread(target=analyze_cv_background, args=(m_link, file_bytes), kwargs={'process_id': new_pid, 'override_role_tag': override_role_tag}).start()
                        else:
                            errors.append(f"Failed to insert/update for {fname}")

                else:
                    errors.append(f"No sourcing match for file {fname}")
            except Exception as e:
                conn.rollback()
                errors.append(f"Failed to process {fname}: {e}")

        # Final Sequence Fix if we manually inserted any IDs
        if did_insert_explicit_id and has_process_id:
             try:
                 # Sync sequence to max(id) to avoid future collisions
                 cur.execute("SELECT setval(pg_get_serial_sequence('process', 'id'), (SELECT MAX(id) FROM process))")
                 conn.commit()
             except Exception as e_seq:
                 conn.rollback()
                 # This can happen if id column is not SERIAL, safe to ignore
                 # logger.warning(f"[Bulk Upload] Sequence sync warning: {e_seq}")

        cur.close(); conn.close()

        result = {"uploaded_count": uploaded_count, "errors": errors, "uploaded_profiles": uploaded_profiles}
        if rejected:
            result["rejected_files"] = rejected

        return jsonify(result), 200

    except Exception as e:
        logger.exception("[Upload Multiple CVs] failed")
        return jsonify({"uploaded_count": 0, "errors": [str(e)]}), 500

@app.get("/process/download_cv")
def process_download_cv():
    linkedin = (request.args.get("linkedin") or "").strip()
    if not linkedin:
        return "LinkedIn URL required", 400

    linkedin_norm = linkedin.split('?')[0].rstrip('/')
    linkedin_path = _normalize_linkedin_to_path(linkedin)
    
    def _standardize_host(url_str: str) -> str:
        s = (url_str or "").strip()
        if not s: return s
        return re.sub(r'^https?://[^/]+', 'https://www.linkedin.com', s, flags=re.I)

    linkedin_norm_www = _standardize_host(linkedin_norm)

    try:
        import psycopg2
        from psycopg2 import sql
        pg_host=os.getenv("PGHOST","localhost")
        pg_port=int(os.getenv("PGPORT","5432"))
        pg_user=os.getenv("PGUSER","postgres")
        pg_password=os.getenv("PGPASSWORD", "")
        pg_db=os.getenv("PGDATABASE","candidate_db")
        conn=psycopg2.connect(host=pg_host, port=pg_port, user=pg_user, password=pg_password, dbname=pg_db)
        cur=conn.cursor()

        cur.execute("SELECT column_name FROM information_schema.columns WHERE table_schema='public' AND table_name='process' AND column_name='cv'")
        if not cur.fetchone():
             cur.close(); conn.close()
             return "CV column not found in database", 404

        row = None
        if linkedin_path:
            cur.execute("SELECT column_name FROM information_schema.columns WHERE table_schema='public' AND table_name='process' AND column_name='normalized_linkedin'")
            if cur.fetchone():
                cur.execute("SELECT cv, name FROM process WHERE normalized_linkedin = %s AND cv IS NOT NULL LIMIT 1", (linkedin_path,))
                row = cur.fetchone()

        if not row:
            cur.execute("SELECT cv, name FROM process WHERE linkedinurl = %s AND cv IS NOT NULL LIMIT 1", (linkedin_norm,))
            row = cur.fetchone()

        if not row and linkedin_norm_www != linkedin_norm:
             cur.execute("SELECT cv, name FROM process WHERE linkedinurl = %s AND cv IS NOT NULL LIMIT 1", (linkedin_norm_www,))
             row = cur.fetchone()
        
        if not row and linkedin_path:
             cur.execute("SELECT cv, name FROM process WHERE LOWER(linkedinurl) LIKE %s AND cv IS NOT NULL LIMIT 1", (f"%{linkedin_path.lower()}",))
             row = cur.fetchone()

        cur.close(); conn.close()

        if row and row[0]:
            pdf_data = row[0]
            candidate_name = (row[1] or "candidate").strip().replace(" ", "_")
            candidate_name = re.sub(r'[^a-zA-Z0-9_]', '', candidate_name)
            
            from flask import make_response
            response = make_response(bytes(pdf_data))
            response.headers.set('Content-Type', 'application/pdf')
            response.headers.set('Content-Disposition', f'attachment; filename="{candidate_name}_CV.pdf"')
            return response
        else:
            return "CV not found", 404

    except Exception as e:
        logger.error(f"[Download CV] {e}")
        return f"Error: {str(e)}", 500

def _strip_level_suffix(seniority: str) -> str:
    """
    Strip '-level' suffix from seniority for database storage.
    Example: 'Mid-level' -> 'Mid', 'Senior-level' -> 'Senior'
    
    This ensures DB stores clean values without '-level' suffix.
    UI layer should add 'level' back when displaying.
    """
    if not seniority:
        return ""
    # Remove '-level' suffix (lowercase only, as normalized values use consistent casing)
    return re.sub(r'-level$', '', seniority).strip()

def _normalize_seniority_to_8_levels(seniority_text: str, total_experience_years=None) -> str:
    """
    Normalize freeform seniority to one of the 8 specified levels:
    1. Junior-level
    2. Mid-level
    3. Senior-level
    4. Lead-level
    5. Manager-level
    6. Expert-level
    7. Director-level
    8. Executive-level
    
    Rules:
    - Map based on keywords in the seniority text
    - Use experience years as fallback if provided
    - Return empty string if cannot determine
    """
    if not seniority_text:
        # Fallback to experience-based mapping
        if total_experience_years is not None:
            try:
                years = float(total_experience_years)
                if years < 2:
                    return "Junior-level"
                elif years < 5:
                    return "Mid-level"
                elif years < 8:
                    return "Senior-level"
                elif years < 12:
                    return "Lead-level"
                else:
                    return "Expert-level"
            except Exception:
                pass
        return ""

    s = str(seniority_text).strip().lower()
    # Pre-convert experience years to float once for all threshold checks below
    _exp_years = None
    if total_experience_years is not None:
        try:
            _exp_years = float(total_experience_years)
        except Exception:
            pass

    # Exact matches first (case-insensitive)
    exact_matches = {
        "junior-level": "Junior-level",
        "mid-level": "Mid-level",
        "senior-level": "Senior-level",
        "lead-level": "Lead-level",
        "manager-level": "Manager-level",
        "expert-level": "Expert-level",
        "director-level": "Director-level",
        "executive-level": "Executive-level",
    }
    if s in exact_matches:
        return exact_matches[s]
    
    def _kw_match(keyword, text):
        """Word-boundary safe match to avoid substring collisions (e.g. 'cto' inside 'director')."""
        return bool(re.search(r'\b' + re.escape(keyword) + r'\b', text))

    # Executive level - highest: VP, Founder, CEO, etc.
    executive_keywords = ["executive", "ceo", "cto", "cfo", "coo", "cxo", "chief", "president", "vp", "vice president", "c-level", "founder"]
    for keyword in executive_keywords:
        if _kw_match(keyword, s):
            return "Executive-level"
    
    # Director level: any title containing "Director"
    director_keywords = ["director", "head of", "group director"]
    for keyword in director_keywords:
        if _kw_match(keyword, s):
            return "Director-level"
    
    # Expert level: Architect, Principal, Staff, etc. (NOT Specialist — that is Mid)
    expert_keywords = ["expert", "principal", "staff", "distinguished", "fellow", "architect"]
    for keyword in expert_keywords:
        if _kw_match(keyword, s):
            return "Expert-level"
    
    # Manager level - check before Lead/Senior to handle "Senior Manager" correctly
    manager_keywords = ["manager", "mgr", "supervisor", "team lead"]
    for keyword in manager_keywords:
        if _kw_match(keyword, s):
            return "Manager-level"
    
    # Lead level - "Lead" as a title designation (not "team lead", already caught above)
    lead_keywords = ["lead"]
    for keyword in lead_keywords:
        if _kw_match(keyword, s):
            return "Lead-level"
    
    # Senior level - "Senior" alone, without higher-level terms (already checked above)
    senior_keywords = ["senior"]
    for keyword in senior_keywords:
        if _kw_match(keyword, s):
            return "Senior-level"
    
    # Mid level: Associate, Specialist (validated domain contributor, not yet independent lead)
    mid_keywords = ["mid", "intermediate", "associate", "specialist"]
    for keyword in mid_keywords:
        if _kw_match(keyword, s):
            return "Mid-level"

    # Coordinator is always treated as Junior regardless of years of experience.
    # This check is placed after Mid but note that the title-based override in the caller
    # takes precedence when a job title (not just seniority text) is available.
    if _kw_match("coordinator", s):
        return "Junior-level"

    # Junior level: Assistant, entry-level titles
    junior_keywords = ["junior", "entry", "trainee", "intern", "graduate", "jr", "assistant"]
    for keyword in junior_keywords:
        if _kw_match(keyword, s):
            return "Junior-level"
    
    # Fallback to experience-based mapping
    if _exp_years is not None:
        if _exp_years < 2:
            return "Junior-level"
        elif _exp_years < 5:
            return "Mid-level"
        elif _exp_years < 8:
            return "Senior-level"
        elif _exp_years < 12:
            return "Lead-level"
        else:
            return "Expert-level"
    
    # Default to empty if cannot determine
    return ""

def _is_internship_role(job_title):
    """
    Check if a job title indicates an internship role.
    Returns True if the title contains 'intern' or 'internship' (case-insensitive).
    """
    if not job_title:
        return False
    return bool(re.search(r'\bintern\b|\binternship\b', job_title, re.IGNORECASE))

def _normalize_company_name(company_name):
    """
    Normalize company name for duplicate detection.
    Removes common suffixes and converts to lowercase for consistent matching.
    
    Note: Currently handles common US/UK company suffixes. International suffixes
    (GmbH, S.A., AG, etc.) are not normalized but can be added if needed.
    
    Returns normalized company name or None if input is empty.
    """
    if not company_name:
        return None
    
    # Convert to lowercase and remove common company suffixes
    normalized = company_name.lower().strip()
    normalized = re.sub(r'\s+(inc\.?|llc\.?|ltd\.?|corp\.?|corporation|company|co\.?|limited|group|plc)$', '', normalized, flags=re.IGNORECASE)
    normalized = normalized.strip()
    
    return normalized if normalized else None

def _recalculate_tenure_and_experience(experience_list):
    """
    Recalculate total_experience_years and tenure from experience list.

    Two-step calculation approach:
      Step 1 – Baseline estimate: current_year − earliest_start_year gives the full
               career span to date, used as a sanity-check reference.
      Step 2 – Detailed sum: Apply a global timeline merge across all employers to
               account for overlapping concurrent roles and employment gaps.  This is
               the authoritative total and is always ≤ baseline (gaps reduce it below
               the career-span ceiling).

    Calculation rules:
    1. Internship roles are excluded from both total_experience and employer count.
    2. total_experience_years uses Step-2 non-inclusive year differences so the result
       aligns with how Gemini and most recruiter tools count experience
       (e.g. "2015 to 2020" = 5 years, not 6).
    3. Tenure = total_experience_years / number_of_non_overlapping_employment_windows.
       When two employers overlap in time (dual/concurrent employment), they count as
       ONE window, not two.  Internships are excluded from both numerator and window
       count.  This reflects how recruiters measure average tenure: a person working
       two jobs simultaneously has one employment period, not two.
    4. Implausible dates (start < 1950, start > current_year, end < start,
       span > 50 yr) are discarded before the calculation to prevent Gemini
       parsing errors from corrupting results.

    Args:
        experience_list: List of experience strings in format
            "Job Title, Company, StartYear to EndYear|present"
            or "Job Title, Company, Month YYYY to Month YYYY|present"

    Returns:
        dict: {
            "total_experience_years": float,  # Step-2 globally-merged inclusive total
            "baseline_years": float,          # Step-1 full career span (max−min+1)
            "tenure": float,  # Average per-employer inclusive tenure (excl. internships)
            "employer_count": int,  # Unique non-intern employer count
            "total_roles": int      # All roles including internships
        }
    """
    if not experience_list or not isinstance(experience_list, list):
        return {
            "total_experience_years": 0.0,
            "baseline_years": 0.0,
            "tenure": 0.0,
            "employer_count": 0,
            "total_roles": 0
        }

    current_year = datetime.now().year
    all_periods = []        # global list for total_experience_years
    employer_periods = {}   # per-company list for tenure
    total_roles = len(experience_list)

    for entry in experience_list:
        if not entry or not isinstance(entry, str):
            continue

        # Expected format: "Job Title, Company, StartYear to EndYear|present"
        # or "Job Title, Company, Month YYYY to Month YYYY|present"
        parts = [p.strip() for p in entry.split(',')]

        if len(parts) < 3:
            continue

        job_title = parts[0]
        company = parts[1]
        # Join all remaining parts so company names with commas (e.g. "Milky Pharmacy
        # Co., Ltd.") don't split the duration string across multiple parts[2+] indices.
        duration_str = ', '.join(parts[2:])

        is_intern = _is_internship_role(job_title)

        # Primary regex: handles optional leading month name ("Aug 2020 to present")
        duration_match = re.search(
            r'(?:\w+\s+)?(\d{4})\s*(?:to|[-–—])\s*(?:\w+\s+)?(present|\d{4})',
            duration_str, re.IGNORECASE
        )
        if not duration_match:
            # Fallback: extract any 4-digit years directly from the duration string.
            # Handles formats like "15 Aug 2015 to 10 Dec 2020" where the primary
            # regex cannot reach past the day number before the end-date.
            years_found = re.findall(r'\b(\d{4})\b', duration_str)
            present_in_str = bool(re.search(r'\bpresent\b', duration_str, re.IGNORECASE))
            if len(years_found) >= 2:
                start_year = int(years_found[0])
                end_year = int(years_found[-1])
            elif len(years_found) == 1 and present_in_str:
                start_year = int(years_found[0])
                end_year = current_year
            else:
                continue
        else:
            start_year = int(duration_match.group(1))
            end_part = duration_match.group(2).lower()
            end_year = current_year if end_part == 'present' else int(end_part)

        # Validation layer: discard implausible dates
        if start_year < 1950 or start_year > current_year:
            continue
        if end_year < start_year or (end_year - start_year) > 50:
            continue

        # (end_year >= start_year is already guaranteed by the validation layer above)
        if not is_intern:
            normalized_company = _normalize_company_name(company)
            if normalized_company:
                employer_periods.setdefault(normalized_company, []).append((start_year, end_year))
                all_periods.append((start_year, end_year))

    # ── Step 1: Career span (earliest start → current year) ─────────────────
    # Total years of experience = full career span from first job start to now.
    # This matches recruiter convention: if someone started working in 2007 and
    # is still active today, their experience is counted as current_year − 2007
    # regardless of any career gaps.
    if all_periods:
        _min_start = min(s for s, _ in all_periods)
        baseline_years = float(current_year - _min_start)
    else:
        baseline_years = 0.0

    # ── Step 2: Active-employment sum (overlaps merged, gaps excluded) ────────
    # Used only for tenure denominator; NOT used for total_experience_years.
    # Global merge ensures cross-employer boundary years (e.g. end 2020 / start 2020)
    # are counted only once, and overlapping concurrent roles don't inflate totals.
    all_periods.sort()
    merged_global = []
    for start, end in all_periods:
        if merged_global and start <= merged_global[-1][1]:
            merged_global[-1] = (merged_global[-1][0], max(merged_global[-1][1], end))
        else:
            merged_global.append((start, end))
    # Non-inclusive: "2015 to 2020" = 2020 − 2015 = 5 years
    # (avoids the +1 per-segment overcount that inflates vs Gemini's calculation)
    total_experience = sum(e - s for s, e in merged_global)

    # ── Tenure: non-overlapping employment windows as divisor ─────────────────
    # When two employers overlap in time (dual/concurrent employment), they form a
    # single "employment window" and count as ONE unit in the divisor — not two.
    # Formula: total_experience_years / number_of_non_overlapping_windows
    #
    # Algorithm:
    #   Step A – Per-employer merge: collapse multiple roles at the SAME company
    #            into that employer's merged periods (non-strict ≤ merge).
    #   Step B – Cross-employer window count: merge across DIFFERENT employers
    #            using STRICT overlap (start < prev_end), so adjacent but non-
    #            concurrent employers (end=2012, start=2012) are separate windows
    #            while truly concurrent employers (IQVIA+Milky overlap) merge.
    #
    # Example: IQVIA (2021-2026) + Milky Pharmacy (2025-2026, overlapping) → 1 window.
    #          Parexel (2012-2015) and TOMOKI (2007-2012) are adjacent → 2 windows.
    #          Result: 3 windows (TOMOKI, Parexel, IQVIA+Milky), tenure = total_exp / 3.
    all_emp_merged = []
    for _emp_periods in employer_periods.values():
        _sorted = sorted(_emp_periods)
        _merged = []
        for _s, _e in _sorted:
            if _merged and _s <= _merged[-1][1]:
                _merged[-1] = (_merged[-1][0], max(_merged[-1][1], _e))
            else:
                _merged.append((_s, _e))
        all_emp_merged.extend(_merged)

    all_emp_merged.sort()
    merged_windows = []
    for start, end in all_emp_merged:
        if merged_windows and start < merged_windows[-1][1]:  # strict overlap
            merged_windows[-1] = (merged_windows[-1][0], max(merged_windows[-1][1], end))
        else:
            merged_windows.append((start, end))
    effective_employer_count = len(merged_windows)
    employer_count = len(employer_periods)  # kept for logging / return value
    tenure = round(total_experience / effective_employer_count, 1) if effective_employer_count > 0 else 0.0

    return {
        "total_experience_years": round(baseline_years, 1),
        "baseline_years": baseline_years,
        "tenure": tenure,
        "employer_count": employer_count,
        "total_roles": total_roles
    }

def _analyze_cv_bytes_sync(pdf_bytes):
    """
    Synchronous helper to parse PDF bytes via Gemini.
    Supports translation for non-English CVs before analysis.
    
    SOURCE OF TRUTH: CV Column
    - Gemini exclusively references the cv column in the process table (Postgres)
    - Parse employment history strictly in format: "Job Title, Company, StartYear to EndYear" 
      OR "Job Title, Company, StartYear to present" (for current positions)
    
    Returns structured dict or None.
    """
    try:
        from pypdf import PdfReader
    except ImportError:
        logger.warning("[CV Sync] pypdf not installed")
        return None

    try:
        reader = PdfReader(io.BytesIO(pdf_bytes))
        text = ""
        for page in reader.pages:
            t = page.extract_text()
            if t: text += t + "\n"
        
        if not text.strip(): return None

        # Detect language and translate if non-English
        original_text = text
        try:
            # Use LLM to detect if CV is in non-English language
            lang_detect_prompt = (
                "Analyze this text and determine if it's primarily in English or another language.\n"
                "Return JSON: {\"language\": \"<language_code>\", \"is_english\": true/false}\n"
                f"Text sample (first {LANG_DETECTION_SAMPLE_LENGTH} chars): {text[:LANG_DETECTION_SAMPLE_LENGTH]}"
            )
            lang_resp_text = unified_llm_call_text(lang_detect_prompt)
            lang_obj = _extract_json_object((lang_resp_text or "").strip())
            
            if lang_obj and not lang_obj.get("is_english", True):
                # Translate to English using the translation pipeline
                logger.info(f"[CV Sync] Detected non-English CV, translating from {lang_obj.get('language', 'unknown')}")
                source_lang = lang_obj.get("language", "")
                # Use translate_text_pipeline for translation
                # Note: Translation is limited to CV_TRANSLATION_MAX_CHARS to balance API limits and processing time
                # If CV exceeds this limit, later portions may not be translated but will still be analyzed
                translated_result = translate_text_pipeline(text[:CV_TRANSLATION_MAX_CHARS], "english", source_lang)
                if translated_result and translated_result.get("translated"):
                    text = translated_result["translated"]
                    logger.info(f"[CV Sync] Translation completed using {translated_result.get('engine', 'unknown')}")
                    if len(original_text) > CV_TRANSLATION_MAX_CHARS:
                        logger.warning(f"[CV Sync] CV truncated for translation ({len(original_text)} > {CV_TRANSLATION_MAX_CHARS} chars). Later portions analyzed in original language.")
        except Exception as e:
            logger.warning(f"[CV Sync] Language detection/translation failed, proceeding with original: {e}")
            text = original_text

        prompt = (
            "SYSTEM:\n"
            "SOURCE OF TRUTH: You are analyzing a CV from the cv column in the process table (Postgres).\n"
            "This CV is the EXCLUSIVE source for all information. Do not infer or add skills not explicitly in the CV.\n\n"
            "Analyze the following CV text.\n"
            "Return STRICT JSON only with these keys:\n"
            "{\n"
            "  \"skillset\": [\"Skill1\", \"Skill2\", ...],\n"
            "  \"total_experience_years\": <number>,\n"
            "  \"tenure\": <number>,\n"
            "  \"experience\": [\"Job Title, Company, StartYear to EndYear|present\", ...],\n"
            "  \"education\": [\"University Name, Degree Type, Discipline\", ...],\n"
            "  \"product_list\": [\"Product1\", \"Product2\", ...],\n"
            "  \"company\": \"<Current/Latest Company Name>\",\n"
            "  \"job_title\": \"<Current/Latest Job Title>\",\n"
            "  \"country\": \"<Country or Location>\",\n"
            "  \"seniority\": \"<Seniority>\",\n"
            "  \"sector\": \"<Sector>\",\n"
            "  \"job_family\": \"<Job Family>\"\n"
            "}\n"
            "Rules:\n"
            "1. Skillset: Extract ONLY skills explicitly mentioned in the CV. Max 15 items. Do not infer or add skills.\n"
            "2. Total Experience: Calculate sum of all employment durations in years, EXCLUDING internships and intern positions. Only count full-time, part-time, and regular employment. Return a number rounded to 1 decimal place.\n"
            "3. Tenure: Calculate average tenure. Formula: total_experience_years / number of NON-OVERLAPPING employment windows. Rules: (a) Treat repeated employment at the same company as ONE employer. (b) When two DIFFERENT employers overlap in time (concurrent/dual employment), count them as ONE employment window in the divisor \u2014 do not count both separately. (c) Exclude internships and intern positions from both the numerator and the window count. Return a number rounded to 1 decimal place.\n"
            "   Example 1: Someone at Google 2015-2017 and again 2019-2021 \u2192 total_exp=4, windows=1 (Google), tenure=4/1=4.0\n"
            "   Example 2: Google 3yr, Amazon 2yr, Intern at Microsoft 1yr (excluded) \u2192 total_exp=5, windows=2, tenure=5/2=2.5\n"
            "   Example 3: IQVIA 2021-present AND Milky Pharmacy Jan-2025-present (concurrent/overlapping) \u2192 they form 1 window, not 2. If total_exp=13 over 3 windows (TOMOKI, Parexel, IQVIA+Milky), tenure=13/3=4.3\n"
            "4. Experience: STRICTLY parse employment history in format 'Job Title, Company, StartYear to EndYear'. If current job, use 'present' instead of EndYear. MANDATORY: Include EVERY SINGLE employment entry from the CV - do not omit any job.\n"
            "5. Education: Format each entry as 'University Name, Degree Type, Discipline'. MANDATORY: Include ALL educational qualifications - degrees, certifications, diplomas. Do not omit any.\n"
            "6. Products: Identify the LATEST company in the employment history. List its specific products, drugs, therapeutics, software platforms, or services. "
            "Use the company name to infer known products if they are not explicitly mentioned in the CV (e.g., AstraZeneca \u2192 Tagrisso, Farxiga; Pfizer \u2192 Eliquis, Xeljanz; Roche \u2192 Herceptin, Avastin; Novartis \u2192 Cosentyx, Entresto). "
            "For non-pharma companies, list their core product lines or service offerings. "
            "MANDATORY: Always return at least 1\u20133 items in product_list. If specific products cannot be identified, return the company's primary service domain (e.g., 'Clinical Trial Management', 'Drug Development', 'Medical Devices').\n"
            "7. Identify CURRENT employment details (company, job_title, country).\n"
            "8. Infer Seniority, Sector, and Job Family based on the profile.\n"
            "9. CRITICAL REQUIREMENT: Parse COMPLETE employment history without ANY omissions. Every job mentioned must be in the experience array.\n"
            "10. CRITICAL REQUIREMENT: Parse COMPLETE education history without ANY omissions. Every degree/certification must be in the education array.\n"
            "11. IMPORTANT: If a field value cannot be determined, return an empty string \"\" instead of 'unknown', 'N/A', or similar placeholders.\n"
            "12. No commentary, no extra keys. Output only valid JSON.\n\n"
            f"CV TEXT:\n{text[:CV_ANALYSIS_MAX_CHARS]}\n\nJSON:"
        )
        resp_text = unified_llm_call_text(prompt)
        raw = (resp_text or "").strip()
        obj = _extract_json_object(raw)
        
        # Clean company and job_title
        if obj:
             for field in ["company", "job_title", "seniority", "sector", "job_family", "country"]:
                 if obj.get(field):
                     # Remove quotes and surrounding whitespace
                     value = re.sub(r'^[\s"\'`]+|[\s"\'`]+$', '', str(obj[field])).strip()
                     # Replace "unknown" or variations with empty string
                     if value.lower() in ['unknown', 'n/a', 'na', 'not specified', 'not available']:
                         obj[field] = ''
                     else:
                         obj[field] = value
             
             # Normalize seniority to one of the 8 specified levels, then strip '-level' suffix for DB storage
             if obj.get('seniority'):
                 normalized_seniority = _normalize_seniority_to_8_levels(obj['seniority'], obj.get('total_experience_years'))
                 obj['seniority'] = _strip_level_suffix(normalized_seniority)

             # Strict title-based seniority override: the job title always wins over Gemini's seniority
             # classification for these well-defined cases.
             if obj.get('job_title'):
                 _jt_lower = str(obj['job_title']).strip().lower()
                 if re.search(r'\bcoordinator\b', _jt_lower):
                     # "Coordinator" in the title always maps to Junior regardless of experience
                     obj['seniority'] = 'Junior'
                 elif re.search(r'\bmanager\b', _jt_lower) and \
                         str(obj.get('seniority', '')).lower() not in ('director', 'expert', 'executive'):
                     # "Manager" in the title always maps to Manager unless already a higher level
                     obj['seniority'] = 'Manager'
             
             # Post-process: Recalculate tenure and total_experience_years from experience list
             # This ensures consistent application of business rules regardless of Gemini's interpretation
             experience_list = obj.get('experience', [])
             if experience_list and isinstance(experience_list, list):
                 # Store original Gemini values before recalculation
                 gemini_total_exp = obj.get('total_experience_years', 0)
                 gemini_tenure = obj.get('tenure', 0)
                 
                 recalc = _recalculate_tenure_and_experience(experience_list)
                 
                 # Update the values with recalculated ones
                 obj['total_experience_years'] = recalc['total_experience_years']
                 obj['tenure'] = recalc['tenure']
                 
                 # Log if there's a significant difference from Gemini's calculation
                 if abs(recalc['total_experience_years'] - float(gemini_total_exp or 0)) > 0.5:
                     logger.info(f"[CV Sync] Recalculated total_experience_years: {recalc['total_experience_years']} (Gemini: {gemini_total_exp})")
                 if abs(recalc['tenure'] - float(gemini_tenure or 0)) > 0.5:
                     logger.info(f"[CV Sync] Recalculated tenure: {recalc['tenure']} (Gemini: {gemini_tenure}, employers: {recalc['employer_count']})")

        return obj
    except Exception as e:
        logger.warning(f"[CV Sync] Analysis failed: {e}")
        return None

def _core_assess_profile(data):
    """
    Core assessment logic separated from endpoint wrapper.
    Args:
        data (dict): contains keys:
            job_title, role_tag, company, country, seniority, sector,
            experience_text, target_skills (list), candidate_skills (list),
            custom_weights (dict, optional), linkedinurl (optional),
            assessment_level (str, optional): 'L1' or 'L2',
            tenure (float, optional): Average tenure per employer
    Returns:
        dict: assessment result object
    """
    job_title = data.get("job_title", "")
    role_tag = data.get("role_tag", "")
    company = data.get("company", "")
    country = data.get("country", "")
    seniority = data.get("seniority", "")
    sector = data.get("sector", "")
    experience_text = data.get("experience_text", "")
    target_skills = data.get("target_skills", []) or []
    candidate_skills = data.get("candidate_skills", []) or []
    process_skills = data.get("process_skills", []) or []
    custom_weights = data.get("custom_weights", {}) or {}
    linkedinurl = data.get("linkedinurl", "")
    assessment_level = data.get("assessment_level", "L2").upper()  # L2 by default
    tenure = data.get("tenure")  # Average tenure per employer
    vskillset_results = data.get("vskillset_results")  # vskillset inference results for scoring
    product = data.get("product", []) or []  # Product list from CV analysis

    # Normalise candidate country: resolve a city name to its country so that Gemini
    # and country_heuristic both see a canonical country string (e.g. "Tokyo" → "Japan").
    # This ensures individual and bulk assessments behave identically when a city is stored.
    if country:
        _ctc_cities  = CITY_TO_COUNTRY_DATA.get("cities", {}) if isinstance(CITY_TO_COUNTRY_DATA, dict) else {}
        _ctc_aliases = CITY_TO_COUNTRY_DATA.get("aliases", {}) if isinstance(CITY_TO_COUNTRY_DATA, dict) else {}
        _ctc_fallback_cities = {
            "tokyo": "Japan", "osaka": "Japan", "beijing": "China", "shanghai": "China",
            "hong kong": "China", "seoul": "South Korea", "mumbai": "India",
            "delhi": "India", "bangalore": "India", "bangkok": "Thailand",
            "jakarta": "Indonesia", "kuala lumpur": "Malaysia", "manila": "Philippines",
            "hanoi": "Vietnam", "taipei": "Taiwan", "sydney": "Australia",
            "melbourne": "Australia", "london": "United Kingdom", "berlin": "Germany",
            "paris": "France", "new york": "United States", "san francisco": "United States",
            "toronto": "Canada", "dubai": "United Arab Emirates", "abu dhabi": "United Arab Emirates",
        }
        _ctc_fallback_aliases = {
            "uk": "United Kingdom", "usa": "United States", "us": "United States",
            "uae": "United Arab Emirates",
        }
        _cv = country.lower().strip()
        if _ctc_aliases:
            _cv = _ctc_aliases.get(_cv, _cv).lower()
        else:
            _cv = _ctc_fallback_aliases.get(_cv, _cv)
        # City lookup: try full value, first comma-separated token, then last token
        # (handles "Tokyo" → "Japan", "Tokyo, Japan" → "Japan", "Yokohama, Kanagawa, Japan" → "Japan")
        if _ctc_cities:
            _cv_parts = [p.strip() for p in _cv.split(",")]
            _resolved = _ctc_cities.get(_cv) or _ctc_cities.get(_cv_parts[0])
            if not _resolved and len(_cv_parts) > 1:
                _resolved = _ctc_cities.get(_cv_parts[-1])
            if _resolved:
                country = _resolved
        else:
            _cv_parts = [p.strip() for p in _cv.split(",")]
            _resolved = _ctc_fallback_cities.get(_cv) or _ctc_fallback_cities.get(_cv_parts[0])
            if not _resolved and len(_cv_parts) > 1:
                _resolved = _ctc_fallback_cities.get(_cv_parts[-1])
            if _resolved:
                country = _resolved
    
    # Log assessment inputs for debugging
    logger.info(f"[CORE_ASSESS] LinkedIn: {linkedinurl}")
    logger.info(f"[CORE_ASSESS] Candidate skills count: {len(candidate_skills)}, first 10: {candidate_skills[:10] if candidate_skills else []}")
    logger.info(f"[CORE_ASSESS] Target skills count: {len(target_skills)}, first 10: {target_skills[:10] if target_skills else []}")

    # Base weights configuration
    # If custom weights are provided and valid, use them. Otherwise default.
    default_weights = {
        "jobtitle_role_tag": 30.0,
        "skillset": 20.0,
        "tenure": 15.0,  # Average tenure per employer
        "country": 10.0,
        "company": 10.0,
        "product": 5.0,  # Product experience
        "seniority": 5.0,
        "sector": 5.0
    }
    
    # Map frontend keys to internal keys
    weights = default_weights.copy()
    if custom_weights:
        try:
            # Safe parsing helper
            def _get_weight(keys, default_val):
                if isinstance(keys, str): keys = [keys]
                for k in keys:
                    if k in custom_weights:
                        return float(custom_weights[k])
                return float(default_val)

            cw = {
                "jobtitle_role_tag": _get_weight(["jobtitle_role_tag", "job_title", "jobTitle"], 30.0),
                "skillset": _get_weight(["skillset", "skills"], 20.0),
                "tenure": _get_weight(["tenure", "avg_tenure"], 15.0),
                "country": _get_weight(["country", "location"], 10.0),
                "company": _get_weight(["company"], 10.0),
                # product excluded from active criteria; omit from cw so sum stays 100
                "seniority": _get_weight(["seniority"], 5.0),
                "sector": _get_weight(["sector", "industry"], 5.0)
            }
            # Verify sum roughly 100
            total_w = sum(cw.values())
            if 99.0 <= total_w <= 101.0:
                weights = cw
            else:
                pass # logger.warning(f"[Assess] Custom weights sum {total_w} out of range (99-101). Using defaults.")
        except Exception as e:
            pass # logger.warning(f"[Assess] Invalid custom weights: {e}")

    
    # Identify active criteria.
    # A criterion is active when EITHER the candidate has explicit data for it OR experience_text
    # is available so Gemini can infer from the CV.  Without this, DB Dockin profiles whose process
    # table fields (job_title, company, country, seniority, sector) are sparse (e.g. name-only XLS
    # upload) would have only "skillset" in active_criteria and Skillset would receive 100% weight.
    _has_exp = bool(experience_text and experience_text.strip())
    active_criteria = []
    # jobtitle_role_tag: role_tag (criteria requirement) is mandatory; job_title or exp for inference
    if role_tag and (job_title or _has_exp):
        active_criteria.append("jobtitle_role_tag")

    if country or _has_exp: active_criteria.append("country")
    if company or _has_exp: active_criteria.append("company")
    if seniority or _has_exp: active_criteria.append("seniority")
    if sector or _has_exp: active_criteria.append("sector")
    # Product is excluded from the assessment category breakdown;
    # Gemini continues to populate it independently via product_list.
    # Tenure active if value is provided and valid
    if tenure is not None and tenure != "":
        try:
            float(tenure)  # Validate it's numeric
            active_criteria.append("tenure")
        except (ValueError, TypeError):
            pass
    # Skillset active if we have target skills AND (candidate skills OR experience text to infer from)
    if target_skills and (candidate_skills or experience_text):
        active_criteria.append("skillset")
    
    if not active_criteria:
        return {
            "assessment_level": "Level 1",
            "is_level2": False,
            "stars": 0,
            "total_score": "0%",
            "criteria": {},
            "comments": "No data available for assessment."
        }

    # Determine Level 2 status
    is_level2 = False
    if "jobtitle_role_tag" in active_criteria and "skillset" in active_criteria and len(active_criteria) >= 5:
        is_level2 = True

    # Distribute missing weights evenly
    total_weight_target = 100.0
    active_base_sum = sum(weights[c] for c in active_criteria)
    
    if active_base_sum > 0:
        missing_weight = total_weight_target - active_base_sum
        bonus_per_active = missing_weight / len(active_criteria)
        
        final_weights = {}
        for c in active_criteria:
            final_weights[c] = weights[c] + bonus_per_active
    else:
        # Fallback if active criteria base weights sum to 0 (unlikely unless configured 0)
        final_weights = {c: (100.0 / len(active_criteria)) for c in active_criteria}

    assessment_results = {}
    
    try:
        # Prepare skills context
        skills_context = ""
        if "skillset" in active_criteria:
            skills_context = (
                f"Target Skills (required): {', '.join(target_skills)}\n"
                f"Candidate Skills (found): {', '.join(candidate_skills) if candidate_skills else '(Analyze from experience text)'}\n"
            )
            if process_skills:
                skills_context += f"Process-provided Skills (hints): {', '.join(process_skills)}\n"

            # Generate appropriate prompt based on assessment level
            if assessment_level == "L2":
                # L2 PROMPT: Contextual inference allowed
                prompt = (
                    "SYSTEM: You are an expert sourcing assessor performing LEVEL 2 ASSESSMENT with contextual inference.\n"
                    "Your task: Assess skills using BOTH explicit evidence AND conservative contextual inference.\n\n"
                    "LEVEL 2 SKILLSET ASSESSMENT RULES:\n"
                    "1. CONFIRMED SKILLS: Skills explicitly mentioned in experience text (quote exact phrase)\n"
                    "2. CONTEXTUAL INFERENCE: Apply high-probability inference based on:\n"
                    "   - Job Title + Company combination (e.g., Game Programmer at Epic Games → Unreal Engine, C++)\n"
                    "   - Industry domain knowledge (e.g., Senior Java Developer → Spring Framework, Maven)\n"
                    "   - Company's core products (e.g., AWS Engineer at Amazon → EC2, S3, Lambda)\n"
                    "   - Sector/Job Family context when direct skills cannot be determined\n"
                    "3. MINIMUM SKILLSET: Generate at least 10 distinct skills (combining confirmed and inferred) related to the profile's job title and company\n"
                    "4. FALLBACK STRATEGY: If direct skillsets cannot be generated, suggest skills based on:\n"
                    "   - Sector (industry-specific skills)\n"
                    "   - Job Family (role-specific skills)\n"
                    "   - Product/Company domain (company/product-specific skills)\n"
                    "5. CONSERVATIVE APPROACH: Only infer skills with >80% probability given the context\n"
                    "6. EVIDENCE REQUIREMENT: For inferred skills, cite the contextual basis (job title + company)\n"
                    "7. LABEL CLEARLY: Mark skills as 'confirmed' (explicit) or 'inferred' (contextual)\n\n"
                    "EXAMPLES OF VALID L2 INFERENCE:\n"
                    "- 'Game Programmer at Epic Games' → Infer: Unreal Engine (Epic's core product), C++ (game engine requirement)\n"
                    "- 'Machine Learning Engineer at Google' → Infer: TensorFlow (Google's ML framework), Python\n"
                    "- 'iOS Developer at Apple' → Infer: Swift, Xcode, iOS SDK\n\n"
                    "For each field present, classify as:\n"
                    "- 'match' (explicit evidence OR high-confidence inference)\n"
                    "- 'related' (partial evidence OR moderate inference)\n"
                    "- 'unrelated' (no evidence and cannot infer)\n\n"
                    "INPUT:\n"
                    f"Target Role Tag: {role_tag}\n"
                    f"Candidate Job Title: {job_title}\n"
                    f"Candidate Company: {company}\n"
                    f"Candidate Country: {country}\n"
                    f"Candidate Seniority: {seniority}\n"
                    f"Candidate Sector: {sector}\n"
                    f"{skills_context}"
                    f"Experience Text:\n{experience_text[:1000]}...\n\n"
                    "OUTPUT JSON with keys: 'jobtitle_role_tag', 'company', 'country', 'seniority', 'sector', 'skillset'.\n"
                    "Each value must be object: {{ \"status\": \"match\"|\"related\"|\"unrelated\", \"comment\": \"...\", \"evidence\": \"quoted text or contextual basis\" }}\n"
                    "IMPORTANT: If a candidate field is empty but experience_text is provided, INFER the status from the experience_text and job history — do NOT default to 'not_assessed'. Only output 'not_assessed' if the experience_text provides no relevant information for that criterion.\n"
                    "SKILLSET ASSESSMENT:\n"
                    "- Match = Most/all target skills found (confirmed OR inferred with high confidence)\n"
                    "- Related = Some target skills found (confirmed OR inferred)\n"
                    "- Unrelated = Few/no target skills found or inferred\n"
                    "- Include 'confirmed_skills' array (explicitly mentioned)\n"
                    "- Include 'inferred_skills' array (contextually inferred with basis)\n"
                    "- Total combined skillset (confirmed + inferred) should contain at least 10 skills when possible\n"
                    "- Include 'missing_skills' array (neither confirmed nor inferable)\n"
                )
            else:
                # L1 PROMPT: Strictly extractive, NO inference (existing prompt)
                prompt = (
                    "SYSTEM: You are a strict, evidence-based sourcing assessor performing LEVEL 1 ASSESSMENT.\n"
                    "Your task is EXTRACTIVE ONLY - confirm skills ONLY if they appear verbatim or as clear synonyms in the parsed CV experience text.\n"
                    "DO NOT hallucinate or infer skills without explicit evidence from the CV.\n\n"
                    "LEVEL 1 SKILLSET ASSESSMENT RULES:\n"
                    "1. ONLY confirm a skill if you can quote the exact phrase or a direct synonym from the parsed CV experience text\n"
                    "2. Use word-boundary matching - 'Python' in text matches target skill 'Python'\n"
                    "3. If a target skill is NOT found in parsed CV experience text, mark as 'unrelated' - DO NOT guess\n"
                    "4. For each confirmed skill, provide the quoted snippet from experience as evidence\n"
                    "5. NO INFERENCE ALLOWED - only explicit mentions from parsed CV data count\n"
                    "6. ONLY use parsed information from the CV column in the database\n\n"
                    "For each field present, classify as:\n"
                    "- 'match' (exact/strong evidence with quote)\n"
                    "- 'related' (partial evidence with quote)\n"
                    "- 'unrelated' (no evidence found)\n\n"
                    "INPUT:\n"
                    f"Target Role Tag: {role_tag}\n"
                    f"Candidate Job Title: {job_title}\n"
                    f"Candidate Company: {company}\n"
                    f"Candidate Country: {country}\n"
                    f"Candidate Seniority: {seniority}\n"
                    f"Candidate Sector: {sector}\n"
                    f"{skills_context}"
                    f"Parsed CV Experience Text (ONLY source of truth for skills):\n{experience_text[:1000]}...\n\n"
                    "OUTPUT JSON with keys: 'jobtitle_role_tag', 'company', 'country', 'seniority', 'sector', 'skillset'.\n"
                    "Each value must be object: {{ \"status\": \"match\"|\"related\"|\"unrelated\", \"comment\": \"...\", \"evidence\": \"quoted text or empty\" }}\n"
                    "IMPORTANT: If a candidate field is empty but experience_text is provided, infer the status from the CV text. Only output 'not_assessed' if no relevant information can be found in the experience_text for that criterion.\n"
                    "SKILLSET ASSESSMENT:\n"
                    "- Match = ALL target skills found with evidence from parsed CV (100%)\n"
                    "- Related = SOME target skills found with evidence from parsed CV (50%)\n"
                    "- Unrelated = NO or FEW target skills found in parsed CV (0%)\n"
                    "- Include 'confirmed_skills' array with skills that have evidence from parsed CV\n"
                    "- Include 'missing_skills' array with skills without evidence in parsed CV\n"
                )
            resp_text = unified_llm_call_text(prompt, temperature=0)
            raw = (resp_text or "").strip()
            parsed = _extract_json_object(raw)
            if isinstance(parsed, dict):
                assessment_results = parsed
                
                # POST-VALIDATION: Validate extracted skills against source text using word-boundary regex
                if "skillset" in active_criteria and "skillset" in assessment_results:
                    skillset_result = assessment_results["skillset"]
                    confirmed_skills = []
                    inferred_skills = []
                    missing_skills = []
                    
                    exp_lower = (experience_text or "").lower()
                    
                    # Guard against zero target skills
                    if not target_skills or len(target_skills) == 0:
                        assessment_results["skillset"] = {
                            "status": "not_assessed",
                            "comment": "No target skills provided",
                            "confirmed_skills": [],
                            "inferred_skills": [],
                            "missing_skills": [],
                            "evidence": "0/0 skills confirmed"
                        }
                    else:
                        # Create lowercased Gemini confirmed list once for efficiency
                        # Guard: Ensure we only process string values (Gemini might return dicts)
                        gemini_confirmed_raw = skillset_result.get("confirmed_skills", [])
                        gemini_confirmed = [s for s in gemini_confirmed_raw if isinstance(s, str)]
                        gemini_confirmed_lower = [s.lower() for s in gemini_confirmed]
                        
                        gemini_inferred_raw = skillset_result.get("inferred_skills", [])
                        gemini_inferred = [s for s in gemini_inferred_raw if isinstance(s, str)] if gemini_inferred_raw else []
                        gemini_inferred_lower = [s.lower() for s in gemini_inferred] if gemini_inferred else []
                        
                        for target_skill in target_skills:
                            # Use word boundary regex to find exact matches
                            pattern = r'\b' + re.escape(target_skill.lower()) + r'\b'
                            if re.search(pattern, exp_lower):
                                confirmed_skills.append(target_skill)
                            else:
                                # Check if Gemini claimed it was found
                                if target_skill in gemini_confirmed or target_skill.lower() in gemini_confirmed_lower:
                                    # Gemini said it found it, but we can't verify - mark as inferred
                                    inferred_skills.append(target_skill)
                                # L2: Also accept Gemini's inferred skills (contextual inference)
                                elif assessment_level == "L2" and gemini_inferred and (target_skill in gemini_inferred or target_skill.lower() in gemini_inferred_lower):
                                    inferred_skills.append(target_skill)
                                else:
                                    missing_skills.append(target_skill)
                        
                        # Update skillset result with validated sets
                        total_skills = len(target_skills)
                        confirmed_count = len(confirmed_skills)
                        inferred_count = len(inferred_skills)
                        
                        # Recalculate status based on assessment level
                        if assessment_level == "L2":
                            # L2: Count both confirmed and inferred skills
                            total_found = confirmed_count + inferred_count
                            if total_found >= total_skills * 0.75:  # 75%+ found (confirmed OR inferred)
                                status = "match"
                            elif total_found >= total_skills * 0.4:  # 40%+ found
                                status = "related"
                            else:
                                status = "unrelated"
                        else:
                            # L1: Count only confirmed skills (strict)
                            if confirmed_count >= total_skills * 0.8:  # 80%+ confirmed
                                status = "match"
                            elif confirmed_count >= total_skills * 0.4:  # 40%+ confirmed
                                status = "related"
                            else:
                                status = "unrelated"
                        
                        # Update comment with validation results
                        comment_parts = []
                        if confirmed_skills:
                            comment_parts.append(f"Confirmed: {', '.join(confirmed_skills[:3])}{'...' if len(confirmed_skills) > 3 else ''}")
                        if inferred_skills:
                            label = "Skills inferred from CV" if assessment_level == "L2" else "Skills suggested from CV (unverified)"
                            comment_parts.append(f"{label}: {', '.join(inferred_skills[:2])}{'...' if len(inferred_skills) > 2 else ''}")
                        if missing_skills:
                            comment_parts.append(f"Missing: {', '.join(missing_skills[:3])}{'...' if len(missing_skills) > 3 else ''}")
                        
                        evidence_text = f"{confirmed_count}/{total_skills} confirmed"
                        if assessment_level == "L2" and inferred_count > 0:
                            evidence_text += f", {inferred_count} inferred"
                        
                        assessment_results["skillset"] = {
                            "status": status,
                            "comment": " | ".join(comment_parts) if comment_parts else "No skills validated",
                            "confirmed_skills": confirmed_skills,
                            "inferred_skills": inferred_skills,
                            "missing_skills": missing_skills,
                            "evidence": evidence_text
                        }
    except Exception as e:
        logger.warning(f"[Gemini Assess Core] {e}")
            
    # Fallback heuristics
    def heuristic_status(key, val, target):
        if not val: return "not_assessed", ""
        v = str(val).lower(); t = str(target).lower()
        if t in v or v in t: return "match", "Confirmed match"
        return "related", "Related to role requirements"

    def jobtitle_heuristic(candidate_title, required_tag):
        """
        Assess job title against role_tag with seniority gate.
        - Exact/near-exact match (substring containment) → 'match' (1.0)
        - Token overlap AND seniority levels match → 'related' (0.5)
        - No token overlap, OR token overlap but seniority mismatch → 'unrelated' (0)
        """
        if not candidate_title: return "not_assessed", ""
        v = str(candidate_title).lower()
        t = str(required_tag).lower()
        if t in v or v in t: return "match", "Job title closely matches the role requirements"
        v_tokens = set(re.findall(r'\b\w+\b', v)) - {"the", "a", "an", "of", "and", "or", "for", "in", "at"}
        t_tokens = set(re.findall(r'\b\w+\b', t)) - {"the", "a", "an", "of", "and", "or", "for", "in", "at"}
        if v_tokens & t_tokens:
            # Seniority gate: conceptual overlap only scores "related" when seniority matches.
            # Close over `seniority` (candidate) and `required_seniority` from _core_assess_profile.
            def _sn(s):
                return re.sub(r'-level$', '', str(s).lower().strip()) if s else ""
            cs = _sn(seniority)
            rs = _sn(required_seniority)
            _ACCEPTABLE_JT = {("lead", "manager"), ("expert", "director")}
            seniority_ok = (not cs or not rs) or (cs == rs) or ((cs, rs) in _ACCEPTABLE_JT)
            if not seniority_ok:
                return "unrelated", f"Relevant title but seniority level does not match (candidate: {cs or 'unknown'}, required: {rs or 'unknown'})"
            return "related", "Relevant job title but partial match"
        return "unrelated", "Job title does not align with the role requirements"

    def seniority_heuristic(candidate_seniority, required_seniority):
        """
        Assess seniority using strict hierarchy defined by Search Criteria.

        Required seniority → acceptable candidate seniorities:
          Associate  → Junior, Mid, Senior
          Manager    → Lead, Manager
          Director   → Director, Expert, Executive

        Any seniority level outside the defined set is not a valid match.
        Candidate seniority must be within the acceptable set for the required level;
        otherwise 0 stars/score (unrelated).
        """
        if not candidate_seniority: return "not_assessed", ""
        if not required_seniority: return "not_assessed", ""
        # Normalize: lowercase, strip whitespace and trailing '-level' suffix
        def _norm(s):
            return re.sub(r'-level$', '', str(s).lower().strip())
        cs = _norm(candidate_seniority)
        rs = _norm(required_seniority)
        # Strict hierarchy: required level → set of acceptable candidate levels
        _HIERARCHY = {
            "associate": {"junior", "mid", "senior"},
            "manager":   {"lead", "manager"},
            "director":  {"director", "expert", "executive"},
        }
        acceptable = _HIERARCHY.get(rs)
        if acceptable is not None:
            if cs in acceptable:
                return "match", f"Seniority level aligns with the role requirements ({candidate_seniority})"
            return "unrelated", "Seniority level does not meet the role requirements"
        # Exact match for levels not covered by the hierarchy mapping
        if cs == rs:
            return "match", "Seniority level matches the role requirements"
        return "unrelated", "Seniority level does not meet the role requirements"

    def country_heuristic(candidate_country, required_country):
        """
        Assess country using binary match logic.
        Country must match; any mismatch yields 'unrelated' (not 'related').
        Recognises major cities as belonging to their countries so that,
        e.g., "Tokyo" matches required "Japan".
        City-to-country mapping is loaded from city_to_country.json; falls back
        to a minimal hardcoded dict if the file is unavailable.
        """
        _json_cities = CITY_TO_COUNTRY_DATA.get("cities", {}) if isinstance(CITY_TO_COUNTRY_DATA, dict) else {}
        _json_aliases = CITY_TO_COUNTRY_DATA.get("aliases", {}) if isinstance(CITY_TO_COUNTRY_DATA, dict) else {}

        # Hardcoded fallback covers the most common cases when the JSON file is absent
        _FALLBACK_CITIES = {
            "tokyo": "japan", "osaka": "japan",
            "beijing": "china", "shanghai": "china", "hong kong": "china",
            "seoul": "south korea",
            "mumbai": "india", "delhi": "india", "bangalore": "india",
            "bangkok": "thailand", "jakarta": "indonesia",
            "kuala lumpur": "malaysia", "manila": "philippines",
            "hanoi": "vietnam", "taipei": "taiwan",
            "sydney": "australia", "melbourne": "australia",
            "london": "united kingdom", "berlin": "germany",
            "paris": "france", "amsterdam": "netherlands",
            "new york": "united states", "san francisco": "united states",
            "toronto": "canada", "vancouver": "canada",
            "dubai": "united arab emirates", "abu dhabi": "united arab emirates",
        }
        _FALLBACK_ALIASES = {
            "uk": "united kingdom", "usa": "united states", "us": "united states",
            "uae": "united arab emirates",
            "south korea": "korea", "republic of korea": "korea", "kr": "korea",
        }
        # All aliases that resolve to "korea" plus "korea" itself — used for post-alias normalisation
        _KOREA_VARIANTS = {k for k, v in _FALLBACK_ALIASES.items() if v == "korea"} | {"korea"}

        def _resolve(val):
            v = str(val).lower().strip()
            # Apply aliases (JSON first, then fallback)
            if _json_aliases:
                v = _json_aliases.get(v, v).lower()  # aliases may be title-cased in JSON
            else:
                v = _FALLBACK_ALIASES.get(v, v)
            # Normalise South Korea / Korea variants regardless of JSON alias source
            if v in _KOREA_VARIANTS:
                v = "korea"
            # Resolve city to country (JSON cities values are capitalised; lower for comparison)
            # Try full value, first comma-separated token, then last token
            # (handles "Tokyo" → "Japan", "Tokyo, JP" → "Japan", "Unknown City, Japan" → "Japan")
            if _json_cities:
                _v_parts = [p.strip() for p in v.split(",")]
                resolved = _json_cities.get(v) or _json_cities.get(_v_parts[0])
                if not resolved and len(_v_parts) > 1:
                    resolved = _json_cities.get(_v_parts[-1])
                if resolved:
                    return resolved.lower()
            else:
                _v_parts = [p.strip() for p in v.split(",")]
                v = _FALLBACK_CITIES.get(v) or _FALLBACK_CITIES.get(_v_parts[0])
                if not v and len(_v_parts) > 1:
                    v = _FALLBACK_CITIES.get(_v_parts[-1])
                v = v or str(val).lower().strip()
            return v

        if not candidate_country: return "not_assessed", ""
        if not required_country: return "not_assessed", ""
        cc = _resolve(candidate_country)
        rc = _resolve(required_country)
        if cc == rc or cc in rc or rc in cc:
            return "match", "Location matches"
        return "unrelated", "Candidate's location does not match the required country"

    # Derive required seniority from role_tag if not explicitly provided
    required_seniority = data.get("required_seniority", "")
    if not required_seniority and role_tag:
        _SENIORITY_KEYWORDS = [
            "director", "vp", "vice president", "president", "head of", "chief",
            "manager", "lead", "senior", "principal", "staff", "junior", "associate",
            "analyst", "coordinator", "executive", "officer", "specialist"
        ]
        rt_lower = role_tag.lower()
        for kw in _SENIORITY_KEYWORDS:
            if re.search(r'\b' + re.escape(kw) + r'\b', rt_lower):
                required_seniority = kw.title()
                break

    # Required country (optional — provided by caller if known)
    required_country = data.get("required_country", "")

    def skill_heuristic(targets, candidates, exp_text):
        """
        Compare skillset with jskillset using ratio-based scoring.
        Returns (status, comment, match_count, total_count, match_ratio) tuple.
        The match_ratio is used for weighted scoring calculation.
        """
        if not targets: 
            return "not_assessed", "", 0, 0, 0.0
        
        t_set = set(t.lower() for t in targets)
        c_source = (candidates or []) + re.split(r'\W+', (exp_text or "").lower())
        c_set = set(c.lower() for c in c_source if c)
        
        matches = t_set.intersection(c_set)
        match_count = len(matches)
        total_count = len(t_set)
        match_ratio = match_count / total_count if total_count else 0.0
        
        # Determine status based on match ratio (for star display and categorization)
        if match_ratio > 0.6: 
            status = "match"
            comment = f"Strong skillset alignment — {match_count} of {total_count} required skills confirmed"
        elif match_ratio > 0.2: 
            status = "related"
            comment = f"Partial skillset match — {match_count} of {total_count} required skills confirmed"
        else:
            status = "unrelated"
            comment = f"Limited skillset match — {match_count} of {total_count} required skills identified"
        
        return status, comment, match_count, total_count, match_ratio

    for c in active_criteria:
        # seniority, country, and jobtitle_role_tag use deterministic heuristics when the
        # candidate's data field is explicitly present (reliable, prevents Gemini drift).
        # When the field is empty but experience_text is available, fall back to Gemini's
        # inference result so the criterion is scored rather than left as "not_assessed".
        if c == "seniority":
            if seniority:
                st, cm = seniority_heuristic(seniority, required_seniority)
            elif c in assessment_results and assessment_results[c].get("status") in ("match", "related", "unrelated"):
                st = assessment_results[c].get("status", "not_assessed")
                cm = assessment_results[c].get("comment", "Inferred from experience")
            else:
                st, cm = "not_assessed", ""
            assessment_results[c] = {"status": st, "comment": cm}
            continue
        if c == "country":
            if required_country and country:
                st, cm = country_heuristic(country, required_country)
            elif required_country and not country and c in assessment_results and assessment_results[c].get("status") in ("match", "related", "unrelated"):
                # Gemini inferred country from experience_text when field was empty
                st = assessment_results[c].get("status", "not_assessed")
                cm = assessment_results[c].get("comment", "Inferred from experience")
            elif country:
                # No required country specified; candidate has a location → no restriction, treat as match
                st, cm = "match", "Location matches"
            else:
                st, cm = "not_assessed", ""
            assessment_results[c] = {"status": st, "comment": cm}
            continue
        if c == "jobtitle_role_tag":
            if job_title:
                st, cm = jobtitle_heuristic(job_title, role_tag)
            elif c in assessment_results and assessment_results[c].get("status") in ("match", "related", "unrelated"):
                # Gemini inferred job title match from experience_text when field was empty
                st = assessment_results[c].get("status", "not_assessed")
                cm = assessment_results[c].get("comment", "Inferred from experience")
            else:
                st, cm = "not_assessed", ""
            assessment_results[c] = {"status": st, "comment": cm}
            continue
        if c not in assessment_results or assessment_results[c].get("status") not in ["match", "related", "unrelated"]:
            if c == "company":
                 # Only default to "match/Present" when the candidate actually has company data.
                 # When the field is empty (criteria included via experience_text inference),
                 # treat as "not_assessed" rather than falsely claiming presence.
                 if company:
                     assessment_results[c] = {"status": "match", "comment": "Present"}
                 else:
                     assessment_results[c] = {"status": "not_assessed", "comment": "Company not found in candidate profile"}
            elif c == "sector":
                 if sector:
                     assessment_results[c] = {"status": "related", "comment": "Sector present"}
                 else:
                     assessment_results[c] = {"status": "not_assessed", "comment": "Sector not found in candidate profile"}
            elif c == "product":
                 # Assess product: if products are listed, it's a match
                 product_count = len(product) if isinstance(product, list) else 0
                 if product_count >= 3:
                     st, cm = "match", f"{product_count} products listed"
                 elif product_count >= 1:
                     st, cm = "related", f"{product_count} product(s) listed"
                 else:
                     st, cm = "unrelated", "No products listed"
                 assessment_results[c] = {"status": st, "comment": cm}
            elif c == "tenure":
                 # Assess tenure: <2 years = weak, 2-4 years = related, >4 years = match
                 try:
                     tenure_val = float(tenure)
                     if tenure_val >= 4.0:
                         st, cm = "match", f"Strong tenure — {tenure_val:.1f} years average per employer"
                     elif tenure_val >= 2.0:
                         st, cm = "related", f"Acceptable tenure — {tenure_val:.1f} years average per employer"
                     else:
                         st, cm = "unrelated", f"Short tenure — {tenure_val:.1f} years average per employer"
                 except (ValueError, TypeError):
                     st, cm = "not_assessed", "Tenure data unavailable"
                 assessment_results[c] = {"status": st, "comment": cm}
            elif c == "skillset":
                 st, cm, match_count, total_count, match_ratio = skill_heuristic(target_skills, candidate_skills, experience_text)
                 assessment_results[c] = {
                     "status": st, 
                     "comment": cm,
                     "match_count": match_count,
                     "total_count": total_count,
                     "match_ratio": match_ratio
                 }

    # Helper function to generate recruiter-style narrative comments based on score and assessment
    def generate_recruiter_narrative(assessment_results, active_criteria, data, total_score_value):
        """Generate professional recruiter-style narrative aligned with assessment score"""
        narrative_parts = []
        
        # Determine tone based on score
        is_high_score = total_score_value >= 70  # Good/excellent range
        is_mid_score = 40 <= total_score_value < 70  # Moderate range
        is_low_score = total_score_value < 40  # Weak range
        
        # Analyze key categories for detailed feedback
        skillset_res = assessment_results.get("skillset", {})
        skillset_status = skillset_res.get("status", "")
        skillset_comment = skillset_res.get("comment", "")
        
        jobtitle_res = assessment_results.get("jobtitle_role_tag", {})
        jobtitle_status = jobtitle_res.get("status", "")
        
        company_res = assessment_results.get("company", {})
        seniority_res = assessment_results.get("seniority", {})
        seniority_status = seniority_res.get("status", "")
        
        sector_res = assessment_results.get("sector", {})
        sector_status = sector_res.get("status", "")
        
        country_res = assessment_results.get("country", {})
        country_status = country_res.get("status", "")
        
        # Build narrative based on score range
        if is_high_score:
            # Positive, recruiter-style appraisal
            narrative_parts.append("Strong alignment with role requirements.")
            
            if skillset_status in ["match", "related"]:
                narrative_parts.append(f"Skillset coverage is {skillset_status} - {skillset_comment}.")
            
            if sector_status in ["match", "related"] and "sector" in active_criteria:
                narrative_parts.append("Sector expertise aligns well.")
            
            if seniority_status in ["match", "related"] and "seniority" in active_criteria:
                tenure_info = data.get("tenure", "")
                if tenure_info:
                    narrative_parts.append(f"Seniority level appropriate with {tenure_info} average tenure.")
                else:
                    narrative_parts.append("Seniority level is appropriate for the role.")
            
            if country_status in ["match", "related"]:
                narrative_parts.append("Geographic alignment supports local market knowledge.")
            
            narrative_parts.append("Well-suited for senior roles in multinational companies. Recommend advancing.")
            
        elif is_mid_score:
            # Balanced, constructive feedback
            narrative_parts.append("Moderate fit with some alignment to requirements.")
            
            if skillset_status == "related":
                narrative_parts.append(f"Skillset shows partial coverage - {skillset_comment}.")
            elif skillset_status == "unrelated":
                narrative_parts.append(f"Skillset coverage is limited - {skillset_comment}.")
            else:
                narrative_parts.append("Skillset overlap exists but may have gaps.")
            
            if sector_status == "unrelated" and "sector" in active_criteria:
                narrative_parts.append("Limited sector-specific experience noted.")
            
            if seniority_status == "unrelated" and "seniority" in active_criteria:
                narrative_parts.append("Seniority level may not fully align with role expectations.")
            
            narrative_parts.append("Consider for further screening to assess specific competencies.")
            
        else:  # is_low_score
            # Constructive, gap-focused feedback
            narrative_parts.append("Limited alignment with role requirements.")
            
            # Identify key gaps
            gaps = []
            if skillset_status == "unrelated":
                gaps.append(f"skillset ({skillset_comment})")
            if sector_status == "unrelated" and "sector" in active_criteria:
                gaps.append("sector experience")
            if seniority_status == "unrelated" and "seniority" in active_criteria:
                gaps.append("seniority level")
            if jobtitle_status == "unrelated":
                gaps.append("job title match")
            
            if gaps:
                narrative_parts.append(f"Key gaps identified in: {', '.join(gaps)}.")
            
            tenure_info = data.get("tenure", "")
            if tenure_info:
                # Handle both string and numeric tenure values
                if isinstance(tenure_info, (int, float)):
                    tenure_val = float(tenure_info)
                else:
                    # String format - remove "Years"/"years" and parse
                    tenure_str = str(tenure_info).replace("Years", "").replace("years", "").strip()
                    try:
                        tenure_val = float(tenure_str)
                    except ValueError:
                        tenure_val = None
                
                if tenure_val is not None and tenure_val < 2:
                    narrative_parts.append("Short average tenure may indicate limited depth.")

            
            narrative_parts.append("Skillset coverage is partial, reducing fit for the role. Not proceeding recommended.")
        
        return " ".join(narrative_parts)

    total_score = 0.0
    breakdown = {}
    comments_list = []
    category_appraisals = {}
    
    missing_fields = [k for k in weights.keys() if k not in active_criteria]
    if missing_fields:
        nice_names = [k.replace("jobtitle_role_tag", "Role").capitalize() for k in missing_fields]
        comments_list.append(f"{', '.join(nice_names)} Not Assessed")

    # Helper function to convert status to qualitative descriptor
    def status_to_descriptor(status):
        """Convert match/related/unrelated to qualitative descriptors"""
        if status == "match":
            return "Strong"
        elif status == "related":
            return "Suitable"
        else:  # unrelated or not_assessed
            return "Weak"
    
    # Helper function to calculate skillset factor using vskillset or fallback to match_ratio
    def calculate_skillset_factor(vskillset_results, target_skills, match_ratio_fallback):
        """
        Calculate skillset scoring factor using vskillset results.
        Scoring rules (extractive-first):
          - confirmed (source=="confirmed"): full credit (1.0)
          - inferred High (probability>=75): full credit (1.0)
          - inferred Medium (probability 40-74): half credit (0.5)
          - inferred Low (<40) / missing: no credit (0.0)
        Falls back to match_ratio if vskillset is not available.
        
        Returns: (factor, log_message)
        """
        if vskillset_results and isinstance(vskillset_results, list) and target_skills:
            jskillset_total_count = len(target_skills)
            if jskillset_total_count == 0:
                return 0.0, "Skillset scoring: No jskills (target_skills) available"

            weighted_sum = 0.0
            for item in vskillset_results:
                if not isinstance(item, dict):
                    continue
                source = item.get("source", "inferred")
                category = item.get("category", "Low")
                prob = item.get("probability", 0)
                if source == "confirmed":
                    weighted_sum += 1.0
                elif category == "High" or prob >= 75:
                    weighted_sum += 1.0
                elif category == "Medium" or prob >= 40:
                    weighted_sum += 0.5
                # Low / missing = 0.0

            factor = weighted_sum / jskillset_total_count
            log_msg = f"Skillset scoring (extractive+inferred): {weighted_sum:.1f}/{jskillset_total_count} = {factor:.2f}"
            return factor, log_msg
        else:
            # Fallback to original ratio-based scoring
            return match_ratio_fallback, f"Skillset scoring (fallback): match_ratio = {match_ratio_fallback:.2f}"
    
    # Map internal category names to display names
    category_display_names = {
        "jobtitle_role_tag": "Job Title",
        "skillset": "Skillset",
        "seniority": "Seniority",
        "company": "Company",
        "sector": "Sector",
        "country": "Country",
        "tenure": "Tenure",
        "product": "Product"
    }

    # Compute integer display weights using the largest-remainder method so that
    # the sum of displayed weight percentages is always exactly 100.
    # Step 1: take the floor of each weight.
    # Step 2: distribute the remaining 1-unit increments to criteria with the
    #         largest fractional parts until the total reaches 100.
    _floor_weights = {c: int(final_weights[c]) for c in active_criteria}
    _floor_sum = sum(_floor_weights.values())
    _extra = 100 - _floor_sum
    _remainders_order = sorted(active_criteria, key=lambda c: final_weights[c] - int(final_weights[c]), reverse=True)
    _int_weights = dict(_floor_weights)
    for _i in range(max(0, _extra)):
        if _i < len(_remainders_order):
            _int_weights[_remainders_order[_i]] += 1

    for c in active_criteria:
        res = assessment_results.get(c, {})
        st = res.get("status", "unrelated")
        cm = res.get("comment", "")
        
        # Calculate factor based on status OR ratio (for skillset)
        factor = 0.0
        
        if c == "skillset":
            # NEW: Use vskillset-based scoring formula
            # Formula: (vskillset_high_count / jskillset_total_count) * weight = actual_points
            # jskillset = target_skills (from login/process table)
            # vskillset = inference results with High probability
            match_ratio = res.get("match_ratio", 0.0)
            factor, log_msg = calculate_skillset_factor(vskillset_results, target_skills, match_ratio)
            logger.info(f"[CORE_ASSESS] {log_msg}")
        elif c in ("seniority", "country"):
            # Binary scoring: seniority and country must be an exact match.
            # A "related" status (partial match) scores 0 — mismatched seniority or
            # country should never earn partial credit.
            if st == "match": factor = 1.0
            # related / unrelated / not_assessed = 0.0
        else:
            # Standard status-based scoring for other categories
            if st == "match": factor = 1.0
            elif st == "related": factor = 0.5
            # unrelated or not_assessed = 0.0
        
        points = final_weights[c] * factor
        total_score += points
        
        breakdown[c] = round(points, 1)
        
        # Calculate stars for this category
        category_stars = 0
        if c == "skillset":
            # Stars based on the actual factor used for scoring (already calculated above)
            skill_factor = factor
            
            if skill_factor >= 0.8: category_stars = 5
            elif skill_factor >= 0.6: category_stars = 4
            elif skill_factor >= 0.4: category_stars = 3
            elif skill_factor >= 0.2: category_stars = 2
            elif skill_factor > 0: category_stars = 1
            # else category_stars = 0
        else:
            # Standard status-based stars for other categories
            if st == "match": category_stars = 5
            elif c in ("seniority", "country"):
                # Binary categories: only exact match earns stars
                category_stars = 0
            elif st == "related": category_stars = 3
            elif st == "unrelated": category_stars = 1
            # not_assessed = 0
        
        # Add category appraisal with weightage, stars, and status
        display_name = category_display_names.get(c, c.capitalize())
        descriptor = status_to_descriptor(st)
        weight_percent = _int_weights.get(c, int(round(final_weights[c])))
        
        # Generate star string — "Unable to Access" when data was unavailable
        if st == "not_assessed":
            star_string = "Unable to Access"
        else:
            star_string = "★" * category_stars + "☆" * (5 - category_stars)
        
        category_appraisals[display_name] = {
            "rating": descriptor,
            "status": st,
            "comment": cm if cm else descriptor,
            "weight_percent": weight_percent,
            "stars": category_stars,
            "star_string": star_string
        }

    final_percent = min(100, max(0, int(round(total_score))))
    stars = int(round(final_percent / 20.0))
    if stars > 5: stars = 5
    
    # Generate recruiter-style narrative aligned with score
    final_comments = generate_recruiter_narrative(assessment_results, active_criteria, data, total_score)
    
    # Generate overall comment (≤MAX_COMMENT_LENGTH chars, professional tone)
    if final_percent >= ASSESSMENT_EXCELLENT_THRESHOLD:
        overall_comment = "Excellent match for the role requirements"
    elif final_percent >= ASSESSMENT_GOOD_THRESHOLD:
        overall_comment = "Good fit with relevant experience"
    elif final_percent >= ASSESSMENT_MODERATE_THRESHOLD:
        overall_comment = "Moderate alignment, some gaps present"
    else:
        overall_comment = "Limited match to role requirements"
    
    # Ensure overall comment is ≤MAX_COMMENT_LENGTH chars
    if len(overall_comment) > MAX_COMMENT_LENGTH:
        overall_comment = overall_comment[:COMMENT_TRUNCATE_LENGTH] + "..."
    
    # Determine assessment level display
    level_display = "Assessment"
    
    # DIAGNOSTIC LOGGING - Log assessment inputs and output
    logger.info(f"[CORE_ASSESS] LinkedIn: {linkedinurl[:50] if linkedinurl else 'N/A'}")
    logger.info(f"[CORE_ASSESS] Profile data keys: {', '.join([k for k, v in data.items() if v])}")
    logger.info(f"[CORE_ASSESS] Active criteria ({len(active_criteria)}): {', '.join(active_criteria)}")
    
    # Log skillset ratio calculation if skillset is in active criteria
    if "skillset" in active_criteria and "skillset" in assessment_results:
        skillset_res = assessment_results["skillset"]
        match_count = skillset_res.get("match_count", 0)
        total_count = skillset_res.get("total_count", 0)
        match_ratio = skillset_res.get("match_ratio", 0.0)
        skillset_weight = final_weights.get("skillset", 0)
        skillset_points = breakdown.get("skillset", 0)
        logger.info(f"[CORE_ASSESS] Skillset: {match_count}/{total_count} matched ({match_ratio:.1%}), weight={skillset_weight}%, points={skillset_points}")
    
    logger.info(f"[CORE_ASSESS] Final score: {final_percent}% | Stars: {stars} | Level: {assessment_level}")
    logger.info(f"[CORE_ASSESS] Breakdown: {breakdown}")
    
    out_obj = {
        "assessment_level": level_display,
        "is_level2": is_level2 or (assessment_level == "L2"),
        "stars": stars,
        "total_score": f"{final_percent}%",
        "criteria": breakdown,
        "comments": final_comments,
        "overall_comment": overall_comment,
        "category_appraisals": category_appraisals
    }

    if linkedinurl:
        _assess_username = _CV_USERNAME_SAFE_RE.sub('', (data.get("username") or ""))
        _uname_suffix = f"_{_assess_username}" if _assess_username else ""
        safe_name = "assessment_" + hashlib.sha256(linkedinurl.encode("utf-8")).hexdigest()[:16] + _uname_suffix + ".json"
        assess_dir = os.path.join(OUTPUT_DIR, "assessments")
        os.makedirs(assess_dir, exist_ok=True)
        out_path = os.path.join(assess_dir, safe_name)
        
        try:
            with open(out_path, "w", encoding="utf-8") as f:
                json.dump(out_obj, f, indent=2, ensure_ascii=False)
            out_obj["file"] = f"output/assessments/{safe_name}"
        except Exception as e:
            logger.warning(f"Failed to write assessment file: {e}")

    return out_obj

def analyze_cv_background(linkedinurl, pdf_bytes, process_id=None, override_role_tag=None):
    """
    Extracts text from PDF bytes using synchronous helper,
    updates DB, and performs automatic assessment.
    
    SOURCE OF TRUTH: CV Column
    - Gemini exclusively references the cv column in the process table (Postgres)
    - No merging with existing process.skillset data
    - No reconciliation with login.jskillset
    - Employment history strictly follows format: "Job Title, Company, StartYear to EndYear"
      OR "Job Title, Company, StartYear to present" (for current positions)
    
    process_id: optional int — process table primary key. Used to identify the record
    when linkedinurl is empty (e.g. candidates docked without a LinkedIn URL).
    override_role_tag: optional str — role_tag selected by the recruiter in Step 3.
    When provided, takes precedence over the role_tag stored in the process table,
    ensuring the criteria-file jskillset for the selected role is written immediately.
    """
    linkedinurl = (linkedinurl or "").strip()
    if not linkedinurl and not process_id:
        logger.warning("[CV BG] Skipping CV analysis: linkedinurl is empty and no process_id provided")
        return
    # Identifier label used in log messages
    _id_label = linkedinurl[:50] if linkedinurl else f"process_id={process_id}"
    _CV_ANALYZE_SEMAPHORE.acquire()
    try:
        obj = _analyze_cv_bytes_sync(pdf_bytes)
        if not obj:
            logger.warning(f"[CV BG] Analysis returned None for {_id_label}")
            return

        skillset = obj.get("skillset", [])
        total_exp = obj.get("total_experience_years", 0)
        tenure = obj.get("tenure", 0.0)  # Extract calculated tenure
        experience = obj.get("experience", [])
        education = obj.get("education", [])
        product_list = obj.get("product_list", [])
        seniority = obj.get("seniority", "")
        sector = obj.get("sector", "")
        job_family = obj.get("job_family", "")
        company = obj.get("company", "")
        job_title = obj.get("job_title", "")
        country = obj.get("country", "")

        # Always look up ML_Master_Company.json first when a company is present.
        # The Master file data is the highest-quality validated source, so it takes
        # priority over Gemini's sector extraction (which can misclassify CROs like
        # IQVIA or Labcorp as Pharmaceutical).
        if company:
            ml_sector = _resolve_sector_from_ml_master(company, log_prefix="[CV BG] ")
            if ml_sector:
                sector = ml_sector

        # Always look up ML_Master_Jobfamily_Seniority.json first when a job title is present.
        # ML Master is the primary source of truth — it overrides Gemini's job_family and
        # seniority values as well as any hardcoded keyword rules.
        if job_title:
            ml_jf, ml_sn = _resolve_jobfamily_seniority_from_ml_master(job_title, log_prefix="[CV BG] ")
            if ml_jf:
                job_family = ml_jf
            if ml_sn:
                seniority = ml_sn

        # Log product extraction for debugging
        if product_list and len(product_list) > 0:
            truncated = '...' if len(product_list) > 3 else ''
            logger.info(f"[CV BG] Extracted {len(product_list)} products for {_id_label}: {product_list[:3]}{truncated}")
        else:
            logger.warning(f"[CV BG] No products extracted for {_id_label} (company: {company})")
        
        # Create skillset string without length limits (DB columns now TEXT type)
        skillset_raw = ",".join([str(s).strip() for s in skillset if str(s).strip()])
        skillset_str = skillset_raw
        experience_str = "\n".join([str(e).strip() for e in experience if str(e).strip()])
        education_str = "\n".join([str(e).strip() for e in education if str(e).strip()])
        product_str = ", ".join([str(p).strip() for p in product_list if str(p).strip()])
        
        import psycopg2
        from psycopg2 import sql as pgsql
        pg_host=os.getenv("PGHOST","localhost"); pg_port=int(os.getenv("PGPORT","5432"))
        pg_user=os.getenv("PGUSER","postgres"); pg_password=os.getenv("PGPASSWORD", "")
        pg_db=os.getenv("PGDATABASE","candidate_db")
        conn=psycopg2.connect(host=pg_host, port=pg_port, user=pg_user, password=pg_password, dbname=pg_db)
        cur=conn.cursor()
        
        normalized = _normalize_linkedin_to_path(linkedinurl) if linkedinurl else None

        # Determine the WHERE clause for all subsequent queries.
        # Priority: process_id (most precise) > sourcing lookup by URL > URL direct
        if process_id:
            where_clause = "id = %s"
            params = [process_id]
        else:
            where_clause = "linkedinurl = %s"
            params = [linkedinurl]
        
        sourcing_id = None
        if linkedinurl and not process_id:
            try:
                cur.execute("SELECT column_name FROM information_schema.columns WHERE table_schema='public' AND table_name='process' AND column_name='id'")
                if cur.fetchone():
                    cur.execute("SELECT id FROM sourcing WHERE linkedinurl = %s LIMIT 1", (linkedinurl,))
                    row_sid = cur.fetchone()
                    if row_sid: sourcing_id = row_sid[0]
            except Exception: pass
            
            if sourcing_id:
                where_clause = "id = %s"; params = [sourcing_id]
        
        # Check available columns and their character limits to ensure 'product' is fetched
        # and to allow safe truncation of VARCHAR fields before writes.
        cur.execute("""
            SELECT column_name, character_maximum_length
            FROM information_schema.columns
            WHERE table_schema='public' AND table_name='process'
        """)
        _col_info = cur.fetchall()
        all_cols = {r[0].lower() for r in _col_info}
        # Map column_name -> max_length (None for TEXT/unlimited columns)
        _process_col_max = {r[0].lower(): r[1] for r in _col_info if r[1] is not None}
        has_product = 'product' in all_cols
        
        # Build SELECT query dynamically based on columns
        select_fields = ["skillset"]
        if has_product: select_fields.append("product")
        # Include username if available so we can reconcile against login.jskillset
        if 'username' in all_cols:
            select_fields.append("username")
        # Include exp and tenure to implement write-once guard (never overwrite once set)
        if 'exp' in all_cols:
            select_fields.append("exp")
        if 'tenure' in all_cols:
            select_fields.append("tenure")
        
        cur.execute(pgsql.SQL("SELECT {} FROM process WHERE {}").format(pgsql.SQL(", ").join(pgsql.Identifier(f) for f in select_fields), pgsql.SQL(where_clause)), tuple(params))
        row = cur.fetchone()
        
        if not row and not sourcing_id and normalized:
             cur.execute("SELECT column_name FROM information_schema.columns WHERE table_schema='public' AND table_name='process' AND column_name='normalized_linkedin'")
             if cur.fetchone():
                 where_clause = "normalized_linkedin = %s"; params = [normalized]
                 cur.execute(pgsql.SQL("SELECT {} FROM process WHERE {}").format(pgsql.SQL(", ").join(pgsql.Identifier(f) for f in select_fields), pgsql.SQL(where_clause)), (normalized,))
                 row = cur.fetchone()
        
        if row:
            # SOURCE OF TRUTH: CV Column - use extracted skillset exclusively (no merge with existing)
            # Gemini must exclusively reference the cv column in the process table (Postgres)
            # skillset_str was already set from CV extraction above - no modification needed
            # Merge product if available (index 1 if product was selected)
            if has_product:
                prod_idx = select_fields.index("product") if "product" in select_fields else None
                if prod_idx is not None and len(row) > prod_idx and row[prod_idx]:
                    curr_prod = [p.strip() for p in row[prod_idx].split(',') if p.strip()]
                    prod_set = set(curr_prod)
                    prod_set.update([str(p).strip() for p in product_list if str(p).strip()])
                    product_str = ", ".join(list(prod_set))
                elif prod_idx is not None and len(row) > prod_idx and not row[prod_idx]:
                    # Existing DB product is empty — use freshly extracted product_str (may also be empty)
                    pass  # product_str already set from CV extraction above
        
        # Write-once guard for total_experience_years / tenure / experience:
        # If exp is already populated in the DB, do not overwrite it.
        # This prevents Gemini's non-deterministic re-runs from changing previously stored values.
        _exp_already_set = False
        if row and 'exp' in select_fields:
            _exp_idx = select_fields.index("exp")
            _existing_exp = row[_exp_idx] if len(row) > _exp_idx else None
            if _existing_exp is not None and str(_existing_exp).strip() not in ('', '0', '0.0'):
                _exp_already_set = True
                logger.info(f"[CV BG] exp already set for {_id_label} ({_existing_exp}) — skipping recalculation")
        
        # If Gemini extracted no products and DB already has a non-empty product value,
        # keep the existing DB value rather than overwriting with empty string.
        if not product_str and row and has_product and 'product' in select_fields:
            prod_idx = select_fields.index("product")
            if len(row) > prod_idx and row[prod_idx]:
                product_str = row[prod_idx]  # Preserve existing products
                logger.info(f"[CV BG] Preserved existing product data for {_id_label} (Gemini extraction returned empty)")
        
        # PATCH: Infer geographic region
        geo_region = _infer_region_from_country(country) if country else ""

        # Use previously fetched columns info
        cols = all_cols
        
        update_parts = []
        update_vals = []
        # NOTE: Deliberately exclude 'skillset' here so we can reconcile against jskillset
        mapping = {
            "exp": str(total_exp), "experience": experience_str,
            "education": education_str, "product": product_str, "seniority": seniority,
            "sector": sector, "job_family": job_family, "company": company, "jobtitle": job_title,
            "country": country, 
            "geographic": geo_region, # Use inferred region, NOT country
            "tenure": float(tenure) # Numeric average tenure
        }
        
        for col_key, val in mapping.items():
            db_col = col_key
            if col_key == "job_family" and "jobfamily" in cols: db_col = "jobfamily"
            if col_key == "jobtitle":
                if "jobtitle" in cols: db_col = "jobtitle"
                elif "role" in cols: db_col = "role"
            
            # Check tenure specifically as it's a numeric type, ensure column exists
            if col_key == "tenure" and "tenure" not in cols:
                continue

            # Write-once guard: skip exp, experience, tenure if already populated
            if col_key in ("exp", "experience", "tenure") and _exp_already_set:
                continue

            if db_col in cols and (val is not None and val != ""):
                # Truncate string values to the column's VARCHAR limit to prevent
                # "value too long for type character varying(N)" errors.
                if isinstance(val, str) and db_col in _process_col_max:
                    val = val[:_process_col_max[db_col]]
                update_parts.append(pgsql.SQL("{} = %s").format(pgsql.Identifier(db_col)))
                update_vals.append(val)
            
        if update_parts:
            update_sql = pgsql.SQL("UPDATE process SET {} WHERE {}").format(pgsql.SQL(", ").join(update_parts), pgsql.SQL(where_clause))
            update_vals.extend(params)
            cur.execute(update_sql, tuple(update_vals))
            conn.commit()

        # Update Sourcing table if fields exist
        if company or job_title or country:
            # Fetch sourcing column VARCHAR limits to prevent "value too long" errors.
            # Failures are non-fatal — worst case, truncation is skipped and the sourcing
            # update may still fail (caught by its own try/except below).
            _sourcing_col_max = {}
            try:
                cur.execute("""
                    SELECT column_name, character_maximum_length
                    FROM information_schema.columns
                    WHERE table_schema='public' AND table_name='sourcing'
                    AND character_maximum_length IS NOT NULL
                """)
                _sourcing_col_max = {r[0].lower(): r[1] for r in cur.fetchall()}
            except Exception as _e_sl:
                logger.debug(f"[CV BG] Could not fetch sourcing column limits: {_e_sl}")

            def _trunc_to_limit(val, col, col_max):
                """Truncate str val to the column's VARCHAR limit if one exists."""
                if isinstance(val, str) and col in col_max:
                    return val[:col_max[col]]
                return val

            s_upd = []; s_vals = []
            if company: s_upd.append("company = %s"); s_vals.append(_trunc_to_limit(company, 'company', _sourcing_col_max))
            if job_title: s_upd.append("jobtitle = %s"); s_vals.append(_trunc_to_limit(job_title, 'jobtitle', _sourcing_col_max))
            if country: s_upd.append("country = %s"); s_vals.append(_trunc_to_limit(country, 'country', _sourcing_col_max))
            
            s_vals.append(linkedinurl) # where clause
            if s_upd:
                try:
                    cur.execute(pgsql.SQL("UPDATE sourcing SET {} WHERE linkedinurl = %s").format(pgsql.SQL(", ".join(s_upd))), tuple(s_vals))
                    conn.commit()
                except Exception as e_src:
                    logger.warning(f"[CV BG] Sourcing update failed: {e_src}")

        # Trigger Auto-Assessment
        fetch_cols_sql = pgsql.SQL("SELECT jobtitle, company, country, role_tag, userid, username, skillset FROM process WHERE {}").format(pgsql.SQL(where_clause))
        cur.execute(fetch_cols_sql, tuple(params))
        ctx_row = cur.fetchone()

        if ctx_row:
            job_title_db, company_db, country_db, role_tag_db, userid_db, username_db, skillset_db = ctx_row
            if not role_tag_db and username_db:
                # Try sourcing table first (authoritative), fallback to process table
                cur.execute("SELECT role_tag FROM sourcing WHERE username=%s AND role_tag IS NOT NULL AND role_tag != '' LIMIT 1", (username_db,))
                rt_row = cur.fetchone()
                if not rt_row or not rt_row[0]:
                    cur.execute("SELECT role_tag FROM process WHERE username=%s AND role_tag IS NOT NULL AND role_tag != '' LIMIT 1", (username_db,))
                    rt_row = cur.fetchone()
                if rt_row and rt_row[0]: role_tag_db = rt_row[0]

            # SOURCE OF TRUTH: CV Column - use extracted skillset exclusively
            # Gemini must exclusively reference the cv column in the process table (Postgres)
            # No reconciliation with login.jskillset - CV is the single source of truth
            final_skillset_str = skillset_str
            
            try:
                # Persist CV-extracted skillset to process table only if not already populated.
                # Guardrail: skillset must not overwrite an existing value — assessment sets it authoritatively.
                if 'skillset' in cols and final_skillset_str:
                    up_where = where_clause
                    up_params = list(params)
                    cur.execute(
                        pgsql.SQL("UPDATE process SET skillset = %s WHERE {} AND (skillset IS NULL OR TRIM(skillset) = '')").format(pgsql.SQL(up_where)),
                        tuple([final_skillset_str] + up_params)
                    )
                    if cur.rowcount:
                        conn.commit()
                        logger.info(f"[CV BG] CV-extracted skillset saved for {_id_label} (exclusive source)")
                    else:
                        logger.info(f"[CV BG] Skillset already set for {_id_label} — skipping CV overwrite")
            except Exception as e_save:
                logger.warning(f"[CV BG] Failed to persist CV skillset: {e_save}")

            if role_tag_db:
                # Use criteria file skillset as target_skills (role-specific), with
                # login.jskillset only as a last resort.  This matches _assess_and_persist
                # behaviour and prevents the L1 assessment from scoring against the wrong role.
                effective_role_for_skills = override_role_tag or role_tag_db
                target_skills_final = []
                if username_db and effective_role_for_skills:
                    _cv_criteria = _read_search_criteria(username_db, effective_role_for_skills)
                    if _cv_criteria:
                        target_skills_final = _cv_criteria.get("Skillset") or []
                if not target_skills_final:
                    target_skills_final = _fetch_jskillset(username_db) if username_db else []
                c_skills_list = [s.strip() for s in (final_skillset_str if 'final_skillset_str' in locals() else (skillset_db or "")).split(',') if s.strip()]
                profile_data = {
                    "job_title": job_title_db or "",
                    "role_tag": role_tag_db,
                    "company": company_db or "",
                    "country": country_db or "",
                    "seniority": seniority or "",
                    "sector": sector or "",
                    "experience_text": experience_str,
                    "target_skills": target_skills_final,
                    "candidate_skills": c_skills_list,
                    "linkedinurl": linkedinurl,
                    "tenure": tenure,
                    "username": username_db or "",  # Include username in assessment filename
                }

                # Idempotency check: background auto-assessment is L1; skip if rating already exists
                _bg_skip = False
                if 'rating' in cols:
                    try:
                        _ensure_rating_metadata_columns(cur, conn)
                        cur.execute(
                            pgsql.SQL("SELECT rating, rating_level FROM process WHERE {}").format(pgsql.SQL(where_clause)),
                            tuple(params)
                        )
                        _bg_meta_row = cur.fetchone()
                        if _bg_meta_row and _bg_meta_row[0]:
                            _bg_existing_meta = {"level": (_bg_meta_row[1] or "").upper()}
                            _bg_allow, _bg_reason = _should_overwrite_existing(_bg_existing_meta, "L1", False)
                            if not _bg_allow:
                                logger.info(f"[CV BG] Skipping auto-assessment for {_id_label}: {_bg_reason}")
                                _bg_skip = True
                    except Exception as _e_bg_idem:
                        logger.warning(f"[CV BG] Idempotency check failed (continuing): {_e_bg_idem}")

                if not _bg_skip:
                    assessment_result = _core_assess_profile(profile_data)
                    if 'rating' in cols:
                        _ensure_rating_metadata_columns(cur, conn)
                        rating_json = json.dumps(assessment_result, ensure_ascii=False)
                        upd_sql = pgsql.SQL(
                            "UPDATE process SET rating = %s, rating_level = %s, "
                            "rating_updated_at = NOW(), "
                            "rating_version = COALESCE(rating_version, 0) + 1 "
                            "WHERE {}"
                        ).format(pgsql.SQL(where_clause))
                        cur.execute(upd_sql, (rating_json, "L1", *params))
                        conn.commit()
                
                # --- NEW: Trigger role_tag -> jskill sync during background CV analysis ---
                if role_tag_db and "jskill" in cols:
                     upd_js_sql = pgsql.SQL("UPDATE process SET jskill = %s WHERE {}").format(pgsql.SQL(where_clause))
                     cur.execute(upd_js_sql, (role_tag_db, *params))
                     conn.commit()
                
                # Trigger role-specific jskillset sync from criteria file (login.jskillset is fallback only)
                # override_role_tag (from Step 3 selection) takes precedence over DB value
                effective_role_tag = override_role_tag or role_tag_db
                if username_db and effective_role_tag:
                    _sync_criteria_jskillset_to_process(username_db, effective_role_tag, linkedinurl, normalized, process_id=process_id)
                elif username_db:
                    _sync_login_jskillset_to_process(username_db, linkedinurl, normalized, process_id=process_id)
                # --------------------------------------------------------------------------

        cur.close(); conn.close()
    except Exception as e:
        logger.error(f"Error in CV analysis background task: {e}")
    finally:
        _CV_ANALYZE_SEMAPHORE.release()

@app.post("/process/parse_cv_and_update")
def process_parse_cv_and_update():
    """
    Manually trigger CV analysis via sync helper.
    
    SOURCE OF TRUTH: CV Column
    - Fetches CV from process.cv column and analyzes with Gemini
    - Returns skillset and employment history strictly from CV
    - Employment history format: "Job Title, Company, StartYear to EndYear"
      OR "Job Title, Company, StartYear to present" (for current positions)
    """
    data = request.get_json(force=True, silent=True) or {}
    linkedinurl = (data.get("linkedinurl") or "").strip()
    if not linkedinurl: return jsonify({"error": "linkedinurl required"}), 400

    try:
        import psycopg2
        pg_host=os.getenv("PGHOST","localhost"); pg_port=int(os.getenv("PGPORT","5432"))
        pg_user=os.getenv("PGUSER","postgres"); pg_password=os.getenv("PGPASSWORD", "")
        pg_db=os.getenv("PGDATABASE","candidate_db")
        conn=psycopg2.connect(host=pg_host, port=pg_port, user=pg_user, password=pg_password, dbname=pg_db)
        cur=conn.cursor()
        
        linkedin_norm = linkedinurl.split('?')[0].rstrip('/')
        linkedin_path = _normalize_linkedin_to_path(linkedinurl)
        
        cur.execute("SELECT column_name FROM information_schema.columns WHERE table_schema='public' AND table_name='process' AND column_name='cv'")
        if not cur.fetchone():
             cur.close(); conn.close()
             return jsonify({"error": "No CV column in DB"}), 500

        row = None
        if linkedin_path:
             cur.execute("SELECT column_name FROM information_schema.columns WHERE table_schema='public' AND table_name='process' AND column_name='normalized_linkedin'")
             if cur.fetchone():
                 cur.execute("SELECT cv FROM process WHERE normalized_linkedin = %s AND cv IS NOT NULL LIMIT 1", (linkedin_path,))
                 row = cur.fetchone()
        
        if not row:
             cur.execute("SELECT cv FROM process WHERE linkedinurl = %s AND cv IS NOT NULL LIMIT 1", (linkedin_norm,))
             row = cur.fetchone()
             
        if not row or not row[0]:
             cur.close(); conn.close()
             return jsonify({"error": "CV not found for this profile"}), 404
             
        pdf_bytes = bytes(row[0])
        cur.close(); conn.close()
        
        obj = _analyze_cv_bytes_sync(pdf_bytes)
        if not obj:
             return jsonify({"error": "Analysis returned no data"}), 500
             
        # Persist synchronously so DB fields are ready before bulk_assess runs.
        # (Background thread approach caused a race: bulk_assess read empty fields
        # because the thread hadn't finished writing when the next request arrived.)
        analyze_cv_background(linkedinurl, pdf_bytes)
        
        return jsonify({
            "skillset": obj.get("skillset", []),
            "total_years": obj.get("total_experience_years", 0),
            "tenure": obj.get("tenure", 0.0), # Includes tenure in API response
            "experience": obj.get("experience", []),
            "education": obj.get("education", []),
            "product": obj.get("product_list", []),
            "company": obj.get("company", ""),
            "job_title": obj.get("job_title", ""),
            "country": obj.get("country", ""),
            "experience_text": "\n".join(obj.get("experience", [])),
            "education_text": "\n".join(obj.get("education", []))
        }), 200

    except Exception as e:
        logger.error(f"[Parse CV Update] {e}")
        return jsonify({"error": str(e)}), 500

@app.get("/process/pending_assessments")
def process_pending_assessments():
    userid = (request.args.get("userid") or "").strip()
    if not userid: return jsonify({"rows": []})
    page_size = int(request.args.get("page_size", 100))
    try:
        import psycopg2
        pg_host=os.getenv("PGHOST","localhost"); pg_port=int(os.getenv("PGPORT","5432"))
        pg_user=os.getenv("PGUSER","postgres"); pg_password=os.getenv("PGPASSWORD", "")
        pg_db=os.getenv("PGDATABASE","candidate_db")
        conn=psycopg2.connect(host=pg_host, port=pg_port, user=pg_user, password=pg_password, dbname=pg_db)
        cur=conn.cursor()
        cur.execute("SELECT column_name FROM information_schema.columns WHERE table_schema='public' AND table_name='process' AND column_name='normalized_linkedin'")
        has_normalized = bool(cur.fetchone())
        cols = ["name", "company", "jobtitle", "country", "linkedinurl", "experience", "rating"]
        if has_normalized: cols.append("normalized_linkedin")
        from psycopg2 import sql
        query = sql.SQL("SELECT {fields} FROM process WHERE userid=%s ORDER BY id DESC LIMIT %s").format(
            fields=sql.SQL(', ').join(map(sql.Identifier, cols)))
        cur.execute(query, (userid, page_size))
        rows = []
        for r in cur.fetchall():
            row_dict = {"name": r[0], "company": r[1], "jobtitle": r[2], "country": r[3], "linkedinurl": r[4], "experience": r[5], "rating": r[6]}
            if has_normalized: row_dict["normalized_linkedin"] = r[7]
            rows.append(row_dict)
        cur.close(); conn.close()
        return jsonify({"rows": rows})
    except Exception as e:
        return jsonify({"error": str(e)}), 500


def _generate_vskillset_for_profile(linkedinurl, target_skills, experience_text="", cv_data=None):
    """
    Generate vskillset for a profile using Gemini inference.
    Returns list of skill evaluation results or None if failed.
    Idempotent: returns existing vskillset from DB without regenerating if already populated.
    """
    if not target_skills or len(target_skills) == 0:
        return None
    
    try:
        import psycopg2
        pg_host = os.getenv("PGHOST", "localhost")
        pg_port = int(os.getenv("PGPORT", "5432"))
        pg_user = os.getenv("PGUSER", "postgres")
        pg_password = os.getenv("PGPASSWORD", "")
        pg_db = os.getenv("PGDATABASE", "candidate_db")

        # --- Idempotency guard: if vskillset already exists in DB, return it without re-running Gemini ---
        if linkedinurl:
            try:
                _norm_gen = linkedinurl.lower().strip().rstrip('/')
                _conn_idem = psycopg2.connect(host=pg_host, port=pg_port, user=pg_user, password=pg_password, dbname=pg_db)
                _cur_idem = _conn_idem.cursor()
                _cur_idem.execute("""
                    SELECT vskillset FROM process
                    WHERE LOWER(TRIM(TRAILING '/' FROM linkedinurl)) = %s
                    LIMIT 1
                """, (_norm_gen,))
                _idem_row = _cur_idem.fetchone()
                _cur_idem.close(); _conn_idem.close()
                if _idem_row and _idem_row[0]:
                    _vs_existing = _idem_row[0]
                    if isinstance(_vs_existing, str):
                        try:
                            _vs_existing = json.loads(_vs_existing)
                        except Exception:
                            _vs_existing = []
                    if isinstance(_vs_existing, list) and len(_vs_existing) > 0:
                        logger.info(f"[vskillset_gen] Returning existing vskillset ({len(_vs_existing)} items) for {linkedinurl[:50]} — skipping Gemini re-inference")
                        return _vs_existing
            except Exception as _e_idem_gen:
                logger.warning(f"[vskillset_gen] Idempotency check failed ({_e_idem_gen}); proceeding with generation")
        
        # Use experience as primary context, cv as fallback
        profile_context = experience_text if experience_text else ""
        if not profile_context and cv_data:
            # Extract text from CV if needed
            try:
                if isinstance(cv_data, bytes):
                    from pypdf import PdfReader
                    import io
                    reader = PdfReader(io.BytesIO(cv_data))
                    text = ""
                    for page in reader.pages:
                        t = page.extract_text()
                        if t: text += t + "\n"
                    profile_context = text[:3000]
            except Exception:
                pass
        
        if not profile_context:
            logger.warning(f"[vskillset_gen] No experience or CV data for {linkedinurl}")
            return None
        
        # Call unified LLM to evaluate skills
        
        prompt = f"""SYSTEM:
You are an expert technical recruiter evaluating candidate skillsets based on their work experience.

TASK:
For each skill in the list below, evaluate the candidate's likely proficiency based on their experience.
Assign a probability score (0-100) and categorize as Low (<40), Medium (40-74), or High (75-100).
Provide clear reasoning based on job titles, companies, and experience patterns.

CANDIDATE PROFILE:
{profile_context[:3000]}

SKILLS TO EVALUATE:
{json.dumps(target_skills, ensure_ascii=False)}

OUTPUT FORMAT (JSON):
{{
  "evaluations": [
    {{
      "skill": "skill_name",
      "probability": 0-100,
      "category": "Low|Medium|High",
      "reason": "Brief explanation based on companies and roles"
    }}
  ]
}}

Return ONLY the JSON object, no other text."""
        
        raw_text = (unified_llm_call_text(prompt) or "").strip()
        
        # Extract JSON from response
        parsed = _extract_json_object(raw_text)
        
        if not parsed or "evaluations" not in parsed:
            logger.warning(f"[vskillset_gen] Gemini returned invalid JSON")
            return None
        
        results = parsed["evaluations"]
        
        # Ensure all required fields are present
        for item in results:
            if "probability" not in item:
                item["probability"] = 50
            if "category" not in item:
                prob = item.get("probability", 50)
                if prob >= 75:
                    item["category"] = "High"
                elif prob >= 40:
                    item["category"] = "Medium"
                else:
                    item["category"] = "Low"
            if "reason" not in item:
                item["reason"] = "No reasoning provided"
        
        # Persist to database
        try:
            conn = psycopg2.connect(host=pg_host, port=pg_port, user=pg_user, password=pg_password, dbname=pg_db)
            cur = conn.cursor()
            
            vskillset_json = json.dumps(results, ensure_ascii=False)
            confirmed_skills = [item["skill"] for item in results if item["category"] == "High"]
            # Ensure all skills are strings before joining
            skillset_str = ", ".join([str(s) for s in confirmed_skills if s])
            
            # Check if vskillset column exists
            cur.execute("""
                SELECT column_name 
                FROM information_schema.columns
                WHERE table_schema='public' AND table_name='process' 
                  AND column_name IN ('vskillset', 'skillset')
            """)
            available_cols = {r[0] for r in cur.fetchall()}
            
            # Use normalized URL for consistent match with idempotency checks elsewhere
            _norm_persist = linkedinurl.lower().strip().rstrip('/')
            # Update vskillset if column exists
            if 'vskillset' in available_cols:
                cur.execute(
                    "UPDATE process SET vskillset = %s WHERE LOWER(TRIM(TRAILING '/' FROM linkedinurl)) = %s",
                    (vskillset_json, _norm_persist)
                )
                logger.info(f"[vskillset_gen] Persisted vskillset for {linkedinurl[:50]}")
            
            # Update skillset: merge new High skills into existing value (add only; never remove or replace)
            if 'skillset' in available_cols and confirmed_skills:
                cur.execute(
                    "SELECT skillset FROM process WHERE LOWER(TRIM(TRAILING '/' FROM linkedinurl)) = %s",
                    (_norm_persist,)
                )
                _sk_row = cur.fetchone()
                _existing_sk = (_sk_row[0] or "") if _sk_row else ""
                _existing_parts = [s.strip() for s in _existing_sk.split(",") if s.strip()]
                _existing_set = {s.lower() for s in _existing_parts}
                _new_high = [s for s in confirmed_skills if s.lower() not in _existing_set]
                if _new_high:
                    _merged_sk = ", ".join(_existing_parts + _new_high)
                    cur.execute(
                        "UPDATE process SET skillset = %s WHERE LOWER(TRIM(TRAILING '/' FROM linkedinurl)) = %s",
                        (_merged_sk, _norm_persist)
                    )
                    logger.info(f"[vskillset_gen] Merged {len(_new_high)} new High skills into skillset for {linkedinurl[:50]}")
                else:
                    logger.info(f"[vskillset_gen] No new High skills for {linkedinurl[:50]} — skillset unchanged")
            
            conn.commit()
            cur.close()
            conn.close()
        except Exception as e_db:
            logger.warning(f"[vskillset_gen] Failed to persist for {linkedinurl}: {e_db}")
        
        return results
        
    except Exception as e:
        logger.error(f"[vskillset_gen] Error for {linkedinurl}: {e}")
        return None

@app.post("/process/bulk_assess")
@_rate(_bulk_assess_flask_limit)
@_check_user_rate("bulk_assess")
def process_bulk_assess():
    """
    Accepts JSON payload like:
    {
      "userid": "9896945",
      "linkedinurls": ["https://...","https://..."],
      "async": true,
      "custom_weights": {...},
      "assessment_level": "L2"
    }
    If async is true returns job_id and starts background worker, otherwise runs synchronously and returns results.
    Uses existing _core_assess_profile to perform per-profile assessment and persists rating into process.rating (if column exists).
    """
    payload = request.get_json(force=True, silent=True) or {}
    userid = (payload.get("userid") or "").strip()
    # Accept either a flat list of LinkedIn URLs or a richer candidates list
    # (dicts with linkedinurl + optional name/company for records with no URL).
    raw_candidates = payload.get("candidates") or []
    linkedinurls = payload.get("linkedinurls") or []
    if isinstance(linkedinurls, str):
        linkedinurls = [linkedinurls]

    # Merge: raw_candidates takes precedence when provided
    if raw_candidates and isinstance(raw_candidates, list):
        candidates_list = [
            c if isinstance(c, dict) else {"linkedinurl": str(c)}
            for c in raw_candidates
        ]
    else:
        candidates_list = [{"linkedinurl": u} for u in (linkedinurls if isinstance(linkedinurls, list) else [])]

    if not candidates_list:
        return jsonify({"error": "linkedinurls list or candidates list required"}), 400

    async_flag = bool(payload.get("async"))
    custom_weights = payload.get("custom_weights") or {}
    assessment_level = (payload.get("assessment_level") or payload.get("assessmentLevel") or "L2").strip().upper()
    # Prefer username from payload; fall back to the cookie set by the login endpoint.
    # The Dock In wizard does not always include username in the JSON body, so the cookie
    # is the reliable source for identifying the recruiter when resolving criteria files.
    username = (payload.get("username") or request.cookies.get("username") or "").strip()
    force_reassess = bool(payload.get("force_reassess") or False)
    # role_tag override: explicitly selected by the user in Step 3 of the Dock In wizard.
    # Used as a high-priority fallback when the record is not in the sourcing table (newly docked).
    override_role_tag = (payload.get("role_tag") or "").strip()
    # require_cv: Normal DB mode sets this to False so records without CV are still assessed.
    # Analytic DB mode (CV upload) keeps it True (default).
    require_cv = bool(payload.get("require_cv", True))

    # helper for single assess + persist
    def _assess_and_persist(linkedinurl, candidate_name="", candidate_company="", candidate_process_id=None):
        try:
            # Fetch profile data from database
            import psycopg2
            from psycopg2 import sql
            pg_host=os.getenv("PGHOST","localhost")
            pg_port=int(os.getenv("PGPORT","5432"))
            pg_user=os.getenv("PGUSER","postgres")
            pg_password=os.getenv("PGPASSWORD", "")
            pg_db=os.getenv("PGDATABASE","candidate_db")
            conn=psycopg2.connect(host=pg_host, port=pg_port, user=pg_user, password=pg_password, dbname=pg_db)
            cur=conn.cursor()
            
            # Fetch profile data from process table
            normalized = None
            try:
                normalized = _normalize_linkedin_to_path(linkedinurl)
            except Exception:
                normalized = None
            # Track the process table PK for no-LinkedIn records so role_tag / rating
            # updates can use WHERE id = X instead of WHERE linkedinurl = '' (which
            # either matches nothing or matches all no-URL rows incorrectly).
            _nolinkedin_id = None

            # --- Backend ineligibility check ---
            # If the candidate was marked ineligible in the sourcing table (because the user
            # proceeded to AutoSourcing with this profile unassessed), block the assessment
            # server-side regardless of what the front end sends.
            try:
                cur.execute("""
                    SELECT assessment_ineligible FROM sourcing
                    WHERE linkedinurl = %s LIMIT 1
                """, (linkedinurl,))
                _inelig_row = cur.fetchone()
                if _inelig_row and _inelig_row[0]:
                    logger.info(f"[BULK_ASSESS] Blocked ineligible candidate (url truncated: {linkedinurl[:40]}…)")
                    try:
                        cur.close()
                    finally:
                        conn.close()
                    return {
                        "linkedinurl": linkedinurl,
                        "result": {
                            "_skipped": True,
                            "error": "Assessment blocked: candidate marked ineligible (AutoSourcing accessed with pending assessments)."
                        }
                    }
            except Exception as _inelig_err:
                # If the column does not exist yet, skip silently – the front end still enforces it.
                # CRITICAL: rollback so the psycopg2 connection is not left in an aborted-transaction
                # state (which would cause every subsequent query on this cursor to fail with
                # "InFailedSqlTransaction").
                try:
                    conn.rollback()
                except Exception:
                    pass
                logger.debug(f"[BULK_ASSESS] Ineligibility check skipped: {_inelig_err}")

            # Fetch by linkedinurl (normalized_linkedin column doesn't exist in all schemas).
            # Always include id so _nolinkedin_id can be set for empty-URL candidates.
            cur.execute("""
                SELECT jobtitle, company, country, seniority, sector, experience, skillset, username, role_tag, cv, tenure, product, id
                FROM process 
                WHERE linkedinurl = %s
                LIMIT 1
            """, (linkedinurl,))
            # Column indices: 0=jobtitle,1=company,2=country,3=seniority,4=sector,5=experience,
            #                 6=skillset,7=username,8=role_tag,9=cv,10=tenure,11=product,12=id
            row = cur.fetchone()

            # When the row was found via the initial query but linkedinurl is empty (''),
            # capture the process PK now so rating/role_tag persists use WHERE id = X.
            if row and not linkedinurl:
                try:
                    _nolinkedin_id = row[12]  # id column, see column index comment above
                except (IndexError, TypeError):
                    _nolinkedin_id = candidate_process_id

            # Fallback: when linkedinurl is empty (profile docked without a LinkedIn URL),
            # try to locate the record by process_id (most reliable), then name+company,
            # then name-only (when exactly one row matches).
            if not row and not linkedinurl and (candidate_process_id or candidate_name):
                # 1. Direct process PK lookup — most reliable, no ambiguity.
                if candidate_process_id:
                    try:
                        cur.execute("""
                            SELECT jobtitle, company, country, seniority, sector, experience, skillset, username, role_tag, cv, tenure, product, id
                            FROM process WHERE id = %s LIMIT 1
                        """, (candidate_process_id,))
                        _pid_row = cur.fetchone()
                        if _pid_row:
                            row = _pid_row
                            try:
                                _nolinkedin_id = row[12]
                            except (IndexError, TypeError):
                                _nolinkedin_id = candidate_process_id
                    except Exception:
                        pass

                # 2. Name + company lookup (only when both are present).
                if not row and candidate_name and candidate_company:
                    try:
                        cur.execute("""
                            SELECT jobtitle, company, country, seniority, sector, experience, skillset, username, role_tag, cv, tenure, product, id
                            FROM process
                            WHERE LOWER(TRIM(name)) = LOWER(%s) AND LOWER(TRIM(company)) = LOWER(%s)
                        """, (candidate_name.strip(), candidate_company.strip()))
                        _name_rows = cur.fetchall()
                        if len(_name_rows) == 1:
                            row = _name_rows[0]
                            try:
                                _nolinkedin_id = row[12]
                            except (IndexError, TypeError):
                                _nolinkedin_id = None
                        elif len(_name_rows) > 1:
                            logger.warning(f"[BULK_ASSESS] Name+company lookup for '{candidate_name}' / '{candidate_company}' returned {len(_name_rows)} rows — skipping ambiguous match")
                    except Exception:
                        pass

                # 3. Name-only lookup — safe only when exactly one row matches globally.
                if not row and candidate_name and not candidate_company:
                    try:
                        cur.execute("""
                            SELECT jobtitle, company, country, seniority, sector, experience, skillset, username, role_tag, cv, tenure, product, id
                            FROM process WHERE LOWER(TRIM(name)) = LOWER(%s)
                        """, (candidate_name.strip(),))
                        _name_rows = cur.fetchall()
                        if len(_name_rows) == 1:
                            row = _name_rows[0]
                            try:
                                _nolinkedin_id = row[12]
                            except (IndexError, TypeError):
                                _nolinkedin_id = None
                        elif len(_name_rows) > 1:
                            logger.warning(f"[BULK_ASSESS] Name-only lookup for '{candidate_name}' returned {len(_name_rows)} rows — skipping ambiguous match")
                    except Exception:
                        pass
            
            # Default values if profile not found
            job_title = ""
            company = ""
            country = ""
            seniority = ""
            sector = ""
            job_family = ""
            experience_text = ""
            candidate_skills = []
            username_db = ""
            role_tag = ""
            cv_data = None
            tenure = None
            product = []
            
            if row:
                job_title = row[0] or ""
                company = row[1] or ""
                country = row[2] or ""
                seniority = row[3] or ""
                sector = row[4] or ""
                experience_text = row[5] or ""
                skillset_str = row[6] or ""
                username_db = row[7] or ""
                role_tag = row[8] or ""
                # Safe extraction of CV data (cv column may not exist in all schemas)
                try:
                    cv_data = row[9] if len(row) >= 10 else None
                    tenure = row[10] if len(row) >= 11 else None
                    product_str = row[11] if len(row) >= 12 else ""
                except (IndexError, TypeError):
                    cv_data = None
                    tenure = None
                    product_str = ""
                
                # Parse skillset
                if skillset_str:
                    candidate_skills = [s.strip() for s in skillset_str.split(',') if s.strip()]
                
                # Parse product
                if product_str:
                    try:
                        # Product could be JSON array or comma-separated string
                        product = json.loads(product_str) if product_str.startswith('[') else [s.strip() for s in product_str.split(',') if s.strip()]
                    except:
                        product = [s.strip() for s in product_str.split(',') if s.strip()]
                    
                    # Log product loading regardless of whether list is empty
                    if product:
                        logger.info(f"[BULK_ASSESS] Loaded {len(product)} products from DB for {linkedinurl[:50]}")
                    else:
                        logger.info(f"[BULK_ASSESS] Product field exists but is empty for {linkedinurl[:50]}")
                else:
                    logger.info(f"[BULK_ASSESS] No product data in DB for {linkedinurl[:50]}")
            
            # Check if CV is uploaded - if not, retry briefly in case the CV was just
            # committed by upload_multiple_cvs moments before this call arrived.
            # Allow up to 5 retries × 3 seconds = 15 seconds to accommodate the
            # analyze_cv_background thread which writes back to the process table.
            if not cv_data and row:
                for _retry in range(5):
                    time.sleep(3)
                    try:
                        if _nolinkedin_id and not linkedinurl:
                            cur.execute("SELECT cv FROM process WHERE id = %s LIMIT 1", (_nolinkedin_id,))
                        else:
                            cur.execute("SELECT cv FROM process WHERE linkedinurl = %s LIMIT 1", (linkedinurl,))
                        _cv_row = cur.fetchone()
                        if _cv_row and _cv_row[0]:
                            cv_data = _cv_row[0]
                            _log_key = f"process_id={_nolinkedin_id}" if (_nolinkedin_id and not linkedinurl) else linkedinurl[:50]
                            logger.info(f"[BULK_ASSESS] CV found after {(_retry+1)*3}s retry for {_log_key}")
                            break
                    except Exception as _e:
                        logger.debug(f"[BULK_ASSESS] CV retry query error (attempt {_retry+1}): {_e}")
                        conn.rollback()
            if not cv_data:
                if require_cv:
                    logger.info(f"[BULK_ASSESS] Skipping {(linkedinurl or candidate_name or '?')[:50]} - No CV uploaded (Assessment pending)")
                    return {
                        "linkedinurl": linkedinurl,
                        "result": {
                            "error": "Assessment pending - No CV uploaded"
                        }
                    }
                else:
                    logger.info(f"[BULK_ASSESS] No CV for {(linkedinurl or candidate_name or '?')[:50]} — proceeding with DB-only assessment (Normal DB mode)")

            # --- Analytic DB: wait for CV parsing to populate profile fields ---
            # For candidates with a CV but no profile fields (name-only import), analyze_cv_background
            # runs concurrently and may not have written jobtitle/company/country/experience yet.
            # Wait up to 100 s (20 × 5 s) for those fields to appear so _core_assess_profile can
            # distribute weights across all criteria instead of scoring only Skillset at 100%.
            if cv_data and not job_title and not company and not country and not experience_text:
                _id_for_poll = _nolinkedin_id if (_nolinkedin_id and not linkedinurl) else None
                _url_for_poll = linkedinurl if linkedinurl else None
                if _id_for_poll or _url_for_poll:
                    _log_key_poll = f"process_id={_id_for_poll}" if _id_for_poll else (_url_for_poll[:50] if _url_for_poll else "?")
                    logger.info(f"[BULK_ASSESS] Waiting for CV parsing to populate profile fields for {_log_key_poll}")
                    _pp_fields_found = False
                    for _ppoll in range(20):
                        try:
                            _pp_conn = psycopg2.connect(host=pg_host, port=pg_port, user=pg_user,
                                                        password=pg_password, dbname=pg_db)
                            _pp_cur = _pp_conn.cursor()
                            try:
                                if _id_for_poll:
                                    _pp_cur.execute(
                                        "SELECT jobtitle, company, country, seniority, sector, experience, skillset, tenure "
                                        "FROM process WHERE id = %s LIMIT 1",
                                        (_id_for_poll,)
                                    )
                                else:
                                    _pp_cur.execute(
                                        "SELECT jobtitle, company, country, seniority, sector, experience, skillset, tenure "
                                        "FROM process WHERE linkedinurl = %s LIMIT 1",
                                        (_url_for_poll,)
                                    )
                                _pp_row = _pp_cur.fetchone()
                            finally:
                                _pp_cur.close(); _pp_conn.close()
                            if _pp_row:
                                _pp_jt, _pp_co, _pp_cn, _pp_se, _pp_sc, _pp_ex, _pp_sk, _pp_tn = [
                                    (_pp_row[i] if i < len(_pp_row) else None) for i in range(8)
                                ]
                                if _pp_jt or _pp_co or _pp_cn or _pp_ex:
                                    job_title = _pp_jt or job_title
                                    company = _pp_co or company
                                    country = _pp_cn or country
                                    seniority = _pp_se or seniority
                                    sector = _pp_sc or sector
                                    experience_text = _pp_ex or experience_text
                                    if _pp_sk:
                                        candidate_skills = [s.strip() for s in _pp_sk.split(',') if s.strip()]
                                    if _pp_tn:
                                        tenure = _pp_tn
                                    logger.info(
                                        f"[BULK_ASSESS] Profile fields ready after {_ppoll * 5}s for {_log_key_poll}: "
                                        f"jobtitle='{job_title}', company='{company}', country='{country}'"
                                    )
                                    _pp_fields_found = True
                                    break
                        except Exception as _pp_err:
                            logger.debug(f"[BULK_ASSESS] Profile field poll attempt {_ppoll+1} failed: {_pp_err}")
                        time.sleep(5)
                    if not _pp_fields_found:
                        logger.warning(
                            f"[BULK_ASSESS] Profile fields still empty after wait for {_log_key_poll} — proceeding with available data"
                        )

            # --- Idempotency pre-check: read existing rating metadata from DB ---
            try:
                _ensure_rating_metadata_columns(cur, conn)
                # Use WHERE id = %s for no-LinkedIn candidates to avoid matching ALL
                # empty-URL rows with WHERE linkedinurl = '' (which could be multiple records).
                if _nolinkedin_id and not linkedinurl:
                    cur.execute("""
                        SELECT rating, rating_level, rating_updated_at, rating_version
                        FROM process WHERE id = %s LIMIT 1
                    """, (_nolinkedin_id,))
                else:
                    cur.execute("""
                        SELECT rating, rating_level, rating_updated_at, rating_version
                        FROM process WHERE linkedinurl = %s LIMIT 1
                    """, (linkedinurl,))
                _bulk_row_meta = cur.fetchone()
                _bulk_existing_meta = None
                if _bulk_row_meta and _bulk_row_meta[0]:
                    _bulk_existing_meta = {
                        "rating": _bulk_row_meta[0],
                        "level": (_bulk_row_meta[1] or "").upper(),
                        "updated_at": _bulk_row_meta[2],
                        "version": _bulk_row_meta[3],
                    }
                _bulk_allow, _bulk_reason = _should_overwrite_existing(_bulk_existing_meta, assessment_level, force_reassess)
                if not _bulk_allow:
                    logger.info(f"[BULK_ASSESS] Skipping {linkedinurl[:50]}: {_bulk_reason}")
                    _bulk_existing_obj = _bulk_existing_meta.get("rating") if _bulk_existing_meta else None
                    if isinstance(_bulk_existing_obj, str):
                        try:
                            _bulk_existing_obj = json.loads(_bulk_existing_obj)
                        except Exception:
                            _bulk_existing_obj = {"raw": _bulk_existing_obj}
                    cur.close(); conn.close()
                    if isinstance(_bulk_existing_obj, dict):
                        _bulk_existing_obj["_skipped"] = True
                        _bulk_existing_obj["_note"] = f"skipped - existing rating present ({_bulk_reason})"
                        return {"linkedinurl": linkedinurl, "result": _bulk_existing_obj}
                    return {"linkedinurl": linkedinurl, "result": {"_skipped": True, "error": "assessment skipped - existing rating", "reason": _bulk_reason}}
            except Exception as _e_bulk_idem:
                logger.warning(f"[BULK_ASSESS] Idempotency pre-check failed for {linkedinurl[:50]} (continuing): {_e_bulk_idem}")

            # ALWAYS look up role_tag from sourcing by linkedinurl first — this is the per-profile
            # authoritative source.  The process table's role_tag may have been set during a previous
            # search with a different role and must NOT override the sourcing record.
            _sourcing_role_tag_direct = None
            if linkedinurl:
                try:
                    cur.execute("SELECT role_tag FROM sourcing WHERE linkedinurl = %s AND role_tag IS NOT NULL AND role_tag != '' LIMIT 1", (linkedinurl,))
                    _src_rt_row = cur.fetchone()
                    if _src_rt_row and _src_rt_row[0]:
                        _sourcing_role_tag_direct = _src_rt_row[0]
                except Exception as _e_src_rt:
                    logger.debug(f"[BULK_ASSESS] sourcing role_tag direct lookup failed: {_e_src_rt}")
            if _sourcing_role_tag_direct:
                role_tag = _sourcing_role_tag_direct
                logger.info(f"[BULK_ASSESS] role_tag='{role_tag}' sourced from sourcing table for {linkedinurl[:50]}")
            elif override_role_tag:
                # Use the role_tag explicitly selected by the user in the Dock In wizard Step 3.
                # This takes priority over username-based DB lookups (which may return a stale
                # role from a previous session) when the record is not in the sourcing table.
                role_tag = override_role_tag
                logger.info(f"[BULK_ASSESS] role_tag='{role_tag}' applied from request override for {linkedinurl[:50]}")
            elif not role_tag:
                # Only fall back to process/username lookups when sourcing table has no role_tag
                # AND no override was provided in the request payload.
                if username_db:
                    try:
                        cur.execute("SELECT role_tag FROM sourcing WHERE username = %s AND role_tag IS NOT NULL AND role_tag != '' LIMIT 1", (username_db,))
                        src_row = cur.fetchone()
                        if src_row and src_row[0]:
                            role_tag = src_row[0]
                    except Exception as _e_src_u:
                        logger.debug(f"[BULK_ASSESS] sourcing role_tag by username lookup failed: {_e_src_u}")
                if not role_tag and username_db:
                    try:
                        cur.execute("SELECT role_tag FROM process WHERE username = %s AND role_tag IS NOT NULL AND role_tag != '' LIMIT 1", (username_db,))
                        proc_row = cur.fetchone()
                        if proc_row and proc_row[0]:
                            role_tag = proc_row[0]
                    except Exception as _e_proc_u:
                        logger.debug(f"[BULK_ASSESS] process role_tag by username lookup failed: {_e_proc_u}")
            
            cur.close()
            conn.close()
            # Do NOT call _sync_login_jskillset_to_process here: that function copies
            # the recruiter's global login.jskillset to every candidate, which overwrites
            # each candidate's role-specific skillset with the wrong global value.
            # Per-role skills are written below (lines 4940+) from the criteria file only.
            normalized_for_sync = _normalize_linkedin_to_path(linkedinurl)
            recruiter_username = username or username_db

            target_skills = []
            # Priority 1: load from the criteria JSON file saved after the AutoSourcing search run.
            # This is the authoritative source — never rely on cached/DB-derived values when the file exists.
            required_seniority_from_criteria = ""
            required_country_from_criteria = ""
            _criteria = _read_search_criteria(recruiter_username, role_tag)
            if _criteria:
                _file_skills = _criteria.get("Skillset") or []
                if _file_skills:
                    target_skills = _file_skills
                    logger.info(f"[BULK_ASSESS] target_skills loaded from criteria file ({len(target_skills)} skills) for {linkedinurl[:50]}")
                    # Write the per-role criteria skills to process.jskillset for this candidate only.
                    # This is the authoritative write — scoped to this candidate's linkedin URL.
                    try:
                        _skill_csv = ",".join(str(s).strip() for s in target_skills if str(s).strip())
                        _norm_link_write = normalized_for_sync or _normalize_linkedin_to_path(linkedinurl) or ""
                        _conn_jsk = psycopg2.connect(host=pg_host, port=pg_port, user=pg_user, password=pg_password, dbname=pg_db)
                        _cur_jsk = _conn_jsk.cursor()
                        try:
                            _cur_jsk.execute("""
                                SELECT column_name FROM information_schema.columns
                                WHERE table_schema='public' AND table_name='process'
                                AND column_name IN ('jskillset','jskills','jskill','normalized_linkedin','linkedinurl')
                            """)
                            _proc_cols_jsk = {r[0].lower() for r in _cur_jsk.fetchall()}
                            _jsk_col = next((c for c in ('jskillset', 'jskills', 'jskill') if c in _proc_cols_jsk), None)
                            if _jsk_col:
                                _updated = 0
                                if _norm_link_write and 'normalized_linkedin' in _proc_cols_jsk:
                                    _cur_jsk.execute(
                                        sql.SQL("UPDATE process SET {} = %s WHERE normalized_linkedin = %s").format(sql.Identifier(_jsk_col)),
                                        (_skill_csv, _norm_link_write)
                                    )
                                    _updated = _cur_jsk.rowcount
                                if _updated == 0 and linkedinurl and 'linkedinurl' in _proc_cols_jsk:
                                    _cur_jsk.execute(
                                        sql.SQL("UPDATE process SET {} = %s WHERE linkedinurl = %s").format(sql.Identifier(_jsk_col)),
                                        (_skill_csv, linkedinurl)
                                    )
                            _conn_jsk.commit()
                            logger.info(f"[BULK_ASSESS] criteria jskillset written to process.{_jsk_col} for {linkedinurl[:50]}")
                        finally:
                            _cur_jsk.close()
                            _conn_jsk.close()
                    except Exception as _e_crit_jsk:
                        logger.warning(f"[BULK_ASSESS] criteria jskillset write-back failed for {linkedinurl}: {_e_crit_jsk}")
                required_seniority_from_criteria = (_criteria.get("Seniority") or "").strip()
                required_country_from_criteria = (_criteria.get("Country") or "").strip()
            # Priority 2: fall back to DB-derived values only if the file is unavailable.
            if not target_skills:
                target_skills = _fetch_jskillset_from_process(linkedinurl) or []
            if not target_skills and recruiter_username:
                target_skills = _fetch_jskillset(recruiter_username) or []

            # Deterministic seniority fallback: derive from job_title keywords when model output is absent.
            # Note: we do not pass tenure (per-employer avg) as total_experience_years since they differ.
            if not seniority:
                try:
                    seniority = _normalize_seniority_to_8_levels(job_title or "")
                    if seniority:
                        logger.info(f"[BULK_ASSESS] Seniority derived from job title: '{seniority}' for {linkedinurl[:50]}")
                except Exception as e_sen:
                    logger.warning(f"[BULK_ASSESS] Seniority derivation failed: {e_sen}")

            # Always check ML_Master_Company.json first — it is the highest-quality
            # validated source and must take priority over DB/Gemini sector values.
            # (E.g. Iqvia = "Clinical Research Organisation", not "Pharmaceutical".)
            if company:
                try:
                    ml_sector = _resolve_sector_from_ml_master(company, log_prefix="[BULK_ASSESS] ")
                    if ml_sector:
                        sector = ml_sector
                        logger.info(f"[BULK_ASSESS] Sector resolved from ML_Master_Company for {linkedinurl[:50]}: '{sector}'")
                except Exception as e_ml:
                    logger.warning(f"[BULK_ASSESS] ML_Master_Company sector lookup failed: {e_ml}")

            # Always check ML_Master_Jobfamily_Seniority.json — it is the primary source of truth
            # for job family and seniority, overriding DB/Gemini/hardcoded-keyword values.
            if job_title:
                try:
                    ml_jf, ml_sn = _resolve_jobfamily_seniority_from_ml_master(job_title, log_prefix="[BULK_ASSESS] ")
                    if ml_jf:
                        job_family = ml_jf
                        logger.info(f"[BULK_ASSESS] job_family resolved from ML_Master_JT for {linkedinurl[:50]}: '{job_family}'")
                    if ml_sn:
                        seniority = ml_sn
                        logger.info(f"[BULK_ASSESS] Seniority resolved from ML_Master_JT for {linkedinurl[:50]}: '{seniority}'")
                except Exception as e_ml_jt:
                    logger.warning(f"[BULK_ASSESS] ML_Master_JT job_family/seniority lookup failed: {e_ml_jt}")

            if not sector and (job_title or role_tag or experience_text):
                try:
                    for probe in [job_title, role_tag, experience_text[:200]]:
                        if probe:
                            mapped = _find_best_sector_match_for_text(probe) or _map_keyword_to_sector_label(probe)
                            if mapped:
                                sector = mapped
                                logger.info(f"[BULK_ASSESS] Sector resolved via fallback for {linkedinurl[:50]}: '{sector}'")
                                break
                except Exception as e_sec:
                    logger.warning(f"[BULK_ASSESS] Sector fallback failed: {e_sec}")

            # For L2 assessment, run full extractive-confirm + inference vskillset generation
            vskillset_results = None  # Initialize to None for passing to profile_data
            if assessment_level == "L2" and target_skills and len(target_skills) > 0:
                _vsk_log_key = f"process_id={_nolinkedin_id}" if _nolinkedin_id and not linkedinurl else (linkedinurl[:50] if linkedinurl else "(unknown)")
                logger.info(f"[BULK_ASSESS] L2 mode - checking/generating vskillset for {_vsk_log_key}")
                try:
                    # Idempotency guard: reuse existing vskillset if already persisted
                    _existing_vs_bulk = None
                    try:
                        conn_vsk_idem = psycopg2.connect(host=pg_host, port=pg_port, user=pg_user, password=pg_password, dbname=pg_db)
                        cur_vsk_idem = conn_vsk_idem.cursor()
                        if _nolinkedin_id and not linkedinurl:
                            # No-LinkedIn record: use primary key for exact match
                            cur_vsk_idem.execute(
                                "SELECT vskillset FROM process WHERE id = %s LIMIT 1",
                                (_nolinkedin_id,)
                            )
                        else:
                            # Use normalized URL comparison to handle trailing slash / case differences
                            _norm_vsk = linkedinurl.lower().strip().rstrip('/')
                            cur_vsk_idem.execute("""
                                SELECT vskillset FROM process
                                WHERE LOWER(TRIM(TRAILING '/' FROM linkedinurl)) = %s
                                LIMIT 1
                            """, (_norm_vsk,))
                        _vs_idem_row = cur_vsk_idem.fetchone()
                        if _vs_idem_row and _vs_idem_row[0]:
                            _vs_idem_val = _vs_idem_row[0]
                            if isinstance(_vs_idem_val, str):
                                try:
                                    _vs_idem_val = json.loads(_vs_idem_val)
                                except Exception:
                                    _vs_idem_val = []
                            if isinstance(_vs_idem_val, list) and len(_vs_idem_val) > 0:
                                _existing_vs_bulk = _vs_idem_val
                        cur_vsk_idem.close(); conn_vsk_idem.close()
                    except Exception as _e_vs_idem:
                        logger.warning(f"[BULK_ASSESS] vskillset idempotency check failed ({_e_vs_idem}); will regenerate")

                    if _existing_vs_bulk is not None:
                        vskillset_results = _existing_vs_bulk
                        candidate_skills = [i["skill"] for i in vskillset_results if isinstance(i, dict) and i.get("category") == "High"]
                        logger.info(f"[BULK_ASSESS] Reusing existing vskillset ({len(vskillset_results)} items) for {_vsk_log_key}")
                    else:
                    # New generation
                        profile_context = experience_text
                        if not profile_context and cv_data:
                            cv_bytes_ctx = cv_data if isinstance(cv_data, bytes) else bytes(cv_data)
                            try:
                                from pypdf import PdfReader
                                import io
                                reader = PdfReader(io.BytesIO(cv_bytes_ctx))
                                profile_context = "".join(
                                    (p.extract_text() or "") + "\n" for p in reader.pages
                                )[:3000]
                            except Exception:
                                pass

                        if profile_context and target_skills:
                            # Step 1: Extractive confirm (fast, no Gemini call)
                            confirmed_skills = _extract_confirmed_skills(profile_context, target_skills)
                            confirmed_set = set(s.lower() for s in confirmed_skills)
                            vskillset_results = [
                                {"skill": s, "probability": 100, "category": "High",
                                 "reason": "Explicitly mentioned in experience text", "source": "confirmed"}
                                for s in confirmed_skills
                            ]

                            # Step 2: Gemini inference for remaining skills
                            unconfirmed = [s for s in target_skills if s.lower() not in confirmed_set]
                            if unconfirmed:
                                try:
                                    prompt = (
                                        "SYSTEM: You are an expert technical recruiter evaluating candidate skillsets.\n"
                                        "For each skill below evaluate the candidate's likely proficiency from their experience.\n"
                                        "These skills were NOT found explicitly, so use contextual inference from job titles, companies, products.\n"
                                        "Assign a probability score (0-100) and categorize as Low (<40), Medium (40-74), or High (75-100).\n\n"
                                        f"CANDIDATE PROFILE:\n{profile_context[:3000]}\n\n"
                                        f"SKILLS TO INFER:\n{json.dumps(unconfirmed, ensure_ascii=False)}\n\n"
                                        "OUTPUT FORMAT (JSON):\n"
                                        '{"evaluations": [{"skill": "skill_name", "probability": 0-100, "category": "Low|Medium|High", "reason": "..."}]}\n'
                                        "Return ONLY the JSON object."
                                    )
                                    raw_inf = (unified_llm_call_text(prompt) or "").strip()
                                    _increment_gemini_query_count(username)
                                    parsed = _extract_json_object(raw_inf)
                                    if parsed and "evaluations" in parsed:
                                        for item in parsed["evaluations"]:
                                            item.setdefault("probability", 50)
                                            item.setdefault("reason", "Inferred from context")
                                            item.setdefault("source", "inferred")
                                            prob = item["probability"]
                                            item["category"] = "High" if prob >= 75 else ("Medium" if prob >= 40 else "Low")
                                        vskillset_results += parsed["evaluations"]
                                except Exception as e_inf:
                                    logger.warning(f"[BULK_ASSESS] vskillset inference step failed for {linkedinurl}: {e_inf}")

                            # Persist vskillset to DB
                            if vskillset_results:
                                try:
                                    conn_vsk = psycopg2.connect(host=pg_host, port=pg_port, user=pg_user, password=pg_password, dbname=pg_db)
                                    cur_vsk = conn_vsk.cursor()
                                    cur_vsk.execute("SELECT column_name FROM information_schema.columns WHERE table_schema='public' AND table_name='process' AND column_name IN ('vskillset','skillset')")
                                    avail_cols = {r[0] for r in cur_vsk.fetchall()}
                                    vsk_json = json.dumps(vskillset_results, ensure_ascii=False)
                                    high_skills_str = ", ".join(
                                        i["skill"] for i in vskillset_results if i.get("category") == "High"
                                    )
                                    updates, vals = [], []
                                    if 'vskillset' in avail_cols:
                                        updates.append("vskillset = %s"); vals.append(vsk_json)
                                    if updates:
                                        if _nolinkedin_id and not linkedinurl:
                                            vals.append(_nolinkedin_id)
                                            cur_vsk.execute(sql.SQL("UPDATE process SET {} WHERE id = %s").format(sql.SQL(", ".join(updates))), tuple(vals))
                                        else:
                                            vals.append(linkedinurl)
                                            cur_vsk.execute(sql.SQL("UPDATE process SET {} WHERE linkedinurl = %s").format(sql.SQL(", ".join(updates))), tuple(vals))
                                    # Skillset: merge new High skills into existing value (add only; never remove or replace)
                                    if 'skillset' in avail_cols and high_skills_str:
                                        _new_highs = [s.strip() for s in high_skills_str.split(',') if s.strip()]
                                        if _nolinkedin_id and not linkedinurl:
                                            cur_vsk.execute("SELECT skillset FROM process WHERE id = %s", (_nolinkedin_id,))
                                        else:
                                            cur_vsk.execute("SELECT skillset FROM process WHERE linkedinurl = %s", (linkedinurl,))
                                        _sk_row = cur_vsk.fetchone()
                                        _existing_sk = (_sk_row[0] or "") if _sk_row else ""
                                        _existing_parts = [s.strip() for s in _existing_sk.split(',') if s.strip()]
                                        _existing_set = {s.lower() for s in _existing_parts}
                                        _new_sk = [s for s in _new_highs if s.lower() not in _existing_set]
                                        if _new_sk:
                                            _merged_sk = ', '.join(_existing_parts + _new_sk)
                                            if _nolinkedin_id and not linkedinurl:
                                                cur_vsk.execute(
                                                    "UPDATE process SET skillset = %s WHERE id = %s",
                                                    (_merged_sk, _nolinkedin_id)
                                                )
                                            else:
                                                cur_vsk.execute(
                                                    "UPDATE process SET skillset = %s WHERE linkedinurl = %s",
                                                    (_merged_sk, linkedinurl)
                                                )
                                            logger.info(f"[BULK_ASSESS] Merged {len(_new_sk)} new High skills into skillset for {_vsk_log_key}")
                                        else:
                                            logger.info(f"[BULK_ASSESS] No new High skills for {_vsk_log_key} — skillset unchanged")
                                    conn_vsk.commit()
                                    cur_vsk.close(); conn_vsk.close()
                                except Exception as e_vdb:
                                    logger.warning(f"[BULK_ASSESS] vskillset persist failed for {linkedinurl}: {e_vdb}")

                            # Update candidate_skills from High vskillset results
                            candidate_skills = [i["skill"] for i in vskillset_results if i.get("category") == "High"]
                            logger.info(f"[BULK_ASSESS] vskillset: {len(confirmed_skills)} confirmed, {len(vskillset_results)-len(confirmed_skills)} inferred, {len(candidate_skills)} High total")
                        else:
                            # Fallback to _generate_vskillset_for_profile when no profile context
                            vskillset_results = _generate_vskillset_for_profile(linkedinurl, target_skills, experience_text, cv_data)
                            if vskillset_results:
                                candidate_skills = [i["skill"] for i in vskillset_results if i.get("category") == "High"]
                except Exception as e_vsk:
                    logger.warning(f"[BULK_ASSESS] vskillset generation failed for {linkedinurl}: {e_vsk}")
            
            # Build profile_data consistent with gemini_assess_profile expectations
            profile_data = {
                "job_title": job_title,
                "role_tag": role_tag,
                "company": company,
                "country": country,
                "seniority": seniority,
                "sector": sector,
                "job_family": job_family,
                "experience_text": experience_text,
                "target_skills": target_skills,
                "candidate_skills": candidate_skills,
                "process_skills": [],
                "custom_weights": custom_weights,
                "linkedinurl": linkedinurl,
                "assessment_level": assessment_level,
                "tenure": tenure,
                "product": product,
                "vskillset_results": vskillset_results,  # Pass vskillset_results for scoring
                "required_seniority": required_seniority_from_criteria,
                "required_country": required_country_from_criteria,
                "username": username,  # Include username so assessment filename contains it
            }

            # run core assessment
            try:
                result = _core_assess_profile(profile_data)
                logger.info(f"[BULK_ASSESS] Assessment completed for {linkedinurl[:50]}: Score={result.get('total_score', 'N/A')}, Stars={result.get('stars', 0)}")
                # Mirror gemini_assess_profile: include vskillset in result so the frontend
                # can render the vskillset category breakdown (data.vskillset check in UI).
                if vskillset_results and isinstance(result, dict):
                    result["vskillset"] = vskillset_results
                    # Update the individual assessment file so download_report always returns
                    # a result that includes vskillset — _core_assess_profile writes this file
                    # *before* vskillset is added, so we re-write it here with the full result.
                    if linkedinurl:
                        try:
                            _uname_suffix = _CV_USERNAME_SAFE_RE.sub('', username or "")
                            _uname_suffix = f"_{_uname_suffix}" if _uname_suffix else ""
                            _safe_aname = "assessment_" + hashlib.sha256(linkedinurl.encode("utf-8")).hexdigest()[:16] + _uname_suffix + ".json"
                            _assess_adir = os.path.join(OUTPUT_DIR, "assessments")
                            os.makedirs(_assess_adir, exist_ok=True)
                            _out_apath = os.path.join(_assess_adir, _safe_aname)
                            with open(_out_apath, "w", encoding="utf-8") as _fh:
                                json.dump(result, _fh, indent=2, ensure_ascii=False)
                        except Exception as _e_file:
                            logger.warning(f"[BULK_ASSESS] Failed to update assessment file with vskillset: {_e_file}")
            except Exception as e:
                logger.error(f"[BULK_ASSESS] Assessment error for {linkedinurl}: {e}")
                result = {"error": f"assessment_error: {str(e)}"}

            # persist rating JSON into process.rating if column exists
            try:
                conn=psycopg2.connect(host=pg_host, port=pg_port, user=pg_user, password=pg_password, dbname=pg_db)
                cur=conn.cursor()

                cur.execute("""
                    SELECT column_name FROM information_schema.columns
                    WHERE table_schema='public' AND table_name='process' AND column_name='rating'
                """)
                if cur.fetchone():
                    # store JSON as text (exclude vskillset to keep rating payload compact,
                    # mirroring gemini_assess_profile which strips vskillset before persisting)
                    rating_obj = {k: v for k, v in result.items() if k != "vskillset"} if isinstance(result, dict) else result
                    rating_payload = json.dumps(rating_obj, ensure_ascii=False)
                    logger.info(f"[BULK_PERSIST] Persisting rating for {(linkedinurl or candidate_name or '?')[:50]}, payload size: {len(rating_payload)} bytes")

                    # Ensure metadata columns exist then persist rating + metadata atomically
                    _ensure_rating_metadata_columns(cur, conn)
                    if linkedinurl:
                        cur.execute(
                            "UPDATE process SET rating = %s, rating_level = %s, rating_updated_at = NOW(), "
                            "rating_version = COALESCE(rating_version, 0) + 1 WHERE linkedinurl = %s",
                            (rating_payload, assessment_level, linkedinurl)
                        )
                    elif _nolinkedin_id is not None:
                        # Precise update by process PK (captured during name+company lookup).
                        cur.execute(
                            "UPDATE process SET rating = %s, rating_level = %s, rating_updated_at = NOW(), "
                            "rating_version = COALESCE(rating_version, 0) + 1 WHERE id = %s",
                            (rating_payload, assessment_level, _nolinkedin_id)
                        )
                    elif candidate_name and candidate_company:
                        # Fallback for records without LinkedIn URL: only update when name+company
                        # uniquely identifies one record (avoids corrupting unrelated candidates).
                        cur.execute(
                            "SELECT COUNT(*) FROM process WHERE LOWER(TRIM(name)) = LOWER(%s) AND LOWER(TRIM(company)) = LOWER(%s)",
                            (candidate_name.strip(), candidate_company.strip())
                        )
                        _name_match_count = (cur.fetchone() or [0])[0]
                        if _name_match_count == 1:
                            cur.execute(
                                "UPDATE process SET rating = %s, rating_level = %s, rating_updated_at = NOW(), "
                                "rating_version = COALESCE(rating_version, 0) + 1 "
                                "WHERE LOWER(TRIM(name)) = LOWER(%s) AND LOWER(TRIM(company)) = LOWER(%s)",
                                (rating_payload, assessment_level, candidate_name.strip(), candidate_company.strip())
                            )
                        elif _name_match_count > 1:
                            logger.warning(f"[BULK_PERSIST] Skipping rating persist for '{candidate_name}' — {_name_match_count} rows match (ambiguous)")
                    updated = cur.rowcount
                    logger.info(f"[BULK_PERSIST] Updated {updated} rows by {'linkedinurl' if linkedinurl else 'name'}: {(linkedinurl or candidate_name or '?')[:50]}")

                    if updated > 0:
                        logger.info(f"[BULK_PERSIST] Successfully persisted rating for {(linkedinurl or candidate_name or '?')[:50]}")
                    else:
                        logger.warning(f"[BULK_PERSIST] No rows updated for {(linkedinurl or candidate_name or '?')[:50]} - profile may not exist in process table")

                    # Also sync jskill if result contains role_tag (safe best-effort)
                    role_tag_val = result.get("role_tag") if isinstance(result, dict) else None
                    if role_tag_val:
                        cur.execute("""
                            SELECT column_name FROM information_schema.columns
                            WHERE table_schema='public' AND table_name='process' AND column_name='jskill'
                        """)
                        if cur.fetchone():
                            cur.execute("UPDATE process SET jskill = %s WHERE linkedinurl = %s", (role_tag_val, linkedinurl))

                    conn.commit()

                # Sync role_tag from sourcing → process.
                # When override_role_tag is provided (user's Step 3 selection), it takes absolute
                # priority and the sourcing re-read is skipped so the selected role is not overwritten
                # by a stale sourcing value from a previous session.
                if role_tag:
                    try:
                        if not override_role_tag:
                            # No override — re-read the authoritative value from sourcing.
                            _sr_rt = None
                            cur.execute(
                                "SELECT role_tag FROM sourcing WHERE linkedinurl=%s AND role_tag IS NOT NULL AND role_tag != '' LIMIT 1",
                                (linkedinurl,)
                            )
                            _sr = cur.fetchone()
                            if _sr and _sr[0]:
                                _sr_rt = _sr[0]
                            if not _sr_rt and username_db:
                                cur.execute(
                                    "SELECT role_tag FROM sourcing WHERE username=%s AND role_tag IS NOT NULL AND role_tag != '' LIMIT 1",
                                    (username_db,)
                                )
                                _sr = cur.fetchone()
                                if _sr and _sr[0]:
                                    _sr_rt = _sr[0]
                            if _sr_rt:
                                role_tag = _sr_rt
                    except Exception as _e_src:
                        logger.warning(f"[BULK_PERSIST] Failed to re-read role_tag from sourcing: {_e_src}")
                    cur.execute("""
                        SELECT column_name FROM information_schema.columns
                        WHERE table_schema='public' AND table_name='process' AND column_name='role_tag'
                    """)
                    if not cur.fetchone():
                        cur.execute("ALTER TABLE process ADD COLUMN role_tag TEXT DEFAULT ''")
                    if linkedinurl:
                        cur.execute(
                            "UPDATE process SET role_tag = %s WHERE linkedinurl = %s",
                            (role_tag, linkedinurl)
                        )
                        _rt_id_label = linkedinurl[:50]
                    elif _nolinkedin_id is not None:
                        # No LinkedIn URL — use the process PK captured during name+company lookup
                        cur.execute(
                            "UPDATE process SET role_tag = %s WHERE id = %s",
                            (role_tag, _nolinkedin_id)
                        )
                        _rt_id_label = f"process_id={_nolinkedin_id}"
                    else:
                        _rt_id_label = None
                    if _rt_id_label and cur.rowcount:
                        src = "override" if override_role_tag else "sourcing→process"
                        logger.info(f"[BULK_PERSIST] Synced role_tag='{role_tag}' from {src} for {_rt_id_label}")
                    conn.commit()

                # Sync userid and username from sourcing → process (mirrors gemini_assess_profile owner update).
                # Reads the recruiter's userid/username stored in the sourcing row for this candidate.
                try:
                    cur.execute(
                        "SELECT userid, username FROM sourcing WHERE linkedinurl=%s AND (userid IS NOT NULL OR username IS NOT NULL) LIMIT 1",
                        (linkedinurl,)
                    )
                    _src_owner = cur.fetchone()
                    if _src_owner:
                        _src_userid, _src_username = _src_owner
                        # Build update: only set columns that exist and have values; don't overwrite existing
                        cur.execute("""
                            SELECT column_name FROM information_schema.columns
                            WHERE table_schema='public' AND table_name='process'
                              AND column_name IN ('userid','username')
                        """)
                        _owner_cols = {r[0] for r in cur.fetchall()}
                        _owner_parts = []
                        _owner_vals = []
                        if 'userid' in _owner_cols and _src_userid:
                            _owner_parts.append("userid = COALESCE(NULLIF(userid, ''), %s)")
                            _owner_vals.append(_src_userid)
                        if 'username' in _owner_cols and _src_username:
                            _owner_parts.append("username = COALESCE(NULLIF(username, ''), %s)")
                            _owner_vals.append(_src_username)
                        if _owner_parts:
                            _owner_vals.append(linkedinurl)
                            cur.execute(
                                sql.SQL("UPDATE process SET {} WHERE linkedinurl=%s").format(sql.SQL(", ".join(_owner_parts))),
                                tuple(_owner_vals)
                            )
                            if cur.rowcount:
                                logger.info(f"[BULK_PERSIST] Synced userid/username from sourcing→process for {linkedinurl[:50]}")
                            conn.commit()
                except Exception as _e_owner:
                    try:
                        conn.rollback()
                    except Exception:
                        pass
                    logger.warning(f"[BULK_PERSIST] Failed to sync userid/username from sourcing: {_e_owner}")

                cur.close(); conn.close()
            except Exception as e_db:
                logger.warning(f"[BulkAssess->DB] Failed to write rating for {linkedinurl}: {e_db}")

            # Build return record; include process_id for no-LinkedIn candidates so callers
            # can augment the bulk results JSON and _find_assessment_for_candidate_by_id works.
            _return_rec = {"linkedinurl": linkedinurl, "result": result}
            if _nolinkedin_id and not linkedinurl:
                _return_rec["process_id"] = _nolinkedin_id
            return _return_rec
        except Exception as e:
            logger.error(f"[BulkAssess] Error assessing {linkedinurl}: {e}")
            return {"linkedinurl": linkedinurl, "error": str(e)}

    # If async, spin thread worker and return job id
    if async_flag:
        job_id = "bulk_" + uuid.uuid4().hex[:10]
        
        # Initialize job status — store username so status endpoint can reconstruct filename
        with JOBS_LOCK:
            JOBS[job_id] = {'status': 'running', 'processed': 0, 'total': len(candidates_list),
                            'messages': [], 'errors': [], 'username': username, 'results_fname': ''}
        persist_job(job_id)
        
        def _bg_worker(cands, job_id):
            results = []
            processed = 0
            for cand in cands:
                clink = cand.get("linkedinurl", "") if isinstance(cand, dict) else str(cand)
                cname = cand.get("name", "") if isinstance(cand, dict) else ""
                ccomp = cand.get("company", "") if isinstance(cand, dict) else ""
                cpid = cand.get("process_id") if isinstance(cand, dict) else None
                try:
                    out = _assess_and_persist(clink, cname, ccomp, candidate_process_id=cpid)
                    # Augment record with name and process_id so _find_assessment_for_candidate_by_id
                    # and _find_assessment_for_candidate_by_name can locate it via file scan.
                    if isinstance(out, dict):
                        if cname and "name" not in out:
                            out["name"] = cname
                        if cpid and "process_id" not in out:
                            out["process_id"] = cpid
                    results.append(out)
                except Exception as e:
                    results.append({"linkedinurl": clink, "error": str(e)})
                finally:
                    processed += 1
                    with JOBS_LOCK:
                        JOBS[job_id]['processed'] = processed
                    persist_job(job_id)
                    
            # persist results file for later retrieval — include safe username in filename
            results_file_ok = False
            try:
                _safe_uname = _CV_USERNAME_SAFE_RE.sub('', username or "")
                _uname_suffix = f"_{_safe_uname}" if _safe_uname else ""
                fname = f"{job_id}_results{_uname_suffix}.json"
                path = os.path.join(OUTPUT_DIR, fname)
                with open(path, "w", encoding="utf-8") as fh:
                    json.dump(results, fh, ensure_ascii=False, indent=2)
                    # fsync ensures the data is committed to the physical disk before we mark
                    # the job as 'done'. On Windows, delayed write buffering can cause the file
                    # to appear non-existent to a concurrent reader even after the write() call.
                    os.fsync(fh.fileno())
                # Store the results filename so the status endpoint can read it
                with JOBS_LOCK:
                    if job_id in JOBS:
                        JOBS[job_id]['results_fname'] = fname
                logger.info(f"[BulkAssess] Completed job {job_id} results saved to {path}")
                results_file_ok = True
            except Exception as e:
                logger.warning(f"[BulkAssess] Failed to write results for {job_id}: {e}")
            
            # Only mark job as 'done' after the results file is confirmed on disk.
            # The status endpoint will keep reporting 'running' until it can read the file.
            if results_file_ok:
                with JOBS_LOCK:
                    JOBS[job_id]['status'] = 'done'
                persist_job(job_id)

        threading.Thread(target=_bg_worker, args=(candidates_list, job_id), daemon=True).start()
        log_approval(action="bulk_assessment_triggered",
                     username=username, userid=userid,
                     detail=f"Async bulk assessment started for {len(candidates_list)} profile(s); job_id={job_id}")
        return jsonify({"ok": True, "job_id": job_id}), 202

    # synchronous: do all and return results
    all_results = []
    for cand in candidates_list:
        clink = cand.get("linkedinurl", "") if isinstance(cand, dict) else str(cand)
        cname = cand.get("name", "") if isinstance(cand, dict) else ""
        ccomp = cand.get("company", "") if isinstance(cand, dict) else ""
        cpid = cand.get("process_id") if isinstance(cand, dict) else None
        out = _assess_and_persist(clink, cname, ccomp, candidate_process_id=cpid)
        if isinstance(out, dict):
            if cname and "name" not in out:
                out["name"] = cname
            if cpid and "process_id" not in out:
                out["process_id"] = cpid
        all_results.append(out)
    log_approval(action="bulk_assessment_completed",
                 username=username, userid=userid,
                 detail=f"Synchronous bulk assessment completed for {len(candidates_list)} profile(s)")
    return jsonify({"ok": True, "results": all_results}), 200


@app.get("/process/role_skills")
@_rate(_make_flask_limit("candidates"))
def get_role_skills():
    """Return the target skills from the criteria file for a given role_tag.
    Used by the Dock In Step 3 wizard to display the authoritative job skillset for each role.
    """
    username = request.cookies.get("username", "").strip()
    role_tag = (request.args.get("role_tag") or "").strip()
    if not role_tag or not username:
        return jsonify({"skills": [], "found": False})
    try:
        criteria = _read_search_criteria(username, role_tag)
        if criteria:
            skills = criteria.get("Skillset") or []
            if skills:
                return jsonify({"skills": skills, "found": True})
        return jsonify({"skills": [], "found": False})
    except Exception as e:
        logger.debug(f"[role_skills] Error for role_tag='{role_tag}': {e}")
        return jsonify({"skills": [], "found": False})


@app.get("/process/bulk_assess_status/<job_id>")
@_rate(_make_flask_limit("bulk_assess_status"))
def process_bulk_status(job_id):
    with JOBS_LOCK:
        job = JOBS.get(job_id)
        if not job:
            # Fallback: load from persisted job file when not in memory (e.g., after server restart).
            # Done inside the lock to prevent concurrent requests from racing to restore the same job.
            job_file = os.path.join(OUTPUT_DIR, f"job_{job_id}.json")
            if os.path.exists(job_file):
                try:
                    with open(job_file, "r", encoding="utf-8") as fh:
                        job = json.load(fh)
                    JOBS[job_id] = job
                    logger.info(f"[BulkAssessStatus] Restored job {job_id} from persisted file")
                except Exception as e:
                    logger.warning(f"[BulkAssessStatus] Failed to load persisted job {job_id}: {e}")
    if not job:
        return jsonify({"error": "Job not found"}), 404
    
    # Calculate progress percentage
    job_response = dict(job)
    if 'processed' in job and 'total' in job and job['total'] > 0:
        job_response['progress'] = int((job['processed'] / job['total']) * 100)
    else:
        job_response['progress'] = 0
    
    # If job is done, try to load and include the actual results from JSON file.
    # Only report status='done' to the client once the results file is confirmed readable —
    # this prevents the frontend from closing the progress bar before the backend has fully
    # persisted all assessment results to disk.
    if job.get('status') == 'done':
        try:
            # Use the stored results filename (includes username suffix) with fallback to legacy name.
            fname = job.get('results_fname') or f"{job_id}_results.json"
            path = os.path.join(OUTPUT_DIR, fname)
            if not os.path.exists(path):
                # Fallback: try legacy filename without username suffix
                legacy_fname = f"{job_id}_results.json"
                legacy_path = os.path.join(OUTPUT_DIR, legacy_fname)
                if os.path.exists(legacy_path):
                    fname, path = legacy_fname, legacy_path
            if os.path.exists(path):
                with open(path, "r", encoding="utf-8") as fh:
                    results = json.load(fh)
                # Include results in the response
                job_response['results'] = results
                job_response['progress'] = 100  # Ensure progress is 100% when done
                logger.info(f"[BulkAssessStatus] Loaded {len(results)} results from {fname}")
                return jsonify(job_response)
            else:
                # Defensive: status='done' is set only after results_file_ok=True in the normal
                # flow. This branch handles edge cases such as a persisted job loaded from the
                # job_{id}.json file (e.g. after a server restart) whose status was already
                # 'done' but whose results file was deleted in the interim. Returning
                # 'running' keeps the client polling until the file reappears or the client
                # times out, rather than closing the wizard with incomplete data.
                job_response['status'] = 'running'
                job_response['progress'] = 99
        except Exception as e:
            logger.warning(f"[BulkAssessStatus] Failed to load results for {job_id}: {e}")
    
    return jsonify(job_response)

@app.get("/process/bulk_assess_stream/<job_id>")
def process_bulk_assess_stream(job_id):
    """
    Server-Sent Events (SSE) endpoint for real-time bulk assessment progress.
    Streams progress updates to the client instead of requiring polling.
    """
    def generate_events():
        """Generator function that yields SSE-formatted messages."""
        last_status = None
        last_progress = -1
        
        while True:
            with JOBS_LOCK:
                job = JOBS.get(job_id)
            
            if not job:
                # Retry once in case of a brief race between job creation and first SSE connection.
                time.sleep(1)
                with JOBS_LOCK:
                    job = JOBS.get(job_id)
                if not job:
                    # Send as a regular data event (not a named event) so onmessage handles it.
                    yield f"data: {json.dumps({'status': 'error', 'error': 'Job not found'})}\n\n"
                    break
            
            # Calculate current progress
            current_progress = 0
            if 'processed' in job and 'total' in job and job['total'] > 0:
                current_progress = int((job['processed'] / job['total']) * 100)
            
            current_status = job.get('status', 'pending')
            
            # Only send update if something changed
            if current_status != last_status or current_progress != last_progress:
                event_data = {
                    'status': current_status,
                    'progress': current_progress,
                    'processed': job.get('processed', 0),
                    'total': job.get('total', 0)
                }
                
                # If job is done, include results
                if current_status == 'done':
                    try:
                        fname = f"{job_id}_results.json"
                        path = os.path.join(OUTPUT_DIR, fname)
                        if os.path.exists(path):
                            with open(path, "r", encoding="utf-8") as fh:
                                results = json.load(fh)
                            event_data['results'] = results
                            logger.info(f"[SSE] Loaded {len(results)} results for {job_id}")
                    except Exception as e:
                        logger.warning(f"[SSE] Failed to load results for {job_id}: {e}")
                
                yield f"data: {json.dumps(event_data)}\n\n"
                
                last_status = current_status
                last_progress = current_progress
                
                # If job is done or failed, close the stream
                if current_status in ('done', 'failed'):
                    break
            
            # Wait before checking again (reduce CPU usage)
            time.sleep(0.5)
    
    return Response(
        stream_with_context(generate_events()),
        mimetype='text/event-stream',
        headers={
            'Cache-Control': 'no-cache',
            'X-Accel-Buffering': 'no',  # Disable nginx buffering
            'Connection': 'keep-alive'
        }
    )


@app.get("/process/assessment_status")
def process_assessment_status():
    """Return whether any bulk assessment job is currently in progress.

    Used by SourcingVerify.html to drive the loading bar state when the page
    and the backend run on different ports (where cross-origin SSE is not
    viable). The client polls this lightweight endpoint every few seconds;
    when ``ongoing`` flips to ``false`` the loading bar is hidden and all
    action buttons are re-enabled.

    Returns: ``{ "ongoing": <bool> }``
    """
    with JOBS_LOCK:
        ongoing = any(
            str(jid).startswith("bulk_") and j.get("status") == "running"
            for jid, j in JOBS.items()
        )
    return jsonify({"ongoing": ongoing})

@app.patch("/process/profile_assessment/<path:linkedinurl>")
@_csrf_required
def patch_profile_assessment(linkedinurl):
    """
    HTTP PATCH endpoint for updating individual profile assessments.
    Faster than full POST as it only updates specific fields.
    """
    import psycopg2
    from psycopg2 import sql
    
    try:
        data = request.get_json(force=True, silent=True) or {}
        if 'rating' not in data:
            return jsonify({"error": "rating field required"}), 400
        
        rating = data.get('rating')
        
        # Normalize LinkedIn URL
        normalized = _normalize_linkedin_to_path(linkedinurl)
        
        # Update only the rating field in database
        pg_host = os.getenv("PGHOST", "localhost")
        pg_port = int(os.getenv("PGPORT", "5432"))
        pg_user = os.getenv("PGUSER", "postgres")
        pg_password = os.getenv("PGPASSWORD", "")
        pg_db = os.getenv("PGDATABASE", "candidate_db")
        
        conn = None
        cur = None
        try:
            conn = psycopg2.connect(
                host=pg_host, port=pg_port, user=pg_user, 
                password=pg_password, dbname=pg_db
            )
            cur = conn.cursor()
            
            # Check if rating column exists
            cur.execute("""
                SELECT column_name FROM information_schema.columns 
                WHERE table_schema='public' AND table_name='process' AND column_name='rating'
            """)
            
            if cur.fetchone():
                rating_json = json.dumps(rating) if isinstance(rating, dict) else rating
                cur.execute(
                    sql.SQL("UPDATE process SET rating = %s WHERE linkedinurl = %s"),
                    (rating_json, normalized)
                )
                conn.commit()
                
                updated = cur.rowcount
                
                logger.info(f"[PATCH] Updated assessment for {normalized}")
                return jsonify({"success": True, "updated": updated}) if updated > 0 else (jsonify({"error": "Profile not found"}), 404)
            else:
                return jsonify({"error": "rating column does not exist"}), 500
        finally:
            if cur:
                cur.close()
            if conn:
                conn.close()
                
    except Exception as e:
        logger.error(f"[PATCH] Error updating assessment: {e}")
        return jsonify({"error": str(e)}), 500

@app.post("/user/upload_jd")
def user_upload_jd():
    try:
        username = request.form.get("username", "").strip()
        if not username: return jsonify({"error": "username required"}), 400
        if 'file' not in request.files: return jsonify({"error": "No file part"}), 400
        file = request.files['file']
        if file.filename == '': return jsonify({"error": "No selected file"}), 400
        filename = file.filename.lower()
        if (request.content_length or 0) > _SINGLE_FILE_MAX:
            return jsonify({"error": "File too large (max 6 MB)"}), 413
        file_bytes = file.read()
        if len(file_bytes) > _SINGLE_FILE_MAX:
            return jsonify({"error": "File too large (max 6 MB)"}), 413
        extracted_text = ""
        if filename.endswith('.pdf'):
            if not _is_pdf_bytes(file_bytes):
                return jsonify({"error": "Uploaded file is not a valid PDF"}), 400
            import io
            try:
                from pypdf import PdfReader
                reader = PdfReader(io.BytesIO(file_bytes))
                for page in reader.pages: extracted_text += (page.extract_text() or "") + "\n"
            except ImportError: return jsonify({"error": "pypdf not installed, cannot process PDF"}), 500
            except Exception as e: return jsonify({"error": f"PDF parsing error: {e}"}), 500
        elif filename.endswith('.docx'):
            import io
            try:
                import docx
                doc = docx.Document(io.BytesIO(file_bytes))
                for para in doc.paragraphs: extracted_text += para.text + "\n"
            except ImportError: 
                return jsonify({"error": "python-docx library not installed. Please install it with: pip install python-docx"}), 500
            except Exception as e: 
                return jsonify({"error": f"DOCX parsing error: {e}"}), 500
        elif filename.endswith('.doc'):
            # Note: Legacy .doc format requires the python-docx library (or alternatives like antiword, textract)
            # python-docx primarily supports .docx but can sometimes read .doc files
            import io
            try:
                import docx
                doc = docx.Document(io.BytesIO(file_bytes))
                for para in doc.paragraphs: extracted_text += para.text + "\n"
            except ImportError:
                return jsonify({"error": "python-docx library not installed. Please install it with: pip install python-docx"}), 500
            except Exception as e:
                # Legacy .doc format may not be fully supported by python-docx
                return jsonify({"error": "Legacy .doc format could not be processed. Please convert to .docx or .pdf format."}), 400
        else:
            try: extracted_text = file_bytes.decode('utf-8', errors='ignore')
            except Exception as e: return jsonify({"error": f"Text decoding error: {e}"}), 500
        extracted_text = extracted_text.strip()
        if not extracted_text: return jsonify({"error": "Could not extract text from file"}), 400
        try:
            import psycopg2
            pg_host=os.getenv("PGHOST","localhost"); pg_port=int(os.getenv("PGPORT","5432"))
            pg_user=os.getenv("PGUSER","postgres"); pg_password=os.getenv("PGPASSWORD", "")
            pg_db=os.getenv("PGDATABASE","candidate_db")
            conn=psycopg2.connect(host=pg_host, port=pg_port, user=pg_user, password=pg_password, dbname=pg_db)
            cur=conn.cursor()
            cur.execute("UPDATE login SET jd = %s WHERE username = %s", (extracted_text, username))
            updated = cur.rowcount
            conn.commit(); cur.close(); conn.close()
            if updated == 0: return jsonify({"error": "Username not found"}), 404
        except Exception as e: return jsonify({"error": f"DB error storing JD: {e}"}), 500
        try:
            from chat_gemini_review import analyze_job_description
            analysis_result = analyze_job_description(extracted_text)
            parsed = analysis_result.get("parsed", {})
            skills = parsed.get("skills", [])
            if skills: _persist_jskillset(username, skills)
        except Exception as e: logger.warning(f"Failed to auto-extract skills after upload: {e}")
        return jsonify({"status": "ok", "message": "JD uploaded and stored", "length": len(extracted_text)}), 200
    except Exception as e:
        logger.error(f"[Upload JD] {e}")
        return jsonify({"error": str(e)}), 500

@app.post("/gemini/analyze_jd")
@_rate(_make_flask_limit("gemini"))
@_check_user_rate("gemini")
def gemini_jd_analyze():
    data = request.get_json(force=True, silent=True) or {}
    username = (data.get("username") or "").strip()
    text_input = (data.get("text") or "").strip()
    sectors_data = data.get("sectors") or []
    jd_text = text_input
    if not jd_text and username:
        try:
            import psycopg2
            pg_host=os.getenv("PGHOST","localhost"); pg_port=int(os.getenv("PGPORT","5432"))
            pg_user=os.getenv("PGUSER","postgres"); pg_password=os.getenv("PGPASSWORD", "")
            pg_db=os.getenv("PGDATABASE","candidate_db")
            conn=psycopg2.connect(host=pg_host, port=pg_port, user=pg_user, password=pg_password, dbname=pg_db)
            cur=conn.cursor()
            cur.execute("SELECT jd FROM login WHERE username = %s", (username,))
            row = cur.fetchone()
            cur.close(); conn.close()
            if row and row[0]: jd_text = row[0]
        except Exception as e: return jsonify({"error": f"DB fetch error: {e}"}), 500
    if not jd_text: return jsonify({"error": "No JD text provided or found for user"}), 400
    try:
        from chat_gemini_review import analyze_job_description
        result = analyze_job_description(jd_text, sectors_data)
        parsed = result.get("parsed", {})
        skills = parsed.get("skills", [])
        if username and skills: _persist_jskillset(username, skills)
        response_obj = {
            "seniority": parsed.get("seniority"),
            "job_title": parsed.get("job_title"),
            "sectors": parsed.get("sectors") or ([parsed.get("sector")] if parsed.get("sector") else []),
            "country": parsed.get("country"),
            "summary": result.get("summary"),
            "skills": skills
        }
        return jsonify(response_obj), 200
    except Exception as e:
        logger.warning(f"[Gemini JD Analyze] {e}")
        return jsonify({"error": str(e)}), 500

@app.post("/user/update_skills")
def user_update_skills():
    """
    Update user skills in the jskillset field.
    Syncs frontend skill additions/removals to the database immediately.
    """
    try:
        data = request.get_json(force=True, silent=True) or {}
        username = (data.get("username") or "").strip()
        skills = data.get("skills", [])
        
        if not username:
            return jsonify({"error": "username required"}), 400
        
        # Use existing _persist_jskillset function
        success, message = _persist_jskillset(username, skills)
        
        if success:
            return jsonify({"status": "ok", "message": message}), 200
        else:
            return jsonify({"error": message}), 400
    except Exception as e:
        logger.error(f"[Update Skills] {e}")
        return jsonify({"error": str(e)}), 500

@app.post("/process/scan_and_upload_cvs")
@_rate(_make_flask_limit("upload_cv"))
def process_scan_and_upload_cvs():
    data = request.get_json(force=True, silent=True) or {}
    directory_path = (data.get("directory_path") or "").strip()
    if not directory_path or not os.path.isdir(directory_path): return jsonify({"error": "Valid directory path is required"}), 400
    try:
        files = [f for f in os.listdir(directory_path) if f.lower().endswith('.pdf')]
        if not files: return jsonify({"uploaded_count": 0, "message": "No PDF files found in directory"}), 200
        import psycopg2
        pg_host=os.getenv("PGHOST","localhost"); pg_port=int(os.getenv("PGPORT","5432"))
        pg_user=os.getenv("PGUSER","postgres"); pg_password=os.getenv("PGPASSWORD", "")
        pg_db=os.getenv("PGDATABASE","candidate_db")
        conn=psycopg2.connect(host=pg_host, port=pg_port, user=pg_user, password=pg_password, dbname=pg_db)
        cur=conn.cursor()
        cur.execute("SELECT id, name, linkedinurl FROM process WHERE name IS NOT NULL AND name != ''")
        candidates = cur.fetchall()
        uploaded_count = 0; errors = []
        def normalize(s): return re.sub(r'[^a-z0-9]', '', s.lower())
        candidate_map = {}
        for cid, cname, clink in candidates:
            norm = normalize(cname)
            if len(norm) < 3: continue
            if norm not in candidate_map: candidate_map[norm] = []
            candidate_map[norm].append((cid, clink, cname))
        for fname in files:
            fname_norm = normalize(fname)
            matched_candidate = None
            possible_matches = []
            for norm_name, entries in candidate_map.items():
                if norm_name in fname_norm:
                    for entry in entries: possible_matches.append((len(norm_name), entry))
            if possible_matches:
                possible_matches.sort(key=lambda x: x[0], reverse=True)
                matched_candidate = possible_matches[0][1]
            if matched_candidate:
                cid, clink, cname = matched_candidate
                full_path = os.path.join(directory_path, fname)
                try:
                    with open(full_path, "rb") as f: file_bytes = f.read()
                    if not _is_pdf_bytes(file_bytes):
                        errors.append(f"{fname}: not a valid PDF (magic bytes mismatch)")
                        continue
                    binary_cv = psycopg2.Binary(file_bytes)
                    if cid: cur.execute("UPDATE process SET cv = %s WHERE id = %s", (binary_cv, cid))
                    else: cur.execute("UPDATE process SET cv = %s WHERE linkedinurl = %s", (binary_cv, clink))
                    if cur.rowcount > 0:
                        conn.commit()
                        uploaded_count += 1
                        threading.Thread(target=analyze_cv_background, args=(clink, file_bytes), kwargs={'process_id': cid}).start()
                    else: errors.append(f"DB update failed for {fname} (Candidate: {cname})")
                except Exception as e:
                    conn.rollback(); errors.append(f"Error processing {fname}: {e}")
        cur.close(); conn.close()
        return jsonify({"uploaded_count": uploaded_count, "errors": errors, "message": f"Scanned {len(files)} files, matched and uploaded {uploaded_count}."}), 200
    except Exception as e:
        logger.error(f"[Batch Upload] {e}")
        return jsonify({"error": str(e)}), 500